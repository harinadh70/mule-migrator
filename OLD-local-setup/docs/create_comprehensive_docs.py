#!/usr/bin/env python3
"""
Generate comprehensive documentation for the MuleSoft to Spring Boot Migration Platform.
Creates a professional Word document covering architecture, all 7 pages, API endpoints,
setup instructions, and user guide.
"""
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ======================================================================
# CONSTANTS
# ======================================================================
BRAND_BLUE = RGBColor(0x1E, 0x40, 0xAF)
BRAND_INDIGO = RGBColor(0x63, 0x66, 0xF1)
DARK_BG = RGBColor(0x0F, 0x17, 0x2A)
HEADER_BG = "1E40AF"
ROW_ALT_BG = "F0F4FF"
ROW_WHITE = "FFFFFF"
ACCENT_BG = "EEF2FF"
WARN_BG = "FFF7ED"
CODE_BG = "F1F5F9"
DARK_TEXT = RGBColor(0x1E, 0x29, 0x3B)
GRAY_TEXT = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ======================================================================
# HELPERS
# ======================================================================
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_para(doc, text, bold=False, italic=False, size=10, color=None,
             align=None, before=0, after=6, font="Arial", style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    return p


def add_bullets(doc, items, size=10):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc, items, size=10):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        p.paragraph_format.space_after = Pt(2)


def add_code(doc, code, label=""):
    if label:
        add_para(doc, label, bold=True, size=9, before=6)
    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = DARK_TEXT
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.left_indent = Inches(0.3)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
        p._p.get_or_add_pPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run.font.color.rgb = WHITE
        set_cell_shading(cell, HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r, row_data in enumerate(rows):
        bg = ROW_ALT_BG if r % 2 == 0 else ROW_WHITE
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Arial"
            set_cell_shading(cell, bg)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacer
    return table


def add_info_box(doc, text, box_type="info"):
    colors = {
        "info": ("EEF2FF", BRAND_BLUE),
        "warning": ("FFF7ED", RGBColor(0xF5, 0x9E, 0x0B)),
        "success": ("F0FDF4", RGBColor(0x16, 0xA3, 0x4A)),
        "danger": ("FEF2F2", RGBColor(0xDC, 0x26, 0x26)),
    }
    bg, clr = colors.get(box_type, colors["info"])
    icons = {"info": "INFO", "warning": "WARNING", "success": "TIP", "danger": "CAUTION"}
    p = doc.add_paragraph()
    run = p.add_run(f"  {icons.get(box_type, 'NOTE')}: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run.font.color.rgb = clr
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.name = "Arial"
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.left_indent = Inches(0.2)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)


def add_page_break(doc):
    doc.add_page_break()


# ======================================================================
# DOCUMENT CREATION
# ======================================================================
def create_documentation():
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # -- Styles --
    style = doc.styles["Heading 1"]
    style.font.size = Pt(22)
    style.font.color.rgb = BRAND_BLUE
    style.font.name = "Arial"
    style.font.bold = True
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after = Pt(10)

    style2 = doc.styles["Heading 2"]
    style2.font.size = Pt(16)
    style2.font.color.rgb = BRAND_INDIGO
    style2.font.name = "Arial"
    style2.font.bold = True
    style2.paragraph_format.space_before = Pt(14)
    style2.paragraph_format.space_after = Pt(8)

    style3 = doc.styles["Heading 3"]
    style3.font.size = Pt(13)
    style3.font.color.rgb = DARK_TEXT
    style3.font.name = "Arial"
    style3.font.bold = True
    style3.paragraph_format.space_before = Pt(10)
    style3.paragraph_format.space_after = Pt(6)

    # ──────────────────────────────────────────────────────────────
    # TITLE PAGE
    # ──────────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MuleSoft to Spring Boot")
    run.font.size = Pt(36)
    run.font.color.rgb = BRAND_BLUE
    run.font.name = "Arial"
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Migration Platform")
    run2.font.size = Pt(36)
    run2.font.color.rgb = BRAND_BLUE
    run2.font.name = "Arial"
    run2.bold = True

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Comprehensive Platform Documentation")
    run3.font.size = Pt(18)
    run3.font.color.rgb = BRAND_INDIGO
    run3.font.name = "Arial"

    doc.add_paragraph()
    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Architecture | User Guide | API Reference | Setup Instructions")
    run4.font.size = Pt(11)
    run4.font.color.rgb = GRAY_TEXT
    run4.font.name = "Arial"

    doc.add_paragraph()

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run(f"Version 2.0  |  {datetime.now().strftime('%B %Y')}")
    run5.font.size = Pt(11)
    run5.font.color.rgb = GRAY_TEXT
    run5.font.name = "Arial"

    add_page_break(doc)

    # ──────────────────────────────────────────────────────────────
    # TABLE OF CONTENTS
    # ──────────────────────────────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        ("1.", "Executive Summary", 3),
        ("2.", "Platform Architecture", 4),
        ("3.", "Installation & Setup Guide", 6),
        ("4.", "Pages & Features (All 7 Pages)", 8),
        ("5.", "API Reference (31 Endpoints)", 14),
        ("6.", "Core Migration Engine (7 Modules)", 17),
        ("7.", "Backend Services", 21),
        ("8.", "Frontend Architecture", 23),
        ("9.", "Security", 25),
        ("10.", "Testing Guide", 26),
        ("11.", "Deployment Guide", 27),
        ("12.", "Troubleshooting", 29),
        ("13.", "Appendix", 30),
    ]
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title}")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = BRAND_BLUE
        tab = p.add_run(f"{'.' * (60 - len(title))} {page}")
        tab.font.size = Pt(10)
        tab.font.name = "Arial"
        tab.font.color.rgb = GRAY_TEXT
        p.paragraph_format.space_after = Pt(2)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_heading("1.1 Purpose & Motivation", level=2)
    add_para(doc,
        "The MuleSoft to Spring Boot Migration Platform is an enterprise-grade, AI-powered tool "
        "that automates the conversion of MuleSoft Mule 4 applications into production-ready "
        "Spring Boot projects. It eliminates the manual effort of rewriting integration logic, "
        "reducing migration timelines from weeks to minutes while maintaining code quality and "
        "architectural best practices.")

    add_para(doc,
        "The platform provides a complete migration lifecycle: parsing MuleSoft XML configurations, "
        "mapping connectors to Spring Boot dependencies, converting DataWeave transformations to Java, "
        "generating project structures with OpenAPI/Swagger annotations, validating output with "
        "multi-provider LLM support, and deploying via GitHub integration and Docker builds.")

    doc.add_heading("1.2 Key Capabilities", level=2)
    add_bullets(doc, [
        "Automated XML-to-Spring Boot conversion with 30+ connector support",
        "DataWeave 2.0 to Java stream-based transformation conversion",
        "Multi-provider LLM validation (Anthropic Claude, OpenAI GPT-4, Google Gemini, DeepSeek, Groq, Ollama)",
        "Inline code editing with per-file LLM re-validation",
        "OpenAPI 3.0 specification generation from RAML or MuleSoft XML",
        "GitHub integration with organization support and Git Data API push",
        "JAR/WAR/Docker builds with multi-platform support (5 platforms)",
        "Server-Sent Events (SSE) for real-time build streaming output",
        "Enterprise-grade interactive architecture diagrams with zoom/pan",
        "Cross-page state management via localStorage for seamless workflows",
    ])

    doc.add_heading("1.3 Technology Stack", level=2)
    add_table(doc,
        ["Layer", "Technology", "Version"],
        [
            ["Backend Framework", "Flask (Python)", "3.1.0"],
            ["XML Parsing", "lxml", "5.3.1"],
            ["YAML Processing", "PyYAML", "6.0.2"],
            ["GitHub Integration", "PyGithub", "2.1.0+"],
            ["LLM - Anthropic", "anthropic SDK", "0.39.0+"],
            ["LLM - OpenAI", "openai SDK", "1.50.0+"],
            ["LLM - Google", "google-generativeai", "0.8.0+"],
            ["Production Server", "Gunicorn", "23.0.0"],
            ["Frontend", "Vanilla JS + CSS (no framework)", "ES6+"],
            ["Code Highlighting", "CodeMirror", "5.x CDN"],
            ["Template Engine", "Jinja2", "3.x"],
        ],
        col_widths=[2.0, 2.8, 1.5])

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 2. PLATFORM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. Platform Architecture", level=1)

    doc.add_heading("2.1 Multi-Page Architecture (Flask Blueprints)", level=2)
    add_para(doc,
        "The platform uses Flask Blueprints to organize the application into modular, "
        "self-contained components. Each page is served by its own blueprint with dedicated "
        "routes, templates, JavaScript, and CSS files. The blueprints are registered in "
        "app.py through a central register_blueprints() function.")

    add_table(doc,
        ["Blueprint", "Module", "URL Prefix", "Pages Served"],
        [
            ["main_bp", "blueprints/main.py", "/", "Dashboard"],
            ["migration_bp", "blueprints/migration.py", "/migrate", "Migration Tool"],
            ["swagger_bp", "blueprints/swagger.py", "/swagger", "Swagger Generator"],
            ["github_bp", "blueprints/github_bp.py", "/github", "GitHub Integration"],
            ["build_bp", "blueprints/build.py", "/build", "Build & Test"],
            ["settings_bp", "blueprints/settings_bp.py", "/settings", "Settings"],
            ["(app.py)", "app.py", "/architecture", "Architecture (Protected)"],
        ],
        col_widths=[1.5, 2.2, 1.3, 1.5])

    doc.add_heading("2.2 Directory Structure", level=2)
    add_code(doc, """mulesoft-to-springboot-migrator/
+-- backend/
|   +-- app.py                          # Flask application factory
|   +-- utils.py                        # Shared utilities
|   +-- gunicorn.conf.py               # Production server config
|   +-- requirements.txt               # Python dependencies
|   +-- blueprints/
|   |   +-- __init__.py                # Blueprint registration
|   |   +-- main.py                    # Dashboard & health
|   |   +-- migration.py              # Migration endpoints
|   |   +-- swagger.py                # OpenAPI generation
|   |   +-- github_bp.py             # GitHub integration
|   |   +-- build.py                  # Build & test
|   |   +-- settings_bp.py           # Settings page
|   +-- migrator/                      # Core migration engine
|   |   +-- parser.py                 # MuleSoft XML parser
|   |   +-- flow_converter.py        # Flow-to-Spring converter
|   |   +-- connector_mapper.py      # Connector dependency mapper
|   |   +-- dataweave_converter.py   # DataWeave-to-Java converter
|   |   +-- spring_generator.py      # Spring Boot project generator
|   |   +-- swagger_generator.py     # OpenAPI spec generator
|   |   +-- llm_agent.py             # LLM-assisted conversion
|   |   +-- llm_validator.py         # Multi-provider LLM validation
|   +-- services/
|   |   +-- build_service.py         # Maven/Docker build execution
|   |   +-- github_service.py        # GitHub API wrapper
|   +-- templates/                     # Jinja2 HTML templates
|   |   +-- base.html                # Master template (sidebar, nav)
|   |   +-- dashboard.html           # Dashboard page
|   |   +-- migration.html           # Migration tool
|   |   +-- swagger.html             # Swagger generator
|   |   +-- github.html              # GitHub integration
|   |   +-- build.html               # Build & test
|   |   +-- settings.html            # Settings
|   |   +-- architecture.html        # Architecture diagrams
|   +-- static/
|       +-- css/ (7 files)            # Page-specific stylesheets
|       +-- js/  (7 files)            # Page-specific JavaScript
+-- docs/                              # Generated documentation""")

    doc.add_heading("2.3 Base Template System", level=2)
    add_para(doc,
        "All pages extend base.html which provides a consistent layout with a collapsible "
        "sidebar navigation, toast notification system, modal dialogs, loading overlay with "
        "progress stepper, and status bar. Pages inject content through Jinja2 template blocks.")

    add_table(doc,
        ["Block Name", "Purpose", "Example Usage"],
        [
            ["{% block title %}", "Page title in browser tab", "Migration Tool"],
            ["{% block extra_css %}", "Page-specific CSS includes", '<link href="migration.css">'],
            ["{% block page_header %}", "Header area above content", "Page title + breadcrumbs"],
            ["{% block content %}", "Main page body", "Panels, forms, editors"],
            ["{% block extra_js %}", "Page-specific JS includes", '<script src="migration.js">'],
            ["{% block loading_content %}", "Progress/loading display", "Step indicators"],
        ],
        col_widths=[1.8, 2.0, 2.5])

    doc.add_heading("2.4 Cross-Page State Management", level=2)
    add_para(doc,
        "The MigrationStore object in base.js manages cross-page state via localStorage. "
        "When a migration completes on the Migration page, the result is stored and made "
        "available to Swagger, GitHub, Build, and Dashboard pages.")

    add_code(doc, """// MigrationStore (base.js) - Cross-page state
const MigrationStore = {
    save(result)  { localStorage.setItem('msb_migration_result', JSON.stringify(result)); },
    load()        { return JSON.parse(localStorage.getItem('msb_migration_result') || 'null'); },
    getFiles()    { const r = this.load(); return r ? r.files : null; },
    getSummary()  { const r = this.load(); return r ? r.summary : null; },
};""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 3. INSTALLATION & SETUP
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. Installation & Setup Guide", level=1)

    doc.add_heading("3.1 Prerequisites", level=2)
    add_table(doc,
        ["Requirement", "Version", "Purpose", "Required?"],
        [
            ["Python", "3.8+", "Backend runtime", "Yes"],
            ["pip", "Latest", "Package manager", "Yes"],
            ["Java JDK", "17+", "Building JAR/WAR", "Optional"],
            ["Docker", "20.10+", "Building Docker images", "Optional"],
            ["Git", "2.x", "Version control", "Optional"],
        ],
        col_widths=[1.5, 1.0, 2.5, 1.2])

    add_info_box(doc,
        "Java, Maven, and Docker are only needed for the Build & Test page. The migration, "
        "Swagger generation, and GitHub features work without them. The platform shows helpful "
        "install guidance with download links when prerequisites are missing.",
        "info")

    doc.add_heading("3.2 Installation Steps", level=2)
    add_numbered(doc, [
        "Clone the repository:  git clone <repo-url> && cd mulesoft-to-springboot-migrator",
        "Create a virtual environment:  python3 -m venv venv && source venv/bin/activate",
        "Install dependencies:  pip install -r backend/requirements.txt",
        "Set environment variables (optional):",
    ])
    add_code(doc, """# Optional environment variables
export FLASK_ENV=development          # Enable debug mode
export PORT=5001                      # Server port (default: 5000)
export CORS_ORIGINS=http://localhost:3000  # Allowed CORS origins
export ARCH_USERNAME=admin            # Architecture page username
export ARCH_PASSWORD=admin            # Architecture page password""")

    add_numbered(doc, [
        "Start the server:  cd backend && python3 app.py",
        "Open browser:  http://localhost:5001",
    ])

    doc.add_heading("3.3 LLM Provider Configuration", level=2)
    add_para(doc,
        "LLM validation is optional but recommended. Configure providers on the Settings page "
        "or pass API keys directly on the Migration page. All keys are stored client-side in "
        "localStorage and sent per-request.")

    add_table(doc,
        ["Provider", "Models Available", "API Key Source", "Free Tier?"],
        [
            ["Anthropic", "Claude Sonnet 4, Claude 3.5 Sonnet, Claude 3 Opus, Haiku", "console.anthropic.com", "No"],
            ["OpenAI", "GPT-4o, GPT-4 Turbo, GPT-4o Mini, o3-mini", "platform.openai.com", "No"],
            ["Google", "Gemini 2.5 Pro, 2.0 Flash, 1.5 Pro", "aistudio.google.com", "Yes (limited)"],
            ["DeepSeek", "DeepSeek Chat, Coder, Reasoner", "platform.deepseek.com", "Low cost"],
            ["Groq", "Llama 3.3 70B, Mixtral 8x7B", "console.groq.com", "Yes (rate limited)"],
            ["Ollama", "codellama, llama3, mistral, deepseek-coder-v2", "Local installation", "Yes (local)"],
        ],
        col_widths=[1.2, 2.5, 1.8, 1.0])

    doc.add_heading("3.4 Production Deployment", level=2)
    add_code(doc, """# Using Gunicorn (recommended for production)
cd backend
gunicorn app:app --config gunicorn.conf.py

# Gunicorn configuration (gunicorn.conf.py):
# - Workers: CPU cores x 2 + 1
# - Worker class: gthread (for I/O-bound LLM calls)
# - Threads: 4 per worker
# - Timeout: 120s (accommodates LLM API latency)
# - Access log: stdout with Apache combined format""")

    add_code(doc, """# Docker deployment
docker build -t mulesoft-migrator .
docker run -p 5001:5001 -e PORT=5001 mulesoft-migrator""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 4. PAGES & FEATURES
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. Pages & Features (All 7 Pages)", level=1)

    # ── 4.1 Dashboard ──
    doc.add_heading("4.1 Dashboard (/)", level=2)
    add_para(doc,
        "The Dashboard is the landing page providing an overview of the platform's state "
        "and quick access to all major features. It displays migration statistics from the "
        "most recent run and animated stat counters.")

    add_para(doc, "Features:", bold=True, size=10)
    add_bullets(doc, [
        "Quick action cards: Start Migration, Generate Swagger, Push to GitHub, Build Project, Settings",
        "Live statistics: Flows Converted, Files Generated, Connectors Mapped, Code Quality Score",
        "Animated counter transitions using requestAnimationFrame with cubic easing",
        "Last migration timestamp display",
        "Intersection Observer for scroll-triggered animations",
        "Ripple effect on card click interactions",
    ])

    # ── 4.2 Migration ──
    doc.add_heading("4.2 Migration (/migrate)", level=2)
    add_para(doc,
        "The Migration page is the core feature of the platform. It accepts MuleSoft XML "
        "configurations and produces a complete Spring Boot project with controllers, services, "
        "repositories, configuration files, and tests.")

    add_para(doc, "Input Panel (4 Tabs):", bold=True, size=10)
    add_bullets(doc, [
        "MuleSoft XML: Paste XML directly, upload files via drag-and-drop, or load sample XML",
        "DataWeave: Manage multiple DW 2.0 scripts with add/remove/editor support",
        "Settings: Project name, Group ID, artifact ID, Java version (11/17/21), packaging",
        "LLM Settings: Provider selection, model picker, API key, temperature settings",
    ])

    add_para(doc, "Output Panel (3 Tabs):", bold=True, size=10)
    add_bullets(doc, [
        "Generated Code: File tree with syntax-highlighted CodeMirror editor, inline editing toggle",
        "Code Review: LLM validation results with severity levels, issue descriptions, and suggestions",
        "Summary: Migration statistics, connector mapping results, warnings, and next-step action links",
    ])

    add_para(doc, "Migration Pipeline (6 Steps):", bold=True, size=10)
    add_numbered(doc, [
        "Parse: MuleSoft XML is parsed into structured data (flows, configs, connectors)",
        "Map Connectors: MuleSoft connectors are mapped to Spring Boot Maven dependencies",
        "Convert DataWeave: DataWeave 2.0 scripts are converted to Java stream operations",
        "Generate: Complete Spring Boot project is generated with all files",
        "Review: LLM validates the generated code for correctness and best practices",
        "Complete: Results displayed with download, edit, validate, and cross-page action options",
    ])

    add_para(doc, "Inline Code Editing:", bold=True, size=10)
    add_bullets(doc, [
        "Toggle Edit mode per file to switch from read-only to editable CodeMirror",
        "Save Changes updates the in-memory file and localStorage",
        "Validate This File sends a single file for LLM re-validation",
        "Validate All Changes sends all modified files for bulk re-validation",
        "Modified files are tracked with visual indicators in the file tree",
    ])

    add_para(doc, "Cross-Page Actions (in Summary tab):", bold=True, size=10)
    add_bullets(doc, [
        "Generate Swagger: Navigate to /swagger with pre-loaded migration data",
        "Push to GitHub: Navigate to /github with files ready to push",
        "Build JAR / Run Tests: Navigate to /build to compile and test the generated project",
    ])

    # ── 4.3 Swagger ──
    doc.add_heading("4.3 Swagger / OpenAPI Generator (/swagger)", level=2)
    add_para(doc,
        "The Swagger page generates OpenAPI 3.0 specifications from multiple sources. "
        "It can auto-generate from migration results or accept manual RAML/XML input.")

    add_para(doc, "Input Sources:", bold=True, size=10)
    add_bullets(doc, [
        "From Migration: Auto-loads parsed data from the most recent migration (MigrationStore)",
        "From MuleSoft XML: Paste raw MuleSoft XML for direct OpenAPI generation",
        "From RAML: Paste RAML content for conversion to OpenAPI 3.0",
    ])

    add_para(doc, "Output Features:", bold=True, size=10)
    add_bullets(doc, [
        "Live preview in CodeMirror editor with YAML or JSON formatting",
        "Format toggle between YAML and JSON output",
        "Download as openapi.yaml or openapi.json",
        "Auto-generated paths, schemas, request/response models, and security schemes",
    ])

    # ── 4.4 GitHub ──
    doc.add_heading("4.4 GitHub Integration (/github)", level=2)
    add_para(doc,
        "The GitHub page provides full repository management including authentication, "
        "repo creation, branch management, and file push capabilities using the GitHub "
        "Git Data API (no local git required).")

    add_para(doc, "Workflow:", bold=True, size=10)
    add_numbered(doc, [
        "Connect: Enter GitHub Personal Access Token (PAT) to authenticate",
        "Select Organization: Choose personal account or an organization",
        "Browse/Create Repos: View existing repos or create new ones (public or private)",
        "Select Branch: Choose existing branch or create a new one",
        "Push Files: Push all generated migration files with a custom commit message",
    ])

    add_info_box(doc,
        "GitHub tokens are stored only in localStorage and sent per-request via the request body. "
        "Tokens are never stored server-side. Required PAT scopes: repo (full control).",
        "info")

    add_para(doc, "Features:", bold=True, size=10)
    add_bullets(doc, [
        "Organization support: Switch between personal repos and org repos",
        "Repository creation with name, description, and visibility (public/private)",
        "Branch creation from any existing branch",
        "File push via Git Data API (creates tree + commit + updates ref)",
        "Connection status indicator with user avatar and username display",
    ])

    # ── 4.5 Build & Test ──
    doc.add_heading("4.5 Build & Test (/build)", level=2)
    add_para(doc,
        "The Build & Test page compiles the generated Spring Boot project into deployable "
        "artifacts. It supports JAR, WAR, and Docker builds with live streaming output "
        "via Server-Sent Events (SSE).")

    add_para(doc, "Prerequisites Check:", bold=True, size=10)
    add_bullets(doc, [
        "Automatically checks for Java, Maven, and Docker availability on page load",
        "Shows install guidance with platform-specific commands and download links when tools are missing",
        "Disables build buttons when required prerequisites are unavailable",
        "Re-check button to refresh prerequisite status after installation",
    ])

    add_para(doc, "Build Types:", bold=True, size=10)
    add_table(doc,
        ["Build Type", "Command", "Output", "Requires"],
        [
            ["JAR", "./mvnw clean package -DskipTests", "app.jar", "Java 17+"],
            ["WAR", "./mvnw clean package (war packaging)", "app.war", "Java 17+"],
            ["Docker", "docker build --platform=<platform>", "Docker image", "Java 17+ & Docker"],
            ["Test", "./mvnw test", "Test results", "Java 17+"],
        ],
        col_widths=[1.0, 2.5, 1.5, 1.3])

    add_para(doc, "Docker Multi-Platform Support:", bold=True, size=10)
    add_table(doc,
        ["Platform", "Architecture", "Base Image", "Use Case"],
        [
            ["Linux (x86_64)", "linux/amd64", "eclipse-temurin:17-jdk-alpine", "Standard servers"],
            ["Ubuntu", "linux/amd64", "eclipse-temurin:17-jdk-focal", "Ubuntu-based deployments"],
            ["Red Hat (UBI9)", "linux/amd64", "eclipse-temurin:17-jdk-ubi9-minimal", "Enterprise/OpenShift"],
            ["macOS (Apple Silicon)", "linux/arm64", "eclipse-temurin:17-jdk-alpine", "M1/M2/M3 Macs"],
            ["Windows Server", "linux/amd64", "eclipse-temurin:17-jdk-nanoserver", "Windows containers"],
        ],
        col_widths=[1.5, 1.2, 2.5, 1.3])

    add_para(doc, "Live Build Streaming:", bold=True, size=10)
    add_bullets(doc, [
        "Terminal-style output panel with dark background and monospace font",
        "Real-time SSE streaming of build output (stdout + stderr)",
        "Auto-scrolling with color-coded output (errors in red, warnings in yellow)",
        "Build status tracking: pending, running, success, failed",
        "Artifact download button appears on successful build",
        "Build cleanup to remove temporary files",
    ])

    # ── 4.6 Settings ──
    doc.add_heading("4.6 Settings (/settings)", level=2)
    add_para(doc,
        "The Settings page provides configuration management for all platform features. "
        "All settings are stored client-side in localStorage with the msb_ prefix.")

    add_para(doc, "Settings Sections:", bold=True, size=10)
    add_bullets(doc, [
        "LLM Provider: Select provider (6 options), choose model, enter API key, set base URL for self-hosted",
        "GitHub: Personal Access Token management with connection test",
        "Project Defaults: Project name, group ID, artifact ID, Java version, packaging type",
        "Preferences: UI theme selection, auto-save toggle, notification settings",
    ])

    # ── 4.7 Architecture ──
    doc.add_heading("4.7 Architecture (/architecture)", level=2)
    add_para(doc,
        "The Architecture page provides enterprise-grade interactive SVG diagrams of the "
        "platform's internal architecture. It features four switchable diagram views with "
        "zoom/pan controls and animated data flow connections.")

    add_para(doc, "Diagram Views:", bold=True, size=10)
    add_bullets(doc, [
        "Architecture Layers: Frontend, backend, services, and data tiers with component boxes",
        "Migration Pipeline: Step-by-step visualization of the XML-to-Spring Boot transformation",
        "Technology Stack: All frameworks, libraries, and tools displayed in categorized groups",
        "Data Flow: Request routing through blueprints, services, and core modules with animated connections",
    ])

    add_para(doc, "Interactive Controls:", bold=True, size=10)
    add_bullets(doc, [
        "Zoom in/out with mouse wheel or control buttons",
        "Pan by click-and-drag on the diagram canvas",
        "Fit-to-screen to auto-resize the diagram to the viewport",
        "Animated flowing dots along connection paths showing data direction",
    ])

    add_info_box(doc,
        "The Architecture page is protected with HTTP Basic Authentication. Default credentials: "
        "admin-username / admin-password. Set ARCH_USERNAME and ARCH_PASSWORD environment variables "
        "to change these in production.",
        "warning")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 5. API REFERENCE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. API Reference (31 Endpoints)", level=1)
    add_para(doc,
        "All API endpoints use the /api/ prefix convention. Page routes return HTML; "
        "API routes return JSON. All API endpoints accept and return application/json "
        "unless otherwise noted.")

    doc.add_heading("5.1 Main Endpoints", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/", "Render dashboard page"],
            ["GET", "/api/health", "Health check: {status: 'ok', env: '...'}"],
        ],
        col_widths=[0.8, 2.5, 3.0])

    doc.add_heading("5.2 Migration Endpoints", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/migrate", "Render migration page"],
            ["GET", "/api/llm/providers", "List available LLM providers and models"],
            ["POST", "/api/migrate", "Execute migration (MuleSoft XML to Spring Boot)"],
            ["POST", "/api/validate", "LLM code validation (separate from migration)"],
            ["POST", "/api/migrate/download", "Download generated project as ZIP archive"],
            ["POST", "/api/convert/dataweave", "Convert a single DataWeave script to Java"],
        ],
        col_widths=[0.8, 2.5, 3.0])

    add_para(doc, "POST /api/migrate - Request Body:", bold=True, size=9)
    add_code(doc, """{
    "xmlContent": "<mule>...</mule>",       // MuleSoft XML (required)
    "projectName": "my-app",                // Project name (default: migrated-app)
    "groupId": "com.example",               // Maven group ID
    "javaVersion": "17",                    // Java version: 11, 17, or 21
    "dwScripts": {"transform.dwl": "..."},  // DataWeave scripts (optional)
    "llmConfig": {                          // LLM validation config (optional)
        "provider": "anthropic",
        "model": "claude-sonnet-4-20250514",
        "apiKey": "sk-..."
    }
}""")

    doc.add_heading("5.3 Swagger Endpoints", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/swagger", "Render Swagger generator page"],
            ["POST", "/api/swagger/from-xml", "Generate OpenAPI from MuleSoft XML"],
            ["POST", "/api/swagger/from-raml", "Generate OpenAPI from RAML content"],
            ["POST", "/api/swagger/from-migration", "Generate OpenAPI from migration result data"],
            ["POST", "/api/swagger/download", "Download spec as YAML or JSON file"],
        ],
        col_widths=[0.8, 2.8, 2.7])

    doc.add_heading("5.4 GitHub Endpoints", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/github", "Render GitHub integration page"],
            ["POST", "/api/github/connect", "Authenticate GitHub token, return user info"],
            ["GET", "/api/github/orgs", "List user organizations"],
            ["GET", "/api/github/repos", "List repositories (optional ?org=name filter)"],
            ["POST", "/api/github/repos/create", "Create new repository (personal or org)"],
            ["GET", "/api/github/repos/<o>/<r>/branches", "List branches in a repository"],
            ["POST", "/api/github/repos/<o>/<r>/branches/create", "Create a new branch"],
            ["POST", "/api/github/push", "Push files to repository via Git Data API"],
        ],
        col_widths=[0.8, 3.2, 2.3])

    doc.add_heading("5.5 Build & Test Endpoints", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/build", "Render Build & Test page"],
            ["POST", "/api/build/check-prereqs", "Check Java, Maven, Docker availability"],
            ["POST", "/api/build/jar", "Start JAR build (returns buildId)"],
            ["POST", "/api/build/war", "Start WAR build (returns buildId)"],
            ["POST", "/api/build/docker", "Start Docker build with platform selection"],
            ["POST", "/api/test/start", "Start test execution (returns testId)"],
            ["GET", "/api/build/<id>/stream", "SSE stream of build/test output"],
            ["GET", "/api/build/<id>/artifact", "Download built JAR/WAR artifact"],
            ["POST", "/api/build/<id>/cleanup", "Clean up build temporary files"],
        ],
        col_widths=[0.8, 2.8, 2.7])

    doc.add_heading("5.6 Settings Endpoint", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/settings", "Render settings page (all config is client-side)"],
        ],
        col_widths=[0.8, 2.5, 3.0])

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 6. CORE MIGRATION ENGINE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. Core Migration Engine (7 Modules)", level=1)
    add_para(doc,
        "The migration engine in backend/migrator/ contains seven specialized modules "
        "that work together to transform MuleSoft XML into Spring Boot Java projects.")

    doc.add_heading("6.1 MuleSoft XML Parser (parser.py)", level=2)
    add_para(doc,
        "The parser is the first module in the pipeline. It processes MuleSoft Mule 4 "
        "XML configurations and extracts flows, sub-flows, global configurations, "
        "connectors, and all related metadata into a structured dictionary.")

    add_para(doc, "Supported MuleSoft Components (20+ namespaces):", bold=True, size=10)
    add_bullets(doc, [
        "HTTP: Listeners, request configs, response builders",
        "Schedulers: Cron expressions, fixed-frequency/fixed-delay",
        "Messaging: JMS, AMQP/RabbitMQ, Kafka, VM, Anypoint MQ",
        "File I/O: File, SFTP, FTP, Email (IMAP/POP3/SMTP)",
        "Database: Select, insert, update, delete, stored procedures, bulk operations",
        "Cloud: Salesforce, AWS S3, AWS SQS, AWS SNS",
        "NoSQL: MongoDB, Redis, Elasticsearch",
        "Web Services: SOAP/WS consumer, WS security",
        "Processing: Choice routers, scatter-gather, for-each, try-catch, until-successful",
        "Batch: Batch jobs, batch steps, batch aggregators",
        "Security: APIkit configs, secure properties, TLS contexts, OAuth2",
        "Caching: Object store, caching strategies",
    ])

    add_para(doc, "Output Structure:", bold=True, size=10)
    add_code(doc, """{
    "flows": [...],              // HTTP-triggered flows
    "sub_flows": [...],          // Reusable sub-flows
    "global_configs": [...],     // Global configurations
    "connectors": set(),         // Unique connector names found
    "global_properties": {},     // Global property placeholders
    "error_handlers": [...],     // Error handling strategies
    "batch_jobs": [...],         // Batch processing jobs
    "apikit_configs": [...],     // APIkit router configs
    "secure_properties": [...],  // Encrypted property configs
    "tls_contexts": [...],       // TLS/SSL configurations
    "caching_strategies": [...], // Cache configurations
    "warnings": [...]            // Parse warnings
}""")

    doc.add_heading("6.2 Flow Converter (flow_converter.py)", level=2)
    add_para(doc,
        "Converts parsed MuleSoft flows into Spring Boot Java source files. Each HTTP "
        "listener flow becomes a REST controller method; schedulers become @Scheduled "
        "methods; message listeners become JMS/AMQP/Kafka listener methods.")

    add_para(doc, "Conversion Mapping:", bold=True, size=10)
    add_table(doc,
        ["MuleSoft Source", "Spring Boot Output", "Annotations"],
        [
            ["HTTP Listener", "REST Controller method", "@GetMapping, @PostMapping, etc."],
            ["Scheduler (cron)", "@Scheduled method", "@Scheduled(cron = \"...\")"],
            ["Scheduler (fixed)", "@Scheduled method", "@Scheduled(fixedRate = ...)"],
            ["JMS Listener", "JMS consumer method", "@JmsListener(destination = \"...\")"],
            ["AMQP Listener", "RabbitMQ listener", "@RabbitListener(queues = \"...\")"],
            ["Kafka Listener", "Kafka consumer", "@KafkaListener(topics = \"...\")"],
            ["VM Listener", "Spring Event listener", "@EventListener"],
            ["File Listener", "Spring Integration", "@InboundChannelAdapter"],
        ],
        col_widths=[1.8, 2.2, 2.3])

    add_para(doc, "Supported Processors (30+):", bold=True, size=10)
    add_bullets(doc, [
        "Core: logger, set-payload, set-variable, remove-variable, choice, scatter-gather, for-each, try-catch",
        "HTTP: request (outbound HTTP calls via RestTemplate/WebClient)",
        "Database: select, insert, update, delete, bulk operations, stored procedures",
        "Messaging: publish (JMS send, Kafka produce, AMQP publish)",
        "File: read, write, list, delete, move, copy (local and SFTP/FTP)",
        "Validation: is-not-null, is-email, matches-regex, custom validators",
        "Transform: object-to-json, json-to-object, XML transforms",
    ])

    doc.add_heading("6.3 Connector Mapper (connector_mapper.py)", level=2)
    add_para(doc,
        "Maps discovered MuleSoft connectors to their Spring Boot equivalents, providing "
        "Maven dependency coordinates and Spring configuration properties.")

    add_table(doc,
        ["MuleSoft Connector", "Spring Boot Dependency", "Config Properties"],
        [
            ["http", "spring-boot-starter-web", "server.port, server.servlet.context-path"],
            ["db", "spring-boot-starter-data-jpa + driver", "spring.datasource.url, username, password"],
            ["jms", "spring-boot-starter-activemq", "spring.activemq.broker-url"],
            ["amqp", "spring-boot-starter-amqp", "spring.rabbitmq.host, port, username"],
            ["kafka", "spring-kafka", "spring.kafka.bootstrap-servers"],
            ["file/sftp", "spring-integration-file/sftp", "file.upload-dir, sftp.host"],
            ["email", "spring-boot-starter-mail", "spring.mail.host, port, username"],
            ["salesforce", "Custom REST client", "salesforce.instance-url, client-id"],
            ["s3", "AWS SDK v2", "aws.s3.bucket, region, access-key"],
            ["mongodb", "spring-boot-starter-data-mongodb", "spring.data.mongodb.uri"],
            ["redis", "spring-boot-starter-data-redis", "spring.redis.host, port"],
        ],
        col_widths=[1.5, 2.5, 2.3])

    doc.add_heading("6.4 DataWeave Converter (dataweave_converter.py)", level=2)
    add_para(doc,
        "Converts MuleSoft DataWeave 2.0 transformation scripts into equivalent Java code "
        "using Java Streams API, Jackson for JSON processing, and standard library utilities.")

    add_para(doc, "Supported DataWeave Features:", bold=True, size=10)
    add_bullets(doc, [
        "Collection operators: map, filter, reduce, pluck, groupBy, orderBy, flatMap, flatten",
        "String functions: upper, lower, trim, capitalize, replace, split, join, contains, startsWith",
        "Array functions: sizeOf, indexOf, min, max, avg, sum, first, last, distinctBy",
        "Object functions: keys, values, merge (++), remove (--), mapObject",
        "Type coercion: as String, as Number, as Boolean, as Date",
        "Null handling: default operator, if-else, when/otherwise",
        "Date/Time: now(), period arithmetic, date formatting",
        "Lambda expressions and variable declarations",
        "MEL compatibility: #[payload], flowVars, sessionVars",
    ])

    doc.add_heading("6.5 Spring Boot Generator (spring_generator.py)", level=2)
    add_para(doc,
        "Generates a complete, production-ready Spring Boot project. The generated code includes "
        "OpenAPI/Swagger annotations (@Tag, @Operation, @ApiResponse), Spring Boot Actuator, "
        "structured logging (logback-spring.xml), database drivers, comprehensive @WebMvcTest "
        "tests, and multi-profile application properties.")

    add_para(doc, "Generated Project Structure:", bold=True, size=10)
    add_code(doc, """generated-project/
+-- pom.xml                             # Maven POM with all dependencies
+-- src/main/java/com/example/project/
|   +-- Application.java               # @SpringBootApplication + @OpenAPIDefinition
|   +-- controller/
|   |   +-- *Controller.java           # @RestController + @Tag + @Operation + @ApiResponse
|   +-- service/
|   |   +-- *Service.java             # Business logic from MuleSoft flows
|   +-- repository/
|   |   +-- *Repository.java          # JPA/Spring Data repositories
|   +-- model/
|   |   +-- *Model.java              # Data models / DTOs
|   +-- config/
|   |   +-- AppConfig.java           # Bean configurations
|   +-- exception/
|   |   +-- GlobalExceptionHandler.java  # @ControllerAdvice error handler
|   +-- util/
|       +-- DataWeaveHelper.java      # Converted DW transformations
+-- src/main/resources/
|   +-- application.properties         # Spring Boot configuration
|   +-- application.yml                # YAML config (multi-profile)
|   +-- logback-spring.xml            # Structured logging configuration
+-- src/test/java/com/example/project/
    +-- ApplicationTests.java          # @SpringBootTest context load test
    +-- controller/
        +-- *ControllerTest.java      # @WebMvcTest unit tests""")

    add_para(doc, "Key Dependencies Added:", bold=True, size=10)
    add_bullets(doc, [
        "springdoc-openapi-starter-webmvc-ui: Auto-generates Swagger UI at /swagger-ui.html",
        "spring-boot-starter-actuator: Health, metrics, info endpoints",
        "Database drivers: MySQL, PostgreSQL, H2 (based on connector detection)",
        "spring-boot-starter-test: JUnit 5, Mockito, Spring Test",
    ])

    doc.add_heading("6.6 Swagger Generator (swagger_generator.py)", level=2)
    add_para(doc,
        "Generates OpenAPI 3.0.3 specifications from parsed MuleSoft data or RAML content. "
        "Extracts paths from HTTP listener flows, generates request/response schemas, and "
        "adds appropriate security schemes.")

    add_para(doc, "Features:", bold=True, size=10)
    add_bullets(doc, [
        "Accepts multiple data shapes: flows, endpoints, summary wrappers",
        "Auto-generates path parameters from URI templates ({id}, {name})",
        "Extracts query parameters from set-variable processors",
        "Creates request/response schemas based on flow processors",
        "Adds security schemes based on detected OAuth2/Basic Auth configs",
        "Output: OpenAPI 3.0.3 compliant YAML or JSON",
    ])

    doc.add_heading("6.7 LLM Agent (llm_agent.py)", level=2)
    add_para(doc,
        "Handles unknown or complex MuleSoft elements that cannot be automatically converted. "
        "Uses a configurable LLM provider to generate Java code for unrecognized processors, "
        "DataWeave scripts, or connector configurations.")

    add_para(doc, "Conversion Strategy (Triple Fallback):", bold=True, size=10)
    add_numbered(doc, [
        "LLM-generated Java code (if LLM is enabled and configured)",
        "TODO comment placeholder (if LLM is unavailable or fails)",
        "Warning in migration summary (always recorded for visibility)",
    ])

    doc.add_heading("6.8 LLM Validator (llm_validator.py)", level=2)
    add_para(doc,
        "Validates generated Spring Boot code quality using any of six supported LLM providers. "
        "Performs comprehensive checks including compilation correctness, Spring Boot best practices, "
        "security vulnerabilities, missing imports, and performance suggestions.")

    add_para(doc, "Validation Output:", bold=True, size=10)
    add_code(doc, """{
    "overallScore": 85,          // 0-100 quality score
    "summary": "...",            // Human-readable summary
    "issues": [...],             // Detected problems
    "improvements": [...],       // Suggested improvements
    "missingItems": [...],       // Missing imports/annotations
    "securityIssues": [...],     // Security vulnerabilities
    "bestPractices": [...]       // Best practice recommendations
}""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 7. BACKEND SERVICES
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("7. Backend Services", level=1)

    doc.add_heading("7.1 GitHub Service (github_service.py)", level=2)
    add_para(doc,
        "Wraps the PyGithub library to provide repository management capabilities. "
        "The service is stateless - the GitHub token is passed per-request from the frontend.")

    add_table(doc,
        ["Method", "Parameters", "Returns"],
        [
            ["get_user_info()", "None", "User login, name, avatar, email"],
            ["list_orgs()", "None", "List of {login, avatar_url, description}"],
            ["list_repos(org, sort)", "org name, sort field", "List of repo objects (50 max)"],
            ["list_branches(owner, repo)", "owner, repo name", "List of {name, sha, protected}"],
            ["create_repo(name, desc, private, org)", "Repo details", "New repo object"],
            ["create_branch(owner, repo, name, from)", "Branch details", "Branch ref object"],
            ["push_files(owner, repo, files, branch, msg)", "File dict + commit info", "Commit SHA"],
        ],
        col_widths=[2.5, 2.0, 2.0])

    doc.add_heading("7.2 Build Service (build_service.py)", level=2)
    add_para(doc,
        "Manages the build lifecycle: writes generated files to a temporary directory, "
        "generates Maven Wrapper (mvnw), executes builds asynchronously with threading, "
        "and streams output via SSE.")

    add_para(doc, "Build Flow:", bold=True, size=10)
    add_numbered(doc, [
        "Write all generated files to a temp directory with correct paths",
        "Generate Maven Wrapper (mvnw + .mvn/wrapper/maven-wrapper.properties)",
        "For WAR: modify pom.xml packaging and add Tomcat provided scope",
        "For Docker: generate platform-specific Dockerfile",
        "Execute build command in subprocess with stdout/stderr capture",
        "Stream output lines via SSE to the frontend in real-time",
        "Track build status (pending, running, success, failed) in memory",
        "Provide artifact download path on success",
    ])

    add_para(doc, "SSE Streaming Pattern:", bold=True, size=10)
    add_code(doc, """def generate():
    process = subprocess.Popen(
        cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
        text=True, bufsize=1
    )
    for line in iter(process.stdout.readline, ""):
        yield f"data: {json.dumps({'line': line.rstrip()})}\\n\\n"
    process.wait()
    status = 'success' if process.returncode == 0 else 'failed'
    yield f"data: {json.dumps({'status': status})}\\n\\n"

return Response(generate(), mimetype="text/event-stream",
                headers={"X-Accel-Buffering": "no"})""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 8. FRONTEND ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("8. Frontend Architecture", level=1)

    doc.add_heading("8.1 Base Layer (base.js / base.css)", level=2)
    add_para(doc,
        "The base layer provides shared functionality across all pages: toast notifications, "
        "modal dialogs, sidebar navigation, MigrationStore for cross-page state, keyboard "
        "shortcuts, and HTML escaping utilities.")

    add_para(doc, "Design System (CSS Variables):", bold=True, size=10)
    add_code(doc, """:root {
    --bg-primary: #0f172a;     --bg-secondary: #1e293b;
    --bg-tertiary: #334155;    --bg-card: #1e293b;
    --border: #334155;         --border-hover: #475569;
    --text-primary: #f1f5f9;   --text-secondary: #94a3b8;
    --accent: #6366f1;         --accent-hover: #818cf8;
    --success: #22c55e;        --warning: #f59e0b;
    --error: #ef4444;          --info: #38bdf8;
}""")

    add_para(doc, "Shared Components:", bold=True, size=10)
    add_table(doc,
        ["Component", "Function", "Description"],
        [
            ["Toast System", "showToast(msg, type, duration)", "Auto-dismiss notifications with progress bar"],
            ["Modal System", "openModal(title, body, footer)", "Overlay dialogs with close-on-click"],
            ["MigrationStore", "save/load/getFiles/getSummary", "Cross-page localStorage state management"],
            ["Sidebar", "toggleSidebar()", "Collapsible navigation with active page highlighting"],
            ["Utilities", "escapeHtml(), showLoading()", "HTML escaping, loading overlays"],
        ],
        col_widths=[1.3, 2.5, 2.5])

    doc.add_heading("8.2 Page-Specific Files", level=2)
    add_table(doc,
        ["Page", "JavaScript", "CSS", "Key Functionality"],
        [
            ["Dashboard", "dashboard.js (6.6 KB)", "dashboard.css (5.9 KB)", "Stat animation, ripple effects"],
            ["Migration", "migration.js (54 KB)", "migration.css (14.1 KB)", "Full migration workflow, inline editing"],
            ["Swagger", "swagger.js (16.1 KB)", "swagger.css (9.3 KB)", "Spec generation, CodeMirror editor"],
            ["GitHub", "github.js (24.2 KB)", "github.css (10.3 KB)", "Auth, repos, branches, push"],
            ["Build", "build.js (24.6 KB)", "build.css (14.3 KB)", "SSE streaming, terminal output"],
            ["Settings", "settings.js (16.5 KB)", "settings.css (12.4 KB)", "Form persistence, validation"],
        ],
        col_widths=[1.2, 1.8, 1.8, 1.8])

    doc.add_heading("8.3 Professional UI Features", level=2)
    add_bullets(doc, [
        "Micro-animations: Button hover scale, card float effects, gradient text animations",
        "Skeleton loading: Placeholder animations while content loads",
        "Ripple effects: Material-style click feedback on interactive elements",
        "Smooth transitions: All state changes use CSS transitions (0.2-0.3s ease)",
        "Glassmorphism: Translucent backgrounds with backdrop-filter blur on modals",
        "Responsive design: Mobile-first breakpoints with collapsible sidebar",
        "Dark theme: Full dark mode with carefully chosen contrast ratios",
        "CodeMirror integration: Syntax highlighting for Java, XML, YAML, JSON, Properties",
    ])

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 9. SECURITY
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("9. Security", level=1)

    doc.add_heading("9.1 API Key & Token Management", level=2)
    add_info_box(doc,
        "All sensitive credentials (GitHub tokens, LLM API keys) are stored ONLY in the "
        "browser's localStorage. They are never persisted on the server. Keys are sent "
        "per-request in the request body and immediately discarded after use.",
        "info")

    add_bullets(doc, [
        "GitHub PATs stored in localStorage with msb_ prefix, sent per-request",
        "LLM API keys stored in localStorage, sent only when validation is triggered",
        "No server-side session storage or database of credentials",
        "CORS configured with allowlist of trusted origins (CORS_ORIGINS env var)",
    ])

    doc.add_heading("9.2 Architecture Page Protection", level=2)
    add_bullets(doc, [
        "HTTP Basic Authentication on /architecture endpoint",
        "Credentials set via ARCH_USERNAME and ARCH_PASSWORD environment variables",
        "Returns 401 with WWW-Authenticate header on failed authentication",
        "Default credentials should be changed in production",
    ])

    doc.add_heading("9.3 Input Validation & Safety", level=2)
    add_bullets(doc, [
        "XML content validated before parsing (lxml with secure defaults)",
        "File upload size limit: 50MB (MAX_CONTENT_LENGTH)",
        "HTML output escaped to prevent XSS (escapeHtml utility in base.js)",
        "Build commands executed in isolated temp directories",
        "No shell injection: subprocess commands use list form, not shell=True",
    ])

    # ══════════════════════════════════════════════════════════════
    # 10. TESTING GUIDE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("10. Testing Guide", level=1)

    doc.add_heading("10.1 Platform Verification Checklist", level=2)
    add_numbered(doc, [
        "Start server: cd backend && PORT=5001 python3 app.py",
        "Verify all 7 pages return HTTP 200: /, /migrate, /swagger, /github, /build, /settings, /architecture",
        "Run a migration with sample XML on /migrate",
        "Verify generated code includes Swagger annotations (@Tag, @Operation, @ApiResponse)",
        "Edit a generated file inline and re-validate with LLM",
        "Generate Swagger from migration results on /swagger page",
        "Connect GitHub with PAT on /github, list repos, create a test repo, push files",
        "Check prerequisites on /build page (verify install guidance when tools missing)",
        "Build JAR on /build (requires JDK 17+), verify artifact download",
        "Build Docker image with platform selection (requires Docker)",
        "Run tests on /build, verify SSE streaming output in terminal panel",
        "Check architecture page with Basic Auth credentials",
        "Verify cross-page state: migration result available on dashboard, swagger, github, build pages",
    ])

    doc.add_heading("10.2 Generated Code Testing", level=2)
    add_para(doc,
        "The generated Spring Boot project includes JUnit 5 tests. After building:")
    add_code(doc, """# Run generated project tests
cd /path/to/generated/project
./mvnw test

# Expected tests:
# - ApplicationTests: @SpringBootTest context load verification
# - *ControllerTest: @WebMvcTest for each generated controller
#   - Tests GET/POST/PUT/DELETE endpoints
#   - Mocks service layer with @MockBean
#   - Verifies HTTP status codes and response content""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 11. DEPLOYMENT GUIDE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("11. Deployment Guide", level=1)

    doc.add_heading("11.1 Development", level=2)
    add_code(doc, """cd backend
export FLASK_ENV=development
export PORT=5001
python3 app.py""")

    doc.add_heading("11.2 Production with Gunicorn", level=2)
    add_code(doc, """cd backend
export FLASK_ENV=production
export PORT=5001
export ARCH_USERNAME=your-admin-username
export ARCH_PASSWORD=your-secure-password
export CORS_ORIGINS=https://your-domain.com
gunicorn app:app --config gunicorn.conf.py""")

    doc.add_heading("11.3 Docker Deployment", level=2)
    add_code(doc, """# Build the platform image
docker build -t mulesoft-migrator:latest .

# Run with environment configuration
docker run -d \\
    --name mulesoft-migrator \\
    -p 5001:5001 \\
    -e PORT=5001 \\
    -e FLASK_ENV=production \\
    -e ARCH_USERNAME=admin \\
    -e ARCH_PASSWORD=secure-password \\
    mulesoft-migrator:latest""")

    doc.add_heading("11.4 Production Checklist", level=2)
    add_bullets(doc, [
        "Set FLASK_ENV=production to disable debug mode",
        "Change default Architecture page credentials (ARCH_USERNAME, ARCH_PASSWORD)",
        "Configure CORS_ORIGINS to restrict allowed origins",
        "Set SECRET_KEY to a strong random value",
        "Use Gunicorn or equivalent WSGI server (never Flask dev server in production)",
        "Configure HTTPS/TLS termination at the load balancer or reverse proxy",
        "Set up log aggregation for Gunicorn access and error logs",
        "Monitor /api/health endpoint for uptime checks",
    ])

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 12. TROUBLESHOOTING
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("12. Troubleshooting", level=1)

    add_table(doc,
        ["Issue", "Cause", "Solution"],
        [
            ["Port 5001 already in use", "Previous server instance running", "lsof -ti:5001 | xargs kill -9"],
            ["Import error on startup", "Missing Python dependency", "pip install -r backend/requirements.txt"],
            ["LLM validation fails", "Invalid API key or provider", "Check API key on Settings page, verify provider status"],
            ["GitHub push fails", "Invalid PAT or insufficient scopes", "Generate new PAT with 'repo' scope at github.com/settings/tokens"],
            ["JAR build fails", "Java not installed", "Install JDK 17+ (see /build page for download links)"],
            ["Docker build fails", "Docker not running", "Start Docker Desktop or Docker daemon"],
            ["Architecture page 401", "Wrong credentials", "Set ARCH_USERNAME and ARCH_PASSWORD env vars"],
            ["CORS error in browser", "Origin not in allowlist", "Set CORS_ORIGINS env var to include your frontend URL"],
            ["Large XML timeout", "XML exceeds processing time", "Split into smaller files or increase Gunicorn timeout"],
            ["DataWeave conversion incomplete", "Unsupported DW feature", "Enable LLM agent for better coverage of complex scripts"],
        ],
        col_widths=[1.8, 1.8, 2.7])

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # 13. APPENDIX
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("13. Appendix", level=1)

    doc.add_heading("13.1 Supported MuleSoft Connectors (30+)", level=2)
    add_table(doc,
        ["Category", "Connectors"],
        [
            ["Transport", "HTTP, HTTPS, WebSocket"],
            ["Messaging", "JMS, AMQP (RabbitMQ), Kafka, VM, Anypoint MQ"],
            ["Database", "JDBC (MySQL, PostgreSQL, Oracle, SQL Server), MongoDB, Redis, Elasticsearch"],
            ["File", "File, SFTP, FTP, FTPS"],
            ["Email", "IMAP, POP3, SMTP"],
            ["Cloud", "Salesforce, AWS S3, AWS SQS, AWS SNS"],
            ["Integration", "SOAP/WS Consumer, REST, APIkit"],
            ["Security", "OAuth2, Basic Auth, TLS/SSL, Secure Properties"],
            ["Processing", "Batch, Scheduler, Object Store, Cache"],
        ],
        col_widths=[1.5, 5.0])

    doc.add_heading("13.2 Environment Variables Reference", level=2)
    add_table(doc,
        ["Variable", "Default", "Description"],
        [
            ["FLASK_ENV", "production", "Flask environment (development enables debug)"],
            ["PORT", "5000", "Server port number"],
            ["SECRET_KEY", "(auto-generated)", "Flask secret key for sessions"],
            ["CORS_ORIGINS", "*", "Comma-separated list of allowed CORS origins"],
            ["ARCH_USERNAME", "admin-username", "Architecture page HTTP Basic Auth username"],
            ["ARCH_PASSWORD", "admin-password", "Architecture page HTTP Basic Auth password"],
            ["MAX_CONTENT_LENGTH", "50MB", "Maximum upload file size"],
            ["GUNICORN_BIND", "0.0.0.0:5000", "Gunicorn bind address"],
            ["GUNICORN_WORKERS", "(CPU*2+1)", "Number of Gunicorn worker processes"],
            ["GUNICORN_THREADS", "4", "Threads per Gunicorn worker"],
        ],
        col_widths=[2.0, 1.5, 3.0])

    doc.add_heading("13.3 localStorage Keys", level=2)
    add_table(doc,
        ["Key", "Used By", "Content"],
        [
            ["msb_migration_result", "All pages", "Full migration result (files, summary, validation)"],
            ["msb_llm_provider", "Migration, Settings", "Selected LLM provider name"],
            ["msb_llm_model", "Migration, Settings", "Selected LLM model ID"],
            ["msb_llm_api_key", "Migration, Settings", "LLM provider API key"],
            ["msb_github_token", "GitHub, Settings", "GitHub Personal Access Token"],
            ["msb_project_name", "Migration, Settings", "Default project name"],
            ["msb_group_id", "Migration, Settings", "Default Maven group ID"],
            ["msb_java_version", "Migration, Settings", "Default Java version (11/17/21)"],
        ],
        col_widths=[2.0, 1.5, 3.0])

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("End of Documentation")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_TEXT
    run.font.name = "Arial"
    run.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY_TEXT
    run2.font.name = "Arial"

    # ── Save ──
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "MuleSoft_SpringBoot_Platform_Documentation.docx")
    doc.save(output_path)
    size = os.path.getsize(output_path)
    print(f"Documentation generated successfully!")
    print(f"  File: {output_path}")
    print(f"  Size: {size / 1024:.1f} KB ({size:,} bytes)")
    return output_path


if __name__ == "__main__":
    create_documentation()
