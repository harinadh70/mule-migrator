#!/usr/bin/env python3
"""
Generate comprehensive documentation for MuleSoft to Spring Boot Migrator.
Creates both .docx and .pdf formats.
"""
import os
import sys

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ════════════════════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False,
                            font_size=None, color=None, alignment=None,
                            space_before=None, space_after=None, font_name=None):
    """Add a formatted paragraph."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1B3A5C", stripe=True):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Arial'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if stripe and row_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    # Set column widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
    return table

def add_code_block(doc, code, language=""):
    """Add a code block with gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    # Add shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(code)
    run.font.size = Pt(8.5)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_info_box(doc, text, box_type="info"):
    """Add a colored info/warning/tip box."""
    colors = {
        "info": ("E8F4FD", "1565C0"),
        "warning": ("FFF8E1", "F57F17"),
        "tip": ("E8F5E9", "2E7D32"),
        "important": ("FCE4EC", "C62828"),
    }
    bg, fg = colors.get(box_type, colors["info"])
    labels = {"info": "INFO", "warning": "WARNING", "tip": "TIP", "important": "IMPORTANT"}

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, bg)

    # Label
    p = cell.paragraphs[0]
    label_run = p.add_run(f"  {labels[box_type]}: ")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(*[int(fg[i:i+2], 16) for i in (0, 2, 4)])
    label_run.font.name = 'Arial'

    text_run = p.add_run(text)
    text_run.font.size = Pt(9)
    text_run.font.name = 'Arial'
    text_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Set width
    for row in table.rows:
        for c in row.cells:
            c.width = Inches(6.5)
    return table

def add_bullet_list(doc, items, level=0):
    """Add bullet list items."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'Arial'

def add_numbered_list(doc, items):
    """Add numbered list items."""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'Arial'

def add_page_break(doc):
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# MAIN DOCUMENT CREATION
# ════════════════════════════════════════════════════════════════

def create_documentation():
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Configure Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15

    for i in range(1, 5):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Arial'
        h_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.styles['Heading 1'].font.size = Pt(22)
    doc.styles['Heading 2'].font.size = Pt(16)
    doc.styles['Heading 3'].font.size = Pt(13)
    doc.styles['Heading 4'].font.size = Pt(11)

    # ════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()

    # Title block
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = title_table.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, "1B3A5C")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n")
    run.font.size = Pt(6)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("MuleSoft to Spring Boot Migrator")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Arial'

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run("Complete Technical Documentation")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)
    run.font.name = 'Arial'

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p4.add_run("A Beginner-Friendly Guide")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)
    run.font.name = 'Arial'

    p5 = cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p5.add_run("\n")
    run.font.size = Pt(6)

    cell.width = Inches(6.5)

    doc.add_paragraph()
    doc.add_paragraph()

    # Meta info
    meta_info = [
        ("Version", "1.0"),
        ("Date", "March 2026"),
        ("Author", "Harinadh"),
        ("Technology Stack", "Python Flask + Spring Boot 3.2"),
    ]
    for label, value in meta_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run2 = p.add_run(value)
        run2.font.size = Pt(10)
        run2.font.name = 'Arial'
        run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph()

    toc_items = [
        ("1.", "Introduction", "What is this tool and why does it exist?"),
        ("2.", "Key Concepts for Beginners", "Understanding MuleSoft, Spring Boot, and APIs"),
        ("3.", "System Requirements", "What you need before starting"),
        ("4.", "Installation Guide", "Step-by-step setup instructions"),
        ("5.", "Quick Start Tutorial", "Your first migration in 5 minutes"),
        ("6.", "Understanding the User Interface", "Every button and panel explained"),
        ("7.", "How the Migration Works", "The complete pipeline from XML to Java"),
        ("8.", "Supported MuleSoft Components", "Every connector, processor, and pattern"),
        ("9.", "DataWeave to Java Conversion", "How expressions are translated"),
        ("10.", "AI-Powered Code Validation", "Using LLM models to review code"),
        ("11.", "Multi-File Migration", "Working with multiple XML files"),
        ("12.", "Generated Spring Boot Project", "Understanding the output"),
        ("13.", "API Reference", "All endpoints documented"),
        ("14.", "Production Deployment", "Docker, Nginx, and cloud deployment"),
        ("15.", "Configuration Reference", "Every setting explained"),
        ("16.", "Troubleshooting Guide", "Common problems and solutions"),
        ("17.", "Glossary", "Technical terms explained simply"),
    ]

    for num, title, desc in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.4)
        run2 = p2.add_run(desc)
        run2.font.size = Pt(9)
        run2.font.name = 'Arial'
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run2.italic = True

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 1: INTRODUCTION
    # ════════════════════════════════════════════════════════════
    doc.add_heading('1. Introduction', level=1)

    doc.add_heading('1.1 What is This Tool?', level=2)
    doc.add_paragraph(
        'The MuleSoft to Spring Boot Migrator is a web application that automatically converts '
        'MuleSoft 4 integration projects into Spring Boot 3.2 Java applications. Think of it as a '
        '"translator" that reads your MuleSoft XML configuration files and produces a complete, '
        'ready-to-run Java project.'
    )

    doc.add_heading('1.2 Why Would You Need This?', level=2)
    doc.add_paragraph(
        'Many organizations use MuleSoft (owned by Salesforce) for building APIs and integrations. '
        'However, there are several reasons why a company might want to move to Spring Boot:'
    )
    add_bullet_list(doc, [
        'Cost Reduction: MuleSoft licenses can be expensive. Spring Boot is free and open-source.',
        'Control: With Spring Boot, you own and control your entire codebase.',
        'Flexibility: Spring Boot has a massive ecosystem of libraries and tools.',
        'Talent Pool: More developers know Java/Spring Boot than MuleSoft.',
        'Performance: Spring Boot applications can be highly optimized for your specific use case.',
    ])

    doc.add_heading('1.3 What Does It Actually Do?', level=2)
    doc.add_paragraph('Here is exactly what happens when you use this tool:')
    add_numbered_list(doc, [
        'You upload your MuleSoft XML files (the files that define your APIs and integrations)',
        'The tool reads and understands every element in those XML files',
        'It identifies all connectors (HTTP, Database, JMS, Kafka, etc.)',
        'It converts MuleSoft flows into Java classes with Spring Boot annotations',
        'It translates DataWeave expressions into equivalent Java code',
        'It generates a complete Maven project with all dependencies',
        'It creates configuration files (application.properties) with correct settings',
        'Optionally, an AI model reviews the generated code for quality and suggestions',
        'You download a ZIP file containing a complete, runnable Spring Boot project',
    ])

    doc.add_heading('1.4 Who Is This Documentation For?', level=2)
    doc.add_paragraph(
        'This documentation is written for everyone, from complete beginners to experienced developers. '
        'Every technical term is explained. Every step includes detailed instructions. '
        'You do NOT need prior experience with MuleSoft, Spring Boot, or programming to understand this guide.'
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 2: KEY CONCEPTS
    # ════════════════════════════════════════════════════════════
    doc.add_heading('2. Key Concepts for Beginners', level=1)

    doc.add_heading('2.1 What is MuleSoft?', level=2)
    doc.add_paragraph(
        'MuleSoft is a platform for building application programming interfaces (APIs) and integrations. '
        'It allows businesses to connect different software systems together. For example, connecting '
        'your website to a database, or connecting your payment system to your inventory system.'
    )
    doc.add_paragraph(
        'MuleSoft uses XML configuration files to define how data flows between systems. '
        'These XML files contain "flows" which are like step-by-step instructions: '
        '"When someone sends a request to this URL, do this, then do that, then send back a response."'
    )

    doc.add_heading('2.2 What is Spring Boot?', level=2)
    doc.add_paragraph(
        'Spring Boot is a popular Java framework for building web applications and APIs. '
        'It is open-source (free to use) and is the most widely used Java framework in the world. '
        'Spring Boot makes it easy to create production-ready applications with minimal configuration.'
    )

    doc.add_heading('2.3 What is an API?', level=2)
    doc.add_paragraph(
        'An API (Application Programming Interface) is a way for two software programs to talk to each other. '
        'Think of it like a waiter in a restaurant: you (the client) tell the waiter (the API) what you want, '
        'the waiter goes to the kitchen (the server), and brings back your food (the response).'
    )
    doc.add_paragraph('Common API concepts:')
    add_styled_table(doc,
        ["Term", "What It Means", "Example"],
        [
            ["Endpoint", "A specific URL that accepts requests", "/api/customers"],
            ["GET", "A request to retrieve data", "Get list of all customers"],
            ["POST", "A request to create new data", "Create a new customer"],
            ["PUT", "A request to update existing data", "Update customer address"],
            ["DELETE", "A request to remove data", "Delete a customer record"],
            ["JSON", "A text format for sending/receiving data", '{"name": "John", "age": 30}'],
            ["XML", "Another text format (used by MuleSoft)", "<name>John</name>"],
        ],
        col_widths=[1.0, 2.5, 3.0]
    )

    doc.add_heading('2.4 What is Docker?', level=2)
    doc.add_paragraph(
        'Docker is a tool that packages your application and everything it needs (libraries, settings, etc.) '
        'into a "container." Think of a container like a shipping container: no matter what ship (computer) '
        'carries it, the contents are always the same and always work. This means your application will '
        'run exactly the same way on any computer that has Docker installed.'
    )

    doc.add_heading('2.5 What is DataWeave?', level=2)
    doc.add_paragraph(
        'DataWeave is MuleSoft\'s own programming language for transforming data. For example, if you '
        'receive customer data in one format but need to send it in a different format, DataWeave handles '
        'that transformation. Our tool converts DataWeave code into equivalent Java code.'
    )

    doc.add_heading('2.6 What is an LLM (AI Model)?', level=2)
    doc.add_paragraph(
        'LLM stands for Large Language Model. These are AI systems (like ChatGPT, Claude, Gemini) that '
        'can understand and generate text, including code. Our tool can optionally send the generated '
        'Spring Boot code to an LLM for review, getting suggestions for improvements and catching potential issues.'
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 3: SYSTEM REQUIREMENTS
    # ════════════════════════════════════════════════════════════
    doc.add_heading('3. System Requirements', level=1)

    doc.add_heading('3.1 Minimum Requirements', level=2)
    add_styled_table(doc,
        ["Requirement", "Minimum", "Recommended"],
        [
            ["Operating System", "macOS 10.15+, Windows 10+, Linux", "macOS 12+ or Ubuntu 22.04+"],
            ["Python", "3.9 or higher", "3.12"],
            ["RAM", "4 GB", "8 GB or more"],
            ["Disk Space", "500 MB", "2 GB (for Docker)"],
            ["Browser", "Any modern browser", "Chrome or Firefox (latest)"],
            ["Internet", "Required for LLM validation", "Broadband connection"],
        ],
        col_widths=[1.5, 2.2, 2.8]
    )

    doc.add_heading('3.2 Optional Requirements', level=2)
    add_styled_table(doc,
        ["Component", "When Needed", "Purpose"],
        [
            ["Docker", "Production deployment", "Run the app in a container"],
            ["LLM API Key", "AI code validation", "Review generated code quality"],
            ["Java 17+", "To run generated project", "Build the Spring Boot output"],
            ["Maven 3.9+", "To build generated project", "Compile and package the Java project"],
        ],
        col_widths=[1.5, 2.2, 2.8]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 4: INSTALLATION GUIDE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('4. Installation Guide', level=1)

    doc.add_heading('4.1 Method 1: Local Installation (Recommended for Development)', level=2)

    doc.add_heading('Step 1: Verify Python is Installed', level=3)
    doc.add_paragraph('Open your terminal (on Mac: search for "Terminal" in Spotlight) and type:')
    add_code_block(doc, 'python3 --version')
    doc.add_paragraph('You should see something like "Python 3.9.6" or higher. If you get an error, install Python from python.org.')

    doc.add_heading('Step 2: Navigate to the Project Folder', level=3)
    add_code_block(doc, 'cd "/path/to/mulesoft-to-springboot-migrator/backend"')

    doc.add_heading('Step 3: Install Dependencies', level=3)
    doc.add_paragraph('This installs all the Python packages the tool needs:')
    add_code_block(doc, 'pip3 install -r requirements.txt')
    doc.add_paragraph(
        'This will install Flask (the web framework), lxml (for reading XML files), Gunicorn (the production server), '
        'and optionally the AI/LLM libraries (anthropic, openai, google-generativeai).'
    )

    doc.add_heading('Step 4: Start the Server', level=3)
    doc.add_paragraph('For development (simple, with auto-reload on code changes):')
    add_code_block(doc, 'python3 app.py')
    doc.add_paragraph('For production (faster, handles many users at once):')
    add_code_block(doc, 'python3 -m gunicorn -c gunicorn.conf.py app:app')

    doc.add_heading('Step 5: Open in Browser', level=3)
    doc.add_paragraph('Open your web browser and go to:')
    add_code_block(doc, 'http://localhost:5000')
    add_info_box(doc, 'You should see the MuleSoft to Spring Boot Migrator interface with a dark-themed UI.', "tip")

    doc.add_paragraph()
    doc.add_heading('4.2 Method 2: Docker Installation (Recommended for Production)', level=2)

    doc.add_heading('Step 1: Install Docker', level=3)
    doc.add_paragraph(
        'Download Docker Desktop from docker.com and install it. After installation, '
        'verify it works by opening terminal and typing:'
    )
    add_code_block(doc, 'docker --version')

    doc.add_heading('Step 2: Create Environment File', level=3)
    doc.add_paragraph('Copy the example environment file and fill in your API keys:')
    add_code_block(doc, 'cp .env.example .env\n# Edit .env with your favorite text editor')

    doc.add_heading('Step 3: Start with Docker Compose', level=3)
    add_code_block(doc, '# Basic start (just the app)\ndocker compose up -d\n\n# With Nginx reverse proxy (for production)\ndocker compose --profile with-nginx up -d')

    doc.add_heading('Step 4: Verify', level=3)
    add_code_block(doc, '# Check health endpoint\ncurl http://localhost:5000/api/health\n\n# Expected response:\n# {"status": "ok", "env": "production"}')

    doc.add_heading('4.3 Stopping the Server', level=2)
    add_styled_table(doc,
        ["Method", "Command"],
        [
            ["Local (Gunicorn)", "kill $(pgrep -f 'gunicorn.*app:app')"],
            ["Local (Flask dev)", "Press Ctrl+C in the terminal"],
            ["Docker", "docker compose down"],
        ],
        col_widths=[2.0, 4.5]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 5: QUICK START TUTORIAL
    # ════════════════════════════════════════════════════════════
    doc.add_heading('5. Quick Start Tutorial', level=1)
    doc.add_paragraph(
        'Let\'s walk through your very first migration. This tutorial takes about 5 minutes.'
    )

    doc.add_heading('Step 1: Open the Application', level=2)
    doc.add_paragraph(
        'Start the server (see Chapter 4) and open http://localhost:5000 in your browser. '
        'You will see a dark-themed interface split into two panels: Input (left) and Output (right).'
    )

    doc.add_heading('Step 2: Load Sample Data', level=2)
    doc.add_paragraph(
        'Click the "Load Sample" button in the top-right corner of the header. This fills in '
        'a sample MuleSoft XML configuration that demonstrates HTTP endpoints, database queries, '
        'and DataWeave transformations.'
    )

    doc.add_heading('Step 3: Click "Migrate to Spring Boot"', level=2)
    doc.add_paragraph(
        'Click the large blue "Migrate to Spring Boot" button at the bottom of the left panel. '
        'The tool will process the XML and generate a complete Spring Boot project. '
        'You\'ll see a loading spinner for 1-2 seconds.'
    )

    doc.add_heading('Step 4: Explore the Output', level=2)
    doc.add_paragraph('The right panel will now show three tabs:')
    add_bullet_list(doc, [
        'Generated Files: A file tree showing all generated Java files, configuration files, and build files. '
        'Click any file to see its contents with syntax highlighting.',
        'Summary: Statistics about what was migrated (number of flows, connectors found, dependencies added, etc.)',
        'AI Review: If LLM validation was enabled, shows an AI-generated code review with score and suggestions.',
    ])

    doc.add_heading('Step 5: Download the Project', level=2)
    doc.add_paragraph(
        'Click the "Download ZIP" button to download the complete Spring Boot project as a ZIP file. '
        'Unzip it, open it in your favorite Java IDE (IntelliJ IDEA, Eclipse, VS Code), and run it!'
    )

    add_info_box(doc,
        'The generated project includes a Dockerfile and docker-compose.yml, so you can also '
        'run it with "docker compose up -d" without installing Java locally!', "tip")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 6: UNDERSTANDING THE USER INTERFACE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('6. Understanding the User Interface', level=1)

    doc.add_heading('6.1 Overall Layout', level=2)
    doc.add_paragraph(
        'The application uses a split-panel layout with a dark slate theme. Here is every section explained:'
    )

    add_styled_table(doc,
        ["Section", "Location", "Purpose"],
        [
            ["Header Bar", "Top of page", "Contains logo, title, Load Sample button, and Clear button"],
            ["Input Panel", "Left side", "Where you provide MuleSoft XML, DataWeave scripts, and settings"],
            ["Output Panel", "Right side", "Shows generated code, migration summary, and AI review"],
            ["Status Bar", "Bottom of page", "Shows current status messages and progress"],
            ["Loading Overlay", "Center (when active)", "Animated spinner shown during migration processing"],
        ],
        col_widths=[1.3, 1.5, 3.7]
    )

    doc.add_heading('6.2 Input Panel Tabs', level=2)

    doc.add_heading('Tab 1: MuleSoft XML', level=3)
    doc.add_paragraph(
        'This is the main input tab where you provide your MuleSoft configuration files.'
    )
    add_bullet_list(doc, [
        'Upload Zone: Drag and drop XML files here, or click to browse. Accepts .xml, .raml, .yaml, .yml files.',
        'Multiple File Support: You can upload multiple XML files. Each appears as a "chip" showing the filename and size.',
        'Code Editor: A syntax-highlighted text editor (CodeMirror) where you can paste XML content directly.',
        'File Preview: Click any uploaded file chip to preview its contents in the editor.',
        'Remove Files: Click the X button on any file chip to remove it.',
    ])

    doc.add_heading('Tab 2: DataWeave', level=3)
    doc.add_paragraph(
        'For adding standalone DataWeave 2.0 scripts that exist outside your XML flows.'
    )
    add_bullet_list(doc, [
        'Script List: Shows all DataWeave scripts you\'ve added, displayed as tabs.',
        'Add Script: Click "+" to create a new DataWeave script. Give it a name.',
        'Editor: Write or paste DataWeave code in the syntax-highlighted editor.',
        'Convert: Click "Convert" to see the Java equivalent of just this DataWeave script.',
    ])

    doc.add_heading('Tab 3: Settings', level=3)
    doc.add_paragraph('Configure the output project settings:')
    add_styled_table(doc,
        ["Setting", "Default", "Description"],
        [
            ["Project Name", "my-spring-app", "The name for your Spring Boot project (used in pom.xml)"],
            ["Group ID", "com.example", "Java package namespace (like com.yourcompany)"],
            ["Java Version", "17", "Which Java version to target (17 or 21 recommended)"],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    doc.add_heading('Tab 4: AI Validation', level=3)
    doc.add_paragraph('Configure AI-powered code review:')
    add_bullet_list(doc, [
        'Enable/Disable Toggle: Turn AI validation on or off with the toggle switch.',
        'Provider Dropdown: Select which AI provider to use (Anthropic, OpenAI, Google, etc.).',
        'Model Dropdown: Choose a specific model (e.g., Claude Sonnet 4, GPT-4o, Gemini Pro).',
        'API Key Field: Enter your API key for the selected provider (never stored on disk).',
        'Base URL: Only needed for Ollama (local AI) - defaults to http://localhost:11434.',
        'Provider Info Card: Shows documentation link and model tier (Free, Standard, Premium).',
        'Test Connection: Click to verify your API key works before migrating.',
    ])

    doc.add_heading('6.3 Output Panel Tabs', level=2)

    doc.add_heading('Tab 1: Generated Files', level=3)
    doc.add_paragraph(
        'Shows a hierarchical file tree of all generated files. The tree structure mirrors '
        'a standard Maven project. Click any file to view its contents with syntax highlighting.'
    )

    doc.add_heading('Tab 2: Summary', level=3)
    doc.add_paragraph('Statistics about the migration:')
    add_bullet_list(doc, [
        'Flows Converted: Number of MuleSoft flows successfully migrated',
        'Sub-Flows Converted: Number of reusable sub-flows migrated',
        'Connectors Found: List of all detected connectors (HTTP, Database, JMS, etc.)',
        'DataWeave Scripts Converted: Number of DataWeave transformations translated to Java',
        'Dependencies: Maven dependencies added to pom.xml',
        'Warnings: Any issues or items that may need manual review',
        'XML Files Processed: Count and names of all input XML files',
    ])

    doc.add_heading('Tab 3: AI Review', level=3)
    doc.add_paragraph('If AI validation was enabled, this tab shows:')
    add_bullet_list(doc, [
        'Overall Score: A number from 1 to 10 rating the code quality, displayed as a color-coded circle',
        'Summary: A brief text description of the overall quality',
        'Issues: Problems found, categorized by severity (Critical, Warning, Info)',
        'Improvements: Specific code suggestions with before/after examples',
        'Missing Items: Things that should be added (error handling, logging, etc.)',
        'Security Issues: Potential security vulnerabilities',
        'Best Practices: Recommendations for following Spring Boot conventions',
    ])

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 7: HOW THE MIGRATION WORKS
    # ════════════════════════════════════════════════════════════
    doc.add_heading('7. How the Migration Works', level=1)
    doc.add_paragraph(
        'The migration follows a pipeline of 6 steps. Understanding this pipeline helps you '
        'understand the output and troubleshoot any issues.'
    )

    doc.add_heading('7.1 The Migration Pipeline', level=2)

    steps = [
        ("Step 1: Parse", "parser.py",
         "The XML parser reads each MuleSoft XML file and extracts all meaningful elements: "
         "flows, sub-flows, connectors, configurations, error handlers, DataWeave transformations, "
         "batch jobs, and more. It understands 30+ MuleSoft namespaces. If you provide multiple "
         "XML files, each is parsed independently to avoid namespace conflicts."),
        ("Step 2: Merge", "app.py",
         "If you provided multiple XML files, the parsed results are merged into a single unified "
         "structure. Duplicate flows (same name in different files) are detected and handled: the "
         "first occurrence is kept and a warning is generated. Properties are merged (last wins). "
         "Connectors are combined (union of all detected types)."),
        ("Step 3: Map Connectors", "connector_mapper.py",
         "Each detected connector (HTTP, Database, JMS, Kafka, etc.) is mapped to its Spring Boot "
         "equivalent. This determines which Maven dependencies to include, which Spring annotations "
         "to use, and which configuration properties to generate. The mapper knows about 30+ "
         "connector types."),
        ("Step 4: Convert Flows", "flow_converter.py",
         "Each MuleSoft flow is converted to a Java class. HTTP listener flows become REST "
         "controllers with @GetMapping, @PostMapping, etc. Scheduler flows become @Scheduled "
         "methods. Message listener flows become JMS/Kafka/AMQP listeners. All processors within "
         "flows (choice, for-each, try-catch, etc.) are converted to equivalent Java code."),
        ("Step 5: Generate Project", "spring_generator.py",
         "The complete Spring Boot project is assembled: pom.xml with all dependencies, the main "
         "Application class with required annotations, configuration classes for each connector type, "
         "application.properties with all settings, exception classes, utility classes, Dockerfile, "
         "docker-compose.yml, tests, and .gitignore."),
        ("Step 6: Validate (Optional)", "llm_validator.py",
         "If AI validation is enabled, the generated code is sent to an LLM (AI model) for review. "
         "The AI acts as a senior Java developer reviewing the code, scoring it from 1-10 and providing "
         "actionable feedback on issues, improvements, security, and best practices."),
    ]

    for title, module, desc in steps:
        doc.add_heading(title, level=3)
        p = doc.add_paragraph()
        run = p.add_run(f"Module: {module}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph(desc)

    doc.add_heading('7.2 Visual Pipeline Flow', level=2)

    pipeline_table = doc.add_table(rows=1, cols=6)
    pipeline_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["XML Files", "Parser", "Merger", "Converter", "Generator", "Output ZIP"]
    colors_list = ["E3F2FD", "BBDEFB", "90CAF9", "64B5F6", "42A5F5", "1E88E5"]
    for i, (label, clr) in enumerate(zip(labels, colors_list)):
        cell = pipeline_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Arial'
        set_cell_shading(cell, clr)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 8: SUPPORTED MULESOFT COMPONENTS
    # ════════════════════════════════════════════════════════════
    doc.add_heading('8. Supported MuleSoft Components', level=1)

    doc.add_heading('8.1 Supported Connectors (30+)', level=2)
    doc.add_paragraph(
        'A "connector" in MuleSoft is a pre-built module for connecting to external systems. '
        'Below is every connector this tool supports and what it maps to in Spring Boot:'
    )

    connector_rows = [
        ["HTTP Listener", "Receives API requests", "Spring Web (@RestController)", "spring-boot-starter-web"],
        ["HTTP Request", "Sends API requests", "RestTemplate / WebClient", "spring-boot-starter-webflux"],
        ["Database", "MySQL, Oracle, PostgreSQL, MSSQL", "Spring Data JPA / JdbcTemplate", "spring-boot-starter-data-jpa"],
        ["JMS", "Java message queues", "JmsTemplate / @JmsListener", "spring-boot-starter-activemq"],
        ["AMQP", "RabbitMQ messaging", "RabbitTemplate / @RabbitListener", "spring-boot-starter-amqp"],
        ["Kafka", "Event streaming", "KafkaTemplate / @KafkaListener", "spring-kafka"],
        ["VM", "In-memory messaging", "Spring Events (ApplicationEvent)", "Built-in (no extra dependency)"],
        ["File", "Local file operations", "java.nio.file.Files", "Built-in"],
        ["SFTP", "Secure file transfer", "Spring Integration SFTP", "spring-integration-sftp"],
        ["FTP", "File transfer protocol", "Spring Integration FTP", "spring-integration-ftp"],
        ["Email", "IMAP, POP3, SMTP", "JavaMail / Spring Mail", "spring-boot-starter-mail"],
        ["SOAP/WS", "SOAP web services", "Spring Web Services", "spring-boot-starter-web-services"],
        ["Salesforce", "CRM integration", "RestTemplate + OAuth", "Custom configuration"],
        ["Amazon S3", "Cloud file storage", "AWS SDK v2 S3Client", "software.amazon.awssdk:s3"],
        ["Amazon SQS", "Cloud message queue", "AWS SqsClient", "spring-cloud-aws-messaging"],
        ["Amazon SNS", "Notification service", "AWS SnsClient", "software.amazon.awssdk:sns"],
        ["MongoDB", "NoSQL document DB", "Spring Data MongoDB", "spring-boot-starter-data-mongodb"],
        ["Redis", "In-memory cache/store", "Spring Data Redis", "spring-boot-starter-data-redis"],
        ["Elasticsearch", "Search engine", "Spring Data Elasticsearch", "spring-boot-starter-data-elasticsearch"],
        ["Object Store", "Key-value storage", "Redis (Spring Cache)", "spring-boot-starter-cache"],
        ["Batch", "Batch job processing", "Spring Batch", "spring-boot-starter-batch"],
        ["OAuth/Security", "Authentication", "Spring Security", "spring-boot-starter-security"],
        ["Validation", "Input validation", "Jakarta Bean Validation", "spring-boot-starter-validation"],
        ["Anypoint MQ", "Cloud messaging", "AMQP Provider", "spring-boot-starter-amqp"],
        ["APIkit", "API routing", "Spring MVC routing", "spring-boot-starter-web"],
    ]

    add_styled_table(doc,
        ["MuleSoft Connector", "Purpose", "Spring Boot Equivalent", "Maven Dependency"],
        connector_rows,
        col_widths=[1.3, 1.5, 1.8, 1.9]
    )

    add_page_break(doc)

    doc.add_heading('8.2 Supported Processors (50+)', level=2)
    doc.add_paragraph(
        'A "processor" in MuleSoft is an individual step within a flow. Here are all supported processors:'
    )

    doc.add_heading('Core Logic Processors', level=3)
    add_styled_table(doc,
        ["MuleSoft Processor", "What It Does", "Java/Spring Boot Equivalent"],
        [
            ["choice", "If/else branching based on conditions", "if-else statements"],
            ["for-each", "Loop through a list of items one by one", "Java Stream .forEach() or for loop"],
            ["parallel-for-each", "Process items simultaneously (in parallel)", "parallelStream() / CompletableFuture"],
            ["scatter-gather", "Run multiple tasks at the same time, wait for all", "CompletableFuture.allOf()"],
            ["try", "Attempt something, handle errors if it fails", "try-catch block"],
            ["until-successful", "Retry an operation until it works", "@Retryable annotation (Spring Retry)"],
            ["async", "Run something in the background", "@Async annotation (Spring Async)"],
            ["flow-ref", "Call another flow (like calling a function)", "Service method call"],
            ["first-successful", "Try options until one works", "try-catch chain or fallback pattern"],
            ["round-robin", "Distribute across options evenly", "Load balancing logic"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    doc.add_heading('Data Processors', level=3)
    add_styled_table(doc,
        ["MuleSoft Processor", "What It Does", "Java/Spring Boot Equivalent"],
        [
            ["set-payload", "Set the response body", "Return value or ResponseEntity"],
            ["set-variable", "Store a value for later use", "Local variable declaration"],
            ["logger", "Write a message to the log", "log.info() / log.debug() (SLF4J)"],
            ["raise-error", "Deliberately cause an error", "throw new Exception()"],
            ["parse-template", "Process a text template", "Template engine processing"],
            ["ee:transform", "DataWeave transformation", "Java stream operations (see Chapter 9)"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    doc.add_heading('Connector-Specific Processors', level=3)
    add_styled_table(doc,
        ["MuleSoft Processor", "What It Does", "Java/Spring Boot Equivalent"],
        [
            ["http:request", "Send an HTTP request to another API", "restTemplate.exchange() or webClient.get()"],
            ["db:select", "Query a database (read data)", "jdbcTemplate.query() or repository.find*()"],
            ["db:insert", "Add a new record to database", "jdbcTemplate.update() or repository.save()"],
            ["db:update", "Modify existing database records", "jdbcTemplate.update()"],
            ["db:delete", "Remove records from database", "jdbcTemplate.update() with DELETE SQL"],
            ["db:stored-procedure", "Call a database stored procedure", "jdbcTemplate.call()"],
            ["jms:publish", "Send a message to a JMS queue", "jmsTemplate.convertAndSend()"],
            ["jms:consume", "Read a message from a JMS queue", "jmsTemplate.receiveAndConvert()"],
            ["file:read", "Read a file from disk", "Files.readString(Path)"],
            ["file:write", "Write content to a file", "Files.writeString(Path, content)"],
            ["file:list", "List files in a directory", "Files.list(Path)"],
            ["file:delete", "Delete a file", "Files.deleteIfExists(Path)"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    doc.add_heading('8.3 Supported Error Handling', level=2)
    add_styled_table(doc,
        ["MuleSoft Pattern", "Behavior", "Spring Boot Equivalent"],
        [
            ["on-error-propagate", "Handle error and re-throw it", "catch block that re-throws"],
            ["on-error-continue", "Handle error and continue normally", "catch block that returns fallback"],
            ["when condition", "Handle only specific error types", "catch (SpecificException e)"],
            ["error-handler (global)", "Default handler for all flows", "@ControllerAdvice class"],
        ],
        col_widths=[1.8, 2.3, 2.4]
    )

    doc.add_paragraph()
    doc.add_heading('8.4 Error Type Mapping (100+ types)', level=2)
    doc.add_paragraph('MuleSoft error types are mapped to specific Java exceptions:')
    add_styled_table(doc,
        ["MuleSoft Error Type", "Java Exception", "HTTP Status"],
        [
            ["HTTP:NOT_FOUND", "ResourceNotFoundException", "404"],
            ["HTTP:BAD_REQUEST", "BadRequestException", "400"],
            ["HTTP:UNAUTHORIZED", "UnauthorizedException", "401"],
            ["HTTP:FORBIDDEN", "AccessDeniedException", "403"],
            ["HTTP:TIMEOUT", "SocketTimeoutException", "408"],
            ["HTTP:TOO_MANY_REQUESTS", "TooManyRequestsException", "429"],
            ["DB:CONNECTIVITY", "DataAccessResourceFailureException", "503"],
            ["DB:BAD_SQL_SYNTAX", "BadSqlGrammarException", "500"],
            ["VALIDATION:INVALID_*", "ConstraintViolationException", "400"],
            ["ANY (catch-all)", "Exception", "500"],
        ],
        col_widths=[2.0, 2.5, 1.0]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 9: DATAWEAVE TO JAVA CONVERSION
    # ════════════════════════════════════════════════════════════
    doc.add_heading('9. DataWeave to Java Conversion', level=1)
    doc.add_paragraph(
        'DataWeave is MuleSoft\'s proprietary data transformation language. This tool converts '
        'DataWeave 2.0 expressions into equivalent Java code using Java Streams and standard libraries.'
    )

    doc.add_heading('9.1 Supported DataWeave Operations', level=2)

    doc.add_heading('Collection Operations', level=3)
    add_styled_table(doc,
        ["DataWeave", "What It Does", "Java Equivalent"],
        [
            ["map", "Transform each item in a list", ".stream().map(...).collect(toList())"],
            ["filter", "Keep only items matching a condition", ".stream().filter(...).collect(toList())"],
            ["reduce", "Combine all items into one value", ".stream().reduce(...)"],
            ["flatMap", "Map and flatten nested lists", ".stream().flatMap(...).collect(toList())"],
            ["groupBy", "Group items by a key", ".stream().collect(groupingBy(...))"],
            ["orderBy", "Sort items by a field", ".stream().sorted(...).collect(toList())"],
            ["distinctBy", "Remove duplicate items", ".stream().distinct().collect(toList())"],
            ["flatten", "Flatten nested arrays", ".stream().flatMap(Collection::stream)"],
            ["pluck", "Extract values from an object", "map.entrySet().stream()..."],
            ["sizeOf", "Count items in a list", ".size() or .length()"],
            ["isEmpty", "Check if list is empty", ".isEmpty()"],
            ["first / last", "Get first or last item", ".get(0) / .get(size()-1)"],
            ["min / max", "Find minimum or maximum", ".stream().min(...) / .max(...)"],
            ["sum / avg", "Calculate sum or average", ".stream().mapToDouble(...).sum()/.average()"],
        ],
        col_widths=[1.2, 2.3, 3.0]
    )

    doc.add_heading('String Operations', level=3)
    add_styled_table(doc,
        ["DataWeave", "Java Equivalent"],
        [
            ["upper()", ".toUpperCase()"],
            ["lower()", ".toLowerCase()"],
            ["trim()", ".trim()"],
            ["capitalize()", "StringUtils.capitalize()"],
            ["contains()", ".contains()"],
            ["startsWith()", ".startsWith()"],
            ["endsWith()", ".endsWith()"],
            ["replace ... with ...", ".replace(old, new)"],
            ["splitBy()", ".split()"],
            ["joinBy()", "String.join()"],
            ['++ (concatenation)', '+ (string concat)'],
            ["matches(regex)", ".matches(regex)"],
            ["substringBefore()", ".substring(0, indexOf(...))"],
            ["substringAfter()", ".substring(indexOf(...) + 1)"],
        ],
        col_widths=[2.5, 4.0]
    )

    doc.add_heading('Type Coercion', level=3)
    add_styled_table(doc,
        ["DataWeave", "Java Equivalent", "Example"],
        [
            ["as String", "String.valueOf(x)", 'num as String -> String.valueOf(num)'],
            ["as Number", "Double.parseDouble(x)", '"42" as Number -> Double.parseDouble("42")'],
            ["as Boolean", "Boolean.parseBoolean(x)", '"true" as Boolean -> Boolean.parseBoolean(...)'],
            ["as Date", "LocalDate.parse(x)", "str as Date -> LocalDate.parse(str)"],
            ["as DateTime", "LocalDateTime.parse(x)", "str as DateTime -> LocalDateTime.parse(str)"],
        ],
        col_widths=[1.2, 2.3, 3.0]
    )

    doc.add_heading('Null Handling', level=3)
    add_styled_table(doc,
        ["DataWeave Pattern", "Java Equivalent"],
        [
            ['value default "fallback"', 'value != null ? value : "fallback"'],
            ["isEmpty(value)", "value == null || value.isEmpty()"],
            ["isBlank(value)", "value == null || value.trim().isEmpty()"],
        ],
        col_widths=[3.0, 3.5]
    )

    doc.add_heading('9.2 DataWeave Expression Examples', level=2)

    doc.add_paragraph('Example 1: Simple mapping', style='Heading 4')
    doc.add_paragraph('DataWeave:')
    add_code_block(doc, '%dw 2.0\noutput application/json\n---\npayload map (item) -> {\n    fullName: item.firstName ++ " " ++ item.lastName,\n    age: item.age as Number\n}')
    doc.add_paragraph('Generated Java:')
    add_code_block(doc, 'List<Map<String,Object>> result = ((List<Map<String,Object>>) payload)\n    .stream()\n    .map(item -> {\n        Map<String,Object> obj = new LinkedHashMap<>();\n        obj.put("fullName", item.get("firstName") + " " + item.get("lastName"));\n        obj.put("age", Double.parseDouble(String.valueOf(item.get("age"))));\n        return obj;\n    })\n    .collect(Collectors.toList());')

    doc.add_paragraph('Example 2: Filter and transform', style='Heading 4')
    doc.add_paragraph('DataWeave:')
    add_code_block(doc, 'payload filter (item) -> item.status == "active"\n    map (item) -> item.email')
    doc.add_paragraph('Generated Java:')
    add_code_block(doc, 'List<Object> result = ((List<Map<String,Object>>) payload)\n    .stream()\n    .filter(item -> "active".equals(item.get("status")))\n    .map(item -> item.get("email"))\n    .collect(Collectors.toList());')

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 10: AI-POWERED CODE VALIDATION
    # ════════════════════════════════════════════════════════════
    doc.add_heading('10. AI-Powered Code Validation', level=1)
    doc.add_paragraph(
        'The tool can optionally send the generated Spring Boot code to an AI language model '
        'for a professional code review. This is like having a senior Java developer look at your '
        'code and suggest improvements.'
    )

    doc.add_heading('10.1 Supported AI Providers', level=2)

    add_styled_table(doc,
        ["Provider", "Models Available", "Pricing", "Best For"],
        [
            ["Anthropic Claude", "Claude Sonnet 4, Claude 3.5 Sonnet, Claude 3 Opus, Claude 3.5 Haiku", "Paid (API key required)", "Best overall code review quality"],
            ["OpenAI GPT", "GPT-4o, GPT-4 Turbo, GPT-4o-mini, o3-mini", "Paid (API key required)", "Wide availability, good quality"],
            ["Google Gemini", "Gemini 2.5 Pro, Gemini 2.0 Flash, Gemini 1.5 Pro", "Free tier available", "Cost-effective with good quality"],
            ["DeepSeek", "DeepSeek Chat, DeepSeek Coder, DeepSeek Reasoner", "Very affordable", "Budget-friendly code analysis"],
            ["Groq", "Llama 3.3 70B, Mixtral 8x7B, Llama 3.1 8B", "Free!", "Fast and free, great for testing"],
            ["Ollama (Local)", "CodeLlama, Llama 3, DeepSeek Coder, Mistral, Qwen", "Free (runs locally)", "Complete privacy, no internet needed"],
        ],
        col_widths=[1.2, 2.0, 1.3, 2.0]
    )

    doc.add_heading('10.2 How to Set Up AI Validation', level=2)
    add_numbered_list(doc, [
        'Go to the "AI Validation" tab in the input panel',
        'Toggle the switch to "Enable AI Validation"',
        'Select your preferred provider from the dropdown',
        'Choose a model (each has different speed/quality tradeoffs)',
        'Enter your API key (get one from the provider\'s website - links are shown)',
        'Click "Test Connection" to verify everything works',
        'Now when you click "Migrate", the AI review will be included automatically',
    ])

    doc.add_heading('10.3 Understanding the AI Review', level=2)

    doc.add_heading('Score (1-10)', level=3)
    add_styled_table(doc,
        ["Score Range", "Color", "Meaning"],
        [
            ["8-10", "Green", "Excellent - Code is production-ready with minor suggestions"],
            ["6-7", "Yellow", "Good - Code works but has notable improvements to make"],
            ["4-5", "Orange", "Fair - Significant issues that should be addressed"],
            ["1-3", "Red", "Poor - Major problems that need fixing before use"],
        ],
        col_widths=[1.2, 1.0, 4.3]
    )

    doc.add_heading('Issue Severity Levels', level=3)
    add_styled_table(doc,
        ["Severity", "Meaning", "Action Required"],
        [
            ["Critical", "Could cause bugs, crashes, or security vulnerabilities", "Must fix before deploying"],
            ["Warning", "Will work but not following best practices", "Should fix for production quality"],
            ["Info", "Suggestions for improvement, not problems", "Nice to have, can fix later"],
        ],
        col_widths=[1.0, 3.0, 2.5]
    )

    doc.add_heading('10.4 Getting Free API Keys', level=2)
    add_styled_table(doc,
        ["Provider", "How to Get API Key", "Free Tier"],
        [
            ["Groq", "Sign up at console.groq.com", "Completely free with generous limits"],
            ["Google Gemini", "Sign up at aistudio.google.com", "Free tier with rate limits"],
            ["Ollama", "Install from ollama.com (no key needed)", "Completely free, runs on your computer"],
            ["DeepSeek", "Sign up at platform.deepseek.com", "Very cheap, sometimes free credits"],
        ],
        col_widths=[1.2, 2.8, 2.5]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 11: MULTI-FILE MIGRATION
    # ════════════════════════════════════════════════════════════
    doc.add_heading('11. Multi-File Migration', level=1)
    doc.add_paragraph(
        'Real MuleSoft projects typically have multiple XML files. This tool handles that scenario seamlessly.'
    )

    doc.add_heading('11.1 How Multi-File Works', level=2)
    add_numbered_list(doc, [
        'Upload multiple XML files using the upload zone (drag multiple files at once or click to select multiple)',
        'Each file appears as a chip showing its name and size',
        'Click any chip to preview the file contents',
        'When you click "Migrate", ALL uploaded files are processed together',
        'The tool parses each file independently, then merges the results into a single unified project',
    ])

    doc.add_heading('11.2 Merge Rules', level=2)
    add_styled_table(doc,
        ["Component", "Merge Rule", "On Conflict"],
        [
            ["Flows", "Combined from all files", "Duplicate names: first kept, warning generated"],
            ["Sub-Flows", "Combined from all files", "Duplicate names: first kept, warning generated"],
            ["Configurations", "Combined by name", "Duplicate names: first kept"],
            ["Properties", "Merged into single dict", "Same key: last value wins"],
            ["Connectors", "Union of all detected", "No conflicts possible (sets)"],
            ["Error Handlers", "All collected", "No conflicts (all kept)"],
            ["Batch Jobs", "All collected", "No conflicts (all kept)"],
        ],
        col_widths=[1.3, 2.3, 2.9]
    )

    doc.add_heading('11.3 Tips for Multi-File Migration', level=2)
    add_bullet_list(doc, [
        'Give each flow a unique name across all files to avoid merge conflicts',
        'Check the Summary tab for any duplicate warnings after migration',
        'Global configurations should be defined in only one file',
        'The generated output is always a single unified Spring Boot project',
    ])

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 12: GENERATED SPRING BOOT PROJECT
    # ════════════════════════════════════════════════════════════
    doc.add_heading('12. Generated Spring Boot Project', level=1)
    doc.add_paragraph(
        'When migration completes, you get a complete, ready-to-build Spring Boot project. '
        'Here is what each generated file does.'
    )

    doc.add_heading('12.1 Project Structure', level=2)
    add_code_block(doc,
        'my-spring-app/\n'
        '+-- pom.xml                              # Maven build file (dependencies)\n'
        '+-- Dockerfile                           # Container image definition\n'
        '+-- docker-compose.yml                   # Container orchestration\n'
        '+-- src/\n'
        '|   +-- main/\n'
        '|   |   +-- java/com/example/myspringapp/\n'
        '|   |   |   +-- Application.java         # Main entry point\n'
        '|   |   |   +-- controller/              # REST API controllers\n'
        '|   |   |   +-- service/                 # Business logic services\n'
        '|   |   |   +-- config/                  # Configuration classes\n'
        '|   |   |   +-- exception/               # Custom exception classes\n'
        '|   |   |   +-- util/                    # Utility classes\n'
        '|   |   +-- resources/\n'
        '|   |       +-- application.properties    # App configuration\n'
        '|   |       +-- application-dev.properties # Dev profile settings\n'
        '|   |       +-- application-prod.properties# Prod profile settings\n'
        '|   +-- test/\n'
        '|       +-- java/.../ApplicationTests.java# Unit tests\n'
        '+-- .gitignore'
    )

    doc.add_heading('12.2 Key Generated Files', level=2)

    add_styled_table(doc,
        ["File", "Purpose", "When Generated"],
        [
            ["pom.xml", "Maven build file listing all project dependencies", "Always"],
            ["Application.java", "Main class with @SpringBootApplication and conditional annotations", "Always"],
            ["*Controller.java", "REST endpoints converted from HTTP listener flows", "When HTTP flows exist"],
            ["*Service.java", "Business logic converted from flow processors", "When flows have service logic"],
            ["*Listener.java", "Message listeners for JMS/Kafka/AMQP", "When messaging flows exist"],
            ["*Scheduler.java", "Scheduled tasks from cron/fixed-delay flows", "When scheduler flows exist"],
            ["SchedulingConfig.java", "Enables Spring scheduling", "When schedulers detected"],
            ["JmsConfig.java", "JMS connection factory configuration", "When JMS connector used"],
            ["AmqpConfig.java", "RabbitMQ configuration", "When AMQP connector used"],
            ["KafkaConfig.java", "Kafka producer/consumer configuration", "When Kafka connector used"],
            ["SecurityConfig.java", "OAuth2/JWT security configuration", "When OAuth connector used"],
            ["AsyncConfig.java", "Thread pool for async operations", "When async processors used"],
            ["CacheConfig.java", "Redis cache configuration", "When ObjectStore/Redis used"],
            ["RestTemplateConfig.java", "HTTP client beans", "When HTTP request connector used"],
            ["WebClientConfig.java", "Reactive HTTP client beans", "When HTTP request connector used"],
            ["ResourceNotFoundException.java", "404 exception class", "Always"],
            ["BadRequestException.java", "400 exception class", "Always"],
            ["JsonUtil.java", "JSON serialization utility", "Always"],
            ["application.properties", "Database URLs, server port, logging, etc.", "Always"],
            ["Dockerfile", "Multi-stage Docker build for the generated app", "Always"],
            ["docker-compose.yml", "Container setup with required services (DB, MQ, etc.)", "Always"],
        ],
        col_widths=[2.0, 2.5, 2.0]
    )

    doc.add_heading('12.3 Running the Generated Project', level=2)
    doc.add_paragraph('After downloading and unzipping the generated project:')

    doc.add_heading('Option A: With Maven (requires Java 17+ and Maven installed)', level=3)
    add_code_block(doc, 'cd my-spring-app\nmvn spring-boot:run')

    doc.add_heading('Option B: With Docker (requires Docker installed)', level=3)
    add_code_block(doc, 'cd my-spring-app\ndocker compose up -d')

    add_info_box(doc,
        'The generated docker-compose.yml includes all required services (MySQL, RabbitMQ, Kafka, etc.) '
        'that your migrated application needs. Just run docker compose up and everything starts automatically!', "tip")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 13: API REFERENCE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('13. API Reference', level=1)
    doc.add_paragraph(
        'The migrator exposes several HTTP API endpoints. You can use these programmatically '
        '(with tools like curl, Postman, or your own code) instead of the web UI.'
    )

    endpoints = [
        {
            "method": "GET",
            "path": "/api/health",
            "desc": "Check if the server is running and healthy.",
            "request": "No body needed",
            "response": '{"status": "ok", "env": "production"}'
        },
        {
            "method": "GET",
            "path": "/api/llm/providers",
            "desc": "Get all available LLM providers with their models, tiers, and documentation links.",
            "request": "No body needed",
            "response": '{"anthropic": {"name": "Anthropic Claude", "models": [...], ...}, ...}'
        },
        {
            "method": "POST",
            "path": "/api/migrate",
            "desc": "The main migration endpoint. Accepts MuleSoft XML and returns a complete Spring Boot project.",
            "request": '{\n  "muleXmlFiles": [{"name": "flow.xml", "content": "<mule>...</mule>"}],\n  "dataweaveScripts": {"script1": "%dw 2.0 ..."},\n  "projectName": "my-app",\n  "groupId": "com.example",\n  "javaVersion": "17",\n  "llmConfig": {"enabled": true, "provider": "anthropic", "model": "...", "apiKey": "..."}\n}',
            "response": '{\n  "success": true,\n  "files": {"pom.xml": "...", "src/main/java/.../...": "..."},\n  "summary": {"flowsConverted": 5, "connectorsFound": [...], ...},\n  "llmValidation": {"overallScore": 8, ...}\n}'
        },
        {
            "method": "POST",
            "path": "/api/validate",
            "desc": "Standalone validation endpoint. Send already-generated files for AI review without re-migrating.",
            "request": '{\n  "files": {...generated files...},\n  "summary": {...migration summary...},\n  "llmConfig": {"provider": "openai", "model": "gpt-4o", "apiKey": "..."}\n}',
            "response": '{"success": true, "validation": {"overallScore": 7, "issues": [...], ...}}'
        },
        {
            "method": "POST",
            "path": "/api/migrate/download",
            "desc": "Download the generated project as a ZIP file.",
            "request": '{"files": {...}, "projectName": "my-app"}',
            "response": "Binary ZIP file download"
        },
        {
            "method": "POST",
            "path": "/api/convert/dataweave",
            "desc": "Convert a standalone DataWeave script to Java code.",
            "request": '{"script": "%dw 2.0\\noutput application/json\\n---\\npayload map ..."}',
            "response": '{"success": true, "result": {"java_code": "...", "imports": [...], "warnings": []}}'
        },
    ]

    for ep in endpoints:
        doc.add_heading(f'{ep["method"]} {ep["path"]}', level=2)
        doc.add_paragraph(ep["desc"])
        doc.add_paragraph('Request:', style='Heading 4')
        add_code_block(doc, ep["request"])
        doc.add_paragraph('Response:', style='Heading 4')
        add_code_block(doc, ep["response"])
        doc.add_paragraph()

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 14: PRODUCTION DEPLOYMENT
    # ════════════════════════════════════════════════════════════
    doc.add_heading('14. Production Deployment', level=1)
    doc.add_paragraph(
        'This section explains how to deploy the migrator tool itself for production use. '
        'This is different from running the generated Spring Boot project.'
    )

    doc.add_heading('14.1 Architecture Overview', level=2)
    doc.add_paragraph('The production deployment uses three layers:')
    add_styled_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Reverse Proxy", "Nginx", "SSL/TLS termination, rate limiting, static file caching, security headers"],
            ["Application Server", "Gunicorn", "Multi-worker Python WSGI server handling concurrent requests"],
            ["Application", "Flask", "The actual migrator web application and API"],
        ],
        col_widths=[1.3, 1.5, 3.7]
    )

    doc.add_heading('14.2 Gunicorn Configuration', level=2)
    doc.add_paragraph(
        'Gunicorn is the production server that replaces Flask\'s built-in development server. '
        'It can handle many requests simultaneously.'
    )
    add_styled_table(doc,
        ["Setting", "Default", "Description"],
        [
            ["workers", "CPU cores x 2 + 1", "Number of worker processes handling requests"],
            ["worker_class", "gthread", "Thread-based workers (good for I/O like LLM API calls)"],
            ["threads", "4", "Threads per worker for concurrent request handling"],
            ["timeout", "120 seconds", "Max time for a single request (LLM calls can be slow)"],
            ["bind", "0.0.0.0:5000", "Network address and port to listen on"],
            ["graceful_timeout", "30 seconds", "Time to finish in-progress requests on shutdown"],
        ],
        col_widths=[1.5, 1.8, 3.2]
    )

    doc.add_paragraph('All settings can be overridden with environment variables:')
    add_code_block(doc, 'GUNICORN_WORKERS=8\nGUNICORN_THREADS=4\nGUNICORN_TIMEOUT=120\nGUNICORN_BIND=0.0.0.0:5000')

    doc.add_heading('14.3 Nginx Configuration', level=2)
    doc.add_paragraph('Nginx sits in front of Gunicorn and provides:')
    add_bullet_list(doc, [
        'Rate Limiting: Prevents abuse by limiting requests per second (2/s for migrations, 10/s for APIs)',
        'Gzip Compression: Compresses responses to save bandwidth (level 6)',
        'Security Headers: X-Frame-Options, Content-Security-Policy, X-XSS-Protection, etc.',
        'Static File Caching: Serves CSS/JS files with 7-day browser cache',
        'SSL/TLS: HTTPS support with certificate configuration (template included)',
        'Request Size Limit: 50 MB maximum upload size',
        'Proxy Timeouts: 120 seconds for migration/validation endpoints',
    ])

    doc.add_heading('14.4 Docker Deployment', level=2)
    doc.add_paragraph('The Dockerfile creates a secure, minimal production image:')
    add_bullet_list(doc, [
        'Base Image: Python 3.12 slim (minimal footprint, around 500MB total)',
        'Non-root User: Runs as "appuser" for security (cannot modify system files)',
        'Health Check: Automatic monitoring every 30 seconds',
        'Optimized Layers: Dependencies installed before code for fast rebuilds',
    ])

    doc.add_heading('Docker Commands', level=3)
    add_styled_table(doc,
        ["Command", "What It Does"],
        [
            ["docker compose up -d", "Start the application in background"],
            ["docker compose --profile with-nginx up -d", "Start with Nginx reverse proxy"],
            ["docker compose logs -f app", "View real-time application logs"],
            ["docker compose down", "Stop all containers"],
            ["docker compose restart", "Restart all containers"],
            ["docker compose ps", "Show running container status"],
        ],
        col_widths=[3.5, 3.0]
    )

    doc.add_heading('14.5 Environment Variables Reference', level=2)
    add_styled_table(doc,
        ["Variable", "Default", "Description"],
        [
            ["FLASK_ENV", "production", "Flask environment mode"],
            ["PORT", "5000", "Application port number"],
            ["SECRET_KEY", "(must set)", "Random string for session security"],
            ["CORS_ORIGINS", "*", "Allowed origins for API requests"],
            ["GUNICORN_WORKERS", "4", "Number of Gunicorn worker processes"],
            ["GUNICORN_THREADS", "4", "Threads per worker"],
            ["GUNICORN_TIMEOUT", "120", "Request timeout in seconds"],
            ["NGINX_PORT", "80", "Nginx HTTP port"],
            ["NGINX_SSL_PORT", "443", "Nginx HTTPS port"],
            ["ANTHROPIC_API_KEY", "(empty)", "API key for Anthropic Claude"],
            ["OPENAI_API_KEY", "(empty)", "API key for OpenAI GPT models"],
            ["GOOGLE_API_KEY", "(empty)", "API key for Google Gemini"],
            ["DEEPSEEK_API_KEY", "(empty)", "API key for DeepSeek models"],
            ["GROQ_API_KEY", "(empty)", "API key for Groq (free)"],
        ],
        col_widths=[1.8, 1.2, 3.5]
    )

    add_info_box(doc,
        'API keys can be set either in the .env file (for Docker) or entered directly in the UI. '
        'The UI never stores API keys on disk for security.', "important")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 15: CONFIGURATION REFERENCE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('15. Configuration Reference', level=1)

    doc.add_heading('15.1 Project Files', level=2)
    add_styled_table(doc,
        ["File", "Location", "Purpose"],
        [
            ["app.py", "backend/", "Main Flask application with all routes and logic"],
            ["gunicorn.conf.py", "backend/", "Gunicorn production server configuration"],
            ["requirements.txt", "backend/", "Python package dependencies"],
            ["parser.py", "backend/migrator/", "MuleSoft XML parser (30+ namespace support)"],
            ["connector_mapper.py", "backend/migrator/", "Connector to Spring Boot dependency mapper"],
            ["flow_converter.py", "backend/migrator/", "MuleSoft flow to Java class converter"],
            ["spring_generator.py", "backend/migrator/", "Complete Spring Boot project generator"],
            ["dataweave_converter.py", "backend/migrator/", "DataWeave 2.0 to Java code converter"],
            ["llm_validator.py", "backend/migrator/", "Multi-provider LLM integration (6 providers)"],
            ["index.html", "backend/templates/", "Single-page web application HTML"],
            ["app.js", "backend/static/", "Frontend JavaScript application logic"],
            ["style.css", "backend/static/", "Dark-themed UI stylesheet"],
            ["Dockerfile", "project root", "Docker container image definition"],
            ["docker-compose.yml", "project root", "Docker Compose service orchestration"],
            ["nginx.conf", "nginx/", "Nginx reverse proxy configuration"],
            [".env.example", "project root", "Environment variable template"],
        ],
        col_widths=[1.8, 1.5, 3.2]
    )

    doc.add_heading('15.2 Frontend Settings Storage', level=2)
    doc.add_paragraph('The frontend stores non-sensitive settings in your browser\'s localStorage:')
    add_styled_table(doc,
        ["Key", "What It Stores", "Security"],
        [
            ["llm_enabled", "Whether AI validation is on or off", "Safe (not sensitive)"],
            ["llm_provider", "Selected AI provider name", "Safe (not sensitive)"],
            ["llm_model", "Selected AI model name", "Safe (not sensitive)"],
            ["llm_base_url", "Custom API URL (for Ollama)", "Safe (not sensitive)"],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    add_info_box(doc,
        'API keys are NEVER stored in localStorage, cookies, or anywhere on disk. '
        'You must re-enter your API key each time you open the application. '
        'This is a deliberate security measure.', "important")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 16: TROUBLESHOOTING GUIDE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('16. Troubleshooting Guide', level=1)

    doc.add_heading('16.1 Installation Issues', level=2)

    problems = [
        ("Problem: 'pip' command not found",
         "On macOS, use pip3 instead of pip. Similarly, use python3 instead of python.",
         "pip3 install -r requirements.txt\npython3 -m gunicorn -c gunicorn.conf.py app:app"),
        ("Problem: 'gunicorn' command not found",
         "Gunicorn might be installed in a directory not on your PATH. Use python3 -m gunicorn instead.",
         "python3 -m gunicorn -c gunicorn.conf.py app:app"),
        ("Problem: Address already in use (port 5000)",
         "Another process is using port 5000. Kill it first, or use a different port.",
         "# Kill existing process on port 5000\nkill $(lsof -t -i :5000)\n\n# Or start on a different port\nPORT=8080 python3 -m gunicorn -c gunicorn.conf.py app:app"),
        ("Problem: lxml installation fails",
         "lxml requires system libraries. Install them first.",
         "# macOS\nbrew install libxml2 libxslt\n\n# Ubuntu/Debian\nsudo apt-get install libxml2-dev libxslt1-dev"),
    ]

    for title, desc, solution in problems:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        doc.add_paragraph(desc)
        doc.add_paragraph('Solution:')
        add_code_block(doc, solution)
        doc.add_paragraph()

    doc.add_heading('16.2 Migration Issues', level=2)

    add_styled_table(doc,
        ["Issue", "Cause", "Solution"],
        [
            ["'Invalid XML' error", "XML has syntax errors or unsupported encoding", "Validate XML at xmlvalidation.com first"],
            ["Missing connectors in output", "Namespace not declared in XML root element", "Ensure all xmlns declarations are present"],
            ["Empty generated files", "XML has no flows or processors", "Check that your XML contains <flow> elements"],
            ["DataWeave warnings", "Complex patterns not fully converted", "Review the Java output and adjust manually"],
            ["Duplicate flow warnings", "Same flow name in multiple XML files", "Rename flows to be unique across files"],
            ["Large file timeout", "File exceeds processing capacity", "Split into smaller XML files or increase timeout"],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    doc.add_heading('16.3 LLM Validation Issues', level=2)
    add_styled_table(doc,
        ["Issue", "Cause", "Solution"],
        [
            ["'Missing API key' error", "No API key provided", "Enter API key in the AI Validation tab"],
            ["'Connection test failed'", "Invalid API key or network issue", "Verify key at provider's website, check internet"],
            ["Validation takes too long", "Large project or slow model", "Use a faster model (e.g., Groq Llama, Gemini Flash)"],
            ["'Module not installed' error", "LLM library not installed", "pip3 install anthropic openai google-generativeai"],
            ["Truncated validation results", "Context too large for model", "This is normal - model sees summarized code"],
        ],
        col_widths=[1.7, 2.0, 2.8]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # CHAPTER 17: GLOSSARY
    # ════════════════════════════════════════════════════════════
    doc.add_heading('17. Glossary', level=1)
    doc.add_paragraph('Every technical term used in this documentation, explained in plain English:')

    glossary = [
        ("API", "Application Programming Interface. A way for two programs to talk to each other over the internet."),
        ("API Key", "A secret password that lets your application use a service (like an AI model)."),
        ("Annotation", "In Java, a special marker starting with @ that adds behavior to code (e.g., @GetMapping means 'handle GET requests')."),
        ("Container", "A lightweight, portable package containing your application and everything it needs to run."),
        ("CORS", "Cross-Origin Resource Sharing. A security feature that controls which websites can access your API."),
        ("DataWeave", "MuleSoft's proprietary language for transforming data between formats."),
        ("Dependency", "A library or package your project needs to work (listed in pom.xml for Java projects)."),
        ("Docker", "A platform for building and running containers (packaged applications)."),
        ("Docker Compose", "A tool for defining and running multi-container Docker applications with a single command."),
        ("Endpoint", "A specific URL path in an API that accepts requests (e.g., /api/customers)."),
        ("Flask", "A lightweight Python web framework used to build the migrator application."),
        ("Flow", "In MuleSoft, a sequence of processing steps triggered by an event (like an HTTP request)."),
        ("Gunicorn", "A production-grade Python web server that can handle many concurrent requests."),
        ("HTTP", "Hypertext Transfer Protocol. The standard protocol for web communication."),
        ("IDE", "Integrated Development Environment. A program for writing code (e.g., VS Code, IntelliJ)."),
        ("JMS", "Java Message Service. A standard for sending messages between applications."),
        ("JSON", "JavaScript Object Notation. A lightweight text format for data exchange."),
        ("JPA", "Java Persistence API. A standard for connecting Java applications to databases."),
        ("Kafka", "A distributed event streaming platform for high-throughput messaging."),
        ("LLM", "Large Language Model. An AI system that can understand and generate text and code."),
        ("Maven", "A build tool for Java projects that manages dependencies and compilation."),
        ("Namespace", "In XML, a way to avoid naming conflicts by grouping elements under a unique prefix."),
        ("Nginx", "A high-performance web server commonly used as a reverse proxy."),
        ("OAuth", "An authentication standard that lets applications access resources on behalf of a user."),
        ("pom.xml", "Project Object Model. The Maven build configuration file listing all dependencies."),
        ("Processor", "In MuleSoft, an individual processing step within a flow (like a logger or database query)."),
        ("REST", "Representational State Transfer. A standard architecture for building web APIs."),
        ("Reverse Proxy", "A server that sits in front of your application to handle SSL, caching, and load balancing."),
        ("Spring Boot", "A popular Java framework for building production-ready web applications."),
        ("SSL/TLS", "Security protocols that encrypt data in transit (the 'S' in HTTPS)."),
        ("Sub-Flow", "In MuleSoft, a reusable group of processors that can be called from multiple flows."),
        ("XML", "Extensible Markup Language. A text format for structured data, used by MuleSoft configurations."),
        ("YAML", "Yet Another Markup Language. A human-readable text format for configuration files."),
        ("ZIP", "A compressed file format that bundles multiple files into one downloadable package."),
    ]

    for term, definition in glossary:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run2 = p.add_run(definition)
        run2.font.size = Pt(10)
        run2.font.name = 'Arial'

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # FINAL PAGE
    # ════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    end_table = doc.add_table(rows=1, cols=1)
    end_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = end_table.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, "1B3A5C")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n")

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("End of Documentation")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Arial'

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run("MuleSoft to Spring Boot Migrator v1.0")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)
    run.font.name = 'Arial'

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p4.add_run("March 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)
    run.font.name = 'Arial'

    p5 = cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p5.add_run("\n")

    cell.width = Inches(6.5)

    # ── Save ──
    output_dir = "/Users/harinadh/Documents/My code/mulesoft-to-springboot-migrator/docs"
    os.makedirs(output_dir, exist_ok=True)

    docx_path = os.path.join(output_dir, "MuleSoft-to-SpringBoot-Migrator-Documentation.docx")
    doc.save(docx_path)
    print(f"Word document saved: {docx_path}")
    return docx_path


if __name__ == "__main__":
    path = create_documentation()
    print(f"\nDocumentation created successfully at:\n{path}")
