#!/usr/bin/env python3
"""
Generate comprehensive platform documentation for MuleSoft to Spring Boot Migration Platform.
Creates a professional Word document covering all platform features, architecture, and usage.
"""
import os
import sys

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ════════════════════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc, text, style=None, bold=False, italic=False,
                            font_size=None, color=None, alignment=None,
                            space_before=None, space_after=None, font_name=None):
    """Add a formatted paragraph."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1B3A5C", stripe=True):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Arial'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if stripe and row_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    # Set column widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
    return table


def add_code_block(doc, code, language=""):
    """Add a code block with gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    # Add shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(code)
    run.font.size = Pt(8.5)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_info_box(doc, text, box_type="info"):
    """Add a colored info/warning/tip box."""
    colors = {
        "info": ("E8F4FD", "1565C0"),
        "warning": ("FFF8E1", "F57F17"),
        "tip": ("E8F5E9", "2E7D32"),
        "important": ("FCE4EC", "C62828"),
    }
    bg, fg = colors.get(box_type, colors["info"])
    labels = {"info": "INFO", "warning": "WARNING", "tip": "TIP", "important": "IMPORTANT"}

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, bg)

    p = cell.paragraphs[0]
    label_run = p.add_run(f"  {labels[box_type]}: ")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(*[int(fg[i:i+2], 16) for i in (0, 2, 4)])
    label_run.font.name = 'Arial'

    text_run = p.add_run(text)
    text_run.font.size = Pt(9)
    text_run.font.name = 'Arial'
    text_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for row in table.rows:
        for c in row.cells:
            c.width = Inches(6.5)
    return table


def add_bullet_list(doc, items, level=0):
    """Add bullet list items."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'Arial'


def add_numbered_list(doc, items):
    """Add numbered list items."""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'Arial'


def add_page_break(doc):
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# MAIN DOCUMENT CREATION
# ════════════════════════════════════════════════════════════════

def create_documentation():
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Configure Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15

    for i in range(1, 5):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Arial'
        h_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.styles['Heading 1'].font.size = Pt(22)
    doc.styles['Heading 2'].font.size = Pt(16)
    doc.styles['Heading 3'].font.size = Pt(13)
    doc.styles['Heading 4'].font.size = Pt(11)

    # ════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()

    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = title_table.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, "1B3A5C")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n")
    run.font.size = Pt(6)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("MuleSoft to Spring Boot Migration Platform")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Arial'

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run("Complete Technical Documentation & User Guide")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)
    run.font.name = 'Arial'

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p4.add_run("\n")
    run.font.size = Pt(6)

    # Version info below title box
    doc.add_paragraph()
    add_formatted_paragraph(doc, "Version 2.0", bold=True, font_size=14,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_formatted_paragraph(doc, "March 2026", font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            color=(0x66, 0x66, 0x66), font_name='Arial')
    add_formatted_paragraph(doc, "Multi-Page Web Platform for Automated Migration",
                            font_size=11, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            color=(0x88, 0x88, 0x88), font_name='Arial')

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════
    doc.add_heading('Table of Contents', level=1)
    add_formatted_paragraph(doc, "This document contains the following sections:",
                            font_size=10, space_after=6, font_name='Arial')

    toc_items = [
        "1. Executive Summary",
        "2. Platform Architecture",
        "3. Installation & Setup Guide",
        "4. Pages & Features",
        "5. API Reference",
        "6. Core Migration Modules",
        "7. Backend Services",
        "8. Frontend Architecture",
        "9. Security",
        "10. Testing Guide",
        "11. Deployment Guide",
        "12. Troubleshooting",
        "13. Appendices",
    ]
    for item in toc_items:
        add_formatted_paragraph(doc, item, font_size=11, space_before=4,
                                space_after=4, font_name='Arial', bold=True,
                                color=(0x1B, 0x3A, 0x5C))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_heading('1.1 Purpose & Motivation', level=2)
    add_formatted_paragraph(
        doc,
        "The MuleSoft to Spring Boot Migration Platform is a comprehensive web-based tool designed to "
        "automate the conversion of MuleSoft ESB applications into Spring Boot microservices. "
        "Organizations frequently face the challenge of migrating from proprietary integration platforms "
        "to open-source frameworks to reduce licensing costs, increase flexibility, and leverage modern "
        "cloud-native architectures.",
        font_size=10, space_after=6, font_name='Arial'
    )
    add_formatted_paragraph(
        doc,
        "This platform addresses the migration challenge by providing an end-to-end workflow: upload "
        "MuleSoft XML configurations, convert DataWeave transformation scripts, generate Spring Boot "
        "projects with proper dependency management, validate the generated code using Large Language "
        "Model (LLM) agents, generate OpenAPI/Swagger documentation, push to GitHub repositories, "
        "and build deployable artifacts (JAR, WAR, Docker images) -- all from a single web interface.",
        font_size=10, space_after=6, font_name='Arial'
    )

    doc.add_heading('1.2 Key Features', level=2)
    add_bullet_list(doc, [
        "Multi-page web platform built with Flask Blueprints for modular, scalable architecture",
        "Inline code editing with toggle edit mode, save, and discard capabilities",
        "GitHub integration for pushing generated projects directly to repositories via Git Data API",
        "Build system supporting JAR, WAR, and Docker image generation with multi-platform support",
        "Swagger/OpenAPI 3.0 specification generation from MuleSoft XML, RAML, or migration results",
        "LLM-powered code review and validation with multi-provider support (OpenAI, Anthropic, Google, Ollama)",
        "Drag-and-drop multi-file XML upload with project settings management",
        "DataWeave to Java conversion engine for transformation script migration",
        "Per-file and bulk validation with scoring and detailed feedback",
        "Cross-page state management using browser localStorage for seamless workflow",
        "Server-Sent Events (SSE) for real-time streaming build output",
        "Download generated projects as ZIP archives",
    ])

    doc.add_heading('1.3 Technology Stack', level=2)
    add_styled_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Backend", "Python 3.8+ / Flask", "Web framework and REST API"],
            ["Frontend", "HTML5 / CSS3 / JavaScript", "Responsive single-page-like UI"],
            ["Templating", "Jinja2", "Server-side HTML rendering with template inheritance"],
            ["Code Generation", "Custom Python modules", "MuleSoft XML parsing and Spring Boot code generation"],
            ["LLM Integration", "OpenAI, Anthropic, Google, Ollama", "Code review, validation, and enhancement"],
            ["GitHub API", "PyGithub / REST API", "Repository management and file push"],
            ["Build Tools", "Maven / Docker", "JAR, WAR, and Docker image builds"],
            ["API Docs", "OpenAPI 3.0 / Swagger", "API specification generation"],
            ["Package Mgmt", "pip / requirements.txt", "Python dependency management"],
        ],
        col_widths=[1.2, 2.0, 3.3]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 2. PLATFORM ARCHITECTURE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('2. Platform Architecture', level=1)

    doc.add_heading('2.1 Multi-Page Architecture (Flask Blueprints)', level=2)
    add_formatted_paragraph(
        doc,
        "The platform is organized using Flask Blueprints, which provide a modular structure where "
        "each major feature area is encapsulated in its own blueprint with dedicated routes, templates, "
        "and static assets. This design allows independent development, testing, and maintenance of each "
        "feature module.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_styled_table(doc,
        ["Blueprint", "URL Prefix", "Module File", "Responsibility"],
        [
            ["main", "/", "blueprints/main.py", "Dashboard, health check, static pages"],
            ["migration", "/migrate, /api/migrate", "blueprints/migration.py", "Core migration workflow, file upload, code generation, validation, download"],
            ["swagger", "/swagger, /api/swagger", "blueprints/swagger.py", "OpenAPI/Swagger generation from XML, RAML, or migration results"],
            ["github_bp", "/github, /api/github", "blueprints/github_bp.py", "GitHub connectivity, repo browsing, branch management, file push"],
            ["build", "/build, /api/build", "blueprints/build.py", "Build pipeline for JAR, WAR, Docker; prerequisites check; SSE streaming"],
            ["settings_bp", "/settings", "blueprints/settings_bp.py", "Project defaults, LLM configuration, token management"],
        ],
        col_widths=[1.0, 1.5, 1.7, 2.3]
    )

    add_info_box(doc, "Each blueprint registers both page routes (returning HTML) and API routes "
                 "(returning JSON) under its own URL prefix, keeping concerns cleanly separated.", "info")

    doc.add_heading('2.2 Directory Structure', level=2)
    add_code_block(doc, """mulesoft-to-springboot-migrator/
|-- app.py                      # Flask application factory & blueprint registration
|-- blueprints/
|   |-- __init__.py
|   |-- main.py                 # Dashboard & health endpoints
|   |-- migration.py            # Migration workflow API
|   |-- swagger.py              # Swagger/OpenAPI generation
|   |-- github_bp.py            # GitHub integration
|   |-- build.py                # Build pipeline & SSE streaming
|   |-- settings_bp.py          # Settings page
|-- services/
|   |-- github_service.py       # GitHub API wrapper (PyGithub)
|   |-- build_service.py        # Maven/Docker build orchestration
|-- migrator/
|   |-- parser.py               # MuleSoft XML parser (30+ namespaces)
|   |-- connector_mapper.py     # Connector & dependency mapping
|   |-- dataweave_converter.py  # DataWeave to Java conversion
|   |-- spring_generator.py     # Spring Boot project generation
|   |-- llm_agent.py            # LLM agent with triple fallback
|   |-- llm_validator.py        # Multi-provider code validation
|   |-- swagger_generator.py    # OpenAPI 3.0 spec generation
|-- templates/
|   |-- base.html               # Base template with sidebar & nav
|   |-- dashboard.html          # Dashboard page
|   |-- migrate.html            # Migration page
|   |-- swagger.html            # Swagger generator page
|   |-- github.html             # GitHub integration page
|   |-- build.html              # Build & test page
|   |-- settings.html           # Settings page
|   |-- architecture.html       # Architecture diagrams (protected)
|-- static/
|   |-- css/
|   |   |-- base.css            # Core styles, design system
|   |   |-- dashboard.css       # Dashboard-specific styles
|   |   |-- migration.css       # Migration page styles
|   |   |-- swagger.css         # Swagger page styles
|   |   |-- github.css          # GitHub page styles
|   |   |-- build.css           # Build page styles
|   |   |-- settings.css        # Settings page styles
|   |-- js/
|   |   |-- base.js             # Core utilities, stores, API helpers
|   |   |-- dashboard.js        # Dashboard logic
|   |   |-- migration.js        # Migration workflow logic
|   |   |-- swagger.js          # Swagger generation logic
|   |   |-- github.js           # GitHub integration logic
|   |   |-- build.js            # Build pipeline logic
|   |   |-- settings.js         # Settings management logic
|-- requirements.txt            # Python dependencies
|-- Dockerfile                  # Docker build configuration
|-- docs/                       # Documentation""", "text")

    doc.add_heading('2.3 Base Template System', level=2)
    add_formatted_paragraph(
        doc,
        "The platform uses Jinja2 template inheritance with a single base template (base.html) "
        "that provides the shared layout, navigation sidebar, loading screen, and common scripts. "
        "All page templates extend this base and override specific blocks.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_styled_table(doc,
        ["Block Name", "Purpose", "Override Required"],
        [
            ["title", "Sets the browser tab title for the page", "Yes"],
            ["extra_css", "Includes page-specific CSS stylesheets", "Optional"],
            ["loading_content", "Custom loading animation text shown during page load", "Optional"],
            ["page_header", "Page title and description displayed in the header area", "Yes"],
            ["content", "Main page content area", "Yes"],
            ["extra_js", "Includes page-specific JavaScript files", "Optional"],
        ],
        col_widths=[1.5, 3.5, 1.5]
    )

    add_code_block(doc, """<!-- Example: migrate.html extending base.html -->
{% extends "base.html" %}

{% block title %}Migration - MuleSoft to Spring Boot{% endblock %}

{% block extra_css %}
<link rel="stylesheet" href="/static/css/migration.css">
{% endblock %}

{% block page_header %}
<h1>MuleSoft Migration</h1>
<p>Upload XML configurations and generate Spring Boot projects</p>
{% endblock %}

{% block content %}
  <!-- Migration-specific content here -->
{% endblock %}

{% block extra_js %}
<script src="/static/js/migration.js"></script>
{% endblock %}""", "html")

    doc.add_heading('2.4 Cross-Page State Management', level=2)
    add_formatted_paragraph(
        doc,
        "The platform maintains state across pages using browser localStorage through two primary "
        "store patterns. This approach ensures that sensitive data like API keys and tokens never "
        "persist on the server, while migration results and settings are available across all pages.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "MigrationStore", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_formatted_paragraph(
        doc,
        "Stores migration results including generated files, metadata, validation scores, and "
        "conversion history. Used by the Migration, Swagger, GitHub, and Build pages to share data.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "SettingsStore", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_formatted_paragraph(
        doc,
        "Stores user preferences including project defaults (name, group ID, Java version), "
        "LLM configuration (provider, model, API key), and UI preferences. Shared across all pages.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Token Management", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_bullet_list(doc, [
        "GitHub Personal Access Tokens are stored in browser localStorage only",
        "LLM API keys (OpenAI, Anthropic, Google) are stored in browser localStorage only",
        "Tokens are passed per-request via HTTP headers or request body parameters",
        "The server never stores or logs any tokens or API keys",
        "Users can clear all stored tokens from the Settings page",
    ])

    add_info_box(doc, "All sensitive credentials (GitHub tokens, LLM API keys) are stored exclusively "
                 "in the browser's localStorage and are transmitted only when needed for specific API "
                 "calls. The server operates statelessly with respect to credentials.", "important")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 3. INSTALLATION & SETUP GUIDE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('3. Installation & Setup Guide', level=1)

    doc.add_heading('3.1 Prerequisites', level=2)
    add_styled_table(doc,
        ["Prerequisite", "Version", "Required", "Notes"],
        [
            ["Python", "3.8+", "Yes", "Core runtime for the platform"],
            ["pip", "Latest", "Yes", "Python package manager"],
            ["JDK", "17+", "For builds", "Required only for JAR/WAR builds"],
            ["Maven", "3.6+", "For builds", "Required only for JAR/WAR builds"],
            ["Docker", "20.10+", "For Docker builds", "Required only for Docker image builds"],
            ["Git", "2.x", "No", "Optional, GitHub integration uses REST API"],
        ],
        col_widths=[1.2, 1.0, 1.0, 3.3]
    )

    doc.add_heading('3.2 Installation Steps', level=2)
    add_numbered_list(doc, [
        "Clone the repository from your source control system",
        "Navigate to the project root directory",
        "Install Python dependencies using pip",
        "Run the application",
    ])

    add_code_block(doc, """# Clone the repository
git clone <repository-url>
cd mulesoft-to-springboot-migrator

# Install dependencies
pip install -r requirements.txt

# Run the application
python app.py""", "bash")

    add_info_box(doc, "It is recommended to use a Python virtual environment (venv) to avoid "
                 "dependency conflicts with other Python projects on your system.", "tip")

    add_code_block(doc, """# Create and activate virtual environment
python -m venv venv
source venv/bin/activate    # macOS/Linux
venv\\Scripts\\activate       # Windows

# Then install dependencies
pip install -r requirements.txt""", "bash")

    doc.add_heading('3.3 Configuration', level=2)
    add_formatted_paragraph(
        doc,
        "The platform uses environment variables for server-side configuration. All variables have "
        "sensible defaults and are optional for local development.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_styled_table(doc,
        ["Variable", "Default", "Description"],
        [
            ["PORT", "5015", "HTTP port the Flask server listens on"],
            ["FLASK_ENV", "development", "Flask environment mode (development or production)"],
            ["CORS_ORIGINS", "*", "Comma-separated list of allowed CORS origins"],
            ["SECRET_KEY", "auto-generated", "Flask secret key for session management"],
            ["ARCH_USERNAME", "admin", "Username for the Architecture page HTTP Basic Auth"],
            ["ARCH_PASSWORD", "admin", "Password for the Architecture page HTTP Basic Auth"],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    add_code_block(doc, """# Example: Set environment variables before running
export PORT=8080
export FLASK_ENV=production
export SECRET_KEY="your-secret-key-here"
export ARCH_USERNAME="myuser"
export ARCH_PASSWORD="mypassword"
python app.py""", "bash")

    doc.add_heading('3.4 Running the Application', level=2)
    add_formatted_paragraph(
        doc,
        "After installation, start the application with the following command. The server will start "
        "on the configured port (default 5015) and be accessible at http://localhost:5015.",
        font_size=10, space_after=6, font_name='Arial'
    )
    add_code_block(doc, """python app.py""", "bash")
    add_formatted_paragraph(
        doc,
        "You should see output similar to:",
        font_size=10, space_after=4, font_name='Arial'
    )
    add_code_block(doc, """ * Serving Flask app 'app'
 * Debug mode: on
 * Running on http://0.0.0.0:5015
Press CTRL+C to quit""", "text")

    doc.add_heading('3.5 Production Deployment', level=2)
    add_formatted_paragraph(
        doc,
        "For production environments, use Gunicorn as the WSGI server instead of Flask's built-in "
        "development server. Gunicorn provides better performance, process management, and stability.",
        font_size=10, space_after=6, font_name='Arial'
    )
    add_code_block(doc, """# Install Gunicorn
pip install gunicorn

# Run with Gunicorn (4 workers, bind to all interfaces)
gunicorn -w 4 -b 0.0.0.0:5015 app:app --timeout 120""", "bash")

    add_info_box(doc, "The --timeout 120 flag is important because migration and build operations "
                 "can take significant time, especially for large MuleSoft projects.", "warning")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 4. PAGES & FEATURES
    # ════════════════════════════════════════════════════════════
    doc.add_heading('4. Pages & Features', level=1)

    # 4.1 Dashboard
    doc.add_heading('4.1 Dashboard (/)', level=2)
    add_formatted_paragraph(
        doc,
        "The Dashboard serves as the central landing page and navigation hub for the platform. "
        "It provides quick-access cards to all major features and displays real-time migration "
        "statistics pulled from the browser's localStorage.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Features:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Quick action cards linking to Migration, Swagger, GitHub, Build, and Settings pages",
        "Migration status statistics showing total files migrated, success rate, and last migration date",
        "Animated stat counters that increment on page load for visual engagement",
        "Responsive card grid layout that adapts to different screen sizes",
        "Status indicators showing connection state for GitHub and LLM providers",
    ])

    add_info_box(doc, "Dashboard statistics are computed from localStorage data. If the browser's "
                 "localStorage is cleared, the statistics will reset to zero.", "info")

    # 4.2 Migration
    doc.add_heading('4.2 Migration (/migrate)', level=2)
    add_formatted_paragraph(
        doc,
        "The Migration page is the core feature of the platform, providing the complete workflow "
        "for converting MuleSoft XML configurations into Spring Boot projects. It features a "
        "two-panel layout with input controls on the left and generated output on the right.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Input Panel", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_bullet_list(doc, [
        "Multi-file XML upload with drag-and-drop support and file list management",
        "DataWeave script management: add, edit, and remove .dwl scripts for conversion",
        "Project settings: project name, group ID, artifact ID, Java version selection (11, 17, 21)",
        "LLM provider selection dropdown with API key input field",
        "Migrate button to trigger the conversion process",
    ])

    add_formatted_paragraph(doc, "Output Panel", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_bullet_list(doc, [
        "File tree browser displaying the generated Spring Boot project structure",
        "Search/filter functionality to quickly locate files in the generated tree",
        "Code viewer with syntax highlighting for the selected file",
        "Inline code editing: toggle edit mode to modify generated code directly in the browser",
        "Save and discard buttons for edited files, with unsaved change indicators",
        "Per-file LLM validation: validate individual files with score and feedback",
        "Bulk LLM validation: validate all generated files at once",
        "Code review scores displayed as badges (0-100 scale with color coding)",
        "Detailed LLM feedback with suggestions for improvement",
    ])

    add_formatted_paragraph(doc, "Cross-Page Actions", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_bullet_list(doc, [
        "Download as ZIP: packages all generated files into a downloadable archive",
        "Generate Swagger: navigates to the Swagger page with migration data pre-loaded",
        "Push to GitHub: navigates to the GitHub page with migration data ready to push",
        "Build JAR/WAR: navigates to the Build page to compile the generated project",
    ])

    # 4.3 Swagger
    doc.add_heading('4.3 Swagger / OpenAPI Generator (/swagger)', level=2)
    add_formatted_paragraph(
        doc,
        "The Swagger page generates OpenAPI 3.0 specifications from multiple input sources. "
        "The generated specifications document the REST APIs that the migrated Spring Boot "
        "application will expose, including endpoints, request/response schemas, and error codes.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Input Sources:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "From MuleSoft XML: parse HTTP listener and APIkit router configurations directly from XML files",
        "From RAML content: convert existing RAML API definitions into OpenAPI 3.0 format",
        "From migration results: generate specifications from previously completed migration data stored in localStorage",
    ])

    add_formatted_paragraph(doc, "Output Options:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Preview the generated specification in a formatted YAML viewer",
        "Download as YAML file (.yaml)",
        "Download as JSON file (.json)",
        "Copy the specification content to clipboard",
    ])

    # 4.4 GitHub
    doc.add_heading('4.4 GitHub Integration (/github)', level=2)
    add_formatted_paragraph(
        doc,
        "The GitHub page enables direct integration with GitHub repositories, allowing users to "
        "push generated Spring Boot projects without needing Git installed locally. The platform "
        "uses the GitHub Git Data API to create file trees and commits programmatically.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Connection:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Connect using a GitHub Personal Access Token (PAT)",
        "Token is stored in browser localStorage and sent per-request",
        "Displays authenticated user information upon successful connection",
    ])

    add_formatted_paragraph(doc, "Repository Management:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Browse organizations the authenticated user belongs to",
        "List repositories for the user or selected organization",
        "Create new repositories (public or private) directly from the platform",
        "Select existing branches or create new branches for the push target",
    ])

    add_formatted_paragraph(doc, "File Push:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Push all generated migration files to the selected repository and branch",
        "Uses the Git Data API (tree/blob/commit creation) instead of local Git commands",
        "Supports large file sets without local disk operations",
        "Displays push progress and result with commit URL link",
    ])

    add_formatted_paragraph(doc, "Required Token Scopes:", bold=True, font_size=10, font_name='Arial')
    add_styled_table(doc,
        ["Scope", "Required", "Purpose"],
        [
            ["repo", "Yes", "Full control of private and public repositories; required for all operations"],
            ["workflow", "Optional", "Allows push to repositories with GitHub Actions workflows configured"],
            ["delete_repo", "Optional", "Allows deleting repositories from the platform interface"],
        ],
        col_widths=[1.5, 1.0, 4.0]
    )

    # 4.5 Build & Test
    doc.add_heading('4.5 Build & Test (/build)', level=2)
    add_formatted_paragraph(
        doc,
        "The Build page provides a complete build pipeline for compiling and packaging the "
        "generated Spring Boot project. It supports multiple output formats and streams build "
        "output in real-time using Server-Sent Events (SSE) for a terminal-like experience.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Prerequisites Check:", bold=True, font_size=10, font_name='Arial')
    add_formatted_paragraph(
        doc,
        "Before building, the platform verifies the availability of required tools on the server. "
        "Each prerequisite is checked independently with version detection.",
        font_size=10, space_after=4, font_name='Arial'
    )
    add_bullet_list(doc, [
        "Java (JDK 17+): required for JAR and WAR builds",
        "Docker: required for Docker image builds",
        "Maven: required for all build types (JAR, WAR, test execution)",
    ])

    add_formatted_paragraph(doc, "Build Targets:", bold=True, font_size=10, font_name='Arial')

    add_styled_table(doc,
        ["Target", "Command", "Output", "Requirements"],
        [
            ["Build JAR", "mvn clean package", "Executable JAR file", "Java 17+, Maven"],
            ["Build WAR", "mvn clean package (modified pom.xml)", "WAR file for servlet containers", "Java 17+, Maven"],
            ["Build Docker", "docker build", "Docker image", "Docker"],
            ["Run Tests", "mvn test", "Test results and reports", "Java 17+, Maven"],
        ],
        col_widths=[1.2, 2.0, 1.8, 1.5]
    )

    add_formatted_paragraph(doc, "Docker Multi-Platform Support:", bold=True, font_size=10,
                            font_name='Arial', space_before=8)
    add_formatted_paragraph(
        doc,
        "The platform supports building Docker images for multiple target platforms, each with "
        "an appropriate base image optimized for that environment.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_styled_table(doc,
        ["Platform ID", "Architecture", "Base Image", "Use Case"],
        [
            ["linux-amd64", "linux/amd64", "eclipse-temurin:17-jre-alpine", "Standard Linux servers and cloud VMs"],
            ["ubuntu", "linux/amd64", "eclipse-temurin:17-jre-jammy", "Ubuntu-based deployments"],
            ["redhat", "linux/amd64", "registry.access.redhat.com/ubi9/openjdk-17-runtime", "Red Hat / OpenShift environments"],
            ["mac-arm", "linux/arm64", "eclipse-temurin:17-jre-alpine", "Apple Silicon Macs (M1/M2/M3)"],
            ["windows", "linux/amd64", "eclipse-temurin:17-jre-windowsservercore", "Windows container hosts"],
        ],
        col_widths=[1.0, 1.0, 2.8, 1.7]
    )

    add_formatted_paragraph(doc, "SSE Streaming Build Output:", bold=True, font_size=10,
                            font_name='Arial', space_before=8)
    add_formatted_paragraph(
        doc,
        "Build output is streamed to the browser in real-time using Server-Sent Events (SSE). "
        "The build terminal in the UI displays Maven and Docker output as it happens, with "
        "color-coded log levels (INFO, WARNING, ERROR). On successful completion, the built "
        "artifact (JAR, WAR, or Docker image tag) is available for download directly from the browser.",
        font_size=10, space_after=6, font_name='Arial'
    )

    # 4.6 Settings
    doc.add_heading('4.6 Settings (/settings)', level=2)
    add_formatted_paragraph(
        doc,
        "The Settings page allows users to configure project defaults, LLM integration, and "
        "manage stored tokens. All settings are persisted in browser localStorage and applied "
        "as defaults across all other pages.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Project Defaults:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Project Name: default name for generated Spring Boot projects",
        "Group ID: Maven group identifier (e.g., com.example)",
        "Java Version: target Java version (11, 17, or 21)",
        "Spring Boot Version: target Spring Boot framework version",
        "Packaging: default packaging format (JAR or WAR)",
    ])

    add_formatted_paragraph(doc, "LLM Configuration:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Enable/Disable: toggle LLM-powered features on or off globally",
        "Provider: select from OpenAI, Anthropic, Google Gemini, or Ollama (local)",
        "Model: choose the specific model for the selected provider",
        "API Key: enter and store the API key for the selected provider",
        "Base URL: custom endpoint URL (required for Ollama, optional for others)",
        "Test Connection: verify the LLM provider is reachable and the API key is valid",
    ])

    add_formatted_paragraph(doc, "GitHub Token Management:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "View stored GitHub token status (connected or not connected)",
        "Update or replace the stored Personal Access Token",
        "Clear the stored token to disconnect from GitHub",
    ])

    # 4.7 Architecture
    doc.add_heading('4.7 Architecture (/architecture)', level=2)
    add_formatted_paragraph(
        doc,
        "The Architecture page provides enterprise-grade interactive SVG diagrams of the platform's "
        "internal architecture. It features four switchable diagram views with zoom/pan controls, "
        "animated data flow connections, and a professional dark theme. This page is protected with "
        "HTTP Basic Authentication to restrict access to authorized team members.",
        font_size=10, space_after=6, font_name='Arial'
    )
    add_bullet_list(doc, [
        "Protected with HTTP Basic Auth using ARCH_USERNAME and ARCH_PASSWORD environment variables",
        "Architecture Layers view — shows frontend, backend, services, and data tiers",
        "Migration Pipeline view — visualizes the XML-to-Spring Boot transformation stages",
        "Technology Stack view — displays all frameworks, libraries, and tools used",
        "Data Flow view — traces request routing through blueprints, services, and modules",
        "Interactive controls: zoom in/out, pan, fit-to-screen, animated flowing connections",
    ])

    add_info_box(doc, "Default credentials for the Architecture page are admin/admin. Change these "
                 "by setting the ARCH_USERNAME and ARCH_PASSWORD environment variables in production.",
                 "warning")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 5. API REFERENCE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('5. API Reference', level=1)
    add_formatted_paragraph(
        doc,
        "The following table documents all HTTP endpoints exposed by the platform. Page endpoints "
        "return rendered HTML, while API endpoints return JSON responses. All API endpoints use "
        "the /api/ prefix convention.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_styled_table(doc,
        ["Method", "Endpoint", "Description", "Request Body", "Response"],
        [
            ["GET", "/", "Dashboard page", "-", "HTML"],
            ["GET", "/migrate", "Migration page", "-", "HTML"],
            ["POST", "/api/migrate", "Run migration", "multipart/form-data (XML files, settings)", "JSON: generated files, metadata"],
            ["POST", "/api/validate", "LLM code review", "JSON: files, provider, api_key", "JSON: scores, feedback per file"],
            ["POST", "/api/migrate/download", "Download ZIP", "JSON: files object", "application/zip"],
            ["POST", "/api/convert/dataweave", "Convert DataWeave", "JSON: dwl_content", "JSON: java_code"],
            ["GET", "/api/llm/providers", "List LLM providers", "-", "JSON: providers array"],
            ["GET", "/swagger", "Swagger page", "-", "HTML"],
            ["POST", "/api/swagger/from-xml", "Generate from XML", "multipart/form-data (XML file)", "JSON: openapi spec"],
            ["POST", "/api/swagger/from-raml", "Generate from RAML", "JSON: raml_content", "JSON: openapi spec"],
            ["POST", "/api/swagger/from-migration", "Generate from migration", "JSON: migration files", "JSON: openapi spec"],
            ["POST", "/api/swagger/download", "Download spec", "JSON: spec, format", "YAML or JSON file"],
            ["GET", "/github", "GitHub page", "-", "HTML"],
            ["POST", "/api/github/connect", "Connect to GitHub", "JSON: token", "JSON: user info"],
            ["GET", "/api/github/orgs", "List organizations", "Header: X-GitHub-Token", "JSON: orgs array"],
            ["GET", "/api/github/repos", "List repositories", "Header: X-GitHub-Token, Query: org", "JSON: repos array"],
            ["POST", "/api/github/repos/create", "Create repository", "JSON: name, private, org, token", "JSON: repo info"],
            ["GET", "/api/github/repos/<owner>/<repo>/branches", "List branches", "Header: X-GitHub-Token", "JSON: branches array"],
            ["POST", "/api/github/repos/<owner>/<repo>/branches/create", "Create branch", "JSON: branch_name, source, token", "JSON: branch info"],
            ["POST", "/api/github/push", "Push files", "JSON: owner, repo, branch, files, token, message", "JSON: commit info"],
            ["GET", "/build", "Build page", "-", "HTML"],
            ["POST", "/api/build/check-prereqs", "Check prerequisites", "-", "JSON: java, maven, docker status"],
            ["POST", "/api/build/jar", "Build JAR", "JSON: files, project settings", "JSON: build_id"],
            ["POST", "/api/build/war", "Build WAR", "JSON: files, project settings", "JSON: build_id"],
            ["POST", "/api/build/docker", "Build Docker", "JSON: files, platform, settings", "JSON: build_id"],
            ["POST", "/api/test/start", "Run tests", "JSON: files, project settings", "JSON: build_id"],
            ["GET", "/api/build/<id>/stream", "SSE build stream", "-", "text/event-stream"],
            ["GET", "/api/build/<id>/artifact", "Download artifact", "-", "application/octet-stream"],
            ["POST", "/api/build/<id>/cleanup", "Cleanup build", "-", "JSON: status"],
            ["GET", "/settings", "Settings page", "-", "HTML"],
            ["GET", "/api/health", "Health check", "-", "JSON: {status: ok}"],
            ["GET", "/architecture", "Architecture page (Basic Auth)", "-", "HTML"],
        ],
        col_widths=[0.6, 2.2, 1.4, 1.3, 1.0]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 6. CORE MIGRATION MODULES
    # ════════════════════════════════════════════════════════════
    doc.add_heading('6. Core Migration Modules', level=1)
    add_formatted_paragraph(
        doc,
        "The migration engine consists of several specialized Python modules located in the migrator/ "
        "directory. Each module handles a specific aspect of the MuleSoft to Spring Boot conversion "
        "pipeline.",
        font_size=10, space_after=6, font_name='Arial'
    )

    doc.add_heading('6.1 MuleSoft XML Parser (parser.py)', level=2)
    add_formatted_paragraph(
        doc,
        "The XML parser is responsible for reading MuleSoft configuration XML files and extracting "
        "structured data about flows, connectors, transformations, and error handlers. It supports "
        "over 30 MuleSoft XML namespaces covering the full range of MuleSoft components.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Supported Namespaces (30+):", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "http://www.mulesoft.org/schema/mule/core - Core MuleSoft elements (flows, sub-flows, error handlers)",
        "http://www.mulesoft.org/schema/mule/http - HTTP listener and request connectors",
        "http://www.mulesoft.org/schema/mule/db - Database connectors (select, insert, update, delete, stored procedures)",
        "http://www.mulesoft.org/schema/mule/ee/core - Enterprise edition transforms (DataWeave)",
        "http://www.mulesoft.org/schema/mule/apikit - APIkit router and configuration",
        "http://www.mulesoft.org/schema/mule/vm - VM queues for async messaging",
        "http://www.mulesoft.org/schema/mule/jms - JMS messaging connectors",
        "http://www.mulesoft.org/schema/mule/file - File read/write connectors",
        "http://www.mulesoft.org/schema/mule/ftp - FTP/SFTP connectors",
        "http://www.mulesoft.org/schema/mule/email - SMTP/IMAP email connectors",
        "http://www.mulesoft.org/schema/mule/os - ObjectStore for caching/state",
        "http://www.mulesoft.org/schema/mule/validation - Validation module",
    ])

    add_formatted_paragraph(doc, "Extracted Source Types (15+):", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "HTTP Listeners with path, method, port, and response configuration",
        "HTTP Request connectors with URL, method, headers, and query parameters",
        "Database operations (SELECT, INSERT, UPDATE, DELETE, stored procedure calls)",
        "DataWeave transformations with input/output MIME types",
        "APIkit router configurations with flow mappings",
        "Error handlers (on-error-propagate, on-error-continue) with type matching",
        "Flow references and sub-flow definitions",
        "Logger components with level and message extraction",
        "Set-payload and set-variable operations",
        "Choice routers (conditional branching) with when/otherwise clauses",
        "For-each and parallel for-each iterators",
        "Scatter-gather parallel processing blocks",
        "Try-catch blocks with error handler mappings",
        "VM publish/consume for async communication patterns",
        "Object Store operations for caching and distributed state",
    ])

    doc.add_heading('6.2 Connector Mapper (connector_mapper.py)', level=2)
    add_formatted_paragraph(
        doc,
        "The connector mapper translates MuleSoft connectors and their configurations into equivalent "
        "Spring Boot dependencies, annotations, and code patterns. It maintains a comprehensive "
        "mapping database covering 30+ MuleSoft connectors and 50+ error types.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Connector Mapping Examples:", bold=True, font_size=10, font_name='Arial')
    add_styled_table(doc,
        ["MuleSoft Connector", "Spring Boot Equivalent", "Dependencies"],
        [
            ["http:listener", "@RestController + @RequestMapping", "spring-boot-starter-web"],
            ["http:request", "RestTemplate / WebClient", "spring-boot-starter-web / spring-boot-starter-webflux"],
            ["db:select", "JdbcTemplate.query()", "spring-boot-starter-jdbc"],
            ["db:insert / db:update", "JdbcTemplate.update()", "spring-boot-starter-jdbc"],
            ["jms:publish / jms:consume", "JmsTemplate / @JmsListener", "spring-boot-starter-activemq"],
            ["vm:publish / vm:consume", "ApplicationEventPublisher / @EventListener", "spring-boot-starter"],
            ["file:read / file:write", "Files.read / Files.write (java.nio)", "No additional dependency"],
            ["email:send", "JavaMailSender", "spring-boot-starter-mail"],
            ["os:store / os:retrieve", "CacheManager / @Cacheable", "spring-boot-starter-cache"],
        ],
        col_widths=[1.5, 2.3, 2.7]
    )

    add_formatted_paragraph(doc, "Error Type Mapping (50+ types):", bold=True, font_size=10,
                            font_name='Arial', space_before=8)
    add_formatted_paragraph(
        doc,
        "The mapper converts MuleSoft error types to Spring Boot exception classes and HTTP status codes. "
        "For example, HTTP:UNAUTHORIZED maps to HttpClientErrorException with status 401, "
        "DB:CONNECTIVITY maps to DataAccessException, and VALIDATION:INVALID_BOOLEAN maps to "
        "IllegalArgumentException with a 400 Bad Request response.",
        font_size=10, space_after=6, font_name='Arial'
    )

    doc.add_heading('6.3 DataWeave Converter (dataweave_converter.py)', level=2)
    add_formatted_paragraph(
        doc,
        "The DataWeave converter translates MuleSoft DataWeave 2.0 transformation scripts into "
        "equivalent Java code. DataWeave is MuleSoft's proprietary expression language for data "
        "transformation, and converting it to standard Java requires understanding both the DataWeave "
        "syntax and the equivalent Java streaming, mapping, and transformation APIs.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Conversion Capabilities:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "DataWeave map/filter/reduce operations to Java Stream API equivalents",
        "DataWeave object/array construction to Java Map/List builders",
        "DataWeave string interpolation to Java String.format() or concatenation",
        "DataWeave type coercion to Java type casting and parsing",
        "DataWeave conditional expressions to Java ternary operators or if/else",
        "DataWeave null-safe navigation (?.) to Java Optional patterns",
        "DataWeave date/time functions to java.time API equivalents",
        "DataWeave payload references to method parameter or request body access",
    ])

    add_code_block(doc, """// Example: DataWeave to Java conversion
// DataWeave input:
//   %dw 2.0
//   output application/json
//   ---
//   payload.items map ((item) -> {
//       name: upper(item.name),
//       total: item.price * item.quantity
//   })

// Generated Java output:
List<Map<String, Object>> result = items.stream()
    .map(item -> {
        Map<String, Object> mapped = new LinkedHashMap<>();
        mapped.put("name", item.get("name").toString().toUpperCase());
        mapped.put("total",
            ((Number) item.get("price")).doubleValue() *
            ((Number) item.get("quantity")).doubleValue());
        return mapped;
    })
    .collect(Collectors.toList());""", "java")

    doc.add_heading('6.4 Spring Boot Generator (spring_generator.py)', level=2)
    add_formatted_paragraph(
        doc,
        "The Spring Boot generator takes the parsed MuleSoft data and connector mappings to produce "
        "a complete, runnable Spring Boot project. It generates all necessary files including the "
        "main application class, REST controllers with OpenAPI/Swagger annotations (@Tag, @Operation, "
        "@ApiResponse), service classes, repository interfaces, configuration files, pom.xml with "
        "correct dependencies (including springdoc-openapi-starter-webmvc-ui, database drivers, "
        "Spring Boot Actuator), application.properties, and logback-spring.xml for structured logging.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Generated Project Structure:", bold=True, font_size=10, font_name='Arial')
    add_code_block(doc, """generated-project/
|-- pom.xml                          # Maven POM with all required dependencies
|-- src/
|   |-- main/
|   |   |-- java/com/example/project/
|   |   |   |-- Application.java         # @SpringBootApplication main class
|   |   |   |-- controller/
|   |   |   |   |-- *Controller.java      # REST controllers from HTTP listeners
|   |   |   |-- service/
|   |   |   |   |-- *Service.java         # Business logic from MuleSoft flows
|   |   |   |-- repository/
|   |   |   |   |-- *Repository.java      # Data access from DB connectors
|   |   |   |-- model/
|   |   |   |   |-- *Model.java           # Data models / DTOs
|   |   |   |-- config/
|   |   |   |   |-- AppConfig.java        # Bean configurations
|   |   |   |-- exception/
|   |   |   |   |-- GlobalExceptionHandler.java  # Error handler mappings
|   |   |   |-- util/
|   |   |       |-- DataWeaveHelper.java  # Converted DW transformations
|   |   |-- resources/
|   |       |-- application.properties    # Spring Boot configuration
|   |       |-- logback-spring.xml        # Structured logging configuration
|   |-- test/
|       |-- java/com/example/project/
|           |-- ApplicationTests.java     # Context load smoke test
|           |-- controller/
|               |-- *ControllerTest.java  # @WebMvcTest controller tests""", "text")

    doc.add_heading('6.5 LLM Agent (llm_agent.py)', level=2)
    add_formatted_paragraph(
        doc,
        "The LLM agent enhances the migration output by using Large Language Models to improve "
        "code quality, fix potential issues, and add proper documentation. It implements a triple "
        "fallback strategy to maximize reliability across different LLM providers.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Triple Fallback Strategy:", bold=True, font_size=10, font_name='Arial')
    add_numbered_list(doc, [
        "Primary: Use the user-selected LLM provider (OpenAI, Anthropic, Google, or Ollama) with the specified model and API key",
        "Secondary: If the primary provider fails or times out, fall back to an alternative provider from the configured options",
        "Tertiary: If all external LLM providers fail, use rule-based heuristics and templates to generate acceptable output",
    ])

    add_formatted_paragraph(doc, "LLM-Enhanced Operations:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Code improvement: refactoring suggestions, best practice enforcement, and Spring Boot idiom alignment",
        "Documentation generation: adding Javadoc comments, README content, and inline code comments",
        "Error handling enhancement: suggesting proper exception hierarchies and error response patterns",
        "Test generation: creating unit test skeletons for the generated controllers and services",
        "Dependency validation: verifying that all required Maven dependencies are included in pom.xml",
    ])

    doc.add_heading('6.6 LLM Validator (llm_validator.py)', level=2)
    add_formatted_paragraph(
        doc,
        "The LLM validator performs code review on generated files using the configured LLM provider. "
        "It scores each file on a 0-100 scale and provides detailed feedback on code quality, "
        "correctness, and adherence to Spring Boot best practices.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Validation Criteria:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Compilation readiness: will the code compile without errors",
        "Spring Boot conventions: proper use of annotations, dependency injection, and configuration",
        "REST API design: correct HTTP method usage, path variables, request/response handling",
        "Error handling: proper exception handling and error response structure",
        "Code organization: class structure, package organization, and naming conventions",
        "Security practices: input validation, SQL injection prevention, proper authentication patterns",
    ])

    add_formatted_paragraph(doc, "Supported Providers:", bold=True, font_size=10, font_name='Arial')
    add_styled_table(doc,
        ["Provider", "Models", "API Key Required"],
        [
            ["OpenAI", "gpt-4o, gpt-4o-mini, gpt-4-turbo, gpt-3.5-turbo", "Yes"],
            ["Anthropic", "claude-sonnet-4-20250514, claude-3.5-sonnet, claude-3-haiku", "Yes"],
            ["Google Gemini", "gemini-1.5-pro, gemini-1.5-flash, gemini-1.0-pro", "Yes"],
            ["Ollama (Local)", "llama3, codellama, mistral, mixtral (user-configured)", "No (local server)"],
        ],
        col_widths=[1.5, 3.0, 2.0]
    )

    doc.add_heading('6.7 Swagger Generator (swagger_generator.py)', level=2)
    add_formatted_paragraph(
        doc,
        "The Swagger generator creates OpenAPI 3.0 specifications from MuleSoft XML configurations "
        "or RAML API definitions. It extracts endpoint information, request/response schemas, "
        "error responses, and metadata to produce a complete API specification.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Generation Sources:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "MuleSoft XML: parses HTTP listener configurations, APIkit routers, and flow definitions to infer REST endpoints",
        "RAML: converts RAML 1.0 API definitions including resource types, traits, data types, and examples",
        "Migration results: uses the generated Spring Boot controller code to extract endpoint annotations and request/response types",
    ])

    add_formatted_paragraph(doc, "Generated Specification Features:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "OpenAPI 3.0.3 compliant specification",
        "Endpoint paths with HTTP methods, parameters, and request bodies",
        "Response schemas with status codes (200, 400, 404, 500)",
        "Component schemas for reusable data types",
        "Server configuration with base URL and environment variables",
        "Tags for logical endpoint grouping",
    ])

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 7. BACKEND SERVICES
    # ════════════════════════════════════════════════════════════
    doc.add_heading('7. Backend Services', level=1)

    doc.add_heading('7.1 GitHub Service (github_service.py)', level=2)
    add_formatted_paragraph(
        doc,
        "The GitHub service provides a Python wrapper around the GitHub REST API using the PyGithub "
        "library. It handles authentication, repository management, branch operations, and file "
        "pushing using the Git Data API for efficient bulk file uploads.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Key Operations:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Authentication: validate tokens and retrieve authenticated user information",
        "Organization listing: enumerate organizations the user belongs to",
        "Repository browsing: list repositories with filtering by owner or organization",
        "Repository creation: create new public or private repositories with optional initialization",
        "Branch management: list branches, create new branches from a source reference",
        "File push via Git Data API: create blobs, build tree objects, create commits, and update references",
    ])

    add_formatted_paragraph(doc, "Git Data API Push Flow:", bold=True, font_size=10, font_name='Arial')
    add_numbered_list(doc, [
        "Retrieve the current commit SHA and tree SHA for the target branch",
        "Create blob objects for each file to be pushed (base64 encoded content)",
        "Build a new tree object referencing all the created blobs with their file paths",
        "Create a new commit object pointing to the new tree with the parent commit reference",
        "Update the branch reference to point to the new commit SHA",
    ])

    add_info_box(doc, "The Git Data API approach avoids the need for local Git installation, "
                 "clone operations, or disk-based repositories. All operations happen through "
                 "HTTP API calls, making it suitable for server-side and containerized deployments.", "tip")

    doc.add_heading('7.2 Build Service (build_service.py)', level=2)
    add_formatted_paragraph(
        doc,
        "The build service orchestrates the compilation and packaging of generated Spring Boot "
        "projects. It manages temporary build directories, executes Maven and Docker commands, "
        "streams output via SSE, and provides artifact download capabilities.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Build Pipeline:", bold=True, font_size=10, font_name='Arial')
    add_numbered_list(doc, [
        "Receive generated files and project configuration from the frontend",
        "Create a temporary build directory with the complete project structure",
        "Write all source files, pom.xml, and configuration to the temporary directory",
        "For WAR builds: modify pom.xml to change packaging type and add servlet dependencies",
        "Execute the Maven or Docker build command as a subprocess",
        "Capture stdout and stderr and stream each line as an SSE event to the connected client",
        "On completion, store the build artifact path and make it available for download",
        "Cleanup temporary build directories after artifact download or timeout",
    ])

    add_formatted_paragraph(doc, "Docker Multi-Platform Build:", bold=True, font_size=10, font_name='Arial')
    add_formatted_paragraph(
        doc,
        "For Docker builds, the service generates a Dockerfile appropriate for the selected target "
        "platform, configures the base image, copies the built JAR file, and sets the entry point. "
        "The build uses docker buildx for cross-platform compilation when the target architecture "
        "differs from the host.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 8. FRONTEND ARCHITECTURE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('8. Frontend Architecture', level=1)

    doc.add_heading('8.1 Base Layer (base.js / base.css)', level=2)
    add_formatted_paragraph(
        doc,
        "The base layer provides shared functionality and styles used by all pages. It is loaded "
        "on every page through the base.html template and provides core utilities that page-specific "
        "scripts build upon.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "base.js Core Utilities:", bold=True, font_size=10, font_name='Arial')
    add_styled_table(doc,
        ["Utility", "Type", "Description"],
        [
            ["Toast", "UI Component", "Non-blocking notification system (success, error, warning, info) with auto-dismiss"],
            ["Modal", "UI Component", "Reusable modal dialog for confirmations, forms, and content display"],
            ["Sidebar", "UI Component", "Collapsible navigation sidebar with active page highlighting"],
            ["MigrationStore", "Data Store", "localStorage wrapper for migration results with get/set/clear methods"],
            ["SettingsStore", "Data Store", "localStorage wrapper for user settings and preferences"],
            ["apiCall()", "HTTP Helper", "Fetch wrapper with error handling, loading states, and JSON parsing"],
            ["connectSSE()", "HTTP Helper", "Server-Sent Events connection manager for streaming build output"],
            ["formatFileSize()", "Formatter", "Human-readable file size formatting (bytes to KB/MB)"],
            ["debounce()", "Utility", "Function debouncing for search inputs and resize handlers"],
        ],
        col_widths=[1.5, 1.2, 3.8]
    )

    add_formatted_paragraph(doc, "base.css Core Styles:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "CSS custom properties (variables) for consistent theming across all pages",
        "Sidebar layout with collapsible navigation and responsive behavior",
        "Toast notification positioning and animation styles",
        "Modal overlay and content container styles",
        "Loading screen overlay with spinner animation",
        "Typography scale, spacing utilities, and color system",
        "Form element styling (inputs, selects, buttons, toggles)",
        "Card and panel component styles used across multiple pages",
    ])

    doc.add_heading('8.2 Page-Specific JavaScript and CSS', level=2)
    add_formatted_paragraph(
        doc,
        "Each page has dedicated JavaScript and CSS files that handle page-specific logic and "
        "styling. These files are loaded only on their respective pages through the extra_js "
        "and extra_css Jinja2 template blocks.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_styled_table(doc,
        ["Page", "JS File", "CSS File", "Key Responsibilities"],
        [
            ["Dashboard", "dashboard.js", "dashboard.css", "Stat counter animation, action card navigation, status checks"],
            ["Migration", "migration.js", "migration.css", "File upload, code editing, tree browser, validation, download"],
            ["Swagger", "swagger.js", "swagger.css", "Source selection, spec generation, preview, download/copy"],
            ["GitHub", "github.js", "github.css", "Token auth, org/repo browsing, branch management, file push"],
            ["Build", "build.js", "build.css", "Prerequisites check, build triggers, SSE terminal, artifact download"],
            ["Settings", "settings.js", "settings.css", "Form persistence, LLM test connection, token management"],
        ],
        col_widths=[1.0, 1.2, 1.2, 3.1]
    )

    doc.add_heading('8.3 Design System', level=2)
    add_formatted_paragraph(
        doc,
        "The platform uses a comprehensive design system built on CSS custom properties, enabling "
        "consistent theming and easy customization. The design system covers colors, typography, "
        "spacing, breakpoints, and component patterns.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "CSS Custom Properties:", bold=True, font_size=10, font_name='Arial')
    add_code_block(doc, """:root {
    /* Primary Colors */
    --primary: #1B3A5C;
    --primary-light: #2C5282;
    --primary-dark: #102A43;

    /* Accent Colors */
    --accent: #4299E1;
    --success: #48BB78;
    --warning: #ECC94B;
    --danger: #F56565;

    /* Neutral Colors */
    --bg-primary: #FFFFFF;
    --bg-secondary: #F7FAFC;
    --text-primary: #1A202C;
    --text-secondary: #718096;
    --border-color: #E2E8F0;

    /* Spacing */
    --spacing-xs: 0.25rem;
    --spacing-sm: 0.5rem;
    --spacing-md: 1rem;
    --spacing-lg: 1.5rem;
    --spacing-xl: 2rem;

    /* Typography */
    --font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    --font-mono: 'Fira Code', 'Courier New', monospace;
}""", "css")

    add_formatted_paragraph(doc, "Dark Theme:", bold=True, font_size=10, font_name='Arial')
    add_formatted_paragraph(
        doc,
        "The design system includes a dark theme that can be toggled by the user. Dark theme "
        "overrides the CSS custom properties with dark-appropriate values, ensuring all components "
        "automatically adapt without page-specific style changes.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Responsive Breakpoints:", bold=True, font_size=10, font_name='Arial')
    add_styled_table(doc,
        ["Breakpoint", "Min Width", "Target Devices"],
        [
            ["Mobile", "0px", "Phones and small screens"],
            ["Tablet", "768px", "Tablets and small laptops"],
            ["Desktop", "1024px", "Standard desktop monitors"],
            ["Wide", "1440px", "Large monitors and ultrawide displays"],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 9. SECURITY
    # ════════════════════════════════════════════════════════════
    doc.add_heading('9. Security', level=1)
    add_formatted_paragraph(
        doc,
        "The platform implements several security measures to protect user credentials and "
        "ensure safe operation in both development and production environments.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_formatted_paragraph(doc, "Credential Storage:", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_bullet_list(doc, [
        "All tokens (GitHub PAT, LLM API keys) are stored exclusively in the browser's localStorage",
        "The server never persists, logs, or caches any user credentials",
        "Tokens are transmitted per-request via HTTP headers (X-GitHub-Token) or JSON request body fields",
        "Users can view, update, and clear stored tokens from the Settings page at any time",
    ])

    add_formatted_paragraph(doc, "Authentication & Access Control:", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_bullet_list(doc, [
        "The Architecture page (/architecture) is protected with HTTP Basic Authentication",
        "Credentials are configured via ARCH_USERNAME and ARCH_PASSWORD environment variables",
        "All other pages are accessible without authentication (designed for internal/team use)",
    ])

    add_formatted_paragraph(doc, "CORS Configuration:", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_bullet_list(doc, [
        "Cross-Origin Resource Sharing (CORS) is configured via the CORS_ORIGINS environment variable",
        "Default configuration allows all origins (*) for development convenience",
        "In production, restrict CORS_ORIGINS to specific trusted domains",
    ])

    add_formatted_paragraph(doc, "Upload Limits:", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_bullet_list(doc, [
        "Maximum upload size is limited to 50MB to prevent abuse and memory exhaustion",
        "File type validation ensures only XML and DWL files are processed by the migration engine",
    ])

    add_formatted_paragraph(doc, "Output Sanitization:", bold=True, font_size=11,
                            color=(0x1B, 0x3A, 0x5C), font_name='Arial')
    add_bullet_list(doc, [
        "Generated code does not contain references to any external tool names or branding",
        "Output is clean, professional Spring Boot code following standard conventions",
        "Generated comments reference the migration process generically without mentioning specific tools",
    ])

    add_info_box(doc, "For production deployments, always set a strong SECRET_KEY, restrict "
                 "CORS_ORIGINS to your domain, change the default ARCH_USERNAME/ARCH_PASSWORD, "
                 "and serve the application behind HTTPS using a reverse proxy (nginx, Caddy).",
                 "important")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 10. TESTING GUIDE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('10. Testing Guide', level=1)

    doc.add_heading('10.1 Backend Testing with Flask Test Client', level=2)
    add_formatted_paragraph(
        doc,
        "The platform includes backend tests using Flask's built-in test client. Tests cover "
        "API endpoint behavior, migration logic, and service integrations.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_code_block(doc, """# Run all backend tests
python -m pytest tests/ -v

# Run specific test module
python -m pytest tests/test_migration.py -v

# Run with coverage report
python -m pytest tests/ --cov=migrator --cov=blueprints --cov-report=html""", "bash")

    add_formatted_paragraph(doc, "Test Categories:", bold=True, font_size=10, font_name='Arial')
    add_bullet_list(doc, [
        "Unit tests: individual module functions (parser, connector mapper, DataWeave converter)",
        "Integration tests: API endpoint testing with Flask test client",
        "Service tests: GitHub service and build service with mocked external dependencies",
        "End-to-end tests: complete migration workflow from XML upload to ZIP download",
    ])

    doc.add_heading('10.2 Build System Testing', level=2)
    add_formatted_paragraph(
        doc,
        "Testing the build system requires JDK 17 or later installed on the machine. Build tests "
        "verify that generated projects compile correctly and produce valid artifacts.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_code_block(doc, """# Verify Java installation
java -version

# Verify Maven installation
mvn -version

# Run build system tests
python -m pytest tests/test_build.py -v""", "bash")

    add_info_box(doc, "Build system tests create temporary directories, write project files, and "
                 "invoke Maven. They may take several minutes to complete depending on network speed "
                 "(Maven dependency downloads) and machine resources.", "info")

    doc.add_heading('10.3 Docker Testing', level=2)
    add_formatted_paragraph(
        doc,
        "Docker-related tests require Docker Desktop or Docker Engine running on the machine. "
        "These tests verify Docker image builds for each supported platform.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_code_block(doc, """# Verify Docker installation
docker --version

# Verify Docker is running
docker info

# Run Docker-specific tests
python -m pytest tests/test_docker_build.py -v""", "bash")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 11. DEPLOYMENT GUIDE
    # ════════════════════════════════════════════════════════════
    doc.add_heading('11. Deployment Guide', level=1)

    doc.add_heading('11.1 Docker Deployment', level=2)
    add_formatted_paragraph(
        doc,
        "The platform includes a Dockerfile for containerized deployment. The Docker image bundles "
        "the Python application with all dependencies and can optionally include JDK and Maven "
        "for build functionality.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_code_block(doc, """# Build the Docker image
docker build -t mulesoft-migrator:latest .

# Run the container
docker run -d \\
  -p 5015:5015 \\
  -e SECRET_KEY="your-production-secret" \\
  -e FLASK_ENV=production \\
  -e CORS_ORIGINS="https://yourdomain.com" \\
  -e ARCH_USERNAME="admin" \\
  -e ARCH_PASSWORD="secure-password" \\
  --name mulesoft-migrator \\
  mulesoft-migrator:latest""", "bash")

    doc.add_heading('11.2 Gunicorn Production Setup', level=2)
    add_formatted_paragraph(
        doc,
        "For non-containerized production deployments, use Gunicorn as the WSGI server with "
        "a reverse proxy (such as nginx) in front for HTTPS termination and static file serving.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_code_block(doc, """# Install Gunicorn
pip install gunicorn

# Run with Gunicorn
gunicorn \\
  --workers 4 \\
  --bind 0.0.0.0:5015 \\
  --timeout 120 \\
  --access-logfile /var/log/migrator/access.log \\
  --error-logfile /var/log/migrator/error.log \\
  app:app""", "bash")

    add_formatted_paragraph(doc, "Example nginx Reverse Proxy Configuration:", bold=True,
                            font_size=10, font_name='Arial')
    add_code_block(doc, """server {
    listen 443 ssl;
    server_name migrator.yourdomain.com;

    ssl_certificate /etc/ssl/certs/migrator.crt;
    ssl_certificate_key /etc/ssl/private/migrator.key;

    client_max_body_size 50M;

    location / {
        proxy_pass http://127.0.0.1:5015;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }

    # SSE support for build streaming
    location /api/build/ {
        proxy_pass http://127.0.0.1:5015;
        proxy_set_header Host $host;
        proxy_set_header Connection '';
        proxy_http_version 1.1;
        chunked_transfer_encoding off;
        proxy_buffering off;
        proxy_cache off;
    }
}""", "nginx")

    doc.add_heading('11.3 Production Checklist', level=2)
    add_formatted_paragraph(
        doc,
        "Before deploying to production, verify the following configuration items are in place:",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_styled_table(doc,
        ["Item", "Action", "Priority"],
        [
            ["SECRET_KEY", "Set a strong, random secret key (at least 32 characters)", "Critical"],
            ["FLASK_ENV", "Set to 'production' to disable debug mode and detailed errors", "Critical"],
            ["CORS_ORIGINS", "Restrict to your specific domain(s) instead of wildcard (*)", "High"],
            ["ARCH_USERNAME", "Change from default 'admin' to a unique username", "High"],
            ["ARCH_PASSWORD", "Change from default 'admin' to a strong password", "High"],
            ["HTTPS", "Configure SSL/TLS via reverse proxy (nginx, Caddy, or cloud LB)", "Critical"],
            ["Firewall", "Restrict direct access to port 5015; only allow reverse proxy", "High"],
            ["Logging", "Configure Gunicorn access and error log paths", "Medium"],
            ["Monitoring", "Set up health check monitoring on /api/health endpoint", "Medium"],
            ["Backups", "No server-side data to back up (all state is client-side)", "Low"],
        ],
        col_widths=[1.5, 3.5, 1.5]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 12. TROUBLESHOOTING
    # ════════════════════════════════════════════════════════════
    doc.add_heading('12. Troubleshooting', level=1)
    add_formatted_paragraph(
        doc,
        "The following table covers common issues encountered when running the platform and "
        "their recommended solutions.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_styled_table(doc,
        ["Issue", "Possible Cause", "Solution"],
        [
            ["Port already in use",
             "Another process is using port 5015",
             "Change port with PORT=8080 python app.py, or stop the conflicting process with lsof -i :5015"],
            ["Module not found",
             "Python dependencies not installed",
             "Run pip install -r requirements.txt in the project root; ensure virtual environment is activated"],
            ["LLM provider not showing",
             "Provider list API not responding",
             "Check that the Flask server is running; verify /api/llm/providers returns data in the browser"],
            ["GitHub push fails",
             "Invalid or expired Personal Access Token",
             "Generate a new PAT with 'repo' scope from GitHub Settings > Developer settings > Tokens"],
            ["GitHub push fails (403)",
             "Token lacks required permissions",
             "Ensure the token has 'repo' scope; for repos with Actions, also add 'workflow' scope"],
            ["Build fails without Java",
             "JDK not installed or not in PATH",
             "Install JDK 17+ and ensure 'java -version' and 'mvn -version' work from the terminal"],
            ["Docker build fails",
             "Docker daemon not running",
             "Start Docker Desktop (macOS/Windows) or start the Docker service (Linux: sudo systemctl start docker)"],
            ["Migration returns empty results",
             "XML files are not valid MuleSoft configurations",
             "Ensure uploaded files are MuleSoft XML configs with proper namespace declarations"],
            ["Large file upload fails",
             "File exceeds 50MB upload limit",
             "Split large MuleSoft projects into smaller batches of XML files"],
            ["SSE stream disconnects",
             "Network timeout or proxy buffering",
             "If behind a reverse proxy, disable buffering for /api/build/ routes (see nginx config above)"],
            ["DataWeave conversion incomplete",
             "Complex DataWeave expressions not supported",
             "Enable LLM enhancement to improve DataWeave conversion; manually review complex transformations"],
            ["Settings not persisting",
             "Browser localStorage is disabled or full",
             "Check browser privacy settings; ensure localStorage is enabled and not at capacity"],
        ],
        col_widths=[1.5, 1.8, 3.2]
    )

    add_page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 13. APPENDICES
    # ════════════════════════════════════════════════════════════
    doc.add_heading('13. Appendices', level=1)

    doc.add_heading('Appendix A: requirements.txt', level=2)
    add_formatted_paragraph(
        doc,
        "The following Python packages are required by the platform. Install them with "
        "pip install -r requirements.txt.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_code_block(doc, """flask>=2.3.0
flask-cors>=4.0.0
gunicorn>=21.2.0
requests>=2.31.0
PyGithub>=2.1.1
pyyaml>=6.0.1
lxml>=4.9.3
jsonschema>=4.19.0
openai>=1.3.0
anthropic>=0.7.0
google-generativeai>=0.3.0
python-docx>=0.8.11
Jinja2>=3.1.2
Werkzeug>=2.3.0""", "text")

    doc.add_heading('Appendix B: Complete File Listing', level=2)
    add_formatted_paragraph(
        doc,
        "The following table lists all primary source files in the platform with their purpose.",
        font_size=10, space_after=6, font_name='Arial'
    )

    add_styled_table(doc,
        ["File Path", "Purpose"],
        [
            ["app.py", "Flask application factory, blueprint registration, CORS setup"],
            ["blueprints/__init__.py", "Blueprint package initializer"],
            ["blueprints/main.py", "Dashboard routes, health check, architecture page"],
            ["blueprints/migration.py", "Migration API: upload, convert, validate, download"],
            ["blueprints/swagger.py", "Swagger generation: from XML, RAML, migration"],
            ["blueprints/github_bp.py", "GitHub API: connect, repos, branches, push"],
            ["blueprints/build.py", "Build API: JAR, WAR, Docker, tests, SSE streaming"],
            ["blueprints/settings_bp.py", "Settings page route"],
            ["services/github_service.py", "GitHub API wrapper using PyGithub"],
            ["services/build_service.py", "Build orchestration: Maven, Docker, cleanup"],
            ["migrator/parser.py", "MuleSoft XML parser (30+ namespaces)"],
            ["migrator/connector_mapper.py", "Connector-to-Spring-Boot mapping (30+ connectors)"],
            ["migrator/dataweave_converter.py", "DataWeave 2.0 to Java conversion"],
            ["migrator/spring_generator.py", "Spring Boot project file generation"],
            ["migrator/llm_agent.py", "LLM agent with triple fallback strategy"],
            ["migrator/llm_validator.py", "Multi-provider LLM code validation"],
            ["migrator/swagger_generator.py", "OpenAPI 3.0 specification generation"],
            ["templates/base.html", "Base Jinja2 template with sidebar navigation"],
            ["templates/dashboard.html", "Dashboard page template"],
            ["templates/migrate.html", "Migration page template"],
            ["templates/swagger.html", "Swagger generator page template"],
            ["templates/github.html", "GitHub integration page template"],
            ["templates/build.html", "Build and test page template"],
            ["templates/settings.html", "Settings page template"],
            ["templates/architecture.html", "Architecture diagrams page template"],
            ["static/css/base.css", "Core design system and shared styles"],
            ["static/js/base.js", "Core utilities: stores, API helpers, UI components"],
            ["requirements.txt", "Python package dependencies"],
            ["Dockerfile", "Docker build configuration"],
        ],
        col_widths=[2.5, 4.0]
    )

    doc.add_heading('Appendix C: Glossary', level=2)
    add_styled_table(doc,
        ["Term", "Definition"],
        [
            ["MuleSoft", "An integration platform (now part of Salesforce) that uses XML-based configuration for building APIs and integrations"],
            ["Spring Boot", "An opinionated Java framework for building production-ready applications with minimal configuration"],
            ["DataWeave", "MuleSoft's proprietary expression and transformation language for data manipulation"],
            ["Flask", "A lightweight Python web framework used as the backend for this platform"],
            ["Blueprint", "A Flask feature for organizing application routes and views into modular components"],
            ["Jinja2", "A Python templating engine used by Flask for rendering HTML pages"],
            ["OpenAPI / Swagger", "A specification for describing REST APIs, including endpoints, parameters, and schemas"],
            ["RAML", "RESTful API Modeling Language, a YAML-based format for describing REST APIs (used by MuleSoft)"],
            ["PAT", "Personal Access Token, a GitHub authentication credential used instead of passwords"],
            ["Git Data API", "A GitHub REST API for creating Git objects (blobs, trees, commits) programmatically"],
            ["SSE", "Server-Sent Events, a standard for pushing real-time updates from server to browser over HTTP"],
            ["LLM", "Large Language Model, an AI system used for code review and enhancement in this platform"],
            ["CORS", "Cross-Origin Resource Sharing, a browser security mechanism for controlling API access across domains"],
            ["WSGI", "Web Server Gateway Interface, a Python standard for web server to application communication"],
            ["Gunicorn", "A production-grade Python WSGI HTTP server, recommended for deploying Flask applications"],
            ["Maven", "A Java build automation tool used for compiling, testing, and packaging Spring Boot projects"],
            ["pom.xml", "The Maven Project Object Model file that defines project dependencies and build configuration"],
        ],
        col_widths=[1.8, 4.7]
    )

    # ════════════════════════════════════════════════════════════
    # SAVE DOCUMENT
    # ════════════════════════════════════════════════════════════
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "MuleSoft_SpringBoot_Migration_Platform_v2.docx")
    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    size_kb = file_size / 1024
    print(f"Documentation generated successfully!")
    print(f"  File: {output_path}")
    print(f"  Size: {size_kb:.1f} KB ({file_size:,} bytes)")
    return output_path


if __name__ == "__main__":
    create_documentation()
