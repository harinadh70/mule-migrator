#!/usr/bin/env python3
"""
Generate Trial & Error Installation Troubleshooting Guide v2
MuleSoft to Spring Boot Migration Platform
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc, text, bold=False, font_size=None, color=None,
                            alignment=None, space_before=None, space_after=None,
                            font_name='Arial'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(code)
    run.font.size = Pt(8.5)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_info_box(doc, text, box_color="E8F4FD", border_color="2196F3", label="INFO"):
    """Add a styled info/warning box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, box_color)
    # Add border via XML
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:color="{border_color}"/>'
        f'  <w:left w:val="single" w:sz="12" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    p = cell.paragraphs[0]
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_label.font.name = 'Arial'
    run_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_body = p.add_run(text)
    run_body.font.size = Pt(9)
    run_body.font.name = 'Arial'
    run_body.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Spacer after box
    doc.add_paragraph().paragraph_format.space_before = Pt(2)


def add_warning_box(doc, text):
    add_info_box(doc, text, box_color="FFF8E1", border_color="FF9800", label="WARNING")


def add_tip_box(doc, text):
    add_info_box(doc, text, box_color="E8F5E9", border_color="4CAF50", label="TIP")


def add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        run.font.name = 'Arial'
    return h


def add_issue_table(doc, error_text, cause, solution_lines):
    """Add a styled issue table with Error / Cause / Solution rows."""
    tbl = doc.add_table(rows=2 + len(solution_lines), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Set column widths
    for row in tbl.rows:
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(5.1)

    # Row 0 - Error
    set_cell_shading(tbl.cell(0, 0), "FFEBEE")
    set_cell_shading(tbl.cell(0, 1), "FFEBEE")
    _set_cell_text(tbl.cell(0, 0), "Error", bold=True, color=RGBColor(0xC6, 0x28, 0x28))
    _set_cell_text(tbl.cell(0, 1), error_text, font_name='Courier New', font_size=9)

    # Row 1 - Cause
    set_cell_shading(tbl.cell(1, 0), "FFF8E1")
    set_cell_shading(tbl.cell(1, 1), "FFF8E1")
    _set_cell_text(tbl.cell(1, 0), "Cause", bold=True, color=RGBColor(0xE6, 0x51, 0x00))
    _set_cell_text(tbl.cell(1, 1), cause)

    # Solution rows
    for i, sol in enumerate(solution_lines):
        row_idx = 2 + i
        set_cell_shading(tbl.cell(row_idx, 0), "E8F5E9")
        set_cell_shading(tbl.cell(row_idx, 1), "E8F5E9")
        label = "Solution" if i == 0 else ""
        _set_cell_text(tbl.cell(row_idx, 0), label, bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
        _set_cell_text(tbl.cell(row_idx, 1), sol)

    doc.add_paragraph().paragraph_format.space_before = Pt(4)


def _set_cell_text(cell, text, bold=False, color=None, font_name='Arial', font_size=9.5):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def add_env_table(doc, rows_data):
    """Add a table for environment variable reference."""
    tbl = doc.add_table(rows=1 + len(rows_data), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    headers = ["Variable", "Default", "Description"]
    for i, h in enumerate(headers):
        set_cell_shading(tbl.cell(0, i), "1A237E")
        _set_cell_text(tbl.cell(0, i), h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for r_idx, (var, default, desc) in enumerate(rows_data, start=1):
        bg = "F5F5F5" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx in range(3):
            set_cell_shading(tbl.cell(r_idx, c_idx), bg)
        _set_cell_text(tbl.cell(r_idx, 0), var, font_name='Courier New', font_size=9)
        _set_cell_text(tbl.cell(r_idx, 1), default, font_name='Courier New', font_size=9)
        _set_cell_text(tbl.cell(r_idx, 2), desc)

    doc.add_paragraph().paragraph_format.space_before = Pt(4)


def add_prereq_table(doc, rows_data):
    """Add a prerequisite check table."""
    tbl = doc.add_table(rows=1 + len(rows_data), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    headers = ["Prerequisite", "Required?", "Check Command", "Minimum Version"]
    for i, h in enumerate(headers):
        set_cell_shading(tbl.cell(0, i), "1A237E")
        _set_cell_text(tbl.cell(0, i), h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for r_idx, (name, req, cmd, ver) in enumerate(rows_data, start=1):
        bg = "F5F5F5" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx in range(4):
            set_cell_shading(tbl.cell(r_idx, c_idx), bg)
        _set_cell_text(tbl.cell(r_idx, 0), name)
        _set_cell_text(tbl.cell(r_idx, 1), req)
        _set_cell_text(tbl.cell(r_idx, 2), cmd, font_name='Courier New', font_size=9)
        _set_cell_text(tbl.cell(r_idx, 3), ver)

    doc.add_paragraph().paragraph_format.space_before = Pt(4)


# ---------------------------------------------------------------------------
# Main document generation
# ---------------------------------------------------------------------------

def create_document():
    doc = Document()

    # -- Default font styling --
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ===================================================================
    # COVER PAGE
    # ===================================================================
    for _ in range(6):
        doc.add_paragraph()

    add_formatted_paragraph(
        doc, "Trial & Error Installation",
        bold=True, font_size=28, color=(0x1A, 0x23, 0x7E),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0
    )
    add_formatted_paragraph(
        doc, "Troubleshooting Guide",
        bold=True, font_size=28, color=(0x1A, 0x23, 0x7E),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    # Divider line
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "1A237E")
    cell.text = ""
    cell.width = Inches(4)

    add_formatted_paragraph(
        doc, "MuleSoft to Spring Boot Migration Platform v2.0",
        bold=True, font_size=14, color=(0x42, 0x42, 0x42),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=8
    )
    add_formatted_paragraph(
        doc, "March 2026",
        font_size=12, color=(0x75, 0x75, 0x75),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    add_formatted_paragraph(
        doc, "Comprehensive guide to resolving installation, configuration, and runtime issues",
        font_size=10, color=(0x9E, 0x9E, 0x9E),
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    doc.add_page_break()

    # ===================================================================
    # TABLE OF CONTENTS (manual)
    # ===================================================================
    add_section_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1.  Quick Start",
        "2.  Prerequisites",
        "3.  Installation Errors & Fixes",
        "4.  Runtime Errors & Fixes",
        "5.  Platform-Specific Issues",
        "6.  Configuration Reference",
        "7.  Health Checks",
        "8.  Virtual Environment Setup (Recommended)",
        "9.  Docker Deployment",
        "10. Contact & Resources",
    ]
    for item in toc_items:
        p = add_formatted_paragraph(doc, item, font_size=11, color=(0x1A, 0x23, 0x7E), space_before=2, space_after=2)

    doc.add_page_break()

    # ===================================================================
    # 1. QUICK START
    # ===================================================================
    add_section_heading(doc, "1. Quick Start", level=1)
    add_formatted_paragraph(
        doc,
        "Get the MuleSoft to Spring Boot Migration Platform running in under 5 minutes.",
        font_size=10, space_after=6
    )

    add_tip_box(doc, "This is the fastest path. If any step fails, refer to the relevant troubleshooting section below.")

    steps = [
        ("Step 1 - Clone the repository", "git clone <repository-url>\ncd mulesoft-to-springboot-migrator"),
        ("Step 2 - Install dependencies", "pip3 install -r requirements.txt"),
        ("Step 3 - Start the application", "cd backend\npython3 app.py"),
        ("Step 4 - Open in browser", "http://localhost:5000"),
    ]
    for title, code in steps:
        add_formatted_paragraph(doc, title, bold=True, font_size=10, space_before=8, space_after=2)
        add_code_block(doc, code)

    add_info_box(doc, "The application will start on port 5000 by default. If that port is in use, set the PORT environment variable: PORT=5001 python3 app.py")

    doc.add_page_break()

    # ===================================================================
    # 2. PREREQUISITES
    # ===================================================================
    add_section_heading(doc, "2. Prerequisites", level=1)
    add_formatted_paragraph(
        doc,
        "The following software is needed to run the migration platform. Only Python and pip are strictly required.",
        font_size=10, space_after=8
    )

    add_prereq_table(doc, [
        ("Python 3.8+", "Required", "python3 --version", "3.8"),
        ("pip", "Required", "pip3 --version", "20.0+"),
        ("JDK 17+", "Optional", "java -version", "17"),
        ("Docker", "Optional", "docker --version", "20.10+"),
        ("Node.js", "Optional", "node --version", "Not needed for the app"),
    ])

    add_info_box(doc, "JDK is only needed if you want to compile/build the generated Spring Boot projects. Docker is only needed for containerized builds. Node.js is not required by the application itself.")

    doc.add_page_break()

    # ===================================================================
    # 3. INSTALLATION ERRORS & FIXES
    # ===================================================================
    add_section_heading(doc, "3. Installation Errors & Fixes", level=1)
    add_formatted_paragraph(
        doc,
        "Common errors encountered during installation and their solutions.",
        font_size=10, space_after=8
    )

    # 3.1
    add_section_heading(doc, "3.1  pip: command not found", level=2)
    add_issue_table(
        doc,
        error_text="bash: pip: command not found",
        cause="pip is not installed or not in PATH. Many systems ship pip as pip3.",
        solution_lines=[
            "Use pip3 instead of pip:\n  pip3 install -r requirements.txt",
            "Or invoke via Python module:\n  python3 -m pip install -r requirements.txt",
        ]
    )

    # 3.2
    add_section_heading(doc, "3.2  ModuleNotFoundError: No module named 'flask'", level=2)
    add_issue_table(
        doc,
        error_text="ModuleNotFoundError: No module named 'flask'",
        cause="Python dependencies have not been installed yet.",
        solution_lines=[
            "Install all required packages:\n  pip3 install -r requirements.txt",
        ]
    )

    # 3.3
    add_section_heading(doc, "3.3  ModuleNotFoundError: No module named 'lxml'", level=2)
    add_issue_table(
        doc,
        error_text="ModuleNotFoundError: No module named 'lxml'",
        cause="lxml requires native C libraries that may not be present on your system.",
        solution_lines=[
            "macOS: Install Xcode command line tools first:\n  xcode-select --install\n  pip3 install lxml",
            "Linux (Debian/Ubuntu):\n  sudo apt-get install libxml2-dev libxslt-dev python3-dev\n  pip3 install lxml",
            "Linux (RHEL/CentOS):\n  sudo yum install libxml2-devel libxslt-devel python3-devel\n  pip3 install lxml",
        ]
    )

    # 3.4
    add_section_heading(doc, "3.4  Address already in use (Port 5000/5001)", level=2)
    add_issue_table(
        doc,
        error_text="OSError: [Errno 48] Address already in use",
        cause="Another process is already listening on port 5000 (or 5001). On macOS, AirPlay Receiver uses port 5000 by default.",
        solution_lines=[
            "Kill the process occupying the port:\n  lsof -ti:5000 | xargs kill -9",
            "Or start the app on a different port:\n  PORT=5001 python3 app.py",
            "macOS: Disable AirPlay Receiver in System Settings > General > AirDrop & Handoff.",
        ]
    )

    # 3.5
    add_section_heading(doc, "3.5  ImportError: cannot import name 'GithubException'", level=2)
    add_issue_table(
        doc,
        error_text="ImportError: cannot import name 'GithubException' from 'github'",
        cause="The wrong GitHub package is installed. The package named 'github' is not the same as 'PyGithub'.",
        solution_lines=[
            "Uninstall the wrong package and install the correct one:\n  pip3 uninstall github\n  pip3 install PyGithub",
        ]
    )

    add_warning_box(doc, "The correct package is PyGithub (capital P, capital G), not github. Running 'pip install github' installs a completely different, incompatible package.")

    # 3.6
    add_section_heading(doc, "3.6  SSL / OpenSSL Warnings", level=2)
    add_issue_table(
        doc,
        error_text="NotOpenSSLWarning: urllib3 v2 only supports OpenSSL 1.1.1+",
        cause="Your Python installation is linked against LibreSSL or an older OpenSSL version.",
        solution_lines=[
            "This warning is safe to ignore. The application will function correctly.",
            "To suppress the warning, you can downgrade urllib3:\n  pip3 install urllib3<2",
        ]
    )

    # 3.7
    add_section_heading(doc, "3.7  Permission Denied Errors on macOS", level=2)
    add_issue_table(
        doc,
        error_text="ERROR: Could not install packages due to an EnvironmentError: [Errno 13] Permission denied",
        cause="Attempting to install packages into the system Python directory without administrator privileges.",
        solution_lines=[
            "Option A - Install for current user only:\n  pip3 install --user -r requirements.txt",
            "Option B - Use a virtual environment (recommended):\n  python3 -m venv venv\n  source venv/bin/activate\n  pip install -r requirements.txt",
        ]
    )

    add_tip_box(doc, "Using a virtual environment (Section 8) is the recommended approach. It avoids permission issues entirely and keeps your system Python clean.")

    doc.add_page_break()

    # ===================================================================
    # 4. RUNTIME ERRORS & FIXES
    # ===================================================================
    add_section_heading(doc, "4. Runtime Errors & Fixes", level=1)
    add_formatted_paragraph(
        doc,
        "Issues that occur after the application has started successfully.",
        font_size=10, space_after=8
    )

    # 4.1
    add_section_heading(doc, "4.1  LLM Providers List is Empty", level=2)
    add_issue_table(
        doc,
        error_text="Settings page shows no LLM providers / provider dropdown is empty",
        cause="This is expected behavior. LLM providers are loaded dynamically and API keys are stored in the browser's local storage only (never sent to the server for storage).",
        solution_lines=[
            "1. Navigate to the Settings page in the web UI.\n2. Select your LLM provider (OpenAI, Anthropic, Google, etc.).\n3. Enter your API key.\n4. Click Save. The key is stored in your browser only.",
        ]
    )

    # 4.2
    add_section_heading(doc, "4.2  GitHub Push Fails with 401", level=2)
    add_issue_table(
        doc,
        error_text="github.GithubException: 401 {\"message\": \"Bad credentials\"}",
        cause="The GitHub personal access token is invalid, expired, or has insufficient permissions.",
        solution_lines=[
            "1. Go to https://github.com/settings/tokens\n2. Generate a new token (classic) with the 'repo' scope.\n3. Copy the token and paste it in the application's GitHub settings.\n4. Ensure the token has not expired.",
        ]
    )

    # 4.3
    add_section_heading(doc, "4.3  GitHub Push Fails with 422", level=2)
    add_issue_table(
        doc,
        error_text="github.GithubException: 422 {\"message\": \"Reference already exists\"}",
        cause="A file conflict exists or branch protection rules are preventing the push.",
        solution_lines=[
            "Try pushing to a new branch name instead of the existing one.",
            "If the repository has branch protection rules, ensure your token has permission to push to the target branch, or use an unprotected branch.",
        ]
    )

    # 4.4
    add_section_heading(doc, "4.4  Build Fails: Unable to Locate Java Runtime", level=2)
    add_issue_table(
        doc,
        error_text="Error: Unable to locate a Java Runtime.\nBuild failed: Java/Maven not available",
        cause="JDK is not installed on the system. JDK is required only for building the generated Spring Boot projects.",
        solution_lines=[
            "macOS (Homebrew):\n  brew install openjdk@17\n  sudo ln -sfn $(brew --prefix)/opt/openjdk@17/libexec/openjdk.jdk /Library/Java/JavaVirtualMachines/openjdk-17.jdk",
            "Ubuntu/Debian:\n  sudo apt install openjdk-17-jdk",
            "Windows:\n  Download from https://adoptium.net/ and add to PATH.",
        ]
    )

    # 4.5
    add_section_heading(doc, "4.5  Docker Build Fails", level=2)
    add_issue_table(
        doc,
        error_text="Cannot connect to the Docker daemon. Is the docker daemon running?",
        cause="Docker is not installed or the Docker daemon is not running.",
        solution_lines=[
            "1. Install Docker Desktop from https://www.docker.com/products/docker-desktop\n2. Start Docker Desktop and wait for the engine to initialize.\n3. Verify: docker --version && docker info",
        ]
    )

    # 4.6
    add_section_heading(doc, "4.6  Migration Produces Empty Output", level=2)
    add_issue_table(
        doc,
        error_text="Migration completes but the output files are empty or contain only boilerplate.",
        cause="The input XML is invalid, not well-formed, or does not follow MuleSoft schema conventions.",
        solution_lines=[
            "Ensure your XML file:\n  - Is well-formed XML (no syntax errors).\n  - Contains the MuleSoft namespace: http://www.mulesoft.org/schema/mule/core\n  - Has at least one flow or sub-flow element.",
            "Test with a known-good sample file first to verify the tool is working correctly.",
        ]
    )

    # 4.7
    add_section_heading(doc, "4.7  Swagger Generation Fails", level=2)
    add_issue_table(
        doc,
        error_text="Swagger/OpenAPI generation returned empty or errored.",
        cause="The input XML has no HTTP listener endpoints defined. The Swagger generator requires at least one HTTP-based endpoint to produce an API specification.",
        solution_lines=[
            "Ensure your MuleSoft XML contains at least one <http:listener> element within a flow.\nExample:\n  <http:listener config-ref=\"HTTP_Listener\" path=\"/api/resource\" doc:name=\"Listener\"/>",
        ]
    )

    # 4.8
    add_section_heading(doc, "4.8  CORS Errors in Browser", level=2)
    add_issue_table(
        doc,
        error_text="Access to XMLHttpRequest at 'http://localhost:5000/api/...' from origin 'http://localhost:3000' has been blocked by CORS policy.",
        cause="The frontend is running on a different port than the backend, and CORS is not configured to allow cross-origin requests.",
        solution_lines=[
            "Set the CORS_ORIGINS environment variable before starting the backend:\n  CORS_ORIGINS=http://localhost:3000 python3 app.py",
            "Or allow all origins (development only):\n  CORS_ORIGINS=* python3 app.py",
        ]
    )

    add_warning_box(doc, "Never use CORS_ORIGINS=* in production. Always specify the exact allowed origins.")

    doc.add_page_break()

    # ===================================================================
    # 5. PLATFORM-SPECIFIC ISSUES
    # ===================================================================
    add_section_heading(doc, "5. Platform-Specific Issues", level=1)

    # 5.1 macOS
    add_section_heading(doc, "5.1  macOS", level=2)

    add_formatted_paragraph(doc, "Homebrew Python vs System Python", bold=True, font_size=10, space_before=6, space_after=2)
    add_formatted_paragraph(
        doc,
        "macOS ships with a system Python that should not be modified. If you installed Python via Homebrew, ensure your PATH prioritizes Homebrew's Python.",
        font_size=10, space_after=4
    )
    add_code_block(doc, "# Check which Python is active\nwhich python3\n\n# Should show /opt/homebrew/bin/python3 (Apple Silicon)\n# or /usr/local/bin/python3 (Intel Mac)")

    add_formatted_paragraph(doc, "Xcode Command Line Tools", bold=True, font_size=10, space_before=8, space_after=2)
    add_formatted_paragraph(
        doc,
        "Many Python packages with C extensions (lxml, cryptography, etc.) require Xcode command line tools.",
        font_size=10, space_after=4
    )
    add_code_block(doc, "xcode-select --install")

    add_formatted_paragraph(doc, "macOS Sequoia Sandbox Restrictions", bold=True, font_size=10, space_before=8, space_after=2)
    add_formatted_paragraph(
        doc,
        "macOS Sequoia (15.x) and later may impose additional sandbox restrictions on network access and file system operations. If you encounter unexpected permission errors, check System Settings > Privacy & Security for any blocked applications.",
        font_size=10, space_after=4
    )

    # 5.2 Windows
    add_section_heading(doc, "5.2  Windows", level=2)

    add_formatted_paragraph(doc, "Python Command Name", bold=True, font_size=10, space_before=6, space_after=2)
    add_formatted_paragraph(
        doc,
        "On Windows, use 'python' and 'pip' instead of 'python3' and 'pip3'.",
        font_size=10, space_after=4
    )
    add_code_block(doc, "python --version\npip install -r requirements.txt\npython app.py")

    add_formatted_paragraph(doc, "PowerShell vs CMD", bold=True, font_size=10, space_before=8, space_after=2)
    add_formatted_paragraph(
        doc,
        "Environment variable syntax differs between PowerShell and Command Prompt.",
        font_size=10, space_after=4
    )
    add_code_block(doc, "# PowerShell\n$env:PORT = \"5001\"\npython app.py\n\n# Command Prompt (CMD)\nset PORT=5001\npython app.py")

    add_formatted_paragraph(doc, "Line Ending Issues (CRLF)", bold=True, font_size=10, space_before=8, space_after=2)
    add_formatted_paragraph(
        doc,
        "Windows uses CRLF line endings which can cause issues with shell scripts and certain file parsers. Configure Git to handle line endings automatically.",
        font_size=10, space_after=4
    )
    add_code_block(doc, "git config --global core.autocrlf true")

    # 5.3 Linux
    add_section_heading(doc, "5.3  Linux", level=2)

    add_formatted_paragraph(doc, "Package Manager Differences", bold=True, font_size=10, space_before=6, space_after=2)
    add_code_block(doc, "# Debian/Ubuntu\nsudo apt update\nsudo apt install python3 python3-pip python3-venv\n\n# RHEL/CentOS/Fedora\nsudo yum install python3 python3-pip\n# or\nsudo dnf install python3 python3-pip")

    add_formatted_paragraph(doc, "Missing Build Tools", bold=True, font_size=10, space_before=8, space_after=2)
    add_formatted_paragraph(
        doc,
        "Some Python packages require compilation. Install build essentials if you see compilation errors during pip install.",
        font_size=10, space_after=4
    )
    add_code_block(doc, "# Debian/Ubuntu\nsudo apt install build-essential python3-dev\n\n# RHEL/CentOS\nsudo yum groupinstall 'Development Tools'\nsudo yum install python3-devel")

    doc.add_page_break()

    # ===================================================================
    # 6. CONFIGURATION REFERENCE
    # ===================================================================
    add_section_heading(doc, "6. Configuration Reference", level=1)
    add_formatted_paragraph(
        doc,
        "The application can be configured through environment variables. Set these before starting the server.",
        font_size=10, space_after=8
    )

    add_env_table(doc, [
        ("PORT", "5000", "Server port number"),
        ("FLASK_ENV", "production", "Set to 'development' to enable debug mode and auto-reload"),
        ("CORS_ORIGINS", "*", "Comma-separated list of allowed CORS origins"),
        ("SECRET_KEY", "auto-generated", "Session secret key for Flask sessions"),
        ("ARCH_USERNAME", "admin-username", "Username for the Architecture page login"),
        ("ARCH_PASSWORD", "admin-password", "Password for the Architecture page login"),
    ])

    add_formatted_paragraph(doc, "Example usage:", bold=True, font_size=10, space_before=8, space_after=4)
    add_code_block(doc, "# Linux / macOS\nexport PORT=5001\nexport FLASK_ENV=development\nexport CORS_ORIGINS=http://localhost:3000\npython3 app.py")
    add_code_block(doc, "# Or inline (single command)\nPORT=5001 FLASK_ENV=development python3 app.py")

    add_warning_box(doc, "Change ARCH_USERNAME and ARCH_PASSWORD from the defaults before deploying to any shared or production environment.")

    doc.add_page_break()

    # ===================================================================
    # 7. HEALTH CHECKS
    # ===================================================================
    add_section_heading(doc, "7. Health Checks", level=1)
    add_formatted_paragraph(
        doc,
        "Use these endpoints and techniques to verify the application is running correctly.",
        font_size=10, space_after=8
    )

    add_formatted_paragraph(doc, "API Health Endpoint", bold=True, font_size=10, space_before=6, space_after=2)
    add_code_block(doc, "GET /api/health\n\nResponse:\n{\n  \"status\": \"ok\",\n  \"env\": \"production\"\n}")
    add_code_block(doc, "# Test with curl\ncurl http://localhost:5000/api/health")

    add_formatted_paragraph(doc, "Build Prerequisites Check", bold=True, font_size=10, space_before=8, space_after=2)
    add_code_block(doc, "POST /api/build/check-prereqs\n\nResponse:\n{\n  \"java\": {\"available\": true, \"version\": \"17.0.x\"},\n  \"docker\": {\"available\": true, \"version\": \"24.x.x\"},\n  \"maven\": {\"available\": true, \"version\": \"3.9.x\"}\n}")
    add_code_block(doc, "# Test with curl\ncurl -X POST http://localhost:5000/api/build/check-prereqs")

    add_formatted_paragraph(doc, "Browser-Side Debugging", bold=True, font_size=10, space_before=8, space_after=2)
    add_formatted_paragraph(
        doc,
        "Open the browser's Developer Tools (F12 or Cmd+Option+I on macOS) and check the Console tab for JavaScript errors. Common issues include failed API calls, CORS errors, and missing configuration.",
        font_size=10, space_after=4
    )

    add_formatted_paragraph(doc, "Server-Side Debugging", bold=True, font_size=10, space_before=8, space_after=2)
    add_formatted_paragraph(
        doc,
        "Check the terminal where the application is running for Python tracebacks and error messages. Enable debug mode for more verbose output.",
        font_size=10, space_after=4
    )
    add_code_block(doc, "FLASK_ENV=development python3 app.py")

    add_tip_box(doc, "In development mode, Flask will auto-reload when you change source files and provide detailed error pages in the browser.")

    doc.add_page_break()

    # ===================================================================
    # 8. VIRTUAL ENVIRONMENT SETUP
    # ===================================================================
    add_section_heading(doc, "8. Virtual Environment Setup (Recommended)", level=1)
    add_formatted_paragraph(
        doc,
        "A virtual environment isolates the project's Python dependencies from your system Python installation. This prevents version conflicts and permission issues.",
        font_size=10, space_after=8
    )

    add_formatted_paragraph(doc, "Create and Activate", bold=True, font_size=10, space_before=6, space_after=4)
    add_code_block(doc, "# Create a virtual environment\npython3 -m venv venv\n\n# Activate it\n# macOS / Linux:\nsource venv/bin/activate\n\n# Windows (PowerShell):\nvenv\\Scripts\\Activate.ps1\n\n# Windows (CMD):\nvenv\\Scripts\\activate")

    add_formatted_paragraph(doc, "Install Dependencies", bold=True, font_size=10, space_before=8, space_after=4)
    add_code_block(doc, "# After activation, install all dependencies\npip install -r requirements.txt")

    add_formatted_paragraph(doc, "Deactivate", bold=True, font_size=10, space_before=8, space_after=4)
    add_code_block(doc, "# When done, deactivate the virtual environment\ndeactivate")

    add_info_box(doc, "When the virtual environment is active, your terminal prompt will show (venv) at the beginning. You can use 'pip' instead of 'pip3' inside the virtual environment.")

    add_formatted_paragraph(doc, "Verify Installation", bold=True, font_size=10, space_before=8, space_after=4)
    add_code_block(doc, "# Check that packages are installed in the virtual environment\npip list\n\n# Check Flask specifically\npython -c \"import flask; print(flask.__version__)\"")

    doc.add_page_break()

    # ===================================================================
    # 9. DOCKER DEPLOYMENT
    # ===================================================================
    add_section_heading(doc, "9. Docker Deployment", level=1)
    add_formatted_paragraph(
        doc,
        "Deploy the migration platform as a Docker container for consistent, reproducible environments.",
        font_size=10, space_after=8
    )

    add_formatted_paragraph(doc, "Dockerfile", bold=True, font_size=10, space_before=6, space_after=4)
    add_code_block(doc, 'FROM python:3.11-slim\n\nWORKDIR /app\n\nCOPY backend/ .\n\nRUN pip install --no-cache-dir -r requirements.txt\n\nEXPOSE 5000\n\nCMD ["gunicorn", "-w", "4", "-b", "0.0.0.0:5000", "app:app"]')

    add_formatted_paragraph(doc, "Build and Run", bold=True, font_size=10, space_before=8, space_after=4)
    add_code_block(doc, "# Build the Docker image\ndocker build -t mulesoft-migrator .\n\n# Run the container\ndocker run -d -p 5000:5000 --name migrator mulesoft-migrator\n\n# View logs\ndocker logs -f migrator\n\n# Stop the container\ndocker stop migrator")

    add_formatted_paragraph(doc, "Docker Compose (Optional)", bold=True, font_size=10, space_before=8, space_after=4)
    add_code_block(doc, "# docker-compose.yml\nversion: '3.8'\nservices:\n  migrator:\n    build: .\n    ports:\n      - \"5000:5000\"\n    environment:\n      - FLASK_ENV=production\n      - PORT=5000\n      - CORS_ORIGINS=*\n    restart: unless-stopped")

    add_code_block(doc, "# Start with Docker Compose\ndocker-compose up -d\n\n# Stop\ndocker-compose down")

    add_warning_box(doc, "The Docker image uses gunicorn as the production WSGI server. Do not use the Flask development server (python app.py) in production deployments.")

    doc.add_page_break()

    # ===================================================================
    # 10. CONTACT & RESOURCES
    # ===================================================================
    add_section_heading(doc, "10. Contact & Resources", level=1)
    add_formatted_paragraph(
        doc,
        "Additional resources for help and reference.",
        font_size=10, space_after=8
    )

    add_formatted_paragraph(doc, "Report Issues", bold=True, font_size=10, space_before=6, space_after=2)
    add_formatted_paragraph(
        doc,
        "If you encounter an issue not covered in this guide, please open an issue on the project's GitHub repository with:",
        font_size=10, space_after=4
    )
    bullets = [
        "A clear description of the problem",
        "Steps to reproduce the issue",
        "Your operating system and Python version",
        "The full error message or traceback",
        "Any relevant configuration details",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    add_formatted_paragraph(doc, "Useful Links", bold=True, font_size=10, space_before=10, space_after=4)

    links = [
        ("GitHub Issues", "https://github.com/<organization>/mulesoft-to-springboot-migrator/issues"),
        ("Flask Documentation", "https://flask.palletsprojects.com/"),
        ("Spring Boot Documentation", "https://docs.spring.io/spring-boot/docs/current/reference/html/"),
        ("MuleSoft Documentation", "https://docs.mulesoft.com/"),
        ("Docker Documentation", "https://docs.docker.com/"),
    ]

    tbl = doc.add_table(rows=len(links), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for i, (name, url) in enumerate(links):
        bg = "F5F5F5" if i % 2 == 0 else "FFFFFF"
        for c in range(2):
            set_cell_shading(tbl.cell(i, c), bg)
        _set_cell_text(tbl.cell(i, 0), name, bold=True)
        _set_cell_text(tbl.cell(i, 1), url, font_size=9)

    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    add_formatted_paragraph(doc, "Stack Overflow Tags", bold=True, font_size=10, space_before=8, space_after=4)
    add_formatted_paragraph(
        doc,
        "When searching Stack Overflow for related issues, use these tags: flask, python-3.x, spring-boot, mulesoft, lxml, docker, gunicorn",
        font_size=10, space_after=6
    )

    # ===================================================================
    # FOOTER
    # ===================================================================
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "1A237E")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MuleSoft to Spring Boot Migration Platform v2.0  |  Troubleshooting Guide  |  March 2026")
    run.font.size = Pt(8)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True

    return doc


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Trial_Error_Installation_Troubleshooting_Guide_v2.docx"
    )

    print("Generating troubleshooting guide...")
    doc = create_document()
    doc.save(output_path)

    file_size = os.path.getsize(output_path)
    if file_size >= 1024 * 1024:
        size_str = f"{file_size / (1024 * 1024):.1f} MB"
    else:
        size_str = f"{file_size / 1024:.1f} KB"

    print(f"Document saved successfully: {output_path}")
    print(f"File size: {size_str}")
