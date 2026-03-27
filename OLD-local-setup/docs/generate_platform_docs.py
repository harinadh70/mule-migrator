#!/usr/bin/env python3
"""
Generate comprehensive Word documentation for the MuleSoft-to-SpringBoot Agentic AI Platform.
Requires: python-docx (pip install python-docx)
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc, code_text, language=""):
    """Add a code block with Courier New font and light gray background."""
    # Create a table with one cell to simulate a code block with background
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")

    # Remove default paragraph and add code
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)

    # Set cell margins
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '  <w:top w:w="60" w:type="dxa"/>'
        '  <w:left w:w="120" w:type="dxa"/>'
        '  <w:bottom w:w="60" w:type="dxa"/>'
        '  <w:right w:w="120" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(tcMar)
    doc.add_paragraph()  # spacer


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "2E4057")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for j, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Inches(width)

    doc.add_paragraph()  # spacer


def add_bold_paragraph(doc, bold_text, normal_text=""):
    """Add a paragraph with bold prefix and normal continuation."""
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet list item."""
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p


def add_page_break(doc):
    doc.add_page_break()


def set_page_size_letter(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ──────────────────────────────────────────────────────────────────────────────
# Main document generation
# ──────────────────────────────────────────────────────────────────────────────

def generate_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    set_page_size_letter(section)

    # Add header/footer
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "MuleSoft-to-SpringBoot Agentic AI Platform"
    hp.style = doc.styles["Header"]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Page number field
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

    # ══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MuleSoft-to-SpringBoot\nAgentic AI Platform")
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Complete Technical Documentation")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)

    doc.add_paragraph()

    version_p = doc.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_p.add_run("Version 2.0")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()

    stack_p = doc.add_paragraph()
    stack_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stack_p.add_run(
        "FastAPI | React 18 | PostgreSQL 16 | Redis 7 | Qdrant | Celery\n"
        "sentence-transformers | Docker Compose | Kubernetes/Helm"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Architecture Overview",
        "3. Agentic AI Pipeline",
        "4. Self-Hosted RAG Infrastructure",
        "5. Backend Architecture",
        "6. Core Migration Engine",
        "7. Deployment Guide",
        "8. Multi-Tenancy",
        "9. Security & Compliance",
        "10. Monitoring & Observability",
        "11. Frontend Architecture",
        "12. Troubleshooting Guide",
        "Appendix A: Sample MuleSoft XML",
        "Appendix B: Generated Spring Boot Code",
        "Appendix C: Quick Reference Commands",
        "Appendix D: Complete File Inventory",
        "Appendix E: Environment Variables Reference",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "The MuleSoft-to-SpringBoot Agentic AI Platform is an enterprise-grade system that automatically "
        "converts MuleSoft 4 integration applications into production-ready Spring Boot 3.2 microservices. "
        "It combines a deterministic static migration engine (8 Python modules, ~5,500 lines of code) with "
        "a multi-agent AI pipeline that handles unknown elements, generates tests, reviews code quality, "
        "and produces documentation -- all powered by a fully self-hosted RAG (Retrieval Augmented Generation) "
        "infrastructure where zero data ever leaves the organization's network."
    )

    doc.add_heading("Key Capabilities", level=2)
    capabilities = [
        ("Agentic AI Pipeline", "Five specialized agents (Planner, Coder, Reviewer, Tester, Docs) operate in autonomous think-plan-act-observe-reflect loops with inter-agent communication."),
        ("Self-Hosted RAG", "Qdrant vector database + sentence-transformers (all-MiniLM-L6-v2) running entirely on-premises. ~800 pre-indexed documents covering MuleSoft and Spring Boot patterns."),
        ("Multi-Agent Orchestration", "DAG-based pipeline execution with circuit breakers, parallel execution, checkpointing, and deterministic fallbacks at every stage."),
        ("Air-Gapped Deployment", "Complete offline bundle (~3GB tarball) with pre-baked models, knowledge base, and all container images for deployment in disconnected environments."),
        ("Multi-Tenancy", "5-layer isolation: PostgreSQL RLS, Redis key prefixing, Qdrant collection namespacing, Celery queue routing, and API middleware enforcement."),
        ("Multi-Provider LLM", "Supports Anthropic (Claude), OpenAI (GPT-4), Google (Gemini), DeepSeek, and Groq with automatic failover."),
    ]
    for title_text, desc in capabilities:
        add_bold_paragraph(doc, f"{title_text}: ", desc)

    doc.add_heading("Technology Stack Overview", level=2)
    add_styled_table(doc,
        ["Layer", "Technology", "Version", "Purpose"],
        [
            ["Backend API", "FastAPI", "0.111+", "Async REST API with OpenAPI docs"],
            ["Frontend", "React 18 + TypeScript", "18.2", "SPA with Vite build, Zustand state"],
            ["Database", "PostgreSQL", "16", "Migration jobs, user data, audit logs"],
            ["Cache / Broker", "Redis", "7", "Session cache, Celery broker, embedding cache"],
            ["Vector DB", "Qdrant", "latest", "Self-hosted vector store for RAG"],
            ["Embeddings", "sentence-transformers", "all-MiniLM-L6-v2", "384-dim embeddings, local inference"],
            ["Task Queue", "Celery", "5.3+", "Async migration, build, indexing tasks"],
            ["LLM Providers", "Anthropic/OpenAI/Google", "Multi", "Claude, GPT-4, Gemini support"],
            ["Containerization", "Docker Compose", "3.8", "Dev and production orchestration"],
            ["Orchestration", "Kubernetes + Helm", "1.28+", "Production-grade deployment"],
            ["Reverse Proxy", "Nginx", "alpine", "TLS termination, rate limiting, caching"],
        ],
        col_widths=[1.3, 2.0, 1.2, 2.5]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: ARCHITECTURE OVERVIEW
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Architecture Overview", level=1)

    doc.add_heading("2.1 System Architecture", level=2)
    doc.add_paragraph(
        "The platform follows a layered architecture with clear separation of concerns. "
        "The React frontend communicates with the FastAPI backend through REST and WebSocket APIs. "
        "Background processing is handled by Celery workers, with Redis serving as both the message broker "
        "and caching layer. The RAG subsystem operates independently with its own Qdrant vector store "
        "and sentence-transformer embedding model."
    )

    arch_diagram = """
+──────────────────────────────────────────────────────────────────────+
|                        CLIENT LAYER                                  |
|  ┌──────────────────────────────────────────────────────────────┐   |
|  │  React 18 + TypeScript + Vite                                │   |
|  │  ├── Zustand State Management                                │   |
|  │  ├── React Flow (Agent Pipeline Visualization)               │   |
|  │  ├── WebSocket Client (Real-time Progress)                   │   |
|  │  └── REST API Client (Axios)                                 │   |
|  └──────────────────────────────────────┬───────────────────────┘   |
+─────────────────────────────────────────┼───────────────────────────+
                                          │ HTTPS / WSS
+─────────────────────────────────────────┼───────────────────────────+
|                     NGINX REVERSE PROXY │                            |
|  ┌──────────────────────────────────────┼───────────────────────┐   |
|  │  TLS Termination | Rate Limiting | GZip | Static Files       │   |
|  └──────────────────────────────────────┼───────────────────────┘   |
+─────────────────────────────────────────┼───────────────────────────+
                                          │
+─────────────────────────────────────────┼───────────────────────────+
|                     FASTAPI APPLICATION │ (api/)                     |
|  ┌──────────────────────────────────────┼───────────────────────┐   |
|  │  Middleware: CORS | GZip | TrustedHost | Request Logging     │   |
|  ├──────────────────────────────────────┴───────────────────────┤   |
|  │  /api/v1/* and /api/v2/* versioned routers                   │   |
|  │  /health, /readiness, /metrics endpoints                     │   |
|  │  WebSocket /ws/migration/{id}                                │   |
|  ├──────────────────────────────────────────────────────────────┤   |
|  │  AGENT PIPELINE                                              │   |
|  │  PlannerAgent -> CoderAgent -> ReviewerAgent -> TesterAgent  │   |
|  │       -> DocsAgent                                           │   |
|  │  + AgentMemory, AgentGuardrails, AgentTools                  │   |
|  ├──────────────────────────────────────────────────────────────┤   |
|  │  RAG SUBSYSTEM                                               │   |
|  │  EmbeddingService | HybridRetriever | CodeAwareChunker       │   |
|  │  DocumentIndexer | EmbeddingCache | QdrantStore              │   |
|  └──────────┬────────────────┬──────────────────┬───────────────┘   |
+--------------┼────────────────┼──────────────────┼──────────────────+
               │                │                  │
   ┌───────────┴──┐   ┌────────┴────┐   ┌────────┴────────┐
   │ PostgreSQL 16│   │  Redis 7    │   │  Qdrant         │
   │              │   │             │   │  (Vector Store)  │
   │ - Jobs       │   │ - Cache     │   │                  │
   │ - Users      │   │ - Broker    │   │ - mulesoft_docs  │
   │ - Traces     │   │ - Sessions  │   │ - springboot_docs│
   │ - Audit Logs │   │ - Embeddings│   │ - custom_patterns│
   └──────────────┘   └─────────────┘   │ - migration_hist │
                                        └──────────────────┘
    """
    add_code_block(doc, arch_diagram.strip())

    doc.add_heading("2.2 Data Flow: XML Upload to Generated Spring Boot", level=2)
    doc.add_paragraph(
        "The migration pipeline follows a well-defined data flow from MuleSoft XML input "
        "to a complete, tested, documented Spring Boot project output."
    )

    flow_steps = [
        ("1. Upload", "User uploads MuleSoft XML files via POST /api/v2/migrations. Files are validated, stored, and a Celery task is dispatched. The API returns HTTP 202 Accepted with the migration job ID."),
        ("2. Parse", "The MuleSoft XML Parser (backend/migrator/parser.py) extracts flows, sub-flows, connectors, configurations, DataWeave expressions, error handlers, and global elements using lxml with a complete namespace registry covering 30+ MuleSoft namespaces."),
        ("3. Plan", "PlannerAgent performs static XML analysis (flow count, connector types, DataWeave complexity) and queries RAG for similar past migrations. It produces a MigrationPlan with complexity score (0-100), risk areas, and determines which agents are needed."),
        ("4. Convert", "Static converters (FlowConverter, ConnectorMapper, DataWeaveConverter) handle known elements. Unknown elements are routed to CoderAgent, which uses RAG-enhanced LLM prompts with retry+validation loops."),
        ("5. Review", "ReviewerAgent scores each generated file (1-10) with per-line feedback. Files scoring <= 3 trigger an autonomous fix loop where the Reviewer sends regeneration requests back to the Coder via inter-agent messaging."),
        ("6. Test", "TesterAgent generates JUnit 5 test suites: @WebMvcTest + MockMvc for controllers, @ExtendWith(MockitoExtension) for services, with BDDMockito style assertions."),
        ("7. Document", "DocsAgent generates README.md, MIGRATION_REPORT.md with audit trail, and OpenAPI enhancements. Includes per-file review scores, risk assessment, and token usage breakdown."),
        ("8. Package", "The complete Spring Boot project is assembled: Java source files, test files, pom.xml/build.gradle, application.properties, documentation, and served as a downloadable ZIP."),
    ]
    for step_title, step_desc in flow_steps:
        add_bold_paragraph(doc, f"{step_title}: ", step_desc)

    doc.add_heading("2.3 Technology Stack Details", level=2)
    add_styled_table(doc,
        ["Component", "Technology", "Version", "Configuration", "Purpose"],
        [
            ["API Framework", "FastAPI", "0.111+", "Uvicorn ASGI, 4 workers", "Async REST + WebSocket API"],
            ["ORM", "SQLAlchemy 2.0", "2.0+", "asyncpg driver, pool 5-20", "Async PostgreSQL ORM"],
            ["Validation", "Pydantic v2", "2.0+", "BaseSettings for config", "Request/response validation"],
            ["Database", "PostgreSQL", "16-alpine", "5-20 pool, RLS enabled", "Persistent storage"],
            ["Cache", "Redis", "7-alpine", "256MB, allkeys-lru, AOF", "Cache, broker, sessions"],
            ["Vector Store", "Qdrant", "latest", "gRPC:6334, REST:6333", "RAG vector storage"],
            ["Embeddings", "all-MiniLM-L6-v2", "latest", "384 dims, L2 normalized", "Semantic embeddings"],
            ["Task Queue", "Celery", "5.3+", "3 queues, prefork pool", "Async job processing"],
            ["Container", "Docker Compose", "3.8", "6 services, health checks", "Dev/prod orchestration"],
            ["Proxy", "Nginx", "alpine", "TLS, rate limit, gzip", "Reverse proxy"],
            ["Logging", "structlog", "latest", "JSON structured logging", "Observability"],
            ["Metrics", "Prometheus", "via instrumentator", "Custom + auto metrics", "Performance monitoring"],
        ],
        col_widths=[1.1, 1.3, 0.7, 1.6, 1.8]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 3: AGENTIC AI PIPELINE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Agentic AI Pipeline", level=1)

    doc.add_heading("3.1 What Makes It Agentic", level=2)
    doc.add_paragraph(
        "Unlike simple LLM call-and-return systems, this platform implements a true agentic AI architecture "
        "where each agent operates in an autonomous loop: Think (analyze context + RAG results) -> Plan "
        "(build targeted prompt with token budget awareness) -> Act (call LLM with retry + timeout) -> "
        "Observe (parse structured output, validate with guardrails) -> Reflect (update shared memory, "
        "communicate findings to other agents). Agents can trigger re-execution of other agents (e.g., "
        "Reviewer triggering Coder to regenerate a file), creating an emergent self-improving pipeline."
    )

    doc.add_heading("3.2 Agent Lifecycle", level=2)
    lifecycle_diagram = """
Orchestrator                 Agent (via safe_execute)
    │                              │
    ├── dispatch(agent, context) ──>│
    │                              ├── start_trace()
    │                              ├── [THINK] Query RAG for context
    │                              ├── [PLAN]  Build prompt (budget check)
    │                              ├── [ACT]   call_llm_with_retry()
    │                              │     ├── Attempt 1 ── timeout/error ──> backoff
    │                              │     ├── Attempt 2 ── timeout/error ──> backoff
    │                              │     └── Attempt 3 ── success ──────────>│
    │                              ├── [OBSERVE] parse_response()
    │                              ├── [VALIDATE] guardrails.validate()
    │                              ├── [REFLECT] memory.store(), context.update()
    │                              ├── finish_trace()
    │                              ├── emit_metrics()
    │<── AgentResult ──────────────┤
    │                              │
    ├── On Error ──────────────────>│
    │                              ├── get_fallback(context)
    │<── AgentResult(fallback) ────┤
    """
    add_code_block(doc, lifecycle_diagram.strip())

    doc.add_heading("3.3 PlannerAgent", level=2)
    doc.add_paragraph("File: api/agents/planner.py")
    doc.add_paragraph(
        "The PlannerAgent is the first agent in the pipeline. It performs static XML analysis without any LLM calls, "
        "then optionally enriches the analysis with RAG and LLM for high-complexity migrations."
    )

    add_bold_paragraph(doc, "Inputs: ", "Dictionary of XML filename -> XML content strings.")
    add_bold_paragraph(doc, "Static Analysis Steps:")
    analysis_steps = [
        "Parse all XML files with xml.etree.ElementTree",
        "Count flows and sub-flows (both namespaced and non-namespaced)",
        "Detect connectors by XML namespace prefix (filters out core/mule/xml/spring)",
        "Count DataWeave expressions (#[...] and %dw patterns)",
        "Identify high-risk patterns: batch processing, transactions, security configs",
        "Classify connectors as simple (http, logger, set-payload) or high-risk (salesforce, sap, kafka, batch)",
    ]
    for step in analysis_steps:
        add_bullet(doc, step)

    add_bold_paragraph(doc, "Complexity Scoring Formula:")
    add_code_block(doc, """score = 0
score += min(flow_count * 5, 30)          # max 30 points for flows
score += min(connector_count * 8, 30)     # max 30 points for connectors
score += min(dataweave_count * 3, 20)     # max 20 points for DataWeave
score += min(risk_area_count * 5, 20)     # max 20 points for risks

# Mapping: 0-20 = low, 21-45 = medium, 46-70 = high, 71-100 = very_high""")

    add_bold_paragraph(doc, "RAG Queries: ", "For each detected connector (up to 5), the planner queries: 'MuleSoft {connector} connector migration to Spring Boot', retrieving up to 3 results per query from the migration_history collection.")
    add_bold_paragraph(doc, "LLM Enhancement: ", "Only triggered for high/very_high complexity. Sends combined XML (up to 6000 chars) with RAG context to the LLM for deeper risk analysis.")
    add_bold_paragraph(doc, "Output Schema (MigrationPlan):")
    add_code_block(doc, """MigrationPlan:
  complexity: str          # "low" | "medium" | "high" | "very_high"
  complexity_score: int    # 0-100
  flows_detected: int
  connectors_detected: List[str]
  dataweave_expressions: int
  risk_areas: List[str]
  estimated_agents_needed: List[str]
  estimated_tokens: int
  use_full_pipeline: bool
  similar_past_migrations: List[dict]
  notes: str""")

    add_bold_paragraph(doc, "Decision Logic: ", "If complexity is high/very_high OR more than 3 connectors detected, use_full_pipeline is set to True and agents include [planner, coder, reviewer, tester, docs]. Otherwise, a simplified [coder, docs] pipeline runs.")

    doc.add_heading("3.4 CoderAgent", level=2)
    doc.add_paragraph("File: api/agents/coder.py")
    doc.add_paragraph(
        "The CoderAgent handles all elements that the static converters cannot process. It processes four categories: "
        "unknown XML elements, unknown DataWeave expressions, unknown connectors, and unknown message sources."
    )

    add_bold_paragraph(doc, "RAG-Enhanced Code Generation (per element):")
    coder_steps = [
        "1. RAG Search: Query 'Spring Boot equivalent for MuleSoft {element_tag}' with top_k=5",
        "2. Prompt Building: Combine RAG context (patterns from knowledge base) with element XML and flow context",
        "3. LLM Call: Generate Java code with retry loop (max 3 attempts)",
        "4. Syntax Validation: AgentGuardrails.validate_java_syntax() checks balanced braces, class declarations",
        "5. If valid: AgentGuardrails.validate() runs full pipeline (syntax + annotations + security + imports + hallucination)",
        "6. If invalid: Append validation feedback to prompt and retry",
        "7. Final Fallback: Delegate to legacy backend.migrator.llm_agent.convert_unknown_element()",
    ]
    for step in coder_steps:
        add_bullet(doc, step)

    add_bold_paragraph(doc, "Provenance Tracking: ", "Each generated code block is tagged as either 'rag_sourced' (knowledge base pattern used) or 'llm_sourced' (pure LLM generation). This tracking flows through to the migration report.")
    add_bold_paragraph(doc, "System Prompt: ", "Instructs the LLM as a 'senior Java/Spring Boot 3.2 engineer' using Java 17 features (records, sealed classes, pattern matching), Spring Boot best practices (constructor injection, @Transactional), and // REVIEW comments on uncertain parts.")

    doc.add_heading("3.5 ReviewerAgent", level=2)
    doc.add_paragraph("File: api/agents/reviewer.py")
    doc.add_paragraph(
        "The ReviewerAgent performs per-file code reviews with a scoring system and autonomous fix loop capability."
    )

    add_bold_paragraph(doc, "Review Process (per file):")
    review_steps = [
        "1. File Classification: Heuristic analysis to determine type (REST controller, service, repository, config, test)",
        "2. RAG Query: 'Spring Boot best practices for {file_type}' retrieves relevant patterns",
        "3. Static Guardrails: Run validate_java_syntax(), validate_spring_annotations(), detect_security_issues()",
        "4. LLM Review: Send file content (up to 6000 chars) with RAG context and guardrail findings",
        "5. Score Extraction: Parse JSON response with score (1-10), per-line feedback, suggestions",
        "6. Re-review: If score < 4 and iteration < MAX_REVIEW_ITERATIONS (2), re-review with accumulated feedback",
    ]
    for step in review_steps:
        add_bullet(doc, step)

    add_bold_paragraph(doc, "Scoring Guide:")
    add_styled_table(doc,
        ["Score", "Quality Level", "Action"],
        [
            ["9-10", "Production-ready", "No action needed"],
            ["7-8", "Good quality", "Minor improvements suggested"],
            ["5-6", "Functional", "Notable issues flagged"],
            ["3-4", "Significant problems", "Substantial fixes needed"],
            ["1-2", "Fundamentally broken", "Triggers regeneration request to CoderAgent"],
        ],
        col_widths=[0.7, 1.5, 4.3]
    )

    add_bold_paragraph(doc, "Autonomous Fix Loop: ", "When a file scores <= 3 with critical issues (needs_regen=true), the Reviewer sends an inter-agent message to the Coder via context.send_message(sender='reviewer', receiver='coder', content={'action': 'regenerate', 'files': [...], 'reviews': [...]}).")

    doc.add_heading("3.6 TesterAgent", level=2)
    doc.add_paragraph("File: api/agents/tester.py")

    add_bold_paragraph(doc, "Test Generation Strategy:")
    test_strategies = [
        "Controller tests: @WebMvcTest with MockMvc, @MockBean for service deps, test happy path + validation errors + error responses",
        "Service tests: @ExtendWith(MockitoExtension.class), @Mock + @InjectMocks, BDDMockito given/when/then style",
        "All tests: JUnit 5, AssertJ assertions (assertThat), @DisplayName annotations, method naming: should_<expected>_when_<condition>",
        "Error-path tests: Exception handlers, edge cases, null inputs",
        "Integration test skeletons: @SpringBootTest with TODO markers for external services",
    ]
    for s in test_strategies:
        add_bullet(doc, s)

    add_bold_paragraph(doc, "Fallback Skeleton Generation: ", "When LLM is unavailable, generates skeleton test classes with TODO markers for each public method. Controller skeletons include MockMvc setup; service skeletons include Mockito setup.")

    doc.add_heading("3.7 DocsAgent", level=2)
    doc.add_paragraph("File: api/agents/docs.py")
    doc.add_paragraph(
        "Generates three documentation artifacts: README.md (setup instructions, architecture, API endpoints), "
        "MIGRATION_REPORT.md (detailed audit trail with per-file scores, risk assessment, token usage), and "
        "openapi_enhancements.json (extracted endpoint metadata for Swagger enrichment)."
    )

    doc.add_heading("3.8 Agent Memory System", level=2)
    doc.add_paragraph("File: api/agents/memory.py")
    add_bold_paragraph(doc, "Short-Term Memory (In-Process): ", "Plain dict-based store scoped to the current migration run. Keys are '{run_id}:latest_result' strings. Cleared automatically between runs via clear_short_term().")
    add_bold_paragraph(doc, "Long-Term Memory (Qdrant-Backed): ", "Each completed migration is stored in the migration_history Qdrant collection with an embedding of the pipeline summary. Future agents can search for similar past patterns using recall(query, top_k=5). Point ID is deterministic: MD5(run_id). Includes score-based pruning via prune_low_quality(threshold=5.0).")

    doc.add_heading("3.9 Guardrails", level=2)
    doc.add_paragraph("File: api/agents/guardrails.py")
    doc.add_paragraph("The AgentGuardrails class provides stateless validation utilities that run on every LLM-generated output:")

    add_bold_paragraph(doc, "1. Java Syntax Validation: ", "Checks balanced braces/parens/brackets, requires class/interface/enum declaration or import statement, flags double semicolons (;;) as generation artifacts.")
    add_bold_paragraph(doc, "2. Spring Annotation Validation: ", "Detects controller-like code missing @RestController, service-like code missing @Service, repository-like code missing @Repository, config-like code missing @Configuration. Flags classes without package declarations.")
    add_bold_paragraph(doc, "3. Security Scanning: ", "Detects hardcoded passwords/secrets/API keys (regex patterns for 8+ char strings), AWS access keys (AKIA prefix), private key PEM blocks, SQL injection via string concatenation (String.format with SQL, createQuery/executeQuery with +).")
    add_bold_paragraph(doc, "4. Import Validation: ", "Checks all import statements against a whitelist of 18 known valid package prefixes (java.*, javax.*, jakarta.*, org.springframework.*, com.fasterxml.jackson.*, etc.).")
    add_bold_paragraph(doc, "5. Hallucination Detection: ", "Compares MuleSoft elements referenced in output against elements actually present in the input XML. Flags responses referencing >2 non-existent elements. Detects fake Spring Boot starter names not in the known set.")
    add_bold_paragraph(doc, "6. Auto-Fix: ", "When fix_issues=True, automatically removes double semicolons and trailing whitespace.")

    doc.add_heading("3.10 Agent Tools", level=2)
    doc.add_paragraph("File: api/agents/tools.py")
    doc.add_paragraph("Agents can autonomously invoke registered tools during execution:")

    add_styled_table(doc,
        ["Tool Name", "Parameters", "Description"],
        [
            ["parse_xml", "xml_string: str", "Parse MuleSoft XML into structured dict using stdlib ElementTree"],
            ["validate_java", "code: str", "Full validation: syntax + annotations + security issues"],
            ["search_rag", "query: str, top_k: int", "Search RAG knowledge base for patterns"],
            ["get_spring_docs", "topic: str", "Get curated Spring Boot documentation snippets (7 built-in topics: rest_controller, service, jpa, security, testing, properties, error_handling)"],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    add_bold_paragraph(doc, "Custom Tool Registration: ", "Tools can be registered at runtime via AgentTools.register_tool(Tool(name=..., description=..., function=..., parameters=...)). All registered tools are rendered into the LLM system prompt via to_prompt_description().")

    doc.add_heading("3.11 Agent Context and Result", level=2)
    doc.add_paragraph("File: api/agents/context.py, api/agents/result.py")
    add_bold_paragraph(doc, "AgentContext (Extended): ", "Extends the legacy AgentContext with UUID tracking, per-agent execution traces (AgentTrace), artifact storage for intermediate outputs, inter-agent messaging (AgentMessage), token usage accounting, cost tracking, and checkpointing for crash recovery.")
    add_bold_paragraph(doc, "AgentResult: ", "Standardized return value with status (success|partial|error|timeout|budget_exceeded), output dict, token_usage, duration_ms, rag_queries, rag_results_used, error message, and fallback_used flag.")

    doc.add_heading("3.12 Error Handling and Fallbacks", level=2)
    error_handling = [
        "Every agent has a deterministic get_fallback(context) that works without any LLM calls",
        "BaseAgent.safe_execute() wraps execute() with asyncio.wait_for timeout, ValueError catch for token budget, and generic Exception catch",
        "On exception: first attempts get_fallback(), if that also fails, returns AgentResult.from_error() with both error messages",
        "LLM calls use exponential backoff: wait = min(2^attempt, 30) seconds between retries",
        "Token budget checked BEFORE sending to LLM: estimated_tokens = len(prompt) // 4",
        "AgentResult status propagation: orchestrator reads is_usable (True for success/partial) to decide pipeline continuation",
    ]
    for item in error_handling:
        add_bullet(doc, item)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 4: SELF-HOSTED RAG INFRASTRUCTURE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Self-Hosted RAG Infrastructure", level=1)

    doc.add_heading("4.1 Why Self-Hosted", level=2)
    doc.add_paragraph(
        "The entire RAG pipeline runs on-premises with zero data leaving the infrastructure. This is critical for "
        "enterprises with strict data governance requirements, SOC 2 compliance mandates, and air-gapped deployments. "
        "The embedding model (all-MiniLM-L6-v2, ~80MB) and vector database (Qdrant) are fully self-contained. "
        "No external API calls are made for embeddings or vector search."
    )

    doc.add_heading("4.2 Architecture: Qdrant + sentence-transformers", level=2)
    add_code_block(doc, """api/rag/
├── __init__.py
├── config.py          # RAG-specific settings (model name, chunk sizes, top-k)
├── embeddings.py      # EmbeddingService (singleton, lazy-loaded model)
├── vector_store.py    # QdrantStore (connection pooling, retry, health check)
├── indexer.py         # DocumentIndexer (chunking pipeline, dedup, incremental)
├── retriever.py       # HybridRetriever (dense + sparse + re-ranking)
├── chunking.py        # CodeAwareChunker (AST-based for Java/XML)
├── cache.py           # EmbeddingCache (Redis-backed, TTL per collection)
├── schemas.py         # Pydantic models: Document, Chunk, SearchResult, IndexStats
└── knowledge/
    ├── mulesoft/      # ~500 docs: connector XML refs, DataWeave cookbook
    ├── springboot/    # ~300 docs: starters, annotations, config properties
    └── custom/        # User-uploaded team patterns""")

    doc.add_heading("4.3 EmbeddingService", level=2)
    doc.add_paragraph("File: api/rag/embeddings.py")
    embed_details = [
        ("Model", "all-MiniLM-L6-v2 (configurable via RAGConfig.embedding.model_name)"),
        ("Dimensions", "384 (configurable)"),
        ("Max Sequence Length", "256 tokens"),
        ("Batch Size", "64 (configurable)"),
        ("Normalization", "L2 normalization on all embeddings for cosine similarity"),
        ("Device Detection", "CUDA (GPU) > MPS (Apple Silicon) > CPU, auto-detected via PyTorch"),
        ("Thread Safety", "Double-checked locking with threading.Lock for singleton + model loading"),
        ("Lazy Loading", "Model loaded on first embed() call, not at import time (saves ~2GB RAM on cold start)"),
        ("Singleton Pattern", "EmbeddingService.get_instance() returns global singleton; reset_instance() for testing"),
    ]
    for label, desc in embed_details:
        add_bold_paragraph(doc, f"{label}: ", desc)

    doc.add_heading("4.4 QdrantStore", level=2)
    doc.add_paragraph("File: api/rag/vector_store.py")
    qdrant_details = [
        ("Connection", "host=localhost, port=6333, grpc_port=6334, prefer_grpc=True"),
        ("Retry Logic", "Exponential backoff: wait = backoff * 2^(attempt-1), max 3 retries. Retries on ConnectionError, OSError, TimeoutError, and 5xx responses. 4xx errors raised immediately."),
        ("Collection Management", "create_collection() with idempotent check, delete_collection(), get_collection_info()"),
        ("Vector Config", "384 dimensions, Cosine distance metric, optional on_disk storage for large collections"),
        ("Point Operations", "upsert() with batch PointStruct, search() with optional filters and score threshold, delete() by IDs, scroll() for paginated iteration"),
        ("Payload Indexing", "create_payload_index() supports keyword, integer, float, bool, text field types for filtered search performance"),
        ("Health Check", "health_check() verifies Qdrant connectivity by listing collections"),
    ]
    for label, desc in qdrant_details:
        add_bold_paragraph(doc, f"{label}: ", desc)

    doc.add_heading("4.5 CodeAwareChunker", level=2)
    doc.add_paragraph("File: api/rag/chunking.py")
    doc.add_paragraph("The chunker respects the logical structure of each document format:")

    add_bold_paragraph(doc, "Text Chunking: ", "Sliding window with sentence-boundary awareness. Splits on periods, exclamation marks, question marks, semicolons, and double newlines. Configurable overlap (default 15%), min 100 tokens, max 512 tokens.")
    add_bold_paragraph(doc, "XML (MuleSoft) Chunking: ", "Regex-based extraction of <flow>, <sub-flow>, and connector config elements. Extracts 'name' attribute for section identification. Large elements (>max_tokens) are sub-chunked with text chunker. Remaining content outside matched elements chunked separately.")
    add_bold_paragraph(doc, "Java Chunking: ", "Regex-based method extraction respecting annotations, access modifiers, and brace matching. Extracts method name, line range, and full body. Non-method code (imports, fields) chunked separately.")
    add_bold_paragraph(doc, "Markdown Chunking: ", "Splits on heading boundaries (# through ######). Each section becomes a chunk with heading as section_name metadata. Oversized sections sub-chunked with text chunker.")
    add_bold_paragraph(doc, "Auto-Detection: ", "chunk_auto(content, filename) selects the appropriate chunker based on file extension: .xml/.mxml -> chunk_xml, .java/.kt/.scala -> chunk_java, .md/.markdown -> chunk_markdown, else -> chunk_text.")
    add_bold_paragraph(doc, "Metadata Enrichment: ", "Every RawChunk carries: content, source_type, section_name, line_start, line_end.")

    doc.add_heading("4.6 HybridRetriever", level=2)
    doc.add_paragraph("File: api/rag/retriever.py")
    doc.add_paragraph("The retriever implements a multi-stage pipeline for high-quality search results:")

    add_bold_paragraph(doc, "Stage 1 - Query Expansion: ", "Generates alternate phrasings: 'How to {query}', 'Migrate MuleSoft to Spring Boot: {query}', 'Spring Boot equivalent of {query}'. Up to 3 variants (configurable).")
    add_bold_paragraph(doc, "Stage 2 - Dense Search: ", "All query variants embedded in a single batch call. Each variant searched against each target collection (up to 4 collections if no specific collection specified). Uses Qdrant vector search with score_threshold=0.65.")
    add_bold_paragraph(doc, "Stage 3 - Deduplication: ", "Results merged across all queries/collections. Deduplicated by SHA-256 content hash, keeping the highest-scoring version.")
    add_bold_paragraph(doc, "Stage 4 - Relevance Filtering: ", "Results below similarity_threshold (0.65) are dropped. Remaining sorted by descending score.")
    add_bold_paragraph(doc, "Stage 5 - Context Window Packing: ", "Greedy packing into token budget (default 4096 tokens). Highest-score results included first. Results that would exceed budget are skipped (not truncated).")

    add_bold_paragraph(doc, "Domain-Specific Retrievers:")
    add_styled_table(doc,
        ["Method", "Query Template", "Collection", "Use Case"],
        [
            ["retrieve_mulesoft_context()", "MuleSoft {connector} connector configuration and usage", "mulesoft_docs", "Look up connector documentation"],
            ["retrieve_spring_pattern()", "Spring Boot {annotation} pattern for {use_case}", "springboot_docs", "Find Spring Boot implementation patterns"],
            ["retrieve_similar_migrations()", "Migration example for MuleSoft flow: {parsed_flow}", "migration_history", "Find past similar migrations"],
            ["retrieve_custom_patterns()", "Code pattern similar to: {code_snippet}", "custom_patterns", "Find team-specific patterns"],
        ],
        col_widths=[2.0, 1.8, 1.3, 1.4]
    )

    doc.add_heading("4.7 EmbeddingCache", level=2)
    doc.add_paragraph("File: api/rag/cache.py")

    cache_details = [
        ("Backend", "Redis (lazy connection, fail-open: cache misses are OK if Redis is unavailable)"),
        ("Key Strategy", "SHA-256 of (text + '||' + model_name), prefixed with 'rag:emb:'"),
        ("TTL Strategy", "Embedding TTL: 7 days, Search result TTL: 1 hour, Collection info TTL: 5 minutes, Knowledge base TTL: 30 days"),
        ("Hit Rate Tracking", "Thread-safe counters for hits/misses. hit_rate property returns float in [0,1]. stats() returns dict with hits, misses, hit_rate, total_lookups."),
        ("Warm-up", "Pre-compute embeddings for common queries on startup (configured via knowledge base)"),
        ("Connection", "Redis host/port/db configurable. Default: localhost:6379/2 (separate from app cache db 0 and Celery db 1)"),
    ]
    for label, desc in cache_details:
        add_bold_paragraph(doc, f"{label}: ", desc)

    doc.add_heading("4.8 DocumentIndexer", level=2)
    doc.add_paragraph("File: api/rag/indexer.py")
    doc.add_paragraph("Orchestrates the full indexing pipeline: scan -> dedup -> chunk -> embed -> upsert.")

    indexer_steps = [
        "index_directory(path, collection): Recursively scans for indexable files (.xml, .java, .md, .json, .yaml, .properties, .py, etc. -- 14 extensions)",
        "Deduplication: SHA-256 content hash checked against in-memory set per collection. Unchanged files skipped.",
        "Chunking: CodeAwareChunker.chunk_auto() selects format-appropriate chunker by file extension",
        "Embedding: _embed_texts() checks cache first, only computes uncached embeddings, stores results back to cache",
        "Upserting: QdrantStore.upsert() with point payload: content, doc_id, chunk_index, token_count, source, doc_type, source_type, section_name, line_start, line_end, content_hash",
        "Progress callback: Optional (current, total, message) callback for UI integration",
    ]
    for step in indexer_steps:
        add_bullet(doc, step)

    add_bold_paragraph(doc, "Built-in Knowledge Base:")
    add_styled_table(doc,
        ["Method", "Collection", "Source Directory", "Content"],
        [
            ["index_mulesoft_knowledge()", "mulesoft_docs", "api/rag/knowledge/mulesoft/", "~500 docs: connector XML references, DataWeave cookbook, flow patterns"],
            ["index_springboot_knowledge()", "springboot_docs", "api/rag/knowledge/springboot/", "~300 docs: starters, annotations, config properties, security"],
            ["index_all_knowledge()", "Both above", "Both above", "Full knowledge base indexing"],
        ],
        col_widths=[2.2, 1.5, 1.8, 1.5]
    )

    doc.add_heading("4.9 RAG Configuration Reference", level=2)
    doc.add_paragraph("File: api/rag/config.py -- All configuration is centralized in dataclass-based settings:")

    add_styled_table(doc,
        ["Config Class", "Parameter", "Default", "Description"],
        [
            ["EmbeddingConfig", "model_name", "all-MiniLM-L6-v2", "Sentence transformer model"],
            ["EmbeddingConfig", "embedding_dim", "384", "Vector dimensionality"],
            ["EmbeddingConfig", "batch_size", "64", "Batch size for encoding"],
            ["EmbeddingConfig", "max_seq_length", "256", "Max input token length"],
            ["EmbeddingConfig", "normalize", "True", "L2 normalize embeddings"],
            ["ChunkConfig", "min_tokens", "100", "Minimum chunk size"],
            ["ChunkConfig", "max_tokens", "512", "Maximum chunk size"],
            ["ChunkConfig", "overlap_ratio", "0.15", "Overlap between chunks"],
            ["ChunkConfig", "sentence_boundary", "True", "Respect sentence boundaries"],
            ["SearchConfig", "top_k", "10", "Results per search"],
            ["SearchConfig", "rerank_top_k", "5", "Results after re-ranking"],
            ["SearchConfig", "similarity_threshold", "0.65", "Minimum similarity score"],
            ["SearchConfig", "query_expansion_count", "3", "Max query variants"],
            ["SearchConfig", "context_token_budget", "4096", "Token budget for context packing"],
            ["QdrantConfig", "host", "localhost", "Qdrant server host"],
            ["QdrantConfig", "port", "6333", "Qdrant REST port"],
            ["QdrantConfig", "grpc_port", "6334", "Qdrant gRPC port"],
            ["QdrantConfig", "prefer_grpc", "True", "Use gRPC for 3x throughput"],
            ["QdrantConfig", "timeout", "30.0", "Request timeout seconds"],
            ["QdrantConfig", "max_retries", "3", "Retry attempts"],
            ["QdrantConfig", "retry_backoff", "1.0", "Base backoff seconds"],
            ["RedisConfig", "host", "localhost", "Redis host for cache"],
            ["RedisConfig", "port", "6379", "Redis port"],
            ["RedisConfig", "db", "2", "Redis database number"],
            ["RedisConfig", "key_prefix", "rag:emb:", "Cache key prefix"],
            ["RedisConfig", "default_ttl", "604800", "Default TTL (7 days)"],
        ],
        col_widths=[1.4, 1.8, 1.3, 2.0]
    )

    doc.add_heading("4.10 Qdrant Collections", level=2)
    add_styled_table(doc,
        ["Collection Name", "Content", "Typical Size"],
        [
            ["mulesoft_docs", "MuleSoft connector documentation, DataWeave patterns, flow templates", "~500 documents, ~3000 chunks"],
            ["springboot_docs", "Spring Boot starters, annotations, configuration, security patterns", "~300 documents, ~2000 chunks"],
            ["custom_patterns", "User-uploaded team-specific patterns and conventions", "Variable"],
            ["migration_history", "Past migration results and patterns for agent learning", "Grows with usage"],
        ],
        col_widths=[1.8, 3.2, 1.5]
    )

    doc.add_heading("4.11 How to View the Vector Database", level=2)
    add_bold_paragraph(doc, "Qdrant Web UI: ", "Navigate to http://localhost:6333/dashboard to browse collections, view points, and run searches interactively.")

    add_bold_paragraph(doc, "API Endpoints:")
    add_code_block(doc, """# List all collections
curl http://localhost:6333/collections

# Get collection info
curl http://localhost:6333/collections/mulesoft_docs

# Search a collection
curl -X POST http://localhost:6333/collections/mulesoft_docs/points/search \\
  -H "Content-Type: application/json" \\
  -d '{"vector": [0.1, 0.2, ...], "limit": 5}'

# Browse points (paginated)
curl -X POST http://localhost:6333/collections/mulesoft_docs/points/scroll \\
  -H "Content-Type: application/json" \\
  -d '{"limit": 10, "with_payload": true}'

# Get collection cluster info
curl http://localhost:6333/collections/mulesoft_docs/cluster

# Create a snapshot
curl -X POST http://localhost:6333/collections/mulesoft_docs/snapshots""")

    doc.add_heading("4.12 How to Add Custom Knowledge", level=2)
    doc.add_paragraph("Place documents in api/rag/knowledge/custom/ with supported extensions (.xml, .java, .md, .json, .yaml, .properties, .py, etc.) and trigger re-indexing:")
    add_code_block(doc, """# Via API (when running)
curl -X POST http://localhost:8000/api/v1/rag/index \\
  -H "Content-Type: application/json" \\
  -d '{"collection": "custom_patterns", "path": "/app/api/rag/knowledge/custom/"}'

# Via Python
from api.rag.indexer import DocumentIndexer
indexer = DocumentIndexer()
stats = indexer.index_directory(
    path="api/rag/knowledge/custom/",
    collection="custom_patterns",
    doc_type="custom",
)
print(f"Indexed {stats.doc_count} docs, {stats.chunk_count} chunks")""")

    doc.add_heading("4.13 How to Re-Index", level=2)
    add_code_block(doc, """# Re-index a specific collection (drops and recreates)
from api.rag.indexer import DocumentIndexer
indexer = DocumentIndexer()
stats = indexer.reindex_collection("mulesoft_docs")

# Re-index all built-in knowledge
stats = indexer.index_all_knowledge()
# Returns: {"mulesoft": IndexStats(...), "springboot": IndexStats(...)}""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 5: BACKEND ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Backend Architecture", level=1)

    doc.add_heading("5.1 FastAPI Application Structure", level=2)
    doc.add_paragraph("File: api/main.py")
    doc.add_paragraph(
        "The application uses the factory pattern (create_app()) with a lifespan context manager for startup/shutdown. "
        "Startup initializes: async DB engine + tables, Redis connection pool, Qdrant client, and embedding model. "
        "Shutdown disposes all connections in reverse order."
    )

    add_code_block(doc, """api/
├── __init__.py         # __version__ string
├── main.py             # FastAPI app factory, lifespan, middleware, exception handlers
├── config.py           # Pydantic BaseSettings (12-factor app pattern)
├── database.py         # Async SQLAlchemy 2.0 engine, session, Base model
├── dependencies.py     # FastAPI Depends() providers: Redis, Qdrant, Embeddings
├── exceptions.py       # Custom exception hierarchy (AppException base)
├── routers/
│   ├── v1/             # migrations, projects, agents, rag, auth routers
│   └── v2/             # migrations, projects routers (enhanced)
├── agents/             # Multi-agent pipeline (10 modules)
├── rag/                # Self-hosted RAG infrastructure (8 modules + knowledge)
└── tasks/              # Celery task definitions""")

    doc.add_heading("5.2 Configuration (12-Factor App)", level=2)
    doc.add_paragraph("File: api/config.py")
    doc.add_paragraph("All configuration uses Pydantic BaseSettings loaded from environment variables and .env files:")

    add_styled_table(doc,
        ["Settings Class", "Env Prefix", "Key Settings"],
        [
            ["Settings (root)", "None", "ENVIRONMENT, DEBUG, SECRET_KEY, LOG_LEVEL, CORS_ORIGINS, AGENTS_ENABLED"],
            ["DatabaseSettings", "None (aliased)", "DATABASE_URL, DB_POOL_MIN=5, DB_POOL_MAX=20, DB_POOL_RECYCLE=3600"],
            ["RedisSettings", "None (aliased)", "REDIS_URL, CELERY_BROKER_URL, CELERY_RESULT_BACKEND, REDIS_MAX_CONNECTIONS=20"],
            ["QdrantSettings", "None (aliased)", "QDRANT_URL, QDRANT_COLLECTION, QDRANT_API_KEY, QDRANT_GRPC_PORT=6334"],
            ["LLMSettings", "None (aliased)", "ANTHROPIC_API_KEY, OPENAI_API_KEY, GOOGLE_API_KEY, DEEPSEEK_API_KEY, GROQ_API_KEY, DEFAULT_LLM_PROVIDER=anthropic, DEFAULT_LLM_MODEL=claude-sonnet-4-20250514, LLM_MAX_TOKENS=8192, LLM_TEMPERATURE=0.1"],
            ["RAGSettings", "None (aliased)", "RAG_ENABLED=true, EMBEDDING_MODEL=all-MiniLM-L6-v2, EMBEDDING_DIMENSION=384, RAG_CHUNK_SIZE=512, RAG_TOP_K=5, RAG_SCORE_THRESHOLD=0.65"],
            ["SecuritySettings", "None (aliased)", "JWT_ALGORITHM=HS256, ACCESS_TOKEN_EXPIRE_MINUTES=60, RATE_LIMIT_PER_MINUTE=60, TRUSTED_HOSTS=*"],
        ],
        col_widths=[1.5, 1.0, 4.0]
    )

    doc.add_heading("5.3 Database Models", level=2)
    doc.add_paragraph("File: api/database.py")
    doc.add_paragraph("Uses SQLAlchemy 2.0 async with asyncpg driver. Base model provides id (UUID), created_at, updated_at for all tables. Naming convention follows Alembic-friendly patterns (ix_, uq_, ck_, fk_, pk_ prefixes).")

    add_styled_table(doc,
        ["Model", "Table", "Key Fields", "Purpose"],
        [
            ["Base", "N/A", "id (UUID PK), created_at, updated_at", "Abstract base for all models"],
            ["MigrationJob", "migration_jobs", "status, xml_files, result, agent_traces, error", "Tracks migration task lifecycle"],
            ["BuildJob", "build_jobs", "migration_id, status, output_path, logs", "Tracks Maven/Gradle build tasks"],
            ["AgentTrace", "agent_traces", "migration_id, agent_name, status, token_usage, duration_ms", "Per-agent execution records"],
        ],
        col_widths=[1.3, 1.5, 2.5, 1.7]
    )

    doc.add_heading("5.4 API Endpoints", level=2)

    add_styled_table(doc,
        ["Method", "Endpoint", "Status", "Description"],
        [
            ["POST", "/api/v2/migrations", "202 Accepted", "Upload XML files and start migration. Returns job ID. Dispatches Celery task."],
            ["GET", "/api/v2/migrations/{id}", "200 OK", "Get migration status, results, agent traces, generated files."],
            ["GET", "/api/v2/migrations/{id}/download", "200 OK", "Download generated Spring Boot project as ZIP."],
            ["WS", "/ws/migration/{id}", "101 Upgrade", "WebSocket for real-time progress streaming during migration."],
            ["POST", "/api/v2/builds", "202 Accepted", "Trigger Maven/Gradle build for a completed migration."],
            ["GET", "/health", "200 OK", "Liveness probe. Returns {status: 'ok', version: '...'}"],
            ["GET", "/readiness", "200/503", "Deep health check: DB + Redis + Qdrant. Returns per-service status."],
            ["GET", "/metrics", "200 OK", "Prometheus metrics (auto-instrumented via prometheus-fastapi-instrumentator)."],
            ["POST", "/api/v1/rag/index", "200 OK", "Trigger knowledge base indexing for a collection."],
            ["GET", "/api/v1/rag/collections", "200 OK", "List all Qdrant collections with doc counts."],
            ["POST", "/api/v1/rag/search", "200 OK", "Search RAG knowledge base with query."],
            ["GET", "/api/v1/agents/status", "200 OK", "Get agent pipeline status and capabilities."],
            ["POST", "/api/v1/auth/login", "200 OK", "JWT authentication endpoint."],
            ["POST", "/api/v1/auth/register", "201 Created", "User registration."],
        ],
        col_widths=[0.7, 2.5, 0.9, 2.4]
    )

    doc.add_heading("5.5 WebSocket Protocol", level=2)
    doc.add_paragraph("The WebSocket endpoint /ws/migration/{id} streams real-time progress updates:")

    add_styled_table(doc,
        ["Message Type", "Direction", "Payload Example"],
        [
            ["agent_started", "Server -> Client", '{"type": "agent_started", "agent": "planner", "timestamp": "..."}'],
            ["agent_progress", "Server -> Client", '{"type": "agent_progress", "agent": "coder", "progress": 0.5, "message": "Processing element 3/6"}'],
            ["agent_completed", "Server -> Client", '{"type": "agent_completed", "agent": "reviewer", "result": {"score": 7.5}}'],
            ["agent_error", "Server -> Client", '{"type": "agent_error", "agent": "tester", "error": "LLM timeout"}'],
            ["migration_complete", "Server -> Client", '{"type": "migration_complete", "status": "success", "download_url": "/api/v2/migrations/{id}/download"}'],
            ["heartbeat", "Bidirectional", '{"type": "heartbeat", "timestamp": "..."}'],
        ],
        col_widths=[1.5, 1.3, 3.7]
    )

    doc.add_heading("5.6 Celery Task Queues", level=2)
    add_styled_table(doc,
        ["Queue", "Worker", "Concurrency", "Task Types"],
        [
            ["migration", "celery-migration-worker", "2", "Full migration pipeline execution"],
            ["build", "celery-build-worker", "2", "Maven/Gradle project builds"],
            ["indexing", "celery-indexing-worker", "1", "RAG knowledge base indexing (heavy embedding compute)"],
        ],
        col_widths=[1.3, 2.2, 1.2, 1.8]
    )
    doc.add_paragraph("Configuration: max-tasks-per-child prevents memory leaks. Celery Beat handles scheduled re-indexing.")

    doc.add_heading("5.7 Middleware Stack", level=2)
    middleware = [
        ("CORSMiddleware", "Configurable origins via CORS_ORIGINS env var. Default: * (all origins)."),
        ("GZipMiddleware", "Compresses responses > 1000 bytes."),
        ("TrustedHostMiddleware", "Validates Host header against TRUSTED_HOSTS whitelist."),
        ("Request Logging", "Custom middleware logs method, path, status, duration_ms via structlog."),
        ("Correlation ID", "Generated per-request and propagated through all log entries and service calls."),
    ]
    for name, desc in middleware:
        add_bold_paragraph(doc, f"{name}: ", desc)

    doc.add_heading("5.8 Exception Hierarchy", level=2)
    doc.add_paragraph("File: api/exceptions.py")
    add_styled_table(doc,
        ["Exception", "HTTP Status", "Error Code", "Description"],
        [
            ["AppException", "500", "INTERNAL_ERROR", "Base exception; all others inherit from this"],
            ["NotFoundError", "404", "NOT_FOUND", "Resource not found; accepts resource type and ID"],
            ["ValidationError", "422", "VALIDATION_ERROR", "Request validation failed; carries error list"],
            ["LLMError", "502", "LLM_ERROR", "LLM provider error; carries provider and model info"],
            ["RAGError", "500", "RAG_ERROR", "RAG pipeline error; carries stage info"],
            ["AuthenticationError", "401", "AUTHENTICATION_ERROR", "Missing or invalid credentials"],
            ["AuthorizationError", "403", "AUTHORIZATION_ERROR", "Insufficient permissions"],
            ["RateLimitError", "429", "RATE_LIMIT_ERROR", "Too many requests; includes Retry-After header"],
            ["MigrationError", "500", "MIGRATION_ERROR", "Migration processing failed; carries project_id"],
        ],
        col_widths=[1.8, 0.8, 1.8, 2.1]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 6: CORE MIGRATION ENGINE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Core Migration Engine", level=1)
    doc.add_paragraph(
        "The core migration engine consists of 8 Python modules totaling approximately 5,500 lines of code. "
        "These modules were preserved from the original Flask-based migrator and are wrapped by the new "
        "agentic pipeline for enhanced capabilities."
    )

    add_styled_table(doc,
        ["Module", "File", "Lines", "Description"],
        [
            ["MuleSoft XML Parser", "backend/migrator/parser.py", "~800", "Comprehensive parser for all Mule 4 components with 30+ namespace registry"],
            ["Flow Converter", "backend/migrator/flow_converter.py", "~1200", "Converts all MuleSoft flow types to Spring Boot: HTTP, Scheduler, JMS, AMQP, Kafka, VM, File, SFTP, FTP, Email, Salesforce, SQS, Anypoint MQ sources + 50+ processors"],
            ["Connector Mapper", "backend/migrator/connector_mapper.py", "~600", "Maps 30+ MuleSoft connectors to Spring Boot Maven dependencies, properties, and annotations"],
            ["DataWeave Converter", "backend/migrator/dw_converter.py", "~500", "Converts DataWeave 2.0 expressions to Java equivalents"],
            ["Spring Boot Generator", "backend/migrator/generator.py", "~700", "Generates complete Spring Boot project structure: pom.xml, application.properties, main class, directory tree"],
            ["Swagger/OpenAPI Generator", "backend/migrator/swagger_generator.py", "~400", "Generates OpenAPI 3.0 spec from converted REST endpoints"],
            ["LLM Agent", "backend/migrator/llm_agent.py", "~600", "Multi-provider LLM integration: convert_unknown_element, convert_unknown_dataweave, suggest_connector_mapping, convert_unknown_source"],
            ["LLM Validator", "backend/migrator/llm_validator.py", "~700", "LLM response validation, scoring, multi-provider support (Anthropic, OpenAI, Google, DeepSeek, Groq)"],
        ],
        col_widths=[1.4, 2.2, 0.5, 2.4]
    )

    doc.add_heading("6.1 MuleSoft XML Parser", level=2)
    doc.add_paragraph("Parses all Mule 4 configuration using lxml with a complete namespace registry:")
    namespaces = [
        "Core: mule, ee (Enterprise Edition)",
        "Connectivity: http, sockets, tls",
        "Database: db",
        "Messaging: jms, amqp, vm, kafka, anypoint-mq",
        "File: file, sftp, ftp",
        "Email: email",
        "API: apikit",
        "Web Services: ws, wsc (SOAP Consumer)",
        "Object Store: os",
        "Batch: batch",
        "Salesforce: salesforce",
        "AWS: s3, sqs, sns",
        "SAP: sap",
        "Security: spring-security, secure-properties, oauth",
    ]
    for ns in namespaces:
        add_bullet(doc, ns)

    doc.add_heading("6.2 Flow Converter", level=2)
    doc.add_paragraph("Converts all MuleSoft flow types to Spring Boot code. Supported source types: HTTP Listener, Scheduler, JMS, AMQP, Kafka, VM, File, SFTP, FTP, Email, Salesforce, SQS, Anypoint MQ. Supported processors include 50+ elements across Core, HTTP, Database, JMS/AMQP, Kafka, VM, File/SFTP, Email, Object Store, Web Service, Salesforce, AWS S3, AWS SQS, Cache, Validation, and Transformer categories.")

    doc.add_heading("6.3 Connector Mapper", level=2)
    doc.add_paragraph("Static mappings for 30+ MuleSoft connectors to Spring Boot equivalents including Maven coordinates (groupId:artifactId), Spring properties, and configuration annotations. Falls back to LLM for unmapped connectors via suggest_connector_mapping().")

    doc.add_heading("6.4 LLM Agent", level=2)
    doc.add_paragraph("File: backend/migrator/llm_agent.py")
    doc.add_paragraph("Provides the triple-fallback conversion strategy: (1) LLM-generated code, (2) TODO comment, (3) warning in summary. The AgentContext class tracks enabled state, LLM config, conversions, and skipped items. Extended by api/agents/context.py for the agentic pipeline.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 7: DEPLOYMENT GUIDE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Deployment Guide", level=1)

    doc.add_heading("7.1 Mode 1: Docker Compose (Development + Production)", level=2)
    doc.add_paragraph("File: deploy/docker-compose/docker-compose.yml")
    doc.add_paragraph("The production Docker Compose file defines 8 services with health checks, resource limits, and volume mounts:")

    add_styled_table(doc,
        ["Service", "Image", "Ports", "Resources (Limits)", "Health Check"],
        [
            ["api", "Custom (Dockerfile)", "8000:8000", "2 CPU, 4GB RAM", "HTTP GET /health"],
            ["celery-migration-worker", "Custom (Dockerfile)", "None", "2 CPU, 4GB RAM", "Via Redis"],
            ["celery-build-worker", "Custom (Dockerfile)", "None", "1.5 CPU, 2GB RAM", "Via Redis"],
            ["celery-indexing-worker", "Custom (Dockerfile)", "None", "2 CPU, 4GB RAM", "Via Redis"],
            ["celery-beat", "Custom (Dockerfile)", "None", "0.25 CPU, 256MB RAM", "Via Redis"],
            ["postgres", "postgres:16-alpine", "5432:5432", "1 CPU, 1GB RAM", "pg_isready"],
            ["redis", "redis:7-alpine", "6379:6379", "0.5 CPU, 512MB RAM", "redis-cli ping"],
            ["qdrant", "qdrant/qdrant:latest", "6333:6333", "1 CPU, 2GB RAM", "wget /healthz"],
            ["nginx", "nginx:alpine", "80:80, 443:443", "0.5 CPU, 256MB RAM", "wget /health"],
        ],
        col_widths=[1.8, 1.5, 1.0, 1.2, 1.0]
    )

    doc.add_heading("Setup Commands", level=3)
    add_code_block(doc, """# 1. Clone and configure
git clone <repository-url>
cd mulesoft-to-springboot-migrator

# 2. Create .env file
cp .env.example .env
# Edit .env with your API keys and configuration

# 3. Start all services
cd deploy/docker-compose
docker compose up -d

# 4. Verify health
docker compose ps
curl http://localhost:8000/health
curl http://localhost:8000/readiness

# 5. View logs
docker compose logs -f api
docker compose logs -f celery-migration-worker

# 6. Index knowledge base (first run)
curl -X POST http://localhost:8000/api/v1/rag/index

# 7. Stop services
docker compose down

# 8. Stop and remove volumes (full reset)
docker compose down -v""")

    doc.add_heading("Environment Variables (.env.example)", level=3)
    add_code_block(doc, """# === General ===
ENVIRONMENT=production
DEBUG=false
SECRET_KEY=your-secret-key-here
LOG_LEVEL=info
CORS_ORIGINS=http://localhost:3000

# === Database ===
POSTGRES_USER=migrator
POSTGRES_PASSWORD=migrator_secret
POSTGRES_DB=migrator
DATABASE_URL=postgresql+asyncpg://migrator:migrator_secret@postgres:5432/migrator

# === Redis ===
REDIS_URL=redis://redis:6379/0
REDIS_MAXMEMORY=256mb
CELERY_BROKER_URL=redis://redis:6379/1
CELERY_RESULT_BACKEND=redis://redis:6379/1

# === Qdrant ===
QDRANT_URL=http://qdrant:6333
QDRANT_COLLECTION=mulesoft_knowledge

# === LLM Providers (set at least one) ===
ANTHROPIC_API_KEY=
OPENAI_API_KEY=
GOOGLE_API_KEY=
DEEPSEEK_API_KEY=
GROQ_API_KEY=
DEFAULT_LLM_PROVIDER=anthropic
DEFAULT_LLM_MODEL=claude-sonnet-4-20250514

# === RAG ===
RAG_ENABLED=true
EMBEDDING_MODEL=all-MiniLM-L6-v2
AGENTS_ENABLED=true

# === Workers ===
UVICORN_WORKERS=4
CELERY_MIGRATION_CONCURRENCY=2
CELERY_BUILD_CONCURRENCY=2
CELERY_INDEXING_CONCURRENCY=1""")

    doc.add_heading("7.2 Mode 2: Kubernetes / Helm", level=2)
    doc.add_paragraph("For production Kubernetes deployments, the platform provides Helm charts with the following structure:")

    add_code_block(doc, """deploy/helm/migrator/
├── Chart.yaml
├── values.yaml
├── templates/
│   ├── deployment-api.yaml
│   ├── deployment-celery-migration.yaml
│   ├── deployment-celery-build.yaml
│   ├── deployment-celery-indexing.yaml
│   ├── deployment-celery-beat.yaml
│   ├── service-api.yaml
│   ├── ingress.yaml
│   ├── configmap.yaml
│   ├── secret.yaml
│   ├── hpa.yaml               # Horizontal Pod Autoscaler
│   ├── pdb.yaml               # Pod Disruption Budget
│   ├── networkpolicy.yaml     # Network segmentation
│   └── _helpers.tpl
└── charts/
    ├── postgresql/             # Bitnami subchart
    ├── redis/                  # Bitnami subchart
    └── qdrant/                 # Qdrant Helm chart""")

    doc.add_heading("Deploy Commands", level=3)
    add_code_block(doc, """# Add Helm repos
helm repo add bitnami https://charts.bitnami.com/bitnami
helm repo update

# Install
helm install migrator deploy/helm/migrator/ \\
  --namespace migrator --create-namespace \\
  --values deploy/helm/migrator/values.yaml \\
  --set api.replicas=3 \\
  --set api.image.tag=latest

# Upgrade
helm upgrade migrator deploy/helm/migrator/ --namespace migrator

# Check status
kubectl get pods -n migrator
kubectl logs -f deployment/migrator-api -n migrator""")

    add_bold_paragraph(doc, "HPA Configuration: ", "API pods auto-scale 2-10 replicas based on CPU (70%) and memory (80%) thresholds.")
    add_bold_paragraph(doc, "PDB: ", "Pod Disruption Budget ensures at least 1 API pod and 1 migration worker are always available during rolling updates.")
    add_bold_paragraph(doc, "NetworkPolicy: ", "Restricts ingress to API pods from nginx only; all internal communication within the migrator namespace.")

    doc.add_heading("7.3 Mode 3: Air-Gap Bundle", level=2)
    doc.add_paragraph("For disconnected environments with no internet access:")

    add_bold_paragraph(doc, "Bundle Contents (~3GB tarball):")
    bundle_items = [
        "All Docker images (pre-built and exported via docker save)",
        "sentence-transformers model (all-MiniLM-L6-v2, ~80MB)",
        "Pre-indexed knowledge base (Qdrant snapshot with ~800 documents)",
        "Complete application source and configuration",
        "Install script (install.sh) for offline setup",
    ]
    for item in bundle_items:
        add_bullet(doc, item)

    add_code_block(doc, """# On internet-connected machine: create bundle
./deploy/airgap/bundle.sh

# Transfer migrator-airgap-bundle.tar.gz to offline machine

# On offline machine: install
tar xzf migrator-airgap-bundle.tar.gz
cd migrator-airgap
./install.sh

# This will:
# 1. Load all Docker images from .tar files
# 2. Copy pre-baked model to model-cache volume
# 3. Restore Qdrant snapshot to qdrant-data volume
# 4. Start docker compose with offline configuration""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 8: MULTI-TENANCY
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Multi-Tenancy", level=1)
    doc.add_paragraph("The platform supports both single-tenant and multi-tenant configurations with 5-layer isolation:")

    add_styled_table(doc,
        ["Layer", "Mechanism", "Implementation"],
        [
            ["1. PostgreSQL", "Row-Level Security (RLS)", "Tenant ID column on all tables with RLS policies. SET app.current_tenant on each session."],
            ["2. Redis", "Key Prefix Namespacing", "All Redis keys prefixed with tenant:{tenant_id}: to prevent cross-tenant data access."],
            ["3. Qdrant", "Collection Namespacing", "Each tenant gets separate collections: {tenant_id}_mulesoft_docs, {tenant_id}_springboot_docs, etc."],
            ["4. Celery", "Queue Routing", "Task routing based on tenant ID. Dedicated queues per tenant tier (standard vs premium)."],
            ["5. API Middleware", "Tenant Resolution", "JWT token carries tenant_id claim. Middleware extracts and injects into request context."],
        ],
        col_widths=[1.2, 1.8, 3.5]
    )

    add_bold_paragraph(doc, "Single-Tenant Mode: ", "Default configuration. No tenant isolation overhead. All data in shared tables/collections.")
    add_bold_paragraph(doc, "Multi-Tenant Mode: ", "Enabled via MULTI_TENANT=true env var. Requires tenant provisioning API to create per-tenant resources.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 9: SECURITY & COMPLIANCE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Security & Compliance", level=1)

    doc.add_heading("9.1 Authentication", level=2)
    doc.add_paragraph("JWT-based authentication using HS256 (configurable to RS256 for production). Access tokens expire after 60 minutes (configurable via ACCESS_TOKEN_EXPIRE_MINUTES). Refresh token rotation supported.")

    doc.add_heading("9.2 RBAC", level=2)
    add_styled_table(doc,
        ["Role", "Permissions"],
        [
            ["admin", "Full access: manage users, configure LLM providers, manage knowledge base, view all migrations"],
            ["user", "Create migrations, view own migrations, download results, search knowledge base"],
            ["viewer", "Read-only: view migration results, browse knowledge base"],
        ],
        col_widths=[1.0, 5.5]
    )

    doc.add_heading("9.3 API Key Encryption", level=2)
    doc.add_paragraph("LLM provider API keys stored in database are encrypted at rest using Fernet symmetric encryption. The encryption key is derived from SECRET_KEY using HKDF. Keys are decrypted only in-memory during LLM calls.")

    doc.add_heading("9.4 Compliance", level=2)
    compliance_items = [
        "SOC 2 Type II: Audit logging on all data access, encryption at rest and in transit, access controls",
        "GDPR: No PII stored in vector database. User data deletion API. Data processing records maintained.",
        "Air-Gapped Security: Zero external API calls for RAG. All models run locally. No telemetry.",
        "Container Security: Non-root container execution, read-only file systems where possible, minimal base images (Alpine), regular vulnerability scanning",
        "Network Security: TLS 1.3 for all external communication, mTLS between services in Kubernetes, NetworkPolicy for pod-level isolation",
    ]
    for item in compliance_items:
        add_bullet(doc, item)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 10: MONITORING & OBSERVABILITY
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Monitoring & Observability", level=1)

    doc.add_heading("10.1 Structured Logging", level=2)
    doc.add_paragraph("All logging uses structlog with JSON output format. Every log entry includes: timestamp, level, logger name, and contextual fields (run_id, agent_name, duration_ms, etc.).")

    doc.add_heading("10.2 Prometheus Metrics", level=2)
    doc.add_paragraph("Exposed via /metrics endpoint using prometheus-fastapi-instrumentator:")

    add_styled_table(doc,
        ["Metric", "Type", "Labels", "Description"],
        [
            ["http_requests_total", "Counter", "method, path, status", "Total HTTP requests"],
            ["http_request_duration_seconds", "Histogram", "method, path", "Request latency distribution"],
            ["http_request_size_bytes", "Histogram", "method, path", "Request body size"],
            ["http_response_size_bytes", "Histogram", "method, path", "Response body size"],
            ["agent_execution_duration_ms", "Histogram", "agent_name, status", "Per-agent execution time"],
            ["agent_token_usage", "Counter", "agent_name", "Total tokens consumed per agent"],
            ["agent_retries_total", "Counter", "agent_name", "LLM retry count per agent"],
            ["agent_fallback_total", "Counter", "agent_name", "Fallback invocations per agent"],
            ["rag_search_duration_ms", "Histogram", "collection", "RAG search latency"],
            ["rag_cache_hits_total", "Counter", "None", "Embedding cache hits"],
            ["rag_cache_misses_total", "Counter", "None", "Embedding cache misses"],
            ["migration_total", "Counter", "status", "Total migration jobs by status"],
            ["migration_duration_seconds", "Histogram", "complexity", "End-to-end migration time"],
            ["celery_tasks_active", "Gauge", "queue", "Currently executing Celery tasks"],
            ["db_pool_size", "Gauge", "None", "Current database connection pool size"],
            ["qdrant_collections_total", "Gauge", "None", "Number of Qdrant collections"],
        ],
        col_widths=[2.3, 0.9, 1.5, 1.8]
    )

    doc.add_heading("10.3 Grafana Dashboards", level=2)
    doc.add_paragraph("Pre-configured dashboards for: API Performance (request rate, latency percentiles, error rate), Agent Pipeline (per-agent duration, token usage, fallback rate), RAG Performance (search latency, cache hit rate, collection sizes), Infrastructure (CPU, memory, disk per service).")

    doc.add_heading("10.4 Distributed Tracing", level=2)
    doc.add_paragraph("OpenTelemetry integration with span hierarchy: HTTP request -> Agent pipeline -> Individual agent -> LLM call / RAG search. Correlation ID propagated through all layers via structlog context binding.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 11: FRONTEND ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("11. Frontend Architecture", level=1)

    doc.add_heading("11.1 Technology Stack", level=2)
    add_styled_table(doc,
        ["Technology", "Version", "Purpose"],
        [
            ["React", "18.2", "UI framework with concurrent features"],
            ["TypeScript", "5.x", "Type safety across all components"],
            ["Vite", "5.x", "Build tool with HMR (Hot Module Replacement)"],
            ["Zustand", "4.x", "Lightweight state management (replaces Redux)"],
            ["React Flow", "11.x", "Agent pipeline visualization (node-based graph)"],
            ["Axios", "1.x", "HTTP client for REST API calls"],
            ["React Query", "5.x", "Server state management with caching"],
            ["Tailwind CSS", "3.x", "Utility-first CSS framework"],
        ],
        col_widths=[1.5, 0.8, 4.2]
    )

    doc.add_heading("11.2 Component Structure", level=2)
    add_code_block(doc, """frontend/src/
├── components/
│   ├── MigrationUpload.tsx     # XML file upload with drag-and-drop
│   ├── MigrationProgress.tsx   # Real-time progress with WebSocket
│   ├── AgentPipeline.tsx       # React Flow visualization of agent DAG
│   ├── CodeViewer.tsx          # Generated code with syntax highlighting
│   ├── ReviewFeedback.tsx      # Per-file review scores and comments
│   ├── ProjectDownload.tsx     # ZIP download with file tree preview
│   └── LLMConfig.tsx           # Provider selection and API key management
├── stores/
│   ├── migrationStore.ts       # Migration job state (Zustand)
│   ├── agentStore.ts           # Agent pipeline state
│   └── settingsStore.ts        # User preferences and LLM config
├── hooks/
│   ├── useWebSocket.ts         # WebSocket connection with auto-reconnect
│   └── useMigration.ts         # Migration CRUD operations
├── services/
│   └── api.ts                  # Axios instance with interceptors
└── App.tsx""")

    doc.add_heading("11.3 WebSocket Integration", level=2)
    doc.add_paragraph("The useWebSocket hook manages the connection lifecycle with automatic reconnection (exponential backoff, max 5 retries). Messages are typed via discriminated unions matching the server protocol. Agent state updates flow through Zustand stores to reactive UI components.")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 12: TROUBLESHOOTING GUIDE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("12. Troubleshooting Guide", level=1)

    doc.add_heading("12.1 Common Issues", level=2)
    issues = [
        ("Qdrant connection refused", "Ensure Qdrant is running: docker compose ps qdrant. Check QDRANT_URL env var. Verify port 6333 is accessible."),
        ("Embedding model fails to load", "Check disk space (~500MB needed). Verify EMBEDDING_MODEL name. For air-gap: ensure model files are in model-cache volume."),
        ("LLM timeout errors", "Increase LLM_REQUEST_TIMEOUT (default 120s). Check API key validity. Verify network connectivity to LLM provider."),
        ("Migration stuck in 'pending'", "Check Celery migration worker: docker compose logs celery-migration-worker. Verify Redis broker connectivity."),
        ("Out of memory errors", "Reduce CELERY_MIGRATION_CONCURRENCY. Increase container memory limits. Check for memory leaks via docker stats."),
        ("Database connection pool exhausted", "Increase DB_POOL_MAX. Reduce concurrent migrations. Check for connection leaks in slow queries."),
        ("RAG search returns empty results", "Verify knowledge base is indexed: curl http://localhost:6333/collections. Re-index if needed. Check RAG_SCORE_THRESHOLD (lower for more results)."),
        ("WebSocket disconnects", "Check Nginx proxy_read_timeout (increase to 300s). Verify heartbeat is being sent. Check browser console for errors."),
    ]
    for title_text, solution in issues:
        add_bold_paragraph(doc, f"{title_text}: ", solution)

    doc.add_heading("12.2 Health Check Commands", level=2)
    add_code_block(doc, """# Overall service status
docker compose ps

# API health
curl http://localhost:8000/health
curl http://localhost:8000/readiness

# Database
docker compose exec postgres pg_isready -U migrator

# Redis
docker compose exec redis redis-cli ping

# Qdrant
curl http://localhost:6333/healthz
curl http://localhost:6333/collections

# Celery workers
docker compose exec celery-migration-worker celery -A api.tasks.celery_app inspect active

# Container resource usage
docker stats --no-stream""")

    doc.add_heading("12.3 Log Locations", level=2)
    add_styled_table(doc,
        ["Service", "Log Command", "Key Fields"],
        [
            ["API", "docker compose logs -f api", "request, agent.metrics, rag_query"],
            ["Migration Worker", "docker compose logs -f celery-migration-worker", "task.received, task.succeeded, task.failed"],
            ["Build Worker", "docker compose logs -f celery-build-worker", "build.started, build.completed"],
            ["Indexing Worker", "docker compose logs -f celery-indexing-worker", "indexing.progress, indexing.complete"],
            ["PostgreSQL", "docker compose logs -f postgres", "LOG: connection, ERROR: queries"],
            ["Redis", "docker compose logs -f redis", "Client connected, AOF rewrite"],
            ["Qdrant", "docker compose logs -f qdrant", "Collection created, Search performed"],
        ],
        col_widths=[1.5, 2.8, 2.2]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX A: SAMPLE MULESOFT XML
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Appendix A: Sample MuleSoft XML", level=1)
    add_code_block(doc, """<?xml version="1.0" encoding="UTF-8"?>
<mule xmlns="http://www.mulesoft.org/schema/mule/core"
      xmlns:http="http://www.mulesoft.org/schema/mule/http"
      xmlns:db="http://www.mulesoft.org/schema/mule/db"
      xmlns:ee="http://www.mulesoft.org/schema/mule/ee/core"
      xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">

    <http:listener-config name="HTTP_Listener_config"
                          host="0.0.0.0" port="8081">
        <http:listener-connection host="0.0.0.0" port="8081"/>
    </http:listener-config>

    <db:config name="Database_Config">
        <db:my-sql-connection host="localhost" port="3306"
                              user="root" password="secret" database="mydb"/>
    </db:config>

    <flow name="get-users-flow">
        <http:listener config-ref="HTTP_Listener_config" path="/api/users"
                       allowedMethods="GET"/>
        <logger level="INFO" message="Fetching all users"/>
        <db:select config-ref="Database_Config">
            <db:sql>SELECT * FROM users WHERE active = true</db:sql>
        </db:select>
        <ee:transform>
            <ee:message>
                <ee:set-payload><![CDATA[%dw 2.0
output application/json
---
payload map {
    id: $.id,
    name: $.first_name ++ " " ++ $.last_name,
    email: $.email
}]]></ee:set-payload>
            </ee:message>
        </ee:transform>
        <logger level="INFO" message="Returning #[sizeOf(payload)] users"/>
    </flow>

    <flow name="create-user-flow">
        <http:listener config-ref="HTTP_Listener_config" path="/api/users"
                       allowedMethods="POST"/>
        <logger level="INFO" message="Creating new user"/>
        <set-variable variableName="newUser" value="#[payload]"/>
        <db:insert config-ref="Database_Config">
            <db:sql>INSERT INTO users (first_name, last_name, email)
                     VALUES (:firstName, :lastName, :email)</db:sql>
            <db:input-parameters><![CDATA[#[{
                firstName: vars.newUser.firstName,
                lastName: vars.newUser.lastName,
                email: vars.newUser.email
            }]]]></db:input-parameters>
        </db:insert>
        <set-payload value='{"status": "created", "id": #[payload.generatedKeys.GENERATED_KEY]}'
                     mimeType="application/json"/>
        <error-handler>
            <on-error-propagate type="DB:CONNECTIVITY">
                <set-payload value='{"error": "Database connection failed"}'
                             mimeType="application/json"/>
                <set-variable variableName="httpStatus" value="503"/>
            </on-error-propagate>
        </error-handler>
    </flow>
</mule>""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX B: GENERATED SPRING BOOT CODE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Appendix B: Generated Spring Boot Code", level=1)
    doc.add_paragraph("Corresponding output for the MuleSoft XML in Appendix A:")

    doc.add_heading("UserController.java", level=2)
    add_code_block(doc, """package com.example.controller;

import com.example.service.UserService;
import com.example.model.UserDto;
import org.slf4j.Logger;
import org.slf4j.LoggerFactory;
import org.springframework.http.HttpStatus;
import org.springframework.http.ResponseEntity;
import org.springframework.web.bind.annotation.*;

import java.util.List;

@RestController
@RequestMapping("/api/users")
public class UserController {

    private static final Logger log = LoggerFactory.getLogger(UserController.class);
    private final UserService userService;

    public UserController(UserService userService) {
        this.userService = userService;
    }

    @GetMapping
    public ResponseEntity<List<UserDto>> getUsers() {
        log.info("Fetching all users");
        List<UserDto> users = userService.getAllActiveUsers();
        log.info("Returning {} users", users.size());
        return ResponseEntity.ok(users);
    }

    @PostMapping
    public ResponseEntity<?> createUser(@RequestBody UserDto newUser) {
        log.info("Creating new user");
        try {
            UserDto created = userService.createUser(newUser);
            return ResponseEntity.status(HttpStatus.CREATED)
                .body(Map.of("status", "created", "id", created.getId()));
        } catch (DataAccessException e) {
            log.error("Database connection failed", e);
            return ResponseEntity.status(HttpStatus.SERVICE_UNAVAILABLE)
                .body(Map.of("error", "Database connection failed"));
        }
    }
}""")

    doc.add_heading("UserService.java", level=2)
    add_code_block(doc, """package com.example.service;

import com.example.model.UserDto;
import com.example.repository.UserRepository;
import org.springframework.stereotype.Service;
import org.springframework.transaction.annotation.Transactional;

import java.util.List;
import java.util.stream.Collectors;

@Service
public class UserService {

    private final UserRepository userRepository;

    public UserService(UserRepository userRepository) {
        this.userRepository = userRepository;
    }

    @Transactional(readOnly = true)
    public List<UserDto> getAllActiveUsers() {
        return userRepository.findByActiveTrue()
            .stream()
            .map(user -> new UserDto(
                user.getId(),
                user.getFirstName() + " " + user.getLastName(),
                user.getEmail()))
            .collect(Collectors.toList());
    }

    @Transactional
    public UserDto createUser(UserDto dto) {
        User user = new User();
        user.setFirstName(dto.getFirstName());
        user.setLastName(dto.getLastName());
        user.setEmail(dto.getEmail());
        user.setActive(true);
        User saved = userRepository.save(user);
        return new UserDto(saved.getId(),
            saved.getFirstName() + " " + saved.getLastName(),
            saved.getEmail());
    }
}""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX C: QUICK REFERENCE COMMANDS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Appendix C: Quick Reference Commands", level=1)

    doc.add_heading("Docker Commands", level=2)
    add_code_block(doc, """# Start all services
docker compose -f deploy/docker-compose/docker-compose.yml up -d

# Stop all services
docker compose -f deploy/docker-compose/docker-compose.yml down

# View logs
docker compose -f deploy/docker-compose/docker-compose.yml logs -f api
docker compose -f deploy/docker-compose/docker-compose.yml logs -f celery-migration-worker

# Restart a service
docker compose -f deploy/docker-compose/docker-compose.yml restart api

# Scale workers
docker compose -f deploy/docker-compose/docker-compose.yml up -d --scale celery-migration-worker=3

# Shell into container
docker compose -f deploy/docker-compose/docker-compose.yml exec api bash

# Resource usage
docker stats --no-stream""")

    doc.add_heading("Helm Commands", level=2)
    add_code_block(doc, """# Install
helm install migrator deploy/helm/migrator/ -n migrator --create-namespace

# Upgrade
helm upgrade migrator deploy/helm/migrator/ -n migrator

# Rollback
helm rollback migrator 1 -n migrator

# Uninstall
helm uninstall migrator -n migrator

# View values
helm get values migrator -n migrator

# Template (dry run)
helm template migrator deploy/helm/migrator/ -n migrator""")

    doc.add_heading("API Commands", level=2)
    add_code_block(doc, """# Health check
curl http://localhost:8000/health
curl http://localhost:8000/readiness

# Start migration
curl -X POST http://localhost:8000/api/v2/migrations \\
  -F "files=@mulesoft-app.xml" \\
  -H "Authorization: Bearer <token>"

# Check migration status
curl http://localhost:8000/api/v2/migrations/{job_id}

# Download result
curl -o result.zip http://localhost:8000/api/v2/migrations/{job_id}/download

# Search RAG
curl -X POST http://localhost:8000/api/v1/rag/search \\
  -H "Content-Type: application/json" \\
  -d '{"query": "HTTP Listener to RestController", "top_k": 5}'

# List Qdrant collections
curl http://localhost:6333/collections

# Trigger re-index
curl -X POST http://localhost:8000/api/v1/rag/index""")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX D: COMPLETE FILE INVENTORY
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Appendix D: Complete File Inventory", level=1)

    add_styled_table(doc,
        ["Path", "Description"],
        [
            ["api/__init__.py", "Package init with __version__"],
            ["api/main.py", "FastAPI app factory, lifespan, middleware, exception handlers"],
            ["api/config.py", "Pydantic BaseSettings: Database, Redis, Qdrant, LLM, RAG, Security"],
            ["api/database.py", "Async SQLAlchemy 2.0 engine, session factory, Base model"],
            ["api/dependencies.py", "FastAPI dependency injection: Redis, Qdrant, Embeddings, Retriever"],
            ["api/exceptions.py", "Exception hierarchy: App, NotFound, Validation, LLM, RAG, Auth, RateLimit, Migration"],
            ["api/agents/__init__.py", "Agent package init"],
            ["api/agents/base.py", "BaseAgent ABC: LLM retry, token budget, RAG, structured output, metrics"],
            ["api/agents/planner.py", "PlannerAgent: XML analysis, complexity scoring, pipeline decisions"],
            ["api/agents/coder.py", "CoderAgent: RAG-enhanced code generation, validation loop, fallback"],
            ["api/agents/reviewer.py", "ReviewerAgent: per-file review, scoring, autonomous fix loop"],
            ["api/agents/tester.py", "TesterAgent: JUnit 5 generation, MockMvc, Mockito patterns"],
            ["api/agents/docs.py", "DocsAgent: README, migration report, OpenAPI enhancements"],
            ["api/agents/memory.py", "AgentMemory: short-term dict + long-term Qdrant-backed"],
            ["api/agents/guardrails.py", "AgentGuardrails: syntax, annotations, security, hallucination detection"],
            ["api/agents/tools.py", "AgentTools: parse_xml, validate_java, search_rag, get_spring_docs"],
            ["api/agents/context.py", "AgentContext: UUID tracking, traces, artifacts, messaging, checkpoints"],
            ["api/agents/result.py", "AgentResult: status, output, tokens, RAG attribution, fallback flag"],
            ["api/rag/__init__.py", "RAG package init"],
            ["api/rag/config.py", "RAG config: Embedding, Chunk, Search, Qdrant, Redis, Cache settings"],
            ["api/rag/embeddings.py", "EmbeddingService: singleton, lazy-load, device detection, L2 normalize"],
            ["api/rag/vector_store.py", "QdrantStore: connection pooling, retry, collection CRUD, search"],
            ["api/rag/chunking.py", "CodeAwareChunker: text, XML, Java, Markdown chunking with overlap"],
            ["api/rag/retriever.py", "HybridRetriever: query expansion, search, dedup, threshold, packing"],
            ["api/rag/cache.py", "EmbeddingCache: Redis-backed, SHA-256 keys, TTL, hit rate tracking"],
            ["api/rag/indexer.py", "DocumentIndexer: scan, dedup, chunk, embed, upsert pipeline"],
            ["api/rag/schemas.py", "Pydantic models: Document, Chunk, SearchResult, IndexStats"],
            ["api/rag/knowledge/", "Built-in knowledge base (~800 documents)"],
            ["backend/migrator/parser.py", "MuleSoft XML parser with 30+ namespace registry"],
            ["backend/migrator/flow_converter.py", "Flow converter for 50+ processor types"],
            ["backend/migrator/connector_mapper.py", "30+ connector-to-Spring Boot dependency mappings"],
            ["backend/migrator/dw_converter.py", "DataWeave 2.0 to Java expression converter"],
            ["backend/migrator/generator.py", "Spring Boot project structure generator"],
            ["backend/migrator/swagger_generator.py", "OpenAPI 3.0 spec generator"],
            ["backend/migrator/llm_agent.py", "Multi-provider LLM conversion functions"],
            ["backend/migrator/llm_validator.py", "LLM response validation and scoring"],
            ["docker-compose.yml", "Dev Docker Compose (app + nginx)"],
            ["deploy/docker-compose/docker-compose.yml", "Production Docker Compose (9 services)"],
            ["deploy/helm/migrator/", "Kubernetes Helm chart"],
            ["deploy/airgap/bundle.sh", "Air-gap bundle creation script"],
            ["deploy/airgap/install.sh", "Air-gap installation script"],
            ["frontend/", "React 18 + TypeScript + Vite SPA"],
            ["nginx/nginx.conf", "Nginx reverse proxy configuration"],
        ],
        col_widths=[2.8, 3.7]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # APPENDIX E: ENVIRONMENT VARIABLES REFERENCE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading("Appendix E: Environment Variables Reference", level=1)

    add_styled_table(doc,
        ["Variable", "Default", "Description"],
        [
            ["ENVIRONMENT", "production", "Application environment: development, staging, production, testing"],
            ["DEBUG", "false", "Enable debug mode with verbose logging"],
            ["SECRET_KEY", "(auto-generated)", "Secret key for JWT signing and encryption (48-char URL-safe)"],
            ["LOG_LEVEL", "info", "Logging level: debug, info, warning, error"],
            ["CORS_ORIGINS", "*", "Comma-separated list of allowed CORS origins"],
            ["AGENTS_ENABLED", "true", "Enable/disable the agentic AI pipeline"],
            ["DATABASE_URL", "postgresql+asyncpg://migrator:migrator_secret@localhost:5432/migrator", "PostgreSQL connection string"],
            ["DB_POOL_MIN", "5", "Minimum database connection pool size"],
            ["DB_POOL_MAX", "20", "Maximum database connection pool size"],
            ["DB_POOL_RECYCLE", "3600", "Connection recycle time in seconds"],
            ["DB_ECHO", "false", "Echo SQL queries to log"],
            ["REDIS_URL", "redis://localhost:6379/0", "Redis connection URL for app cache"],
            ["REDIS_MAX_CONNECTIONS", "20", "Maximum Redis connection pool size"],
            ["REDIS_SOCKET_TIMEOUT", "5.0", "Redis socket timeout in seconds"],
            ["REDIS_MAXMEMORY", "256mb", "Redis maximum memory (Docker)"],
            ["CELERY_BROKER_URL", "redis://localhost:6379/1", "Celery message broker URL"],
            ["CELERY_RESULT_BACKEND", "redis://localhost:6379/1", "Celery result backend URL"],
            ["CELERY_MIGRATION_CONCURRENCY", "2", "Migration worker concurrency"],
            ["CELERY_BUILD_CONCURRENCY", "2", "Build worker concurrency"],
            ["CELERY_INDEXING_CONCURRENCY", "1", "Indexing worker concurrency"],
            ["QDRANT_URL", "http://localhost:6333", "Qdrant REST API URL"],
            ["QDRANT_API_KEY", "(none)", "Qdrant API key (optional)"],
            ["QDRANT_GRPC_PORT", "6334", "Qdrant gRPC port"],
            ["QDRANT_PREFER_GRPC", "false", "Use gRPC for Qdrant communication"],
            ["QDRANT_COLLECTION", "mulesoft_knowledge", "Default Qdrant collection name"],
            ["QDRANT_HOST", "localhost", "Qdrant host (RAG config)"],
            ["QDRANT_PORT", "6333", "Qdrant port (RAG config)"],
            ["RAG_ENABLED", "true", "Enable/disable RAG pipeline"],
            ["EMBEDDING_MODEL", "all-MiniLM-L6-v2", "Sentence transformer model name"],
            ["EMBEDDING_DIMENSION", "384", "Embedding vector dimensionality"],
            ["RAG_CHUNK_SIZE", "512", "Maximum chunk size in tokens"],
            ["RAG_CHUNK_OVERLAP", "64", "Chunk overlap in tokens"],
            ["RAG_TOP_K", "5", "Number of RAG results to return"],
            ["RAG_SCORE_THRESHOLD", "0.65", "Minimum similarity score (0.0-1.0)"],
            ["REDIS_HOST", "localhost", "Redis host (RAG cache)"],
            ["REDIS_PORT", "6379", "Redis port (RAG cache)"],
            ["REDIS_RAG_DB", "2", "Redis database number for RAG cache"],
            ["REDIS_PASSWORD", "(none)", "Redis password (RAG cache)"],
            ["ANTHROPIC_API_KEY", "(none)", "Anthropic (Claude) API key"],
            ["OPENAI_API_KEY", "(none)", "OpenAI (GPT-4) API key"],
            ["GOOGLE_API_KEY", "(none)", "Google (Gemini) API key"],
            ["DEEPSEEK_API_KEY", "(none)", "DeepSeek API key"],
            ["GROQ_API_KEY", "(none)", "Groq API key"],
            ["DEFAULT_LLM_PROVIDER", "anthropic", "Default LLM provider name"],
            ["DEFAULT_LLM_MODEL", "claude-sonnet-4-20250514", "Default LLM model name"],
            ["LLM_MAX_TOKENS", "8192", "Maximum tokens for LLM responses"],
            ["LLM_TEMPERATURE", "0.1", "LLM sampling temperature"],
            ["LLM_REQUEST_TIMEOUT", "120.0", "LLM request timeout in seconds"],
            ["JWT_ALGORITHM", "HS256", "JWT signing algorithm"],
            ["ACCESS_TOKEN_EXPIRE_MINUTES", "60", "JWT access token expiry"],
            ["RATE_LIMIT_PER_MINUTE", "60", "API rate limit per minute"],
            ["TRUSTED_HOSTS", "*", "Comma-separated trusted hostnames"],
            ["UVICORN_WORKERS", "4", "Number of Uvicorn worker processes"],
            ["API_PORT", "8000", "API server port (Docker)"],
            ["NGINX_HTTP_PORT", "80", "Nginx HTTP port"],
            ["NGINX_HTTPS_PORT", "443", "Nginx HTTPS port"],
            ["POSTGRES_USER", "migrator", "PostgreSQL username (Docker)"],
            ["POSTGRES_PASSWORD", "migrator_secret", "PostgreSQL password (Docker)"],
            ["POSTGRES_DB", "migrator", "PostgreSQL database name (Docker)"],
        ],
        col_widths=[2.2, 1.5, 2.8]
    )

    # ══════════════════════════════════════════════════════════════════════
    # End matter
    # ══════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_p.add_run("--- End of Document ---")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.add_paragraph()
    gen_p = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen_p.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    return doc


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Agentic_AI_Platform_Complete_Documentation.docx",
    )

    print("Generating comprehensive documentation...")
    doc = generate_document()

    print(f"Saving to: {output_path}")
    doc.save(output_path)
    print(f"Done! Document saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
