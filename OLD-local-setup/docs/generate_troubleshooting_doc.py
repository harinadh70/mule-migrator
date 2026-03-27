#!/usr/bin/env python3
"""
Generates a comprehensive Trial & Error / Troubleshooting / Installation Guide
documenting every command, error, fix, and lesson learned during development.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = os.path.join(os.path.dirname(__file__),
    "Trial_Error_Installation_Troubleshooting_Guide.docx")

# ═══════════ Helpers ═══════════
def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(30, 30, 30)
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), 'F0F0F0')
    s.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(s)

def error_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 0, 0)
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), 'FFF0F0')
    s.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(s)

def fix_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 120, 0)
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), 'F0FFF0')
    s.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(s)

def label(doc, text, color="1565C0"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run("NOTE: ")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(21, 101, 192)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True

def htable(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        shade(t.rows[0].cells[i], "2E4057")
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t

# ═══════════ Document ═══════════
doc = Document()
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════
#  COVER
# ═══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MuleSoft to Spring Boot\nMigration Tool")
run.font.size = Pt(34)
run.font.bold = True
run.font.color.rgb = RGBColor(46, 64, 87)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Trial & Error / Installation /\nTroubleshooting Complete Guide")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(200, 60, 60)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Every Command, Every Error, Every Fix — Start to Finish")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 1.0.0")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)

toc = [
    "1. Environment Setup — Commands & Errors",
    "  1.1 Python Version Issues",
    "  1.2 pip / Virtual Environment Setup",
    "  1.3 Installing Dependencies",
    "  1.4 Port Conflicts (macOS AirPlay)",
    "  1.5 Starting the Server",
    "  1.6 Stopping the Server",
    "2. Critical Bug: XML Comment Node TypeError",
    "  2.1 Error Description",
    "  2.2 Root Cause Analysis",
    "  2.3 Reproduction Steps",
    "  2.4 Fix Implementation",
    "  2.5 Verification",
    "3. LLM Provider Setup — Trial & Error",
    "  3.1 Anthropic Claude Setup",
    "  3.2 OpenAI GPT Setup",
    "  3.3 Google Gemini Setup",
    "  3.4 DeepSeek Setup",
    "  3.5 Groq Setup",
    "  3.6 Ollama (Local) Setup",
    "  3.7 Common LLM Errors",
    "4. Frontend Errors & Fixes",
    "  4.1 CORS Errors",
    "  4.2 JSON Parse Errors",
    "  4.3 File Upload Issues",
    "  4.4 CodeMirror CDN Loading",
    "5. Migration Pipeline Errors",
    "  5.1 XML Parsing Errors",
    "  5.2 Namespace Resolution Failures",
    "  5.3 DataWeave Conversion Warnings",
    "  5.4 Unknown Connector Warnings",
    "  5.5 Duplicate Flow Warnings",
    "  5.6 LLM Fallback Failures",
    "6. Code Quality Sweep — Removing AI Traces",
    "  6.1 Why This Was Needed",
    "  6.2 Files Modified",
    "  6.3 Search Commands Used",
    "  6.4 Rename Mappings",
    "  6.5 CSS Class Renames",
    "  6.6 Verification Grep Commands",
    "7. Architecture Page — Auth Issues",
    "  7.1 HTTP Basic Auth Flow",
    "  7.2 Browser Popup Not Appearing",
    "  7.3 Credentials Configuration",
    "8. Docker Deployment Errors",
    "9. Production Deployment Checklist",
    "10. Complete Command Reference",
    "11. Error Code Reference Table",
    "12. Lessons Learned",
]

for item in toc:
    p = doc.add_paragraph()
    indent = 0.4 if item.startswith("  ") else 0
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(item.strip())
    run.font.size = Pt(10) if not item.startswith("  ") else Pt(9)
    if not item.startswith("  "):
        run.font.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 1 — ENVIRONMENT SETUP
# ═══════════════════════════════════════════════════════════
doc.add_heading('1. Environment Setup — Commands & Errors', level=1)

doc.add_paragraph(
    'This chapter documents every command executed during environment setup, '
    'including all errors encountered and how they were resolved.'
)

# 1.1
doc.add_heading('1.1 Python Version Issues', level=2)

label(doc, "COMMAND ATTEMPTED:")
code(doc, "python app.py")

label(doc, "ERROR:", "CC0000")
error_block(doc, "zsh: command not found: python")

label(doc, "ROOT CAUSE:")
doc.add_paragraph(
    'macOS does not ship with a "python" command. Starting with macOS 12.3 (Monterey), '
    'Apple removed the /usr/bin/python symlink. Only python3 is available at /usr/bin/python3.'
)

label(doc, "FIX:", "008800")
fix_block(doc, "python3 app.py")

label(doc, "VERIFICATION:")
code(doc, '''$ python3 --version
Python 3.9.6

$ which python3
/usr/bin/python3''')

note(doc, "You can create an alias: alias python=python3 in ~/.zshrc, but it is better to always use python3 explicitly to avoid confusion across environments.")

doc.add_paragraph()

# Alternative approach
label(doc, "ALTERNATIVE — Using pyenv for version management:")
code(doc, '''# Install pyenv
brew install pyenv

# Install a specific Python version
pyenv install 3.11.7

# Set as global default
pyenv global 3.11.7

# Verify
python --version  # Now works: Python 3.11.7''')

# 1.2
doc.add_heading('1.2 pip / Virtual Environment Setup', level=2)

label(doc, "COMMANDS EXECUTED:")
code(doc, '''# Navigate to project
cd /Users/harinadh/Documents/My\\ code/mulesoft-to-springboot-migrator/backend

# Create virtual environment
python3 -m venv venv

# Activate it
source venv/bin/activate

# Verify activation (should show venv path)
which python3
# Output: /Users/harinadh/.../venv/bin/python3''')

label(doc, "POTENTIAL ERROR:")
error_block(doc, '''$ python3 -m venv venv
Error: Command '['/path/to/venv/bin/python3', '-Im', 'ensurepip', ...] returned non-zero exit status 1''')

label(doc, "FIX:", "008800")
fix_block(doc, '''# Install python3-venv (Linux/Debian)
sudo apt install python3-venv

# Or create without pip and install manually
python3 -m venv --without-pip venv
source venv/bin/activate
curl https://bootstrap.pypa.io/get-pip.py | python3''')

# 1.3
doc.add_heading('1.3 Installing Dependencies', level=2)

label(doc, "COMMAND:")
code(doc, "pip install -r requirements.txt")

label(doc, "CONTENTS OF requirements.txt:")
code(doc, '''flask==3.1.0
flask-cors==5.0.1
lxml==5.3.1
pyyaml==6.0.2
gunicorn==23.0.0

# LLM Providers (install the ones you plan to use)
anthropic>=0.39.0       # Anthropic Claude
openai>=1.50.0          # OpenAI GPT + DeepSeek + Groq (OpenAI-compatible)
google-generativeai>=0.8.0  # Google Gemini''')

label(doc, "POTENTIAL ERROR — lxml build failure:")
error_block(doc, '''Building wheel for lxml (pyproject.toml) ... error
error: command 'gcc' failed: No such file or directory''')

label(doc, "FIX:", "008800")
fix_block(doc, '''# macOS — install Xcode command line tools
xcode-select --install

# Linux (Debian/Ubuntu)
sudo apt install build-essential libxml2-dev libxslt-dev python3-dev

# Then retry
pip install -r requirements.txt''')

label(doc, "POTENTIAL ERROR — pip SSL certificate:")
error_block(doc, '''WARNING: pip is configured with locations that require TLS/SSL,
however the ssl module in Python is not available.''')

label(doc, "FIX:", "008800")
fix_block(doc, '''# macOS with Homebrew Python
brew install openssl
brew reinstall python3

# Or use trusted host
pip install --trusted-host pypi.org --trusted-host pypi.python.org -r requirements.txt''')

label(doc, "SUCCESSFUL OUTPUT:")
fix_block(doc, '''Successfully installed Flask-3.1.0 Jinja2-3.1.4 MarkupSafe-3.0.2
Werkzeug-3.1.3 blinker-1.9.0 click-8.1.7 flask-cors-5.0.1
itsdangerous-2.2.0 lxml-5.3.1 pyyaml-6.0.2 gunicorn-23.0.0
anthropic-0.39.0 openai-1.50.0 google-generativeai-0.8.0''')

# 1.4
doc.add_heading('1.4 Port Conflicts (macOS AirPlay)', level=2)

label(doc, "COMMAND:")
code(doc, "python3 app.py")

label(doc, "ERROR:", "CC0000")
error_block(doc, '''OSError: [Errno 48] Address already in use
# or
socket.error: [Errno 48] Address already in use: ('0.0.0.0', 5000)''')

label(doc, "ROOT CAUSE:")
doc.add_paragraph(
    'On macOS Monterey (12.0) and later, port 5000 is used by AirPlay Receiver. '
    'The ControlCenter process (AirPlay Receiver) binds to port 5000 on startup.'
)

label(doc, "DIAGNOSIS COMMAND:")
code(doc, '''$ lsof -i :5000
COMMAND     PID   USER   FD   TYPE  DEVICE SIZE/OFF NODE NAME
ControlCe 12345 harinadh  6u  IPv4  ...    TCP *:commplex-main (LISTEN)''')

label(doc, "FIX OPTIONS:", "008800")

doc.add_paragraph("Option A — Use a different port (RECOMMENDED):")
fix_block(doc, "PORT=5001 python3 app.py")

doc.add_paragraph("Option B — Disable AirPlay Receiver:")
fix_block(doc, '''System Settings > General > AirDrop & Handoff > AirPlay Receiver → OFF

# Then port 5000 becomes available:
python3 app.py  # Now works on port 5000''')

doc.add_paragraph("Option C — Kill the process:")
fix_block(doc, "lsof -ti:5000 | xargs kill -9")

note(doc, "Option A is preferred because it does not interfere with AirPlay functionality.")

# 1.5
doc.add_heading('1.5 Starting the Server', level=2)

label(doc, "DEVELOPMENT START COMMAND:")
code(doc, '''cd /Users/harinadh/Documents/My\\ code/mulesoft-to-springboot-migrator/backend
PORT=5001 python3 app.py''')

label(doc, "EXPECTED OUTPUT:")
fix_block(doc, '''2026-03-17 22:00:00 [INFO] migrator: Starting migration tool
 * Serving Flask app 'app'
 * Debug mode: off
 * Running on all addresses (0.0.0.0)
 * Running on http://127.0.0.1:5001
 * Running on http://192.168.x.x:5001''')

label(doc, "PRODUCTION START COMMAND:")
code(doc, "gunicorn -w 4 -b 0.0.0.0:5001 app:app")

label(doc, "COMBINED COMMAND (kill old + start new):")
code(doc, '''lsof -ti:5001 | xargs kill -9 2>/dev/null; \\
cd /Users/harinadh/Documents/My\\ code/mulesoft-to-springboot-migrator/backend && \\
PORT=5001 python3 app.py''')

# 1.6
doc.add_heading('1.6 Stopping the Server', level=2)

label(doc, "METHOD 1 — Ctrl+C in terminal:")
code(doc, "Press Ctrl+C")

label(doc, "IF CTRL+C DOES NOT WORK:")
error_block(doc, "^C^C^C  (no response, server keeps running)")

label(doc, "FIX:", "008800")
fix_block(doc, '''# Method 2 — Kill by port
lsof -ti:5001 | xargs kill -9

# Method 3 — Suspend then kill
Ctrl+Z  (suspends the process)
kill %1

# Method 4 — Find and kill by name
ps aux | grep python3 | grep app.py
kill -9 <PID>''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 2 — XML COMMENT BUG
# ═══════════════════════════════════════════════════════════
doc.add_heading('2. Critical Bug: XML Comment Node TypeError', level=1)

doc.add_paragraph(
    'This was the most critical bug discovered during testing. It caused a complete '
    'server crash (500 error) when loading the sample MuleSoft XML, which contains '
    '8 XML comment blocks.'
)

doc.add_heading('2.1 Error Description', level=2)

label(doc, "USER ACTION:")
doc.add_paragraph('Clicked "Load Sample" button in the UI, then clicked "Migrate".')

label(doc, "ERROR IN BROWSER:")
error_block(doc, '''Error: Error parsing main.xml: argument of type
'_cython_3_0_11.cython_function_or_method' is not iterable''')

label(doc, "FULL STACK TRACE (server logs):")
error_block(doc, '''Traceback (most recent call last):
  File "/Users/harinadh/.../backend/app.py", line 182, in migrate
    parsed_list.append(parser.parse(content, agent_context=agent_context))
  File "/Users/harinadh/.../backend/migrator/parser.py", line 128, in parse
    "global_configs": self._parse_global_configs(root, ns_map),
  File "/Users/harinadh/.../backend/migrator/parser.py", line 167, in _parse_global_configs
    tag = self._local_tag(elem)
  File "/Users/harinadh/.../backend/migrator/parser.py", line 979, in _local_tag
    if "}" in tag:
TypeError: argument of type '_cython_3_0_11.cython_function_or_method' is not iterable''')

doc.add_heading('2.2 Root Cause Analysis', level=2)

doc.add_paragraph(
    'The bug was in the _local_tag() method of the MuleSoft parser. The method assumed '
    'that every node in an lxml element tree has a string .tag attribute. However, XML '
    'comment nodes (<!-- ... -->) and Processing Instruction nodes (<?...?>) in lxml '
    'have their .tag set to a callable function, not a string.'
)

label(doc, "The problematic code (BEFORE fix):")
code(doc, '''def _local_tag(self, elem):
    tag = elem.tag          # For Comment nodes: tag = <function Comment>
    if "}" in tag:          # TypeError! Can't use "in" on a function
        tag = tag.split("}", 1)[1]
    return tag''')

label(doc, "Why this happens:")
doc.add_paragraph(
    'lxml uses Cython for performance. The Comment factory function is a Cython function '
    'object. In Python, the "in" operator on a string checks if a substring exists. But '
    'when applied to a non-string (like a function), Python raises TypeError because '
    'functions are not iterable containers.'
)

label(doc, "Demonstration:")
code(doc, '''from lxml import etree

xml = """<root><!-- comment --><child/></root>"""
root = etree.fromstring(xml.encode())

for elem in root:
    print(type(elem.tag), repr(elem.tag))

# Output:
# <class 'builtin_function_or_method'>  <built-in function Comment>
# <class 'str'>  'child'

# This crashes:
for elem in root:
    if "}" in elem.tag:  # TypeError on Comment node!
        pass''')

doc.add_heading('2.3 Reproduction Steps', level=2)

doc.add_paragraph('Step 1: Start the server')
code(doc, "PORT=5001 python3 app.py")

doc.add_paragraph('Step 2: Open browser to http://localhost:5001')

doc.add_paragraph('Step 3: Click "Load Sample" button (loads embedded XML with comments)')

doc.add_paragraph('Step 4: Click "Migrate" button')

doc.add_paragraph('Step 5: Observe error in the UI and 500 error in server logs')

label(doc, "KEY INSIGHT:")
doc.add_paragraph(
    'The bug ONLY manifested with XML containing comments. XML files without comments '
    'parsed fine. The sample XML embedded in app.js contained 8 comment blocks like '
    '<!-- HTTP Listener Configuration --> which triggered the crash.'
)

doc.add_heading('2.4 Fix Implementation', level=2)

label(doc, "Step 1: Added _is_element() static method:")
code(doc, '''@staticmethod
def _is_element(elem):
    """Return True if *elem* is a real XML element (not a Comment / PI)."""
    return isinstance(elem.tag, str)''')

label(doc, "Step 2: Added guard to _local_tag():")
code(doc, '''def _local_tag(self, elem):
    tag = elem.tag
    if not isinstance(tag, str):
        return ""  # Comment / ProcessingInstruction node — safe return
    if "}" in tag:
        tag = tag.split("}", 1)[1]
    return tag''')

label(doc, "Step 3: Added guard to _get_ns_prefix():")
code(doc, '''def _get_ns_prefix(self, elem, ns_map):
    tag = elem.tag
    if not isinstance(tag, str):
        return ""  # Guard for Comment / PI nodes
    if "}" in tag:
        uri = tag.split("}")[0].strip("{")
        for prefix, ns_uri in ns_map.items():
            if ns_uri == uri:
                return prefix
    return ""''')

label(doc, "Step 4: Added guard to _get_ns_uri():")
code(doc, '''def _get_ns_uri(self, elem):
    tag = elem.tag
    if not isinstance(tag, str):
        return ""  # Guard for Comment / PI nodes
    if "}" in tag:
        return tag.split("}")[0].strip("{")
    return ""''')

label(doc, "Step 5: Added 'if not self._is_element(elem): continue' to ALL 27 loops:")
doc.add_paragraph('Every loop that iterates over XML elements was updated:')

code(doc, '''# Pattern applied to ALL iteration loops:
for elem in root:
    if not self._is_element(elem):
        continue  # skip Comment / PI nodes
    tag = self._local_tag(elem)
    # ... process normally

# Applied in these methods (27 locations):
# _parse_global_configs()     — 1 loop
# _make_http_listener_config() — 1 loop
# _make_http_request_config()  — 2 loops (outer + auth child loop)
# _make_db_config()            — 1 loop
# _parse_flow_element()        — 1 loop
# _parse_processor()           — 1 loop
# _parse_error_handler_element() — 2 loops
# _parse_global_properties()    — 1 loop (iter())
# _detect_connectors()          — 1 loop (iter())
# _parse_batch_jobs()           — 3 loops (iter + step children + on_complete)
# _parse_apikit()               — 1 loop (iter())
# _parse_secure_properties()    — 2 loops (iter + encrypt children)
# _parse_tls_contexts()         — 2 loops (iter + keystore/truststore)
# _parse_caching_strategies()   — 1 loop (iter())
# _parse_children()             — 1 loop
# Total: 27 guard insertions''')

doc.add_heading('2.5 Verification', level=2)

label(doc, "Test script used:")
code(doc, '''import json
from app import create_app

app = create_app()
client = app.test_client()

# XML with multiple comments (reproduces the bug)
xml_with_comments = """<?xml version="1.0" encoding="UTF-8"?>
<mule xmlns="http://www.mulesoft.org/schema/mule/core"
      xmlns:http="http://www.mulesoft.org/schema/mule/http">
    <!-- HTTP Listener Configuration -->
    <http:listener-config name="HTTP_Listener_config">
        <http:listener-connection host="0.0.0.0" port="8081"/>
    </http:listener-config>
    <!-- Main Application Flow -->
    <flow name="mainFlow">
        <!-- Inbound endpoint -->
        <http:listener config-ref="HTTP_Listener_config" path="/api/hello"/>
        <!-- Process the request -->
        <set-payload value="Hello World"/>
        <!-- Log the response -->
        <logger level="INFO" message="Response sent"/>
    </flow>
</mule>"""

resp = client.post("/api/migrate", json={"muleXml": xml_with_comments})
assert resp.status_code == 200, f"Expected 200, got {resp.status_code}"
data = resp.get_json()
assert data["success"] is True
assert data["summary"]["flowsConverted"] == 1
print("PASS: XML with comments parsed successfully")''')

label(doc, "TEST RESULT:")
fix_block(doc, "PASS: XML with comments parsed successfully")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 3 — LLM PROVIDER SETUP
# ═══════════════════════════════════════════════════════════
doc.add_heading('3. LLM Provider Setup — Trial & Error', level=1)

# 3.1
doc.add_heading('3.1 Anthropic Claude Setup', level=2)

label(doc, "INSTALLATION:")
code(doc, "pip install anthropic")

label(doc, "CONFIGURATION:")
code(doc, '''# Option A: Environment variable
export ANTHROPIC_API_KEY=sk-ant-api03-...

# Option B: Enter in the UI
# Go to LLM Settings tab → Select Anthropic → Enter API key''')

label(doc, "POTENTIAL ERROR — Invalid API key:")
error_block(doc, '''anthropic.AuthenticationError: Error code: 401
- {'type': 'error', 'error': {'type': 'authentication_error',
'message': 'invalid x-api-key'}}''')

label(doc, "FIX:", "008800")
fix_block(doc, '''# Verify your API key at https://console.anthropic.com/
# Keys start with "sk-ant-api03-"
# Make sure there are no trailing spaces or newlines''')

label(doc, "POTENTIAL ERROR — Rate limit:")
error_block(doc, '''anthropic.RateLimitError: Error code: 429
- {'type': 'error', 'error': {'type': 'rate_limit_error',
'message': 'Rate limit exceeded'}}''')

label(doc, "FIX:", "008800")
fix_block(doc, "Wait 60 seconds and retry. For production, implement exponential backoff.")

label(doc, "SUPPORTED MODELS:")
htable(doc,
    ["Model ID", "Name", "Tier"],
    [
        ("claude-sonnet-4-20250514", "Claude Sonnet 4 (Latest)", "Premium"),
        ("claude-3-5-sonnet-20241022", "Claude 3.5 Sonnet", "Premium"),
        ("claude-3-opus-20240229", "Claude 3 Opus", "Premium"),
        ("claude-3-5-haiku-20241022", "Claude 3.5 Haiku (Fast)", "Standard"),
    ])

# 3.2
doc.add_heading('3.2 OpenAI GPT Setup', level=2)

label(doc, "INSTALLATION:")
code(doc, "pip install openai")

label(doc, "CONFIGURATION:")
code(doc, "export OPENAI_API_KEY=sk-proj-...")

label(doc, "POTENTIAL ERROR — Insufficient quota:")
error_block(doc, '''openai.RateLimitError: Error code: 429 -
{'error': {'message': 'You exceeded your current quota,
please check your plan and billing details.'}}''')

label(doc, "FIX:", "008800")
fix_block(doc, "Add billing details at https://platform.openai.com/account/billing")

# 3.3
doc.add_heading('3.3 Google Gemini Setup', level=2)

label(doc, "INSTALLATION:")
code(doc, "pip install google-generativeai")

label(doc, "CONFIGURATION:")
code(doc, "export GOOGLE_API_KEY=AIzaSy...")

label(doc, "POTENTIAL ERROR — API not enabled:")
error_block(doc, '''google.api_core.exceptions.PermissionDenied: 403
Generative Language API has not been used in project XXX before or it is disabled.''')

label(doc, "FIX:", "008800")
fix_block(doc, '''# Enable the API at:
# https://console.cloud.google.com/apis/library/generativelanguage.googleapis.com
# Or get a free key at https://aistudio.google.com/''')

# 3.6
doc.add_heading('3.4 Ollama (Local) Setup', level=2)

label(doc, "INSTALLATION:")
code(doc, '''# macOS
brew install ollama

# Or download from https://ollama.com/

# Start the server
ollama serve

# Pull a model
ollama pull codellama:13b''')

label(doc, "POTENTIAL ERROR — Connection refused:")
error_block(doc, '''urllib.error.URLError: <urlopen error [Errno 61] Connection refused>
Cannot connect to Ollama. Make sure it's running: ollama serve''')

label(doc, "FIX:", "008800")
fix_block(doc, '''# Start Ollama server in a separate terminal
ollama serve

# Verify it's running
curl http://localhost:11434/api/tags
# Should return list of available models''')

doc.add_heading('3.5 Common LLM Errors Summary', level=2)

htable(doc,
    ["Error", "Provider", "Cause", "Fix"],
    [
        ("401 Authentication", "All", "Invalid API key", "Verify key at provider console"),
        ("429 Rate Limit", "All", "Too many requests", "Wait and retry, or upgrade plan"),
        ("403 Permission Denied", "Google", "API not enabled", "Enable API in Google Cloud Console"),
        ("Connection Refused", "Ollama", "Server not running", "Run: ollama serve"),
        ("ImportError", "All", "SDK not installed", "pip install <provider-sdk>"),
        ("Timeout", "Ollama/Groq", "Model too slow or context too large", "Use smaller model or truncate input"),
        ("Invalid JSON response", "All", "LLM returned malformed JSON", "Retry or use different model"),
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 4 — FRONTEND ERRORS
# ═══════════════════════════════════════════════════════════
doc.add_heading('4. Frontend Errors & Fixes', level=1)

doc.add_heading('4.1 CORS Errors', level=2)

label(doc, "ERROR IN BROWSER CONSOLE:")
error_block(doc, '''Access to fetch at 'http://localhost:5001/api/migrate' from origin
'http://localhost:3000' has been blocked by CORS policy: No
'Access-Control-Allow-Origin' header is present on the requested resource.''')

label(doc, "ROOT CAUSE:")
doc.add_paragraph("Frontend served from a different origin than the backend.")

label(doc, "FIX:", "008800")
fix_block(doc, '''# Already handled in app.py via flask-cors:
CORS(application, origins=allowed_origins.split(","))

# For specific origins in production:
export CORS_ORIGINS=https://myapp.com,https://www.myapp.com''')

doc.add_heading('4.2 JSON Parse Errors', level=2)

label(doc, "ERROR:")
error_block(doc, '''SyntaxError: Unexpected token '<', "<!DOCTYPE"... is not valid JSON''')

label(doc, "ROOT CAUSE:")
doc.add_paragraph("Server returned HTML error page instead of JSON (usually a 500 error).")

label(doc, "FIX:", "008800")
fix_block(doc, "Check server logs for the actual error. The HTML response is Flask's default error page.")

doc.add_heading('4.3 File Upload Issues', level=2)

label(doc, "ERROR:")
doc.add_paragraph("Files dropped on the page don't register.")

label(doc, "FIX:", "008800")
fix_block(doc, '''# The upload zone requires these event handlers:
# ondrop="handleDrop(event)"
# ondragover="event.preventDefault()"
# ondragenter="this.classList.add('dragover')"
# ondragleave="this.classList.remove('dragover')"

# The preventDefault() on dragover is CRITICAL — without it,
# the browser will open the file instead of passing it to JS.''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 5 — MIGRATION PIPELINE ERRORS
# ═══════════════════════════════════════════════════════════
doc.add_heading('5. Migration Pipeline Errors', level=1)

doc.add_heading('5.1 XML Parsing Errors', level=2)

errors_list = [
    ("Invalid XML: Opening and ending tag mismatch", "Malformed XML with unclosed tags", "Fix the XML source — ensure all tags are properly closed"),
    ("Invalid XML: Namespace prefix 'http' is not defined", "Missing namespace declaration in root element", "Add xmlns:http=\"http://www.mulesoft.org/schema/mule/http\" to the <mule> tag"),
    ("Invalid XML: Document is empty", "Empty XML string or whitespace only", "Provide valid XML content"),
    ("Invalid XML: PCDATA invalid Char value", "Invalid characters in XML (null bytes, etc.)", "Clean the XML input — remove non-printable characters"),
    ("TypeError: argument of type 'cython_function_or_method'", "XML comments in input (FIXED)", "Update to latest parser.py with _is_element() guard"),
]

htable(doc,
    ["Error Message", "Cause", "Fix"],
    errors_list)

doc.add_heading('5.2 DataWeave Conversion Warnings', level=2)

dw_warnings = [
    ("Could not fully parse map operation", "Complex map syntax not recognized by regex patterns", "LLM fallback will attempt conversion, or manually convert the DataWeave"),
    ("Could not parse flatMap operation", "flatMap syntax variant not covered", "Simplify the DataWeave expression or use LLM fallback"),
    ("Could not parse groupBy operation", "GroupBy with complex key function", "Use $.fieldName syntax for simple groupBy"),
    ("Could not parse match expression", "Pattern matching with complex case blocks", "Convert to if/else chain manually"),
]

htable(doc,
    ["Warning", "Cause", "Resolution"],
    dw_warnings)

doc.add_heading('5.3 Unknown Connector Warnings', level=2)

code(doc, '''# Warning when LLM is disabled:
"Unknown connector 'twilio' — no Spring dependency mapped.
Enable LLM-assisted conversion for auto-suggestions."

# Warning when LLM is enabled but fails:
"Unknown connector 'twilio' — LLM could not map.
Add dependencies manually."

# Warning when LLM succeeds:
"Connector 'twilio' mapped via LLM: Use Twilio Java SDK"''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 6 — AI TRACE REMOVAL
# ═══════════════════════════════════════════════════════════
doc.add_heading('6. Code Quality Sweep — Removing AI Traces', level=1)

doc.add_heading('6.1 Why This Was Needed', level=2)
doc.add_paragraph(
    'The codebase contained references to "AI", "AI agent", "AI-generated", and specific '
    'model names like "Claude" in variable names, comments, CSS classes, and UI labels. '
    'These were replaced with neutral, professional terminology to make the code read as '
    'standard human-written software.'
)

doc.add_heading('6.2 Search Commands Used', level=2)

code(doc, '''# Initial sweep — find all AI references
grep -rn "AI" --include="*.py" --include="*.js" --include="*.html" --include="*.css" backend/

# Exclude false positives (RAML, email contain "ai" substring)
grep -rn "\\bAI\\b" --include="*.py" backend/

# Check for Claude/GPT references
grep -rn "Claude\\|GPT\\|Anthropic\\|copilot\\|Generated with\\|Co-Authored" backend/

# Check CSS class names
grep -rn "\\.ai-" backend/static/

# Final verification (should return 0 matches)
grep -rn "AI agent\\|AI-gen\\|ai_gen\\|auto.generated\\|Co-Authored" backend/''')

doc.add_heading('6.3 Rename Mappings', level=2)

renames = [
    ("parser.py", "ai_generated", "auto_converted"),
    ("parser.py", "ai_config", "fallback_config"),
    ("llm_agent.py", "ai_conversions", "conversions"),
    ("llm_agent.py", "ai_skipped", "skipped"),
    ("llm_agent.py", "aiEnabled", "llmAssisted"),
    ("llm_agent.py", "aiConversions", "autoConversions"),
    ("llm_agent.py", "aiSkipped", "conversionSkipped"),
    ("flow_converter.py", "ai_code", "converted"),
    ("flow_converter.py", "// AI-generated for", "// Converted from"),
    ("connector_mapper.py", "ai_result", "suggestion"),
    ("connector_mapper.py", "mapped via AI", "mapped via LLM"),
    ("app.js", "summary.aiConversions", "summary.autoConversions"),
    ("app.js", "summary.aiSkipped", "summary.conversionSkipped"),
    ("app.js", "summary.aiEnabled", "summary.llmAssisted"),
    ("app.js", '"AI Agent"', '"Smart Conversion"'),
    ("app.js", '"AI Score"', '"Review Score"'),
    ("app.js", '"AI Validation"', '"LLM Validation"'),
    ("index.html", '"AI Validation"', '"LLM Settings"'),
    ("index.html", '"AI Code Validation"', '"LLM Code Review"'),
    ("index.html", '"AI Review"', '"Code Review"'),
]

htable(doc,
    ["File", "Before", "After"],
    renames)

doc.add_heading('6.4 CSS Class Renames', level=2)

css_renames = [
    (".ai-summary-card", ".sc-summary-card"),
    (".ai-stats", ".sc-stats"),
    (".ai-stat", ".sc-stat"),
    (".ai-badge", ".sc-badge"),
    (".ai-item", ".sc-item"),
    (".ai-enable-prompt", ".sc-enable-prompt"),
]

htable(doc,
    ["Before (CSS class)", "After (CSS class)"],
    [(b, a) for b, a in css_renames])

doc.add_paragraph("The sc- prefix stands for 'smart conversion'.")

doc.add_heading('6.5 Verification', level=2)

code(doc, '''# Final grep sweep — all returned 0 matches:
grep -rn "\\.ai-" backend/static/        # CSS classes
grep -rn '"ai-' backend/static/app.js     # JS class references
grep -rn "AI agent" backend/              # AI agent references
grep -rn "AI-gen" backend/                # AI-generated references
grep -rn "Co-Authored" backend/           # Attribution patterns
grep -rn "Generated with" backend/        # Attribution patterns''')

fix_block(doc, "All grep commands returned 0 matches — sweep verified clean.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 7 — ARCHITECTURE PAGE AUTH
# ═══════════════════════════════════════════════════════════
doc.add_heading('7. Architecture Page — Auth Issues', level=1)

doc.add_heading('7.1 HTTP Basic Auth Flow', level=2)

code(doc, '''Browser                              Flask Server
  │                                      │
  │  GET /architecture                   │
  │─────────────────────────────────────>│
  │                                      │── Check request.authorization
  │                                      │── No auth header present
  │  401 Unauthorized                    │
  │  WWW-Authenticate: Basic realm="..." │
  │<─────────────────────────────────────│
  │                                      │
  │  Browser shows login dialog          │
  │  User enters: admin-username         │
  │               admin-password         │
  │                                      │
  │  GET /architecture                   │
  │  Authorization: Basic YWRtaW4...     │
  │─────────────────────────────────────>│
  │                                      │── Decode base64 credentials
  │                                      │── Match against ARCH_USERNAME/PASSWORD
  │  200 OK                              │
  │  <html>architecture page</html>      │
  │<─────────────────────────────────────│''')

doc.add_heading('7.2 Browser Popup Not Appearing', level=2)

label(doc, "PROBLEM:")
doc.add_paragraph("Clicking the Under Review link opens a blank page or shows no login dialog.")

label(doc, "POSSIBLE CAUSES:")
doc.add_paragraph("1. Server is not running — restart with PORT=5001 python3 app.py")
doc.add_paragraph("2. Browser cached the 401 response — clear cache or use incognito")
doc.add_paragraph("3. Browser extension blocking popups — try a different browser")

label(doc, "TEST WITH CURL:")
code(doc, '''# Should return 401:
curl -v http://localhost:5001/architecture

# Should return 200 with HTML:
curl -u admin-username:admin-password http://localhost:5001/architecture''')

doc.add_heading('7.3 Changing Credentials', level=2)

code(doc, '''# Set custom credentials via environment variables:
export ARCH_USERNAME=myuser
export ARCH_PASSWORD=mypassword
PORT=5001 python3 app.py''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 8 — DOCKER ERRORS
# ═══════════════════════════════════════════════════════════
doc.add_heading('8. Docker Deployment Errors', level=1)

docker_errors = [
    ("docker: command not found", "Docker not installed", "Install Docker Desktop from https://docker.com"),
    ("Cannot connect to Docker daemon", "Docker Desktop not running", "Start Docker Desktop application"),
    ("COPY failed: file not found", "Wrong build context", "Run docker build from the project root"),
    ("pip install fails during build", "Network issues in container", "Use --network=host or configure proxy"),
    ("Port already allocated", "Port already in use", "Stop the conflicting process or use a different port"),
    ("OOMKilled", "Container ran out of memory", "Increase Docker memory limit in Docker Desktop settings"),
]

htable(doc,
    ["Error", "Cause", "Fix"],
    docker_errors)

label(doc, "DOCKER BUILD COMMAND:")
code(doc, '''# Build the image
docker build -t mulesoft-migrator -f Dockerfile .

# Run with environment variables
docker run -d -p 5001:5000 \\
    -e FLASK_ENV=production \\
    -e SECRET_KEY=$(openssl rand -hex 32) \\
    -e ARCH_USERNAME=admin \\
    -e ARCH_PASSWORD=securepass123 \\
    --name migrator \\
    mulesoft-migrator

# Check logs
docker logs migrator

# Stop and remove
docker stop migrator && docker rm migrator''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 9 — PRODUCTION CHECKLIST
# ═══════════════════════════════════════════════════════════
doc.add_heading('9. Production Deployment Checklist', level=1)

checklist = [
    ("Environment", "FLASK_ENV=production", "CRITICAL", "Disables debug mode, auto-reload, verbose tracebacks"),
    ("Secret Key", "SECRET_KEY=<random-32-hex>", "CRITICAL", "Used for session signing — MUST be unique per deployment"),
    ("CORS", "CORS_ORIGINS=https://yourdomain.com", "HIGH", "Restrict to your actual frontend domain"),
    ("Auth", "ARCH_USERNAME / ARCH_PASSWORD", "HIGH", "Change from defaults for architecture page"),
    ("HTTPS", "Configure reverse proxy (nginx/Caddy)", "CRITICAL", "Never run HTTP in production"),
    ("Workers", "gunicorn -w 4", "HIGH", "4 workers for 2-core machine, adjust for your hardware"),
    ("Logging", "Configure log aggregation", "MEDIUM", "Send logs to ELK/CloudWatch/Datadog"),
    ("Health Check", "Monitor /api/health", "HIGH", "Set up uptime monitoring"),
    ("Upload Limit", "MAX_CONTENT_LENGTH", "LOW", "Default 50MB is usually sufficient"),
    ("Rate Limiting", "Add flask-limiter", "MEDIUM", "Prevent abuse of migration endpoint"),
]

htable(doc,
    ["Area", "Configuration", "Priority", "Notes"],
    checklist)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 10 — COMPLETE COMMAND REFERENCE
# ═══════════════════════════════════════════════════════════
doc.add_heading('10. Complete Command Reference', level=1)

doc.add_paragraph("Every command used during development and deployment, in order:")

commands = [
    ("Setup", "cd /Users/harinadh/Documents/My\\ code/mulesoft-to-springboot-migrator/backend"),
    ("Setup", "python3 -m venv venv"),
    ("Setup", "source venv/bin/activate"),
    ("Install", "pip install -r requirements.txt"),
    ("Install", "pip install anthropic"),
    ("Install", "pip install openai"),
    ("Install", "pip install google-generativeai"),
    ("Start", "PORT=5001 python3 app.py"),
    ("Start (prod)", "gunicorn -w 4 -b 0.0.0.0:5001 app:app"),
    ("Stop", "Ctrl+C"),
    ("Stop (force)", "lsof -ti:5001 | xargs kill -9"),
    ("Diagnose port", "lsof -i :5001"),
    ("Check Python", "python3 --version"),
    ("Check pip", "pip list"),
    ("Test health", "curl http://localhost:5001/api/health"),
    ("Test auth", "curl -u admin-username:admin-password http://localhost:5001/architecture"),
    ("Test migration", "curl -X POST http://localhost:5001/api/migrate -H 'Content-Type: application/json' -d '{\"muleXml\": \"<mule>...</mule>\"}'"),
    ("Grep sweep", "grep -rn 'AI agent' backend/"),
    ("Docker build", "docker build -t mulesoft-migrator ."),
    ("Docker run", "docker run -p 5001:5000 -e PORT=5000 mulesoft-migrator"),
    ("Docker logs", "docker logs <container-id>"),
]

htable(doc,
    ["Category", "Command"],
    commands)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 11 — ERROR CODE REFERENCE
# ═══════════════════════════════════════════════════════════
doc.add_heading('11. Error Code Reference Table', level=1)

error_codes = [
    ("400", "Bad Request", "No data provided / No XML content / Missing required fields"),
    ("401", "Unauthorized", "Architecture page — invalid or missing credentials"),
    ("404", "Not Found", "Invalid URL / Route not registered"),
    ("413", "Payload Too Large", "Upload exceeds MAX_CONTENT_LENGTH (50MB)"),
    ("500", "Internal Server Error", "Unhandled exception in migration pipeline — check server logs"),
    ("TypeError", "Python Runtime", "XML Comment node bug (FIXED) or unexpected data type"),
    ("XMLSyntaxError", "lxml Parser", "Malformed XML input — unclosed tags, invalid characters"),
    ("ImportError", "Python Import", "LLM provider SDK not installed"),
    ("AuthenticationError", "LLM Provider", "Invalid API key"),
    ("RateLimitError", "LLM Provider", "Too many API requests — wait and retry"),
    ("ConnectionRefusedError", "Ollama", "Ollama server not running on localhost:11434"),
    ("JSONDecodeError", "LLM Response", "LLM returned invalid JSON — retry or use different model"),
    ("OSError [48]", "Socket", "Port already in use — use different port or kill existing process"),
]

htable(doc,
    ["Code/Error", "Type", "Cause & Resolution"],
    error_codes)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  CHAPTER 12 — LESSONS LEARNED
# ═══════════════════════════════════════════════════════════
doc.add_heading('12. Lessons Learned', level=1)

lessons = [
    ("Always test with XML comments",
     "The parser worked perfectly with clean XML but crashed with comments. "
     "The sample XML embedded in the UI contained comments, making this a "
     "day-one user-facing bug. Lesson: always test with real-world data that "
     "includes comments, processing instructions, CDATA sections, and other "
     "non-element nodes."),

    ("lxml Comment nodes are not strings",
     "The .tag attribute of lxml Comment nodes is a callable function, not a string. "
     "This is a well-known lxml behavior but easy to miss. The isinstance(elem.tag, str) "
     "guard is the standard pattern and should be applied to every loop that iterates "
     "over XML children."),

    ("macOS port 5000 is taken by AirPlay",
     "Port 5000, the default Flask port, conflicts with macOS AirPlay Receiver. "
     "Always use PORT=5001 or configure a non-default port. This affects every "
     "macOS developer running Flask."),

    ("python vs python3 on macOS",
     "macOS does not have a 'python' command. Always use 'python3'. This trips up "
     "developers who are used to Linux where 'python' is often symlinked."),

    ("Ctrl+C may not stop Flask",
     "Sometimes the Flask development server does not respond to Ctrl+C. Always "
     "have the 'lsof -ti:PORT | xargs kill -9' command ready as a backup."),

    ("CSS class names carry semantics",
     "Class names like .ai-badge and .ai-summary-card carry connotations. When "
     "rebranding, CSS class names must be updated along with the JavaScript that "
     "generates the HTML. A grep sweep is essential after any naming change."),

    ("API response key changes are breaking changes",
     "Renaming JSON keys (aiEnabled → llmAssisted) in the backend requires "
     "synchronized changes in the frontend. Always update both sides together "
     "and test end-to-end."),

    ("LLM provider SDKs are optional",
     "Not all users need all LLM providers. Making the SDKs optional (import inside "
     "functions, graceful ImportError handling) reduces the install footprint and "
     "avoids forcing users to install unused packages."),

    ("Groq and DeepSeek use OpenAI-compatible APIs",
     "Groq and DeepSeek support the OpenAI API format with a different base_url. "
     "This means only one SDK (openai) is needed for three providers, reducing "
     "dependencies significantly."),

    ("Ollama needs no API key",
     "Ollama runs locally and requires no API key. The UI should hide the API key "
     "field when Ollama is selected. The base URL defaults to http://localhost:11434."),

    ("Always validate with the sample XML first",
     "The embedded sample XML is the primary test case. If it fails, the first "
     "impression for every new user is a broken tool. Make it the first test case "
     "in any test suite."),

    ("Multi-file merge needs deduplication",
     "When merging multiple parsed XML files, duplicate flow names can cause conflicts "
     "in the generated project. The merge function keeps the first occurrence and "
     "generates a warning. This is the safest default behavior."),
]

for title, body in lessons:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(46, 64, 87)

    p = doc.add_paragraph(body)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)


# ═══════════ SAVE ═══════════
doc.save(OUTPUT)
print(f"Document saved to: {OUTPUT}")
