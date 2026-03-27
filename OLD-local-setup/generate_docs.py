#!/usr/bin/env python3
"""
Comprehensive Documentation Generator for MuleSoft to Spring Boot Migration Tool.
Generates a detailed Word document covering every aspect of the codebase.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
    "MuleSoft_to_SpringBoot_Migration_Tool_Documentation.docx")

# ════════════════════════════════════════════════════════════════
#  Helpers
# ════════════════════════════════════════════════════════════════
def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells, header=False, shade=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(text)
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.font.size = Pt(9)
                if header:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        if header:
            set_cell_shading(cell, "2E4057")
        elif shade:
            set_cell_shading(cell, shade)
    return row

def add_code_block(doc, code, language="python"):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(30, 30, 30)
    # Background shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    shading.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_note(doc, text, note_type="NOTE"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    colors = {"NOTE": "1565C0", "WARNING": "E65100", "TIP": "2E7D32", "IMPORTANT": "AD1457"}
    run = p.add_run(f"{note_type}: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(colors.get(note_type, "1565C0"))
    run.font.size = Pt(9)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True

def section_break(doc):
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  Create Document
# ════════════════════════════════════════════════════════════════
doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MuleSoft to Spring Boot\nMigration Tool")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(46, 64, 87)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Complete Technical Documentation")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(99, 102, 241)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 1.0.0")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Enterprise Reference Manual")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(130, 130, 130)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ("1", "Executive Summary", ""),
    ("2", "System Architecture Overview", ""),
    ("2.1", "High-Level Architecture Diagram", ""),
    ("2.2", "Technology Stack", ""),
    ("2.3", "Directory Structure", ""),
    ("2.4", "Request-Response Lifecycle", ""),
    ("2.5", "Data Flow Pipeline", ""),
    ("3", "Installation & Setup Guide", ""),
    ("3.1", "Prerequisites", ""),
    ("3.2", "Installation Steps", ""),
    ("3.3", "Configuration", ""),
    ("3.4", "Running the Application", ""),
    ("3.5", "Docker Deployment", ""),
    ("4", "Flask Application (app.py) — Complete Reference", ""),
    ("4.1", "Application Factory Pattern", ""),
    ("4.2", "Route Registration", ""),
    ("4.3", "API Endpoints — Full Specification", ""),
    ("4.4", "Multi-XML Parsing Pipeline", ""),
    ("4.5", "Error Handling Strategy", ""),
    ("4.6", "Authentication — Architecture Page", ""),
    ("5", "MuleSoft XML Parser (parser.py) — Complete Reference", ""),
    ("5.1", "Namespace Registry — All 30+ Namespaces", ""),
    ("5.2", "Core Processor Tags", ""),
    ("5.3", "Source Tags", ""),
    ("5.4", "MuleSoftParser Class — Every Method", ""),
    ("5.5", "XML Comment Node Handling (Bug Fix)", ""),
    ("5.6", "Global Configuration Parsing", ""),
    ("5.7", "Flow Parsing", ""),
    ("5.8", "Error Handler Parsing", ""),
    ("5.9", "Batch Job Parsing", ""),
    ("5.10", "APIkit Parsing", ""),
    ("5.11", "Secure Properties Parsing", ""),
    ("5.12", "TLS Context Parsing", ""),
    ("5.13", "Caching Strategy Parsing", ""),
    ("6", "Flow Converter (flow_converter.py) — Complete Reference", ""),
    ("6.1", "Public API", ""),
    ("6.2", "Source Conversion — All 15 Source Types", ""),
    ("6.3", "Processor Conversion — All 50+ Processor Types", ""),
    ("6.4", "Error Handler Conversion", ""),
    ("6.5", "Batch Job Conversion", ""),
    ("6.6", "Sub-flow Resolution", ""),
    ("7", "DataWeave Converter (dataweave_converter.py) — Complete Reference", ""),
    ("7.1", "Header Parsing", ""),
    ("7.2", "Expression Conversion Engine", ""),
    ("7.3", "All Supported Operations", ""),
    ("7.4", "Body Conversion — Multi-line Scripts", ""),
    ("7.5", "MEL Compatibility Layer", ""),
    ("8", "Connector Mapper (connector_mapper.py) — Complete Reference", ""),
    ("8.1", "Dependency Map — All 30+ Connectors", ""),
    ("8.2", "HTTP Method Mapping", ""),
    ("8.3", "Error Type Mapping — All 50+ Error Types", ""),
    ("8.4", "Spring Config Property Generation", ""),
    ("9", "LLM Conversion Module (llm_agent.py) — Complete Reference", ""),
    ("9.1", "AgentContext Class", ""),
    ("9.2", "Triple Fallback Strategy", ""),
    ("9.3", "Unknown Element Conversion", ""),
    ("9.4", "Unknown DataWeave Conversion", ""),
    ("9.5", "Connector Mapping Suggestions", ""),
    ("9.6", "Unknown Source Conversion", ""),
    ("10", "LLM Validator (llm_validator.py) — Complete Reference", ""),
    ("10.1", "Provider Registry — All 6 Providers", ""),
    ("10.2", "Anthropic Claude Provider", ""),
    ("10.3", "OpenAI GPT Provider", ""),
    ("10.4", "Google Gemini Provider", ""),
    ("10.5", "DeepSeek Provider", ""),
    ("10.6", "Groq Provider", ""),
    ("10.7", "Ollama Provider (Local)", ""),
    ("10.8", "Validation Prompt Engineering", ""),
    ("10.9", "Response Parsing & Error Handling", ""),
    ("11", "Spring Boot Generator (spring_generator.py) — Complete Reference", ""),
    ("11.1", "Project File Generation", ""),
    ("11.2", "pom.xml Generation", ""),
    ("11.3", "Main Application Class", ""),
    ("11.4", "application.properties Generation", ""),
    ("11.5", "Configuration Classes — All 15 Types", ""),
    ("11.6", "Exception Classes", ""),
    ("11.7", "Utility Classes", ""),
    ("11.8", "Infrastructure Files", ""),
    ("12", "Frontend (app.js) — Complete Reference", ""),
    ("12.1", "Tab Switching System", ""),
    ("12.2", "Multi-File Upload", ""),
    ("12.3", "LLM Provider Management", ""),
    ("12.4", "Migration Workflow", ""),
    ("12.5", "File Tree Rendering", ""),
    ("12.6", "Summary Rendering", ""),
    ("12.7", "Validation Results Rendering", ""),
    ("12.8", "Sample MuleSoft XML", ""),
    ("13", "HTML Templates — Complete Reference", ""),
    ("14", "CSS Styling — Complete Reference", ""),
    ("15", "API Flow Diagrams", ""),
    ("15.1", "Migration API Flow", ""),
    ("15.2", "Validation API Flow", ""),
    ("15.3", "Download API Flow", ""),
    ("15.4", "LLM Provider Selection Flow", ""),
    ("16", "MuleSoft to Spring Boot Mapping Reference", ""),
    ("16.1", "Connector Mapping Table", ""),
    ("16.2", "Processor Mapping Table", ""),
    ("16.3", "Error Type Mapping Table", ""),
    ("16.4", "DataWeave to Java Mapping Table", ""),
    ("17", "Testing Guide", ""),
    ("18", "Deployment Guide", ""),
    ("19", "Troubleshooting Guide", ""),
    ("20", "Appendices", ""),
]

for num, title, _ in toc_items:
    p = doc.add_paragraph()
    indent = 0
    if "." in num:
        indent = 0.4
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{num}  {title}")
    run.font.size = Pt(10) if "." not in num else Pt(9)
    if "." not in num:
        run.font.bold = True

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 1 — EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'The MuleSoft to Spring Boot Migration Tool is a comprehensive, web-based application '
    'designed to automate the conversion of MuleSoft 4 (Anypoint Platform) applications into '
    'fully functional Spring Boot 3.2 projects. The tool parses MuleSoft XML configurations, '
    'converts DataWeave 2.0 expressions to Java equivalents, maps MuleSoft connectors to '
    'Spring Boot dependencies, and generates a complete, runnable Spring Boot project including '
    'pom.xml, configuration classes, REST controllers, service classes, and infrastructure files.'
)

doc.add_heading('1.1 Purpose & Motivation', level=2)
doc.add_paragraph(
    'Organizations migrating from MuleSoft to Spring Boot face significant challenges:\n\n'
    '• Manual conversion is error-prone and time-consuming, often taking weeks per application\n'
    '• MuleSoft XML configurations contain complex namespace handling, DataWeave expressions, '
    'and connector-specific configurations that require deep expertise in both platforms\n'
    '• Ensuring the migrated code follows Spring Boot best practices requires experienced engineers\n'
    '• Testing and validating the migrated code adds additional overhead\n\n'
    'This tool addresses these challenges by providing an automated, intelligent migration pipeline '
    'that handles 30+ MuleSoft connectors, 50+ processor types, comprehensive DataWeave conversion, '
    'and optional LLM-powered code review for quality assurance.'
)

doc.add_heading('1.2 Key Features', level=2)

features = [
    ("Multi-XML Input", "Accepts multiple MuleSoft XML files simultaneously, automatically merging them into a unified project structure with deduplication of flows, sub-flows, and configurations."),
    ("Comprehensive XML Parsing", "Supports all MuleSoft 4 configuration components including core processors, all major connectors (HTTP, Database, JMS, AMQP, Kafka, File, SFTP, Email, Salesforce, AWS S3/SQS, MongoDB, Redis, Elasticsearch), APIkit, Batch processing, error handling, secure properties, TLS contexts, and caching strategies."),
    ("DataWeave 2.0 Conversion", "Converts DataWeave expressions, scripts, operators, and functions to equivalent Java/Spring code. Covers 40+ operations including map, filter, reduce, pluck, mapObject, string operations, date/time handling, type coercion, null handling, and MEL compatibility."),
    ("30+ Connector Mapping", "Maps every MuleSoft connector to its Spring Boot equivalent with correct Maven dependencies, Spring annotations, and application.properties configuration."),
    ("Complete Project Generation", "Generates a runnable Spring Boot 3.2 project with pom.xml, main application class, application.properties/yml, 15+ configuration classes, exception classes, utility classes, Dockerfile, docker-compose.yml, and test classes."),
    ("LLM-Powered Code Review", "Optional integration with 6 LLM providers (Anthropic Claude, OpenAI GPT, Google Gemini, DeepSeek, Groq, Ollama) for automated code review covering correctness, security, best practices, and improvement suggestions."),
    ("Smart Conversion (LLM Fallback)", "When the static converters encounter unknown elements, the tool can optionally delegate to an LLM provider for real-time code generation with a triple fallback strategy: LLM code → TODO comment → warning."),
    ("ZIP Download", "One-click download of the complete migrated project as a ZIP file ready for import into any IDE."),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0].cells
hdr[0].text = "Feature"
hdr[1].text = "Description"
for cell in hdr:
    set_cell_shading(cell, "2E4057")
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)

for name, desc in features:
    row = t.add_row()
    row.cells[0].text = name
    row.cells[1].text = desc
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_heading('1.3 Technology Stack', level=2)

stack = [
    ("Backend", "Python 3.9+, Flask 3.1.0, lxml 5.3.1"),
    ("Frontend", "HTML5, CSS3 (Custom Dark Theme), Vanilla JavaScript, CodeMirror 5.65"),
    ("LLM SDKs", "anthropic >= 0.39.0, openai >= 1.50.0, google-generativeai >= 0.8.0"),
    ("Production", "Gunicorn 23.0.0, Docker, docker-compose"),
    ("Target Output", "Spring Boot 3.2, Java 17/21, Maven"),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
hdr = t.rows[0].cells
hdr[0].text = "Layer"
hdr[1].text = "Technologies"
for cell in hdr:
    set_cell_shading(cell, "2E4057")
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)
for layer, tech in stack:
    row = t.add_row()
    row.cells[0].text = layer
    row.cells[1].text = tech
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 2 — SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════
doc.add_heading('2. System Architecture Overview', level=1)

doc.add_heading('2.1 High-Level Architecture Diagram', level=2)
doc.add_paragraph(
    'The system follows a layered architecture with clear separation of concerns:'
)

arch_diagram = """
┌─────────────────────────────────────────────────────────────────────┐
│                        BROWSER (Frontend)                          │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌───────────────────┐  │
│  │ MuleSoft │  │ DataWeave│  │ Settings │  │  LLM Settings     │  │
│  │ XML Tab  │  │   Tab    │  │   Tab    │  │  (Provider/Key)   │  │
│  └────┬─────┘  └────┬─────┘  └────┬─────┘  └────────┬──────────┘  │
│       │              │             │                  │             │
│       └──────────────┴─────────────┴──────────────────┘             │
│                              │                                      │
│                    POST /api/migrate                                 │
│                    { muleXmlFiles, dataweaveScripts, llmConfig }     │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
                               ▼
┌─────────────────────────────────────────────────────────────────────┐
│                     FLASK APPLICATION (app.py)                       │
│                                                                     │
│  ┌─────────────────────────────────────────────────────────────┐    │
│  │  Route: POST /api/migrate                                   │    │
│  │                                                             │    │
│  │  1. Parse request JSON                                      │    │
│  │  2. Create AgentContext (LLM config)                        │    │
│  │  3. For each XML file → MuleSoftParser.parse()              │    │
│  │  4. Merge parsed results (if multiple files)                │    │
│  │  5. DataWeaveConverter.convert() for each DW script         │    │
│  │  6. ConnectorMapper.map_connectors()                        │    │
│  │  7. FlowConverter.convert()                                 │    │
│  │  8. SpringBootGenerator.generate()                          │    │
│  │  9. (Optional) validate_code() via LLM                      │    │
│  │  10. Return JSON response                                   │    │
│  └─────────────────────────────────────────────────────────────┘    │
└──────────────────────────────┬──────────────────────────────────────┘
                               │
            ┌──────────────────┼──────────────────┐
            │                  │                  │
            ▼                  ▼                  ▼
┌───────────────────┐ ┌───────────────┐ ┌──────────────────┐
│   PARSING LAYER   │ │ CONVERSION    │ │  GENERATION      │
│                   │ │ LAYER         │ │  LAYER           │
│ • MuleSoftParser  │ │ • FlowConv.   │ │ • SpringBoot     │
│ • Namespace       │ │ • DataWeave   │ │   Generator      │
│   Registry        │ │   Converter   │ │ • pom.xml        │
│ • 30+ connector   │ │ • Connector   │ │ • Config classes │
│   configs         │ │   Mapper      │ │ • Properties     │
│ • Flow/SubFlow    │ │ • LLM Agent   │ │ • Dockerfile     │
│ • Error handlers  │ │   (fallback)  │ │ • Tests          │
│ • Batch jobs      │ │               │ │                  │
│ • APIkit          │ │               │ │                  │
└───────────────────┘ └───────────────┘ └──────────────────┘
                               │
                               ▼
                    ┌─────────────────────┐
                    │  LLM VALIDATION     │
                    │  (Optional)         │
                    │                     │
                    │ • Anthropic Claude  │
                    │ • OpenAI GPT       │
                    │ • Google Gemini    │
                    │ • DeepSeek         │
                    │ • Groq             │
                    │ • Ollama (local)   │
                    └─────────────────────┘
"""
add_code_block(doc, arch_diagram, "text")

doc.add_heading('2.2 Directory Structure', level=2)

dir_structure = """
mulesoft-to-springboot-migrator/
└── backend/
    ├── app.py                          # Flask application — routes, factory, helpers
    ├── gunicorn.conf.py                # Gunicorn production config
    ├── requirements.txt                # Python dependencies
    │
    ├── migrator/                       # Core migration engine
    │   ├── __init__.py                 # Package init
    │   ├── parser.py                   # MuleSoft XML parser (1004 lines)
    │   ├── flow_converter.py           # Flow → Spring Boot converter (700+ lines)
    │   ├── dataweave_converter.py      # DataWeave → Java converter (913 lines)
    │   ├── connector_mapper.py         # Connector → Spring dependency mapper (496 lines)
    │   ├── llm_agent.py                # LLM-powered fallback conversion (333 lines)
    │   ├── llm_validator.py            # Multi-provider LLM code review (637 lines)
    │   └── spring_generator.py         # Spring Boot project generator (720 lines)
    │
    ├── static/
    │   ├── app.js                      # Frontend JavaScript (900+ lines)
    │   └── style.css                   # Dark theme CSS (1150+ lines)
    │
    └── templates/
        ├── index.html                  # Main application page
        └── architecture.html           # Protected architecture/flow diagram page
"""
add_code_block(doc, dir_structure, "text")

doc.add_heading('2.3 Request-Response Lifecycle', level=2)

doc.add_paragraph(
    'Every migration request follows this exact sequence of operations:'
)

lifecycle = [
    ("1. HTTP Request", "Browser sends POST /api/migrate with JSON body containing muleXmlFiles array, dataweaveScripts object, projectName, groupId, javaVersion, and optional llmConfig."),
    ("2. Input Validation", "Flask route validates that at least one XML file is provided. Single XML string (backward compat) is wrapped into the multi-file format."),
    ("3. AgentContext Creation", "If LLM is enabled and a provider is configured, an AgentContext object is created to track LLM-assisted conversions throughout the pipeline."),
    ("4. XML Parsing (per file)", "Each MuleSoft XML file is passed to MuleSoftParser.parse(). The parser extracts global configs, flows, sub-flows, error handlers, properties, connectors, batch jobs, APIkit configs, secure properties, TLS contexts, and caching strategies."),
    ("5. Result Merging", "If multiple XML files were parsed, results are merged with deduplication. Duplicate flow/sub-flow names generate warnings."),
    ("6. DataWeave Conversion", "Each DataWeave script is converted to Java code via DataWeaveConverter.convert(). The converter handles headers, variables, functions, and the full expression conversion engine."),
    ("7. Connector Mapping", "ConnectorMapper.map_connectors() maps detected MuleSoft connectors to Spring Boot Maven dependencies and configuration properties."),
    ("8. Flow Conversion", "FlowConverter.convert() transforms parsed flows into Spring Boot Java source files — controllers, services, schedulers, listeners, batch jobs."),
    ("9. Project Generation", "SpringBootGenerator.generate() assembles the complete project: pom.xml, main class, properties, config classes, exception classes, Dockerfile, tests."),
    ("10. LLM Validation", "If enabled, the generated code is sent to the configured LLM provider for review. Returns scores, issues, improvements, security concerns."),
    ("11. Response", "JSON response with success flag, files dict, summary, and optional llmValidation."),
]

for step, desc in lifecycle:
    p = doc.add_paragraph()
    run = p.add_run(step + ": ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run(desc)
    run.font.size = Pt(10)

doc.add_heading('2.4 Data Flow Pipeline', level=2)

pipeline = """
Input XML String
       │
       ▼
┌──────────────────┐
│ lxml.etree.      │
│ fromstring()     │──── XMLSyntaxError → ValueError("Invalid XML")
│                  │
└────────┬─────────┘
         │
         ▼
┌──────────────────┐
│ _build_ns_map()  │ ← Merges XML nsmap with NAMESPACES registry
└────────┬─────────┘
         │
    ┌────┴────┬──────────┬──────────┬─────────────┐
    ▼         ▼          ▼          ▼             ▼
 Configs   Flows    Sub-flows   Error       Connectors
    │         │          │       Handlers       │
    │    ┌────┴────┐     │          │           │
    │    │ Source   │     │          │           │
    │    │ detect   │     │          │           │
    │    ├─────────┤     │          │           │
    │    │Processor │     │          │           │
    │    │ parse    │     │          │           │
    │    └─────────┘     │          │           │
    │         │          │          │           │
    └─────────┴──────────┴──────────┴───────────┘
                         │
                         ▼
              ┌──────────────────┐
              │  Parsed Dict     │
              │  {               │
              │   global_configs │
              │   flows          │
              │   sub_flows      │
              │   error_handlers │
              │   connectors     │
              │   batch_jobs     │
              │   ...            │
              │  }               │
              └────────┬─────────┘
                       │
           ┌───────────┼───────────┐
           ▼           ▼           ▼
      FlowConv    ConnMapper   SpringGen
           │           │           │
           └───────────┴───────────┘
                       │
                       ▼
              Project Files Dict
              { "path": "content" }
"""
add_code_block(doc, pipeline, "text")

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 3 — INSTALLATION & SETUP
# ════════════════════════════════════════════════════════════════
doc.add_heading('3. Installation & Setup Guide', level=1)

doc.add_heading('3.1 Prerequisites', level=2)

prereqs = [
    ("Python", "3.9 or higher", "python3 --version"),
    ("pip", "Latest version", "python3 -m pip --version"),
    ("Git", "Any recent version", "git --version"),
    ("Docker (optional)", "For containerized deployment", "docker --version"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["Requirement", "Version", "Check Command"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for req, ver, cmd in prereqs:
    row = t.add_row()
    row.cells[0].text = req
    row.cells[1].text = ver
    row.cells[2].text = cmd
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_heading('3.2 Installation Steps', level=2)

doc.add_paragraph('Step 1: Clone or download the repository')
add_code_block(doc, 'cd /path/to/your/projects\ngit clone <repository-url>\ncd mulesoft-to-springboot-migrator/backend')

doc.add_paragraph('Step 2: Create a virtual environment (recommended)')
add_code_block(doc, 'python3 -m venv venv\nsource venv/bin/activate  # macOS/Linux\n# or: venv\\Scripts\\activate  # Windows')

doc.add_paragraph('Step 3: Install dependencies')
add_code_block(doc, 'pip install -r requirements.txt')

doc.add_paragraph('Step 4: (Optional) Install LLM provider SDKs')
add_code_block(doc, '# Only install the providers you plan to use:\npip install anthropic       # For Anthropic Claude\npip install openai          # For OpenAI GPT, DeepSeek, Groq\npip install google-generativeai  # For Google Gemini')

doc.add_heading('3.3 Configuration', level=2)

doc.add_paragraph('The application uses environment variables for configuration:')

env_vars = [
    ("PORT", "5000", "Server port number"),
    ("FLASK_ENV", "production", "Environment: 'development' or 'production'"),
    ("SECRET_KEY", "(auto-generated)", "Flask session secret key"),
    ("CORS_ORIGINS", "*", "Comma-separated allowed origins"),
    ("ARCH_USERNAME", "admin-username", "Architecture page login username"),
    ("ARCH_PASSWORD", "admin-password", "Architecture page login password"),
    ("ANTHROPIC_API_KEY", "(none)", "Anthropic Claude API key"),
    ("OPENAI_API_KEY", "(none)", "OpenAI API key"),
    ("GOOGLE_API_KEY", "(none)", "Google Gemini API key"),
    ("DEEPSEEK_API_KEY", "(none)", "DeepSeek API key"),
    ("GROQ_API_KEY", "(none)", "Groq API key"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["Variable", "Default", "Description"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for var, default, desc in env_vars:
    row = t.add_row()
    row.cells[0].text = var
    row.cells[1].text = default
    row.cells[2].text = desc
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_heading('3.4 Running the Application', level=2)

doc.add_paragraph('Development mode:')
add_code_block(doc, 'PORT=5001 python3 app.py')

add_note(doc, "On macOS, port 5000 is used by AirPlay Receiver. Use PORT=5001 or disable AirPlay in System Settings > General > AirDrop & Handoff.", "WARNING")

doc.add_paragraph('Production mode with Gunicorn:')
add_code_block(doc, 'gunicorn -w 4 -b 0.0.0.0:5000 app:app')

doc.add_paragraph('Then open your browser to http://localhost:5001 (or whichever port you configured).')

doc.add_heading('3.5 Stopping the Server', level=2)

doc.add_paragraph('To stop the server:')
doc.add_paragraph('• Press Ctrl+C in the terminal')
doc.add_paragraph('• If Ctrl+C does not work, kill the process on the port:')
add_code_block(doc, 'lsof -ti:5001 | xargs kill -9')

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 4 — FLASK APPLICATION (app.py)
# ════════════════════════════════════════════════════════════════
doc.add_heading('4. Flask Application (app.py) — Complete Reference', level=1)

doc.add_paragraph(
    'The Flask application is the entry point and orchestration layer of the migration tool. '
    'It is organized using the Application Factory pattern for testability and modularity.'
)

doc.add_heading('4.1 Application Factory Pattern', level=2)

doc.add_paragraph(
    'The create_app() function is the application factory. It creates and configures a Flask '
    'instance, sets up CORS, logging, and registers all routes.'
)

doc.add_paragraph('Function: create_app()')
add_code_block(doc, '''def create_app():
    """Application factory — returns a configured Flask app."""
    application = Flask(__name__, static_folder="static", template_folder="templates")

    # Configuration from environment
    env = os.environ.get("FLASK_ENV", "production")
    application.config.update(
        ENV=env,
        DEBUG=(env == "development"),
        SECRET_KEY=os.environ.get("SECRET_KEY", os.urandom(32).hex()),
        MAX_CONTENT_LENGTH=50 * 1024 * 1024,  # 50 MB max upload
        JSON_SORT_KEYS=False,
    )

    # CORS
    allowed_origins = os.environ.get("CORS_ORIGINS", "*")
    CORS(application, origins=allowed_origins.split(","))

    # Logging
    log_level = logging.DEBUG if application.config["DEBUG"] else logging.INFO
    logging.basicConfig(
        level=log_level,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    _register_routes(application)
    return application''')

doc.add_paragraph('Detailed Explanation:')

explanations = [
    ("Line 1-2: Function signature and docstring", "The factory pattern allows creating multiple app instances (useful for testing) and defers initialization until the function is called."),
    ("Line 3: Flask constructor", "static_folder='static' tells Flask to serve static files from the /static directory. template_folder='templates' tells Jinja2 where to find HTML templates."),
    ("Line 6: FLASK_ENV", "Reads the FLASK_ENV environment variable. Defaults to 'production' which disables debug mode, verbose logging, and auto-reload."),
    ("Line 7-12: Config update", "ENV: stores the environment name. DEBUG: enables Flask debug mode in development. SECRET_KEY: used for session signing — auto-generated if not set. MAX_CONTENT_LENGTH: 50MB upload limit prevents memory exhaustion. JSON_SORT_KEYS=False: preserves insertion order in JSON responses."),
    ("Line 15-16: CORS setup", "Cross-Origin Resource Sharing is configured to allow requests from specified origins. In production, this should be restricted to the actual frontend domain."),
    ("Line 19-23: Logging", "Configures Python's logging module with timestamps, log level, and logger name. In development, DEBUG level shows all internal operations. In production, INFO level reduces noise."),
    ("Line 25: Route registration", "Delegates route setup to a separate function for cleaner organization."),
]

for title, explanation in explanations:
    p = doc.add_paragraph()
    run = p.add_run(title + ": ")
    run.font.bold = True
    run.font.size = Pt(9)
    run = p.add_run(explanation)
    run.font.size = Pt(9)

doc.add_heading('4.2 Route Registration', level=2)

doc.add_paragraph(
    'All routes are registered inside _register_routes(application). This keeps the factory '
    'function clean and groups all endpoint definitions together.'
)

doc.add_heading('4.3 API Endpoints — Full Specification', level=2)

endpoints = [
    ("GET /", "index()", "Serves the main HTML page (index.html)", "None", "HTML page"),
    ("GET /architecture", "architecture()", "Protected architecture page (HTTP Basic Auth required)", "Authorization header", "HTML page or 401"),
    ("GET /api/health", "health()", "Health check endpoint", "None", '{"status": "ok", "env": "..."}'),
    ("GET /api/llm/providers", "llm_providers()", "Returns available LLM providers and their models", "None", "Provider registry JSON"),
    ("POST /api/migrate", "migrate()", "Main migration endpoint — converts MuleSoft XML to Spring Boot", "JSON body (see below)", "Migration result JSON"),
    ("POST /api/validate", "validate_endpoint()", "Standalone LLM validation endpoint", "JSON body with files and llmConfig", "Validation result JSON"),
    ("POST /api/migrate/download", "download_project()", "Downloads migrated project as ZIP", "JSON body with files dict", "ZIP file"),
    ("POST /api/convert/dataweave", "convert_dataweave()", "Converts a single DataWeave script to Java", "JSON body with script", "Conversion result JSON"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["Endpoint", "Handler", "Description"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for ep, handler, desc, _, _ in endpoints:
    row = t.add_row()
    row.cells[0].text = ep
    row.cells[1].text = handler
    row.cells[2].text = desc
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_heading('4.3.1 POST /api/migrate — Request Body', level=3)

add_code_block(doc, '''{
    "muleXmlFiles": [
        { "name": "main.xml", "content": "<?xml ...>" },
        { "name": "api.xml", "content": "<?xml ...>" }
    ],
    "muleXml": "<?xml ...>",           // Backward compat (single file)
    "dataweaveScripts": {
        "transform-1": "%dw 2.0\\noutput application/json\\n---\\npayload"
    },
    "projectName": "migrated-app",
    "groupId": "com.example",
    "javaVersion": "17",
    "llmConfig": {
        "enabled": true,
        "provider": "anthropic",
        "model": "claude-sonnet-4-20250514",
        "apiKey": "sk-ant-...",
        "baseUrl": ""                   // Optional, for Ollama/custom
    }
}''')

doc.add_heading('4.3.2 POST /api/migrate — Response Body', level=3)

add_code_block(doc, '''{
    "success": true,
    "files": {
        "pom.xml": "<pom content>",
        "src/main/java/com/example/MigratedAppApplication.java": "...",
        "src/main/resources/application.properties": "...",
        // ... all generated files
    },
    "summary": {
        "flowsConverted": 3,
        "subFlowsConverted": 2,
        "connectorsFound": ["http", "database", "jms"],
        "dataweaveScriptsConverted": 1,
        "dependencies": [
            { "groupId": "org.springframework.boot", "artifactId": "spring-boot-starter-web" }
        ],
        "warnings": ["Unknown config 'custom:config' skipped..."],
        "xmlFilesProcessed": 2,
        "xmlFileNames": ["main.xml", "api.xml"],
        "llmAssisted": true,
        "autoConversions": [...],
        "autoConversionCount": 1,
        "conversionSkipped": [...],
        "conversionSkippedCount": 0
    },
    "llmValidation": {
        "overallScore": 8,
        "summary": "Well-structured migration with minor issues...",
        "issues": [...],
        "improvements": [...],
        "missingItems": [...],
        "securityIssues": [...],
        "bestPractices": [...]
    }
}''')

doc.add_heading('4.4 Multi-XML Parsing Pipeline', level=2)

doc.add_paragraph(
    'The migration endpoint supports multiple XML files through the muleXmlFiles array. '
    'Each file is parsed independently, then results are merged using _merge_parsed_results().'
)

doc.add_paragraph('The merge function:')
doc.add_paragraph('• Combines all flows, sub-flows, and configurations')
doc.add_paragraph('• Deduplicates by name — first occurrence wins')
doc.add_paragraph('• Generates warnings for duplicate flow/sub-flow names')
doc.add_paragraph('• Merges connector sets (union)')
doc.add_paragraph('• Merges global properties (later files override earlier)')
doc.add_paragraph('• Concatenates all warnings')

doc.add_paragraph('Comment-separated XML support:')
doc.add_paragraph(
    'If a single XML input contains <!-- File: filename.xml --> markers, the tool automatically '
    'splits it into multiple files using _split_comment_separated_xml(). This supports pasting '
    'multiple files into a single textarea.'
)

doc.add_heading('4.5 Error Handling Strategy', level=2)

errors = [
    ("No data provided", "400", "Request body is empty or not JSON"),
    ("MuleSoft XML content is required", "400", "No XML files and no single XML string"),
    ("Error parsing {name}: {error}", "400", "XML syntax error in a specific file"),
    ("No valid XML content found", "400", "All files were empty after parsing"),
    ("Migration failed", "500", "Unexpected error — full stack trace logged"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["Error Message", "HTTP Status", "Cause"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for msg, status, cause in errors:
    row = t.add_row()
    row.cells[0].text = msg
    row.cells[1].text = status
    row.cells[2].text = cause
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_heading('4.6 Authentication — Architecture Page', level=2)

doc.add_paragraph(
    'The /architecture route is protected with HTTP Basic Authentication. The credentials '
    'are configured via environment variables ARCH_USERNAME and ARCH_PASSWORD, defaulting '
    'to "admin-username" and "admin-password" respectively.'
)

doc.add_paragraph('Authentication flow:')
doc.add_paragraph('1. Browser requests /architecture')
doc.add_paragraph('2. Server checks request.authorization header')
doc.add_paragraph('3. If missing or incorrect: returns 401 with WWW-Authenticate header')
doc.add_paragraph('4. Browser shows native login dialog')
doc.add_paragraph('5. User enters credentials → browser resends with Authorization header')
doc.add_paragraph('6. Server validates and serves architecture.html')

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 5 — PARSER
# ════════════════════════════════════════════════════════════════
doc.add_heading('5. MuleSoft XML Parser (parser.py) — Complete Reference', level=1)

doc.add_paragraph(
    'The parser is the foundation of the migration pipeline. It takes raw MuleSoft XML '
    'and produces a structured Python dictionary that all downstream converters consume. '
    'At 1004 lines, it is the largest module in the system.'
)

doc.add_heading('5.1 Namespace Registry — All 30+ Namespaces', level=2)

doc.add_paragraph(
    'MuleSoft XML uses XML namespaces extensively. The NAMESPACES dictionary maps short '
    'prefixes to full namespace URIs. This is critical for correctly identifying elements '
    'regardless of how the XML author chose to declare namespace prefixes.'
)

namespaces = [
    ("mule", "http://www.mulesoft.org/schema/mule/core", "Core MuleSoft elements (flow, sub-flow, choice, etc.)"),
    ("ee", "http://www.mulesoft.org/schema/mule/ee/core", "DataWeave transforms (ee:transform)"),
    ("http", "http://www.mulesoft.org/schema/mule/http", "HTTP listener and request configs"),
    ("sockets", "http://www.mulesoft.org/schema/mule/sockets", "TCP/UDP socket connections"),
    ("tls", "http://www.mulesoft.org/schema/mule/tls", "TLS/SSL context configurations"),
    ("db", "http://www.mulesoft.org/schema/mule/db", "Database operations (select, insert, etc.)"),
    ("jms", "http://www.mulesoft.org/schema/mule/jms", "JMS messaging (ActiveMQ, etc.)"),
    ("amqp", "http://www.mulesoft.org/schema/mule/amqp", "AMQP messaging (RabbitMQ)"),
    ("vm", "http://www.mulesoft.org/schema/mule/vm", "In-memory VM queues"),
    ("kafka", "http://www.mulesoft.org/schema/mule/kafka", "Apache Kafka integration"),
    ("anypoint-mq", "http://www.mulesoft.org/schema/mule/anypoint-mq", "Anypoint MQ cloud messaging"),
    ("file", "http://www.mulesoft.org/schema/mule/file", "Local file system operations"),
    ("sftp", "http://www.mulesoft.org/schema/mule/sftp", "SFTP file transfer"),
    ("ftp", "http://www.mulesoft.org/schema/mule/ftp", "FTP file transfer"),
    ("email", "http://www.mulesoft.org/schema/mule/email", "Email (IMAP, POP3, SMTP)"),
    ("apikit", "http://www.mulesoft.org/schema/mule/mule-apikit", "APIkit router and configs"),
    ("ws", "http://www.mulesoft.org/schema/mule/ws", "Web Service consumer (SOAP)"),
    ("wsc", "http://www.mulesoft.org/schema/mule/wsc", "Web Service consumer v2"),
    ("os", "http://www.mulesoft.org/schema/mule/os", "Object Store operations"),
    ("batch", "http://www.mulesoft.org/schema/mule/batch", "Batch processing"),
    ("validation", "http://www.mulesoft.org/schema/mule/validation", "Input validation"),
    ("scripting", "http://www.mulesoft.org/schema/mule/scripting", "Script execution (Groovy, etc.)"),
    ("json", "http://www.mulesoft.org/schema/mule/json", "JSON module"),
    ("xml-module", "http://www.mulesoft.org/schema/mule/xml-module", "XML module"),
    ("oauth", "http://www.mulesoft.org/schema/mule/oauth", "OAuth client"),
    ("salesforce", "http://www.mulesoft.org/schema/mule/salesforce", "Salesforce CRM connector"),
    ("s3", "http://www.mulesoft.org/schema/mule/s3", "AWS S3 storage"),
    ("sqs", "http://www.mulesoft.org/schema/mule/sqs", "AWS SQS queues"),
    ("sns", "http://www.mulesoft.org/schema/mule/sns", "AWS SNS notifications"),
    ("mongo", "http://www.mulesoft.org/schema/mule/mongo", "MongoDB connector"),
    ("redis", "http://www.mulesoft.org/schema/mule/redis", "Redis connector"),
    ("elasticsearch", "http://www.mulesoft.org/schema/mule/elasticsearch", "Elasticsearch connector"),
    ("spring", "http://www.mulesoft.org/schema/mule/spring", "Spring module bridge"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["Prefix", "URI", "Purpose"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for prefix, uri, purpose in namespaces:
    row = t.add_row()
    row.cells[0].text = prefix
    row.cells[1].text = uri
    row.cells[2].text = purpose
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

doc.add_heading('5.2 Core Processor Tags', level=2)

doc.add_paragraph(
    'The CORE_PROCESSOR_TAGS set contains all recognized MuleSoft core processing elements. '
    'These are elements within the mule: namespace that transform, route, or control message flow.'
)

processors = [
    ("logger", "Logs a message at a specified level", "log.info() / log.debug()"),
    ("set-payload", "Sets the message payload", "Return value / response body"),
    ("set-variable", "Creates or updates a flow variable", "Local variable assignment"),
    ("remove-variable", "Removes a flow variable", "Variable = null"),
    ("choice", "Conditional routing (if/else)", "if/else if/else"),
    ("scatter-gather", "Parallel execution of multiple routes", "CompletableFuture.allOf()"),
    ("for-each", "Iterates over a collection", "for loop / stream"),
    ("parallel-for-each", "Parallel iteration", "parallelStream()"),
    ("try", "Error handling scope", "try-catch block"),
    ("until-successful", "Retry with backoff", "@Retryable / RetryTemplate"),
    ("first-successful", "Try routes until one succeeds", "Fallback chain"),
    ("round-robin", "Load-balance across routes", "Round-robin routing"),
    ("async", "Asynchronous execution", "@Async / CompletableFuture"),
    ("flow-ref", "Call another flow/sub-flow", "Method call / @Autowired"),
    ("raise-error", "Throw a typed error", "throw new Exception()"),
    ("parse-template", "Process a template file", "Thymeleaf / template engine"),
    ("idempotent-message-validator", "Prevent duplicate processing", "Redis/DB idempotency check"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["MuleSoft Tag", "Purpose", "Spring Boot Equivalent"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for tag, purpose, spring in processors:
    row = t.add_row()
    row.cells[0].text = tag
    row.cells[1].text = purpose
    row.cells[2].text = spring
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_heading('5.3 XML Comment Node Handling (Critical Bug Fix)', level=2)

doc.add_paragraph(
    'A critical bug was discovered and fixed where XML comments (<!-- ... -->) in MuleSoft '
    'XML caused the parser to crash with a TypeError.'
)

p = doc.add_paragraph()
run = p.add_run('Root Cause: ')
run.font.bold = True
run = p.add_run(
    'lxml represents XML comments as special nodes where the .tag attribute is a callable '
    'function (the Comment factory), not a string. When the parser code executed '
    '"}" in tag, Python raised TypeError: argument of type \'cython_function_or_method\' '
    'is not iterable, because you cannot use the "in" operator on a function object.'
)

p = doc.add_paragraph()
run = p.add_run('Fix: ')
run.font.bold = True
run = p.add_run(
    'Added the _is_element() static method that checks isinstance(elem.tag, str). '
    'This returns False for Comment nodes and ProcessingInstruction nodes. The guard '
    'was applied to all 27 iteration loops in the parser, plus defensive checks in '
    '_local_tag(), _get_ns_prefix(), and _get_ns_uri().'
)

doc.add_paragraph('The guard methods:')
add_code_block(doc, '''@staticmethod
def _is_element(elem):
    """Return True if *elem* is a real XML element (not a Comment / PI)."""
    return isinstance(elem.tag, str)

def _local_tag(self, elem):
    tag = elem.tag
    if not isinstance(tag, str):
        return ""  # Comment / ProcessingInstruction node
    if "}" in tag:
        tag = tag.split("}", 1)[1]
    return tag''')

doc.add_paragraph('Usage pattern in every loop:')
add_code_block(doc, '''for elem in root:
    if not self._is_element(elem):
        continue  # skip Comment / PI nodes
    tag = self._local_tag(elem)
    # ... process element''')

doc.add_heading('5.4 MuleSoftParser Class — Method Reference', level=2)

methods = [
    ("parse(xml_content, agent_context=None)", "Main entry point. Parses XML string into structured dict.", "dict"),
    ("_build_ns_map(root)", "Builds namespace prefix → URI map by merging XML nsmap with NAMESPACES.", "dict"),
    ("_parse_global_configs(root, ns_map)", "Extracts all global configurations (HTTP, DB, JMS, etc.).", "list"),
    ("_make_http_listener_config(elem, ns_map)", "Parses http:listener-config into config dict.", "dict"),
    ("_make_http_request_config(elem, ns_map)", "Parses http:request-config with auth support.", "dict"),
    ("_make_db_config(elem, ns_map)", "Parses db:config for MySQL, Oracle, PostgreSQL, MSSQL.", "dict"),
    ("_parse_flows(root, ns_map)", "Finds and parses all <flow> elements.", "list"),
    ("_parse_sub_flows(root, ns_map)", "Finds and parses all <sub-flow> elements.", "list"),
    ("_parse_flow_element(flow_elem, ns_map, is_sub_flow)", "Parses a single flow/sub-flow element.", "dict"),
    ("_try_parse_source(elem, tag, ns, ns_map)", "Detects and parses flow source (15+ types).", "dict or None"),
    ("_parse_processor(elem, tag, ns, ns_map)", "Parses a single processor element recursively.", "dict"),
    ("_extract_transform_dw(elem, processor)", "Extracts DataWeave scripts from ee:transform.", "None (mutates)"),
    ("_parse_error_handlers(root, ns_map)", "Parses global error handlers.", "list"),
    ("_parse_error_handler_element(eh_elem, ns_map)", "Parses a single error handler.", "dict"),
    ("_parse_global_properties(root, ns_map)", "Extracts global-property and configuration-properties.", "dict"),
    ("_detect_connectors(root, ns_map)", "Scans all elements to detect which connectors are used.", "set"),
    ("_parse_batch_jobs(root, ns_map)", "Parses batch:job elements with steps and on-complete.", "list"),
    ("_parse_apikit(root, ns_map)", "Parses apikit:config and apikit:router.", "list"),
    ("_parse_secure_properties(root, ns_map)", "Parses secure-properties config with encryption.", "list"),
    ("_parse_tls_contexts(root, ns_map)", "Parses TLS context with key-store and trust-store.", "list"),
    ("_parse_caching_strategies(root, ns_map)", "Parses caching-strategy elements.", "list"),
    ("_parse_children(elem)", "Generic child element parser.", "list"),
    ("_is_element(elem)", "Static method — True if elem is a real XML element.", "bool"),
    ("_local_tag(elem)", "Strips namespace URI from tag, returns local name.", "str"),
    ("_get_ns_prefix(elem, ns_map)", "Returns namespace prefix for an element.", "str"),
    ("_get_ns_uri(elem)", "Returns namespace URI for an element.", "str"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["Method", "Description", "Returns"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for method, desc, returns in methods:
    row = t.add_row()
    row.cells[0].text = method
    row.cells[1].text = desc
    row.cells[2].text = returns
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

doc.add_heading('5.5 Source Types — All 15 Supported', level=2)

sources = [
    ("http:listener", "http-listener", "@GetMapping / @PostMapping", "path, method, config_ref"),
    ("scheduler", "scheduler", "@Scheduled", "cron/fixedDelay/fixedRate"),
    ("jms:subscriber", "jms-listener", "@JmsListener", "destination, config_ref, ackMode"),
    ("amqp:listener", "amqp-listener", "@RabbitListener", "queueName, config_ref"),
    ("kafka:listener", "kafka-listener", "@KafkaListener", "topic, config_ref"),
    ("vm:listener", "vm-listener", "@EventListener", "queueName, config_ref"),
    ("file:listener", "file-listener", "WatchService / polling", "directory, config_ref"),
    ("sftp:listener", "sftp-listener", "Spring Integration SFTP", "directory, config_ref"),
    ("ftp:listener", "ftp-listener", "Spring Integration FTP", "directory, config_ref"),
    ("email:on-new-email", "email-listener", "JavaMail polling", "config_ref"),
    ("salesforce:on-new-object", "salesforce-on-new-object", "REST polling / streaming", "config_ref"),
    ("sqs:listener", "sqs-listener", "@SqsListener", "queueUrl, config_ref"),
    ("anypoint-mq:subscriber", "anypoint-mq-listener", "Spring AMQP", "destination, config_ref"),
    ("(unknown)", "(LLM fallback)", "LLM-generated config", "auto_converted flag"),
]

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
for i, h in enumerate(["MuleSoft Source", "Internal Type", "Spring Equivalent", "Key Attributes"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(8)
for ms, internal, spring, attrs in sources:
    row = t.add_row()
    row.cells[0].text = ms
    row.cells[1].text = internal
    row.cells[2].text = spring
    row.cells[3].text = attrs
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 6 — CONNECTOR MAPPER
# ════════════════════════════════════════════════════════════════
doc.add_heading('6. Connector Mapper (connector_mapper.py) — Complete Reference', level=1)

doc.add_heading('6.1 Dependency Map — All 30+ Connectors', level=2)

doc.add_paragraph(
    'The CONNECTOR_DEPENDENCY_MAP is the central registry that maps MuleSoft connector '
    'namespaces to their Spring Boot Maven dependency equivalents.'
)

connectors = [
    ("http", "spring-boot-starter-web, spring-boot-starter-webflux"),
    ("database", "spring-boot-starter-data-jpa, spring-boot-starter-jdbc"),
    ("jms", "spring-boot-starter-activemq"),
    ("amqp", "spring-boot-starter-amqp"),
    ("kafka", "spring-kafka"),
    ("vm", "(no external deps — uses Spring Events)"),
    ("file", "(no external deps — uses java.nio)"),
    ("sftp", "spring-integration-sftp"),
    ("ftp", "spring-integration-ftp"),
    ("email", "spring-boot-starter-mail"),
    ("ws / wsc", "spring-boot-starter-web-services, spring-ws-core"),
    ("objectstore", "spring-boot-starter-data-redis, spring-boot-starter-cache"),
    ("batch", "spring-boot-starter-batch"),
    ("validation", "spring-boot-starter-validation"),
    ("salesforce", "spring-boot-starter-web (REST API via RestTemplate)"),
    ("s3", "software.amazon.awssdk:s3, software.amazon.awssdk:sts"),
    ("sqs", "software.amazon.awssdk:sqs, spring-cloud-aws-messaging"),
    ("sns", "software.amazon.awssdk:sns"),
    ("mongo", "spring-boot-starter-data-mongodb"),
    ("redis", "spring-boot-starter-data-redis"),
    ("elasticsearch", "spring-boot-starter-data-elasticsearch"),
    ("anypoint-mq", "spring-boot-starter-amqp"),
    ("oauth", "spring-boot-starter-security, spring-boot-starter-oauth2-client"),
    ("apikit", "(handled by Spring Web)"),
    ("ee", "(no external deps)"),
    ("scripting", "(no external deps)"),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
for i, h in enumerate(["MuleSoft Connector", "Spring Boot Dependencies"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for conn, deps in connectors:
    row = t.add_row()
    row.cells[0].text = conn
    row.cells[1].text = deps
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_heading('6.2 Error Type Mapping — All 50+ Types', level=2)

doc.add_paragraph(
    'The ERROR_TYPE_MAP maps MuleSoft error types to Java exception classes. This is used '
    'during error handler conversion to generate proper catch blocks.'
)

error_categories = [
    ("HTTP Errors", [
        ("HTTP:NOT_FOUND", "ResourceNotFoundException"),
        ("HTTP:BAD_REQUEST", "BadRequestException"),
        ("HTTP:UNAUTHORIZED", "UnauthorizedException"),
        ("HTTP:FORBIDDEN", "AccessDeniedException"),
        ("HTTP:TIMEOUT", "java.util.concurrent.TimeoutException"),
        ("HTTP:CONNECTIVITY", "java.net.ConnectException"),
    ]),
    ("Database Errors", [
        ("DB:CONNECTIVITY", "CannotGetJdbcConnectionException"),
        ("DB:BAD_SQL_SYNTAX", "BadSqlGrammarException"),
        ("DB:QUERY_EXECUTION", "DataAccessException"),
    ]),
    ("JMS Errors", [
        ("JMS:CONNECTIVITY", "JmsException"),
        ("JMS:PUBLISHING", "JmsException"),
        ("JMS:TIMEOUT", "TimeoutException"),
    ]),
    ("File Errors", [
        ("FILE:FILE_NOT_FOUND", "FileNotFoundException"),
        ("FILE:ACCESS_DENIED", "AccessDeniedException"),
        ("SFTP:CONNECTIVITY", "IOException"),
    ]),
    ("Validation Errors", [
        ("VALIDATION:NULL", "NullPointerException"),
        ("VALIDATION:INVALID_EMAIL", "IllegalArgumentException"),
        ("VALIDATION:INVALID_SIZE", "IllegalArgumentException"),
    ]),
    ("Core Errors", [
        ("ANY", "Exception"),
        ("EXPRESSION", "RuntimeException"),
        ("MULE:SECURITY", "SecurityException"),
        ("MULE:RETRY_EXHAUSTED", "RuntimeException"),
    ]),
]

for category, errors in error_categories:
    p = doc.add_paragraph()
    run = p.add_run(f"{category}:")
    run.font.bold = True

    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    for i, h in enumerate(["MuleSoft Error Type", "Java Exception"]):
        t.rows[0].cells[i].text = h
        set_cell_shading(t.rows[0].cells[i], "3E5060")
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)
    for mule_err, java_exc in errors:
        row = t.add_row()
        row.cells[0].text = mule_err
        row.cells[1].text = java_exc
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 7 — DATAWEAVE CONVERTER
# ════════════════════════════════════════════════════════════════
doc.add_heading('7. DataWeave Converter (dataweave_converter.py) — Complete Reference', level=1)

doc.add_paragraph(
    'The DataWeave Converter transforms MuleSoft DataWeave 2.0 expressions and scripts '
    'into equivalent Java code. At 913 lines, it is one of the most complex modules, '
    'handling 40+ distinct conversion patterns.'
)

doc.add_heading('7.1 Conversion Categories', level=2)

dw_categories = [
    ("Payload/Variable References", "payload.field → payload.get(\"field\")\nvars.x → x\nflowVars.x → x\nattributes.queryParams.x → request.getParameter(\"x\")\nattributes.headers.x → request.getHeader(\"x\")"),
    ("String Operations", "upper(x) → x.toUpperCase()\nlower(x) → x.toLowerCase()\ntrim(x) → x.trim()\ncapitalize(x) → StringUtils.capitalize(x)\nx splitBy \",\" → x.split(\",\")\nx joinBy \",\" → String.join(\",\", x)\nx contains \"y\" → x.contains(\"y\")\nx replace \"a\" with \"b\" → x.replace(\"a\", \"b\")"),
    ("Array Operations", "sizeOf(x) → x.size()\nflatten(x) → stream().flatMap().collect()\nfirst(x) → x.get(0)\nlast(x) → x.get(x.size()-1)\nmin/max/sum/avg → stream().mapToDouble()"),
    ("Collection Operators", "map → .stream().map().collect()\nfilter → .stream().filter().collect()\nreduce → .stream().reduce()\ngroupBy → Collectors.groupingBy()\norderBy → .sorted(Comparator.comparing())\ndistinctBy → filter with ConcurrentHashSet\npluck → entrySet().stream().map()\nmapObject → forEach with new LinkedHashMap\nfilterObject → entrySet().stream().filter()"),
    ("Type Coercion", "x as String → String.valueOf(x)\nx as Number → Double.parseDouble(x)\nx as Boolean → Boolean.parseBoolean(x)\nx as Date → LocalDate.parse(x)\nx as DateTime → LocalDateTime.parse(x)"),
    ("Null Handling", "x default \"y\" → x != null ? x : \"y\"\nisEmpty(x) → x == null || String.valueOf(x).isEmpty()\nisBlank(x) → x == null || String.valueOf(x).isBlank()"),
    ("Logical Operators", "and → &&\nor → ||\nnot → !"),
    ("Date/Time", "now() → LocalDateTime.now()"),
    ("Pattern Matching", "x match { case is String -> ... } → if/else instanceof chain"),
    ("Object Literals", "{ key: value } → Map<String, Object> with put()"),
    ("Concatenation", "++ → + (string concatenation)"),
]

for category, mappings in dw_categories:
    p = doc.add_paragraph()
    run = p.add_run(f"{category}:")
    run.font.bold = True
    run.font.size = Pt(10)
    add_code_block(doc, mappings)

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 8 — LLM AGENT
# ════════════════════════════════════════════════════════════════
doc.add_heading('8. LLM Conversion Module (llm_agent.py) — Complete Reference', level=1)

doc.add_heading('8.1 AgentContext Class', level=2)

doc.add_paragraph(
    'AgentContext is the shared state object threaded through every pipeline stage. It tracks '
    'whether LLM assistance is enabled, stores the provider configuration, and records all '
    'conversion attempts (successful and skipped) for the summary.'
)

add_code_block(doc, '''class AgentContext:
    def __init__(self, enabled=False, llm_config=None):
        self.enabled = enabled          # Whether LLM is turned on
        self.llm_config = llm_config    # {provider, apiKey, model, baseUrl}
        self.conversions = []           # Successful conversions
        self.skipped = []               # Skipped items

    def record_conversion(self, element, prompt_summary, generated_code):
        self.conversions.append({
            "element": element,
            "prompt_summary": prompt_summary,
            "generated_code": generated_code[:500],  # truncate
        })

    def record_skipped(self, element, reason):
        self.skipped.append({"element": element, "reason": reason})

    def to_summary(self):
        return {
            "llmAssisted": self.enabled,
            "autoConversions": self.conversions,
            "autoConversionCount": len(self.conversions),
            "conversionSkipped": self.skipped,
            "conversionSkippedCount": len(self.skipped),
        }''')

doc.add_heading('8.2 Triple Fallback Strategy', level=2)

doc.add_paragraph(
    'When the static converters encounter an unknown element, the LLM agent follows a '
    'triple fallback strategy:'
)
doc.add_paragraph('1. LLM-generated code: Call the configured LLM provider with a specialized prompt')
doc.add_paragraph('2. TODO comment: If LLM is unavailable or fails, insert a TODO comment')
doc.add_paragraph('3. Warning in summary: Record the skip in the migration summary')

doc.add_heading('8.3 Conversion Functions', level=2)

functions = [
    ("convert_unknown_element()", "Unknown XML processors", "Java code string", "Builds a prompt with the element tag and raw XML, asks the LLM to generate Spring Boot equivalent code."),
    ("convert_unknown_dataweave()", "Unparseable DataWeave", "Java code string", "Sends the unparseable DataWeave expression with context hint to the LLM for Java conversion."),
    ("suggest_connector_mapping()", "Unknown connectors", "Dict with maven_dependencies", "Asks the LLM to suggest Maven dependencies and Spring properties for an unknown connector."),
    ("convert_unknown_source()", "Unknown message sources", "Dict with source config", "Asks the LLM to determine the Spring Boot equivalent for an unknown inbound endpoint."),
]

for func, input_type, output, desc in functions:
    p = doc.add_paragraph()
    run = p.add_run(f"{func}")
    run.font.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph(f"  Input: {input_type}")
    doc.add_paragraph(f"  Output: {output}")
    doc.add_paragraph(f"  Behavior: {desc}")

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 9 — LLM VALIDATOR
# ════════════════════════════════════════════════════════════════
doc.add_heading('9. LLM Validator (llm_validator.py) — Complete Reference', level=1)

doc.add_heading('9.1 Provider Registry', level=2)

doc.add_paragraph(
    'The LLM_PROVIDERS dictionary defines all supported LLM providers with their models, '
    'API key environment variable names, and documentation URLs.'
)

providers = [
    ("Anthropic Claude", "claude-sonnet-4-20250514, claude-3-5-sonnet, claude-3-opus, claude-3-5-haiku", "ANTHROPIC_API_KEY", "Direct SDK (anthropic library)"),
    ("OpenAI GPT", "gpt-4o, gpt-4-turbo, gpt-4o-mini, o3-mini", "OPENAI_API_KEY", "Direct SDK (openai library)"),
    ("Google Gemini", "gemini-2.5-pro, gemini-2.0-flash, gemini-1.5-pro", "GOOGLE_API_KEY", "Direct SDK (google-generativeai)"),
    ("DeepSeek", "deepseek-chat, deepseek-coder, deepseek-reasoner", "DEEPSEEK_API_KEY", "OpenAI-compatible (base_url override)"),
    ("Groq", "llama-3.3-70b, mixtral-8x7b, llama-3.1-8b", "GROQ_API_KEY", "OpenAI-compatible (base_url override)"),
    ("Ollama (Local)", "codellama:13b, llama3:8b, deepseek-coder-v2, mistral:7b", "(none)", "HTTP API to localhost:11434"),
]

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
for i, h in enumerate(["Provider", "Models", "API Key Env", "Integration"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for name, models, env, integration in providers:
    row = t.add_row()
    row.cells[0].text = name
    row.cells[1].text = models
    row.cells[2].text = env
    row.cells[3].text = integration
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

doc.add_heading('9.2 Provider Architecture', level=2)

doc.add_paragraph(
    'Each provider implements the BaseLLMProvider abstract class with two methods:'
)
doc.add_paragraph('• validate(files, summary) — Sends generated code for full review, returns structured JSON')
doc.add_paragraph('• chat(system_prompt, user_prompt, max_tokens) — Generic chat for real-time conversion')

doc.add_paragraph('Provider class hierarchy:')
add_code_block(doc, '''BaseLLMProvider (abstract)
    ├── AnthropicProvider   — uses anthropic SDK
    ├── OpenAIProvider      — uses openai SDK
    ├── GoogleProvider      — uses google-generativeai SDK
    ├── DeepSeekProvider    — uses openai SDK with custom base_url
    ├── GroqProvider        — uses openai SDK with custom base_url
    └── OllamaProvider      — uses urllib.request (no SDK needed)''')

doc.add_heading('9.3 Validation Response Schema', level=2)

add_code_block(doc, '''{
    "overallScore": 8,           // 1-10 quality score
    "summary": "...",            // 1-2 sentence assessment
    "issues": [                  // List of problems found
        {
            "severity": "critical|warning|info",
            "file": "Controller.java",
            "line": "42",
            "message": "Missing @Valid annotation",
            "suggestion": "Add @Valid to request body parameter"
        }
    ],
    "improvements": [...],       // Suggested code improvements
    "missingItems": [...],       // Things needing manual implementation
    "securityIssues": [...],     // Security vulnerabilities found
    "bestPractices": [...]       // Spring Boot best practice violations
}''')

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 10 — SPRING GENERATOR
# ════════════════════════════════════════════════════════════════
doc.add_heading('10. Spring Boot Generator (spring_generator.py) — Complete Reference', level=1)

doc.add_paragraph(
    'The generator assembles all converted components into a complete, runnable Spring Boot '
    'project. It produces 20-40+ files depending on the connectors detected.'
)

doc.add_heading('10.1 Generated File Types', level=2)

gen_files = [
    ("pom.xml", "Maven project descriptor with all dependencies"),
    ("Application.java", "Main class with @SpringBootApplication and conditional @Enable annotations"),
    ("application.properties", "Connector-specific configuration from parsed MuleSoft configs"),
    ("application.yml", "YAML equivalent of properties"),
    ("application-dev.properties", "Development profile with DEBUG logging"),
    ("application-prod.properties", "Production profile with WARN logging"),
    ("*Controller.java", "REST controllers from HTTP listener flows"),
    ("*Service.java", "Service classes from sub-flows"),
    ("*Scheduler.java", "Scheduled tasks from scheduler flows"),
    ("*Listener.java", "Message listeners (JMS, AMQP, Kafka, etc.)"),
    ("*BatchJob.java", "Spring Batch jobs from batch:job elements"),
    ("config/*.java", "Up to 15 configuration classes (JMS, AMQP, Kafka, Redis, etc.)"),
    ("exception/*.java", "Custom exception classes (5 types)"),
    ("util/JsonUtil.java", "JSON serialization utility"),
    (".gitignore", "Standard Java/Spring .gitignore"),
    ("Dockerfile", "Multi-stage Docker build"),
    ("docker-compose.yml", "Docker Compose with dependent services"),
    ("*Tests.java", "Basic Spring Boot test class"),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
for i, h in enumerate(["File", "Description"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for file, desc in gen_files:
    row = t.add_row()
    row.cells[0].text = file
    row.cells[1].text = desc
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_heading('10.2 Configuration Class Generation', level=2)

config_classes = [
    ("SchedulingConfig", "scheduler flows detected", "@EnableScheduling annotation support"),
    ("JmsConfig", "jms connector detected", "@EnableJms and connection factory setup"),
    ("AmqpConfig", "amqp connector detected", "@EnableRabbit and exchange/queue bindings"),
    ("KafkaConfig", "kafka connector detected", "@EnableKafka and consumer/producer factories"),
    ("RestTemplateConfig", "RestTemplate used in code", "RestTemplate bean creation"),
    ("WebClientConfig", "WebClient used in code", "WebClient.Builder bean + named WebClient beans per HTTP request config"),
    ("CacheConfig", "redis or objectstore detected", "@EnableCaching with Redis store"),
    ("SecurityConfig", "oauth connector detected", "SecurityFilterChain with JWT resource server"),
    ("AsyncConfig", "CompletableFuture.runAsync used", "ThreadPoolTaskExecutor bean"),
    ("BatchConfig", "batch jobs detected", "Spring Batch configuration placeholder"),
    ("SftpConfig", "sftp connector detected", "DefaultSftpSessionFactory + SftpRemoteFileTemplate beans"),
    ("FtpConfig", "ftp connector detected", "FTP session factory placeholder"),
    ("AwsS3Config", "s3 connector detected", "S3Client bean with region"),
    ("AwsSqsConfig", "sqs connector detected", "SqsClient bean with region"),
    ("WebServiceConfig", "ws/wsc connector detected", "WebServiceTemplate bean"),
    ("CorsConfig", "always generated", "WebMvcConfigurer with global CORS mapping"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["Class", "Generated When", "Purpose"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for cls, when, purpose in config_classes:
    row = t.add_row()
    row.cells[0].text = cls
    row.cells[1].text = when
    row.cells[2].text = purpose
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 11 — API FLOW DIAGRAMS
# ════════════════════════════════════════════════════════════════
doc.add_heading('11. API Flow Diagrams', level=1)

doc.add_heading('11.1 Migration API Flow (POST /api/migrate)', level=2)

migration_flow = """
Browser                    Flask App                    Migrator Modules
  │                           │                              │
  │  POST /api/migrate        │                              │
  │  {muleXmlFiles, ...}      │                              │
  │──────────────────────────>│                              │
  │                           │                              │
  │                           │  Validate request JSON       │
  │                           │──┐                           │
  │                           │  │ Check for XML content     │
  │                           │<─┘                           │
  │                           │                              │
  │                           │  Create AgentContext          │
  │                           │──┐                           │
  │                           │  │ (if LLM enabled)          │
  │                           │<─┘                           │
  │                           │                              │
  │                           │  For each XML file:          │
  │                           │  parser.parse(xml)           │
  │                           │─────────────────────────────>│
  │                           │                              │── lxml.fromstring()
  │                           │                              │── _build_ns_map()
  │                           │                              │── _parse_global_configs()
  │                           │                              │── _parse_flows()
  │                           │                              │── _parse_sub_flows()
  │                           │                              │── _parse_error_handlers()
  │                           │                              │── _detect_connectors()
  │                           │  parsed_data                 │── _parse_batch_jobs()
  │                           │<─────────────────────────────│
  │                           │                              │
  │                           │  _merge_parsed_results()     │
  │                           │──┐ (if multiple files)       │
  │                           │<─┘                           │
  │                           │                              │
  │                           │  dw_converter.convert()      │
  │                           │─────────────────────────────>│
  │                           │  converted_dw                │── _parse_header()
  │                           │<─────────────────────────────│── _convert_body()
  │                           │                              │── _convert_expression()
  │                           │                              │
  │                           │  connector_mapper.map()      │
  │                           │─────────────────────────────>│
  │                           │  connector_info              │── Match connectors
  │                           │<─────────────────────────────│── Resolve Maven deps
  │                           │                              │
  │                           │  flow_converter.convert()    │
  │                           │─────────────────────────────>│
  │                           │  spring_files                │── Convert each flow
  │                           │<─────────────────────────────│── Generate Java code
  │                           │                              │
  │                           │  generator.generate()        │
  │                           │─────────────────────────────>│
  │                           │  project_files               │── pom.xml
  │                           │<─────────────────────────────│── Application.java
  │                           │                              │── properties, configs
  │                           │                              │── Dockerfile, tests
  │                           │                              │
  │                           │  (optional) validate_code()  │
  │                           │─────────────────────────────>│
  │                           │  validation_result           │── Send to LLM
  │                           │<─────────────────────────────│── Parse response
  │                           │                              │
  │  JSON Response            │                              │
  │  {success, files, summary}│                              │
  │<──────────────────────────│                              │
  │                           │                              │
"""
add_code_block(doc, migration_flow, "text")

doc.add_heading('11.2 Download API Flow (POST /api/migrate/download)', level=2)

download_flow = """
Browser                    Flask App
  │                           │
  │  POST /api/migrate/download│
  │  {files: {...}, projectName}│
  │──────────────────────────>│
  │                           │
  │                           │── Create BytesIO buffer
  │                           │── Create ZipFile(buffer)
  │                           │── For each file:
  │                           │     zf.writestr(path, content)
  │                           │── buffer.seek(0)
  │                           │
  │  application/zip          │
  │  (binary download)        │
  │<──────────────────────────│
  │                           │
"""
add_code_block(doc, download_flow, "text")

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 12 — FRONTEND
# ════════════════════════════════════════════════════════════════
doc.add_heading('12. Frontend (app.js) — Complete Reference', level=1)

doc.add_paragraph(
    'The frontend is a single-page application built with vanilla JavaScript (no frameworks). '
    'It provides a split-panel interface with input tabs on the left and output tabs on the right.'
)

doc.add_heading('12.1 Global State Variables', level=2)

add_code_block(doc, '''let migrationResult = null;     // Last successful migration response
let dwScripts = {};             // DataWeave scripts {name: content}
let activeDwScript = null;      // Currently selected DW script name
let llmProviders = {};          // Loaded from /api/llm/providers
let uploadedXmlFiles = [];      // Multi-file uploads [{name, content}]''')

doc.add_heading('12.2 Key Functions', level=2)

js_functions = [
    ("switchInputTab(btn)", "Switches active input tab (XML, DataWeave, Settings, LLM)"),
    ("switchOutputTab(btn)", "Switches active output tab (Files, Summary, Code Review)"),
    ("handleFileUpload(event)", "Handles drag-and-drop or file picker XML uploads"),
    ("handleDrop(event)", "Handles drag-and-drop file uploads"),
    ("loadSample()", "Loads embedded sample MuleSoft XML with 8 comment blocks"),
    ("clearAll()", "Resets all inputs and outputs to initial state"),
    ("migrate()", "Main migration function — calls POST /api/migrate"),
    ("downloadProject()", "Downloads ZIP via POST /api/migrate/download"),
    ("loadLLMProviders()", "Fetches provider list from /api/llm/providers"),
    ("onProviderChange()", "Updates model dropdown when provider is selected"),
    ("toggleLLM()", "Enables/disables LLM settings panel"),
    ("testLLMConnection()", "Tests LLM provider connectivity"),
    ("getLLMConfig()", "Builds llmConfig object from form inputs"),
    ("saveLLMSettings()", "Persists LLM settings to localStorage"),
    ("restoreLLMSettings()", "Restores LLM settings from localStorage on page load"),
    ("renderFiles(files)", "Builds file tree and content viewer"),
    ("renderSummary(summary)", "Renders migration summary with stats"),
    ("renderValidation(validation)", "Renders LLM code review results"),
    ("revalidateWithLLM()", "Re-runs LLM validation on existing results"),
    ("escapeHtml(text)", "XSS-safe HTML escaping"),
    ("convertDataWeave()", "Standalone DW conversion via POST /api/convert/dataweave"),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
for i, h in enumerate(["Function", "Description"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for func, desc in js_functions:
    row = t.add_row()
    row.cells[0].text = func
    row.cells[1].text = desc
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 13 — TESTING
# ════════════════════════════════════════════════════════════════
doc.add_heading('13. Testing Guide', level=1)

doc.add_heading('13.1 Backend Testing with Flask Test Client', level=2)

add_code_block(doc, '''import json
from app import create_app

app = create_app()
client = app.test_client()

# Test 1: Health endpoint
resp = client.get("/api/health")
assert resp.status_code == 200
data = resp.get_json()
assert data["status"] == "ok"

# Test 2: Migration with sample XML
sample_xml = """<?xml version="1.0" encoding="UTF-8"?>
<mule xmlns="http://www.mulesoft.org/schema/mule/core"
      xmlns:http="http://www.mulesoft.org/schema/mule/http">
    <http:listener-config name="HTTP_Listener_config">
        <http:listener-connection host="0.0.0.0" port="8081"/>
    </http:listener-config>
    <flow name="mainFlow">
        <http:listener config-ref="HTTP_Listener_config" path="/api/hello"/>
        <set-payload value="Hello World"/>
    </flow>
</mule>"""

resp = client.post("/api/migrate", json={
    "muleXml": sample_xml,
    "projectName": "test-app",
    "groupId": "com.test",
})
assert resp.status_code == 200
data = resp.get_json()
assert data["success"] is True
assert "pom.xml" in data["files"]
assert data["summary"]["flowsConverted"] == 1

# Test 3: LLM providers endpoint
resp = client.get("/api/llm/providers")
assert resp.status_code == 200
providers = resp.get_json()
assert "anthropic" in providers
assert "openai" in providers''')

doc.add_heading('13.2 Testing XML with Comments', level=2)

add_code_block(doc, '''# This test verifies the XML comment bug fix
xml_with_comments = """<?xml version="1.0" encoding="UTF-8"?>
<mule xmlns="http://www.mulesoft.org/schema/mule/core"
      xmlns:http="http://www.mulesoft.org/schema/mule/http">
    <!-- This is a comment that used to crash the parser -->
    <http:listener-config name="config">
        <http:listener-connection host="0.0.0.0" port="8081"/>
    </http:listener-config>
    <!-- Another comment -->
    <flow name="testFlow">
        <!-- Comment inside flow -->
        <http:listener config-ref="config" path="/test"/>
        <set-payload value="test"/>
    </flow>
</mule>"""

resp = client.post("/api/migrate", json={"muleXml": xml_with_comments})
assert resp.status_code == 200  # Previously would return 500''')

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 14 — DEPLOYMENT
# ════════════════════════════════════════════════════════════════
doc.add_heading('14. Deployment Guide', level=1)

doc.add_heading('14.1 Docker Deployment', level=2)

doc.add_paragraph('Create a Dockerfile for the migration tool:')
add_code_block(doc, '''FROM python:3.11-slim

WORKDIR /app
COPY backend/ .
RUN pip install --no-cache-dir -r requirements.txt

EXPOSE 5000
CMD ["gunicorn", "-w", "4", "-b", "0.0.0.0:5000", "app:app"]''')

doc.add_paragraph('Build and run:')
add_code_block(doc, '''docker build -t mulesoft-migrator .
docker run -p 5001:5000 \\
    -e FLASK_ENV=production \\
    -e SECRET_KEY=your-secret-key \\
    mulesoft-migrator''')

doc.add_heading('14.2 Production Checklist', level=2)

checklist = [
    "Set FLASK_ENV=production",
    "Set a strong SECRET_KEY",
    "Restrict CORS_ORIGINS to your domain",
    "Set ARCH_USERNAME and ARCH_PASSWORD to non-default values",
    "Use gunicorn with 4+ workers",
    "Enable HTTPS via reverse proxy (nginx/Caddy)",
    "Set MAX_CONTENT_LENGTH appropriate for your use case",
    "Configure log aggregation",
    "Set up health check monitoring on /api/health",
]

for item in checklist:
    doc.add_paragraph(f"  {item}")

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  CHAPTER 15 — TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════
doc.add_heading('15. Troubleshooting Guide', level=1)

issues = [
    ("Port 5000 already in use (macOS)", "macOS AirPlay Receiver uses port 5000.", "Use PORT=5001 python3 app.py, or disable AirPlay in System Settings > General > AirDrop & Handoff."),
    ("python: command not found", "macOS ships with python3, not python.", "Use python3 instead of python for all commands."),
    ("Ctrl+C not stopping server", "Sometimes the Flask development server doesn't respond to SIGINT.", "Use: lsof -ti:5001 | xargs kill -9"),
    ("TypeError on XML with comments", "lxml Comment nodes have .tag as a callable function.", "Fixed in parser.py. Ensure you have the latest version with _is_element() guard."),
    ("LLM validation returns score 0", "API key is missing or invalid.", "Check the API key in LLM Settings. Verify with the Test Connection button."),
    ("Missing dependency error", "LLM provider SDK not installed.", "Install the specific SDK: pip install anthropic/openai/google-generativeai"),
    ("Ollama connection refused", "Ollama is not running locally.", "Start Ollama: ollama serve. Ensure it's on port 11434."),
    ("Large XML file causes timeout", "Very complex MuleSoft configurations with many flows.", "Split into multiple smaller XML files and use multi-file upload."),
    ("Duplicate flow warnings", "Multiple XML files contain flows with the same name.", "First occurrence is kept. Rename duplicate flows in the source XML."),
]

for problem, cause, solution in issues:
    p = doc.add_paragraph()
    run = p.add_run(f"Problem: {problem}")
    run.font.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run(f"Cause: ")
    run.font.bold = True
    run.font.size = Pt(9)
    run = p.add_run(cause)
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run(f"Solution: ")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 100, 0)
    run = p.add_run(solution)
    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing

section_break(doc)

# ════════════════════════════════════════════════════════════════
#  APPENDICES
# ════════════════════════════════════════════════════════════════
doc.add_heading('16. Appendices', level=1)

doc.add_heading('Appendix A: requirements.txt', level=2)
add_code_block(doc, '''flask==3.1.0
flask-cors==5.0.1
lxml==5.3.1
pyyaml==6.0.2
gunicorn==23.0.0

# LLM Providers (install the ones you plan to use)
anthropic>=0.39.0       # Anthropic Claude
openai>=1.50.0          # OpenAI GPT + DeepSeek + Groq (OpenAI-compatible)
google-generativeai>=0.8.0  # Google Gemini''')

doc.add_heading('Appendix B: Complete File Listing with Line Counts', level=2)

file_listing = [
    ("app.py", "436", "Flask application, routes, helpers"),
    ("migrator/parser.py", "1004", "MuleSoft XML parser"),
    ("migrator/flow_converter.py", "700+", "Flow to Spring Boot converter"),
    ("migrator/dataweave_converter.py", "913", "DataWeave to Java converter"),
    ("migrator/connector_mapper.py", "496", "Connector dependency mapper"),
    ("migrator/llm_agent.py", "333", "LLM conversion module"),
    ("migrator/llm_validator.py", "637", "Multi-provider LLM validator"),
    ("migrator/spring_generator.py", "720", "Spring Boot project generator"),
    ("static/app.js", "900+", "Frontend JavaScript"),
    ("static/style.css", "1150+", "Dark theme CSS"),
    ("templates/index.html", "256", "Main HTML template"),
    ("templates/architecture.html", "~200", "Architecture page template"),
    ("requirements.txt", "11", "Python dependencies"),
    ("gunicorn.conf.py", "~10", "Gunicorn configuration"),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["File", "Lines", "Purpose"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for file, lines, purpose in file_listing:
    row = t.add_row()
    row.cells[0].text = file
    row.cells[1].text = lines
    row.cells[2].text = purpose
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_heading('Appendix C: Glossary', level=2)

glossary = [
    ("MuleSoft", "Enterprise integration platform by Salesforce for connecting applications, data, and devices."),
    ("Spring Boot", "Java framework for building production-ready applications with minimal configuration."),
    ("DataWeave", "MuleSoft's proprietary data transformation language (version 2.0)."),
    ("MEL", "Mule Expression Language — legacy expression syntax from Mule 3.x."),
    ("APIkit", "MuleSoft module for auto-generating API implementations from RAML/OAS specs."),
    ("LLM", "Large Language Model — used for code review and unknown element conversion."),
    ("AgentContext", "Shared state object tracking LLM-assisted conversions through the pipeline."),
    ("Connector", "MuleSoft module that connects to external systems (HTTP, DB, JMS, etc.)."),
    ("Flow", "MuleSoft's primary processing unit — a sequence of processors with an optional source."),
    ("Sub-flow", "Reusable flow fragment called via flow-ref, has no source."),
    ("Processor", "Individual step in a flow that transforms, routes, or processes messages."),
    ("Source", "Inbound endpoint that triggers a flow (HTTP listener, scheduler, message consumer)."),
    ("pom.xml", "Maven Project Object Model — defines dependencies, build plugins, and project metadata."),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
for i, h in enumerate(["Term", "Definition"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "2E4057")
    for r in t.rows[0].cells[i].paragraphs[0].runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
for term, defn in glossary:
    row = t.add_row()
    row.cells[0].text = term
    row.cells[1].text = defn
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

# ════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════
doc.save(OUTPUT_PATH)
print(f"Documentation saved to: {OUTPUT_PATH}")
print(f"Total sections: 16 chapters")
